"""
NTB BRIEFING SÁNG — Auto Generator
====================================
Chạy mỗi sáng: python ntb_briefing_sang.py
Output: NTB_Briefing_Sang_DDMMYYYY.docx

Nguồn data: file Excel Dashboard (download từ GHN hoặc Google Sheets)
Để đổi nguồn data → xem phần CONFIG bên dưới
"""

import sys
import os
import math
from datetime import datetime, timedelta
from pathlib import Path


import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════════════════════
# CONFIG — chỉnh ở đây
# ══════════════════════════════════════════════════════════════════════════════
CONFIG = {
    # Đường dẫn file Excel dashboard (để trống nếu dùng Google Sheets)
    "excel_file": "",

    # Google Sheets ID (lấy từ URL: /spreadsheets/d/ID/edit)
    # Để trống nếu dùng file Excel
    "gsheet_id": "1JZ1eRerRqrpwjZ4HBevQunjd8VquM_cvPFz12TaJfMQ",

    # File credentials Google Service Account (để dùng Google Sheets)
    "credentials_json": r"C:\Users\lap4all\Desktop\Backlog_Automation\credentials.json",

    # Thư mục output
    "output_dir": r"C:\Users\lap4all\Desktop\Backlog_Automation\output",

    # Ngưỡng cảnh báo
    "thresh": {
        "gtc_warn": 60,       # %GTC dưới mức này → cảnh báo
        "gtc_critical": 40,   # %GTC dưới mức này → nghiêm trọng
        "ca2_warn": 70,       # %Gán ca 2 dưới mức này → cảnh báo
        "ca2_critical": 30,   # %Gán ca 2 dưới mức này → nghiêm trọng
        "fd_warn": 0.08,      # %FD trên mức này → cảnh báo
        "fd_critical": 0.12,  # %FD trên mức này → nghiêm trọng
        "aging_warn": 5,      # số đơn aging tối thiểu để cảnh báo
        "treo_warn": 10,      # số đơn treo LC tối thiểu để cảnh báo
        "opr_warn": 30,       # % OPR trễ tối thiểu để cảnh báo
        "opr_min_vol": 5,     # volume tối thiểu OPR mới tính
    }
}


# ══════════════════════════════════════════════════════════════════════════════
# COLORS & STYLES
# ══════════════════════════════════════════════════════════════════════════════
GHN_ORANGE  = RGBColor(0xF2, 0x65, 0x22)
GHN_BLUE    = RGBColor(0x00, 0x72, 0xBC)
RED         = RGBColor(0xC0, 0x39, 0x2B)
RED_BG      = "FADBD8"
ORANGE_COL  = RGBColor(0xE6, 0x7E, 0x22)
ORANGE_BG   = "FDEBD0"
GREEN_COL   = RGBColor(0x1E, 0x84, 0x49)
GREEN_BG    = "D5F5E3"
GRAY        = RGBColor(0x5D, 0x6D, 0x7E)
GRAY_BG     = "F0F3F4"
WHITE       = "FFFFFF"
BLACK       = RGBColor(0x2C, 0x3E, 0x50)
BORDER_COL  = "BDC3C7"


# ══════════════════════════════════════════════════════════════════════════════
# DOCX HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_borders(cell, color=BORDER_COL, size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)


def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font="Arial"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_para(doc_or_cell, text="", bold=False, italic=False, size=11,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=4):
    if hasattr(doc_or_cell, 'paragraphs') and hasattr(doc_or_cell, 'add_paragraph'):
        para = doc_or_cell.add_paragraph()
    else:
        para = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
        para = doc_or_cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if text:
        add_run(para, text, bold=bold, italic=italic, size=size, color=color)
    return para


def add_hr(doc, color_hex="BDC3C7", space_before=4, space_after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


# ══════════════════════════════════════════════════════════════════════════════
# DATA LOADING
# ══════════════════════════════════════════════════════════════════════════════
def load_from_excel(filepath):
    print(f"  Đọc file: {filepath}")
    xl = pd.read_excel(filepath, sheet_name=None)
    return xl


def load_from_gsheet(sheet_id, credentials_json):
    print(f"  Kết nối Google Sheets: {sheet_id}")
    try:
        import gspread
        from google.oauth2.service_account import Credentials
        scopes = ['https://spreadsheets.google.com/feeds',
                  'https://www.googleapis.com/auth/drive']
        creds = Credentials.from_service_account_file(credentials_json, scopes=scopes)
        gc = gspread.authorize(creds)
        sh = gc.open_by_key(sheet_id)
        xl = {}
        for ws in sh.worksheets():
            data = ws.get_all_values()
            if not data:
                continue
            df = pd.DataFrame(data[1:], columns=data[0])
            xl[ws.title] = df
        return xl
    except Exception as e:
        print(f"  [WARN] Không kết nối được Google Sheets: {e}")
        return None


def load_data():
    xl = None
    if CONFIG["excel_file"] and Path(CONFIG["excel_file"]).exists():
        xl = load_from_excel(CONFIG["excel_file"])
    elif CONFIG["gsheet_id"] and Path(CONFIG["credentials_json"]).exists():
        xl = load_from_gsheet(CONFIG["gsheet_id"], CONFIG["credentials_json"])

    if xl is None:
        # Demo mode: tìm file Excel bất kỳ trong thư mục hiện tại
        excels = list(Path(".").glob("*.xlsx")) + list(Path(".").glob("Dash_Board*.xlsx"))
        if excels:
            xl = load_from_excel(str(excels[0]))
        else:
            raise FileNotFoundError(
                "Không tìm thấy data! Đặt file Excel vào cùng thư mục hoặc config lại credentials."
            )
    return xl


# ══════════════════════════════════════════════════════════════════════════════
# DATA ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
def parse_date_col(series):
    """Try to extract date from 'Time' column like '2026-06-23 - Thứ 3'"""
    def try_parse(v):
        if pd.isna(v):
            return pd.NaT
        s = str(v).strip()[:10]
        try:
            return pd.to_datetime(s)
        except Exception:
            return pd.NaT
    return series.apply(try_parse)


def get_n1_date(df_data):
    """Tìm ngày N-1: ngày gần nhất có data (hôm qua hoặc 1-2 ngày trước)"""
    if 'date' not in df_data.columns:
        return None
    today = pd.Timestamp.now().normalize()
    for delta in range(1, 5):
        candidate = today - timedelta(days=delta)
        if (df_data['date'] == candidate).any():
            return candidate
    # fallback: max date
    return df_data['date'].dropna().max()


def analyze(xl):
    T = CONFIG["thresh"]
    n1_label = "(N-1)"
    n1_date_str = ""

    # ── Data sheet (GTC/Gán) ──────────────────────────────────────────────
    bc_stats = {}  # bc -> {am, pct_gtc, pct_gan, ton, chua_gan, pct_gtc_ca1, pct_gan_ca1, pct_gtc_ca2, pct_gan_ca2, gtc_dod}
    bc_stats_n2 = {}

    if 'Data' in xl:
        df = xl['Data'].copy()
        # numeric
        for c in ['% GTC', '% Gán', 'Sản Lượng Giao Thành Công', 'Sản Lượng Gán',
                  'Sản Lượng Tồn', 'Sản Lượng Chưa Gán', 'Volume']:
            if c in df.columns:
                df[c] = pd.to_numeric(df[c], errors='coerce')

        if 'Time' in df.columns:
            df['date'] = parse_date_col(df['Time'])
            n1 = get_n1_date(df)
            if n1:
                n1_date_str = n1.strftime('%d/%m/%Y')
                n1_label = n1.strftime('%d/%m/%Y')
                n2 = n1 - timedelta(days=1)
                df_n1 = df[df['date'] == n1]
                df_n2 = df[df['date'] == n2]
            else:
                df_n1 = df
                df_n2 = pd.DataFrame()
        else:
            df_n1 = df
            df_n2 = pd.DataFrame()

        loai_col = 'Loại Hàng' if 'Loại Hàng' in df_n1.columns else None
        chi_tiet_col = 'Chi tiết' if 'Chi tiết' in df_n1.columns else None
        am_col = 'AM' if 'AM' in df_n1.columns else None

        if chi_tiet_col and am_col:
            for bc, grp in df_n1.groupby(chi_tiet_col):
                am = str(grp[am_col].iloc[0])
                vol = grp['Volume'].sum()
                gtc = grp['Sản Lượng Giao Thành Công'].sum()
                gan = grp['Sản Lượng Gán'].sum()
                ton = grp['Sản Lượng Tồn'].sum()
                chua_gan = grp['Sản Lượng Chưa Gán'].sum()
                pct_gtc = round(gtc / vol * 100, 1) if vol > 0 else None
                pct_gan = round(gan / vol * 100, 1) if vol > 0 else None

                ca1 = grp[grp[loai_col] == 'Hàng Mới Ca 1'] if loai_col else pd.DataFrame()
                ca2 = grp[grp[loai_col] == 'Hàng Mới Ca 2'] if loai_col else pd.DataFrame()

                bc_stats[bc] = {
                    'am': am, 'vol': vol, 'pct_gtc': pct_gtc, 'pct_gan': pct_gan,
                    'ton': round(ton), 'chua_gan': round(chua_gan),
                    'pct_gtc_ca1': round(ca1['% GTC'].iloc[0]*100, 1) if not ca1.empty and '% GTC' in ca1 else None,
                    'pct_gan_ca1': round(ca1['% Gán'].iloc[0]*100, 1) if not ca1.empty and '% Gán' in ca1 else None,
                    'pct_gtc_ca2': round(ca2['% GTC'].iloc[0]*100, 1) if not ca2.empty and '% GTC' in ca2 else None,
                    'pct_gan_ca2': round(ca2['% Gán'].iloc[0]*100, 1) if not ca2.empty and '% Gán' in ca2 else None,
                    'vol_ca2': round(ca2['Volume'].sum()) if not ca2.empty else 0,
                }

            if not df_n2.empty and chi_tiet_col in df_n2.columns:
                for bc, grp in df_n2.groupby(chi_tiet_col):
                    vol = grp['Volume'].sum()
                    gtc = grp['Sản Lượng Giao Thành Công'].sum()
                    bc_stats_n2[bc] = round(gtc / vol * 100, 1) if vol > 0 else None

    # ── FD sheet ─────────────────────────────────────────────────────────
    fd_data = {}  # bc -> {am, fd_n, fd_n1, fd_delta}
    fd_sheet = None
    for name in xl:
        if 'FD' in name.upper():
            fd_sheet = xl[name]
            break

    if fd_sheet is not None:
        df_fd = fd_sheet.copy()
        # Find header row
        for i, row in df_fd.iterrows():
            if str(row.iloc[0]).strip().startswith('(') or str(row.iloc[1]).strip() in ['AM', 'Trần', 'Nguyễn', 'Huỳnh', 'Lê', 'Phan', 'Phạm']:
                df_fd.columns = ['BC', 'AM', 'fd_n', 'fd_n1', 'fd_delta', 'fd_n7', 'fd_delta7', 'vol', 'fd_vol', 'ty_trong'] + list(df_fd.columns[10:])
                df_fd = df_fd.iloc[i:].reset_index(drop=True)
                break
        else:
            # Try iloc[2:]
            if len(df_fd) > 2:
                df_fd = df_fd.iloc[2:].copy()
                df_fd.columns = ['BC', 'AM', 'fd_n', 'fd_n1', 'fd_delta', 'fd_n7', 'fd_delta7', 'vol', 'fd_vol', 'ty_trong'] + list(range(len(df_fd.columns)-10))

        for col in ['fd_n', 'fd_n1', 'fd_delta']:
            if col in df_fd.columns:
                df_fd[col] = pd.to_numeric(df_fd[col], errors='coerce')

        if 'AM' in df_fd.columns and 'BC' in df_fd.columns:
            df_fd = df_fd.dropna(subset=['AM'])
            for _, row in df_fd.iterrows():
                bc = str(row.get('BC', '')).strip()
                am = str(row.get('AM', '')).strip()
                if not bc or not am or am == 'nan':
                    continue
                # Lọc bỏ các row rác (AM là số hoặc % hoặc quá ngắn)
                if am.replace('.', '').replace('%', '').replace(' ', '').isnumeric():
                    continue
                if len(am) < 4 or am in ['AM', 'Bưu Cục', 'TỔNG']:
                    continue
                # AM phải có dạng họ tên tiếng Việt (ít nhất 2 từ)
                if '%' in am or '(' in am or am[0].isdigit():
                    continue
                fd_n = row.get('fd_n')
                try:
                    fd_n = float(fd_n) if fd_n is not None else None
                except:
                    fd_n = None
                if fd_n is None:
                    continue
                fd_data[bc] = {
                    'am': am,
                    'fd_n': fd_n,
                    'fd_n1': row.get('fd_n1'),
                    'fd_delta': row.get('fd_delta'),
                }

    # ── Aging sheet ───────────────────────────────────────────────────────
    aging_data = {}  # bc -> {count, avg_days}
    for name in xl:
        if 'Aging' in name and '5' in name and 'PIVOT' not in name.upper():
            df_ag = xl[name].copy()
            bc_col = next((c for c in df_ag.columns if 'bc' in c.lower() or 'bưu cục' in c.lower()), None)
            if bc_col:
                for bc, grp in df_ag.groupby(bc_col):
                    ag_col = next((c for c in grp.columns if 'aging' in c.lower()), None)
                    aging_data[str(bc)] = {
                        'count': len(grp),
                        'avg_days': round(pd.to_numeric(grp[ag_col], errors='coerce').mean(), 1) if ag_col else 0,
                    }
            break

    # ── Treo LC sheet ─────────────────────────────────────────────────────
    treo_data = {}  # bc -> {count, count_168h}
    for name in xl:
        if 'Treo' in name and 'PV' not in name.upper() and 'PIVOT' not in name.upper():
            df_tr = xl[name].copy()
            wh_col = next((c for c in df_tr.columns if 'warehouse' in c.lower() or 'bưu cục' in c.lower() or 'name' in c.lower()), None)
            time_col = next((c for c in df_tr.columns if 'tồn đọng' in c.lower() or 'thời gian' in c.lower()), None)
            if wh_col:
                for bc, grp in df_tr.groupby(wh_col):
                    count_168 = 0
                    if time_col:
                        count_168 = grp[time_col].astype(str).str.contains('192|168', na=False).sum()
                    treo_data[str(bc)] = {'count': len(grp), 'count_168h': int(count_168)}
            break

    # ── OPR TTS sheet ─────────────────────────────────────────────────────
    opr_data = {}  # bc -> {am, total, tre, pct}
    for name in xl:
        if 'OPR' in name.upper() and 'rawopr' not in name.lower() and 'ODR' not in name.upper():
            df_opr = xl[name].copy()
            bc_col = next((c for c in df_opr.columns if 'kholay' in c.lower() or 'bưu cục' in c.lower()), None)
            am_col2 = next((c for c in df_opr.columns if c == 'AM'), None)
            vol_col = next((c for c in df_opr.columns if 'vol_ltc' in c.lower() or 'volume' in c.lower()), None)
            tre_col = next((c for c in df_opr.columns if 'trễ' in c.lower() or 'tre' in c.lower() or 'đơn trễ' in c.lower()), None)
            if bc_col and vol_col and tre_col:
                for col in [vol_col, tre_col]:
                    df_opr[col] = pd.to_numeric(df_opr[col], errors='coerce').fillna(0)
                for bc, grp in df_opr.groupby(bc_col):
                    total = grp[vol_col].sum()
                    tre = grp[tre_col].sum()
                    am = str(grp[am_col2].iloc[0]) if am_col2 else ''
                    opr_data[str(bc)] = {
                        'am': am, 'total': total, 'tre': tre,
                        'pct': round(tre / total * 100, 1) if total > 0 else 0,
                    }
            break

    # ── BUILD ISSUES PER AM ───────────────────────────────────────────────
    all_ams = set()
    for v in bc_stats.values(): all_ams.add(v['am'])
    for v in fd_data.values(): all_ams.add(v['am'])
    for v in opr_data.values():
        if v['am']: all_ams.add(v['am'])
    all_ams = {a for a in all_ams if a and a != 'nan'}

    results = {}
    scores = {}

    for am in sorted(all_ams):
        issues = []
        score = 100

        # GTC
        am_bcs = {bc: v for bc, v in bc_stats.items() if v['am'] == am}
        worst_gtc = sorted(
            [(bc, v) for bc, v in am_bcs.items() if v['pct_gtc'] is not None and v['pct_gtc'] < T['gtc_warn']],
            key=lambda x: x[1]['pct_gtc']
        )[:3]
        for bc, v in worst_gtc:
            sev = 'critical' if v['pct_gtc'] < T['gtc_critical'] else 'high'
            dod = None
            if bc in bc_stats_n2 and bc_stats_n2[bc] is not None:
                dod = round(v['pct_gtc'] - bc_stats_n2[bc], 1)
            dod_str = (f" | {'▲' if dod>0 else '▼'}{abs(dod)}% DoD") if dod is not None else ''
            issues.append({
                'type': 'GTC thấp', 'bc': bc, 'severity': sev,
                'detail': f"%GTC {v['pct_gtc']}%{dod_str} | %Gán {v['pct_gan']}% | Tồn {v['ton']} đơn | Chưa gán {v['chua_gan']}",
                'm': {'gtc': v['pct_gtc'], 'gan': v['pct_gan'], 'ton': v['ton'], 'chua_gan': v['chua_gan'], 'dod': dod},
            })
            score -= 20 if sev == 'critical' else 12

        # Gán Ca 2
        low_ca2 = sorted(
            [(bc, v) for bc, v in am_bcs.items() if v.get('pct_gan_ca2') is not None and v['pct_gan_ca2'] < T['ca2_warn']],
            key=lambda x: x[1]['pct_gan_ca2']
        )[:2]
        for bc, v in low_ca2:
            sev = 'critical' if v['pct_gan_ca2'] < T['ca2_critical'] else 'high'
            issues.append({
                'type': 'Gán Ca 2 thấp', 'bc': bc, 'severity': sev,
                'detail': f"%Gán Ca 2: {v['pct_gan_ca2']}% | %GTC Ca 2: {v.get('pct_gtc_ca2','?')}% | Volume ca 2: {v.get('vol_ca2','?')} đơn",
                'm': {'gan_ca2': v['pct_gan_ca2'], 'gtc_ca2': v.get('pct_gtc_ca2', 0)},
            })
            score -= 15 if sev == 'critical' else 8

        # FD
        am_fd = sorted(
            [(bc, v) for bc, v in fd_data.items() if v['am'] == am and v.get('fd_n') is not None and v['fd_n'] >= T['fd_warn']],
            key=lambda x: -x[1]['fd_n']
        )[:3]
        for bc, v in am_fd:
            sev = 'critical' if v['fd_n'] >= T['fd_critical'] else 'high'
            delta = v.get('fd_delta') or 0
            dir_str = f"{'tăng' if delta > 0 else 'giảm'} {abs(round(delta*100,1))}% so N-1"
            issues.append({
                'type': 'FD cao', 'bc': bc, 'severity': sev,
                'detail': f"%FD: {round(v['fd_n']*100,1)}% ({dir_str})",
                'm': {'fd': round(v['fd_n']*100,1), 'fd_n1': round((v.get('fd_n1') or 0)*100,1), 'delta': round(delta*100,1)},
            })
            score -= 15 if sev == 'critical' else 10

        # Aging
        am_bc_names = set(am_bcs.keys())
        am_aging = [(bc, v) for bc, v in aging_data.items() if bc in am_bc_names and v['count'] >= T['aging_warn']]
        am_aging.sort(key=lambda x: -x[1]['count'])
        for bc, v in am_aging[:3]:
            sev = 'high' if v['count'] >= 50 else 'medium'
            issues.append({
                'type': 'Aging >5 ngày', 'bc': bc, 'severity': sev,
                'detail': f"{v['count']} đơn | Aging TB: {v['avg_days']} ngày",
                'm': {'aging_count': v['count'], 'avg_days': v['avg_days']},
            })
            score -= 8 if sev == 'high' else 3

        # Treo LC
        am_treo = [(bc, v) for bc, v in treo_data.items() if bc in am_bc_names and v['count'] >= T['treo_warn']]
        am_treo.sort(key=lambda x: -x[1]['count'])
        for bc, v in am_treo[:2]:
            sev = 'high' if v['count'] >= 30 or v['count_168h'] > 5 else 'medium'
            extra = f" | ≥168h: {v['count_168h']} đơn" if v['count_168h'] > 0 else ''
            issues.append({
                'type': 'Treo LC', 'bc': bc, 'severity': sev,
                'detail': f"{v['count']} đơn treo LC{extra}",
                'm': {'treo_count': v['count'], 'treo_168h': v['count_168h']},
            })
            score -= 8 if sev == 'high' else 4

        # OPR TTS
        am_opr = sorted(
            [(bc, v) for bc, v in opr_data.items() if v['am'] == am and v['pct'] >= T['opr_warn'] and v['total'] >= T['opr_min_vol']],
            key=lambda x: -x[1]['pct']
        )[:2]
        for bc, v in am_opr:
            sev = 'high' if v['pct'] >= 60 else 'medium'
            issues.append({
                'type': 'OPR TTS', 'bc': bc, 'severity': sev,
                'detail': f"Tỷ lệ trễ {v['pct']}% | {int(v['tre'])}/{int(v['total'])} đơn bị trễ",
                'm': {'opr_tre': v['pct'], 'tre': int(v['tre']), 'total': int(v['total'])},
            })
            score -= 10 if sev == 'high' else 5

        results[am] = issues
        scores[am] = {
            'score': max(0, score),
            'n_critical': sum(1 for i in issues if i['severity'] == 'critical'),
            'n_high': sum(1 for i in issues if i['severity'] == 'high'),
            'n_total': len(issues),
        }

    return results, scores, n1_label


# ══════════════════════════════════════════════════════════════════════════════
# QUESTION GENERATOR
# ══════════════════════════════════════════════════════════════════════════════
def make_question(issue, am):
    ln = am.split()[-1]
    bc = issue['bc']
    m = issue.get('m', {})
    t = issue['type']

    if t == 'GTC thấp':
        dod = m.get('dod')
        dod_str = ''
        if dod is not None:
            dod_str = f", {'tăng ▲' if dod > 0 else 'giảm ▼'}{abs(dod)}% so hôm kia"
        return (f"{ln} ơi, {bc} hôm qua %GTC chỉ đạt {m.get('gtc')}%{dod_str} — "
                f"tồn {m.get('ton')} đơn, còn {m.get('chua_gan')} đơn chưa gán. "
                f"{ln} báo cáo em nguyên nhân cụ thể và kế hoạch clear tồn hôm nay "
                f"nhé, ưu tiên gán hết ca 1 trước 10h.")

    if t == 'Gán Ca 2 thấp':
        return (f"{ln} ơi, ca 2 hôm qua tại {bc} gán chỉ {m.get('gan_ca2')}%, "
                f"GTC ca 2 chỉ {m.get('gtc_ca2')}%. "
                f"Lý do là gì — shipper nghỉ, hàng về trễ hay nguyên nhân khác? "
                f"Hôm nay {ln} có kế hoạch bổ sung nhân lực ca 2 không?")

    if t == 'FD cao':
        delta = m.get('delta', 0) or 0
        dir_str = f"đang tăng thêm {abs(delta)}%" if delta > 0 else f"đã giảm {abs(delta)}% nhưng vẫn cao"
        return (f"{ln} ơi, %FD tại {bc} đang ở {m.get('fd')}% — {dir_str} so N-1. "
                f"{ln} cho biết top 3 lý do hoàn trả hôm qua và phương án xử lý "
                f"để kéo FD về dưới 5% trong tuần này nhé.")

    if t == 'Aging >5 ngày':
        return (f"{ln} ơi, {bc} đang có {m.get('aging_count')} đơn aging trên 5 ngày "
                f"(TB {m.get('avg_days')} ngày/đơn). "
                f"{ln} rà soát lại và cho biết đơn nào xử lý được hôm nay, "
                f"đơn nào cần support từ vùng?")

    if t == 'Treo LC':
        extra = f" — đặc biệt có {m.get('treo_168h')} đơn treo ≥168h, cần xử lý gấp!" if m.get('treo_168h', 0) > 0 else "."
        return (f"{ln} ơi, {bc} đang treo {m.get('treo_count')} đơn luân chuyển{extra} "
                f"{ln} check nguyên nhân và push xử lý hôm nay nhé.")

    if t == 'OPR TTS':
        return (f"{ln} ơi, OPR lấy hàng TTS tại {bc} hôm qua tỷ lệ trễ {m.get('opr_tre')}% "
                f"({m.get('tre')}/{m.get('total')} đơn bị trễ). "
                f"{ln} trao đổi lại với shipper phụ trách — nguyên nhân trễ và "
                f"lịch lấy hàng hôm nay như thế nào?")

    return f"{ln} ơi, {bc}: {issue.get('detail', '')}. {ln} cho em biết nguyên nhân và kế hoạch nhé."


# ══════════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT GENERATOR
# ══════════════════════════════════════════════════════════════════════════════
SEV_META = {
    'critical': {'label': 'Nghiêm trọng', 'bg': RED_BG, 'color': RED},
    'high':     {'label': 'Cao',           'bg': ORANGE_BG, 'color': ORANGE_COL},
    'medium':   {'label': 'Trung bình',    'bg': GRAY_BG, 'color': GRAY},
}
TYPE_ICON = {
    'GTC thấp': '[GTC]', 'Gán Ca 2 thấp': '[CA2]', 'FD cao': '[FD]',
    'Aging >5 ngày': '[AGE]', 'Treo LC': '[LC]', 'OPR TTS': '[OPR]',
}


def score_grade(s):
    return ('Ổn định', GREEN_COL, GREEN_BG) if s >= 70 else \
           ('Cần theo dõi', ORANGE_COL, ORANGE_BG) if s >= 40 else \
           ('Báo động', RED, RED_BG)


def build_cover(doc, results, scores, n1_label, today_str):
    sorted_ams = sorted(scores.items(), key=lambda x: x[1]['score'])
    n_crit = sum(1 for _, s in scores.items() if s['score'] < 30)
    n_warn = sum(1 for _, s in scores.items() if 30 <= s['score'] < 60)
    n_ok   = sum(1 for _, s in scores.items() if s['score'] >= 60)
    n_total= sum(s['n_total'] for s in scores.values())

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('BRIEFING SÁNG — GIẢI TRÌNH VẬN HÀNH NTB')
    run.bold = True; run.font.size = Pt(18); run.font.color.rgb = GHN_ORANGE; run.font.name = 'Arial'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(f'Dữ liệu N-1: {n1_label}   |   Ngày lập: {today_str}   |   NỘI BỘ')
    r2.font.size = Pt(10); r2.font.color.rgb = GRAY; r2.font.name = 'Arial'

    add_hr(doc, 'F26522', space_before=2, space_after=6)

    # Summary stats table
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(3.5); tbl.columns[1].width = Cm(3.5)
    tbl.columns[2].width = Cm(3.5); tbl.columns[3].width = Cm(3.5)
    stats = [
        (str(n_crit), 'AM Báo Động', RED_BG, RED),
        (str(n_warn), 'AM Coi Chừng', ORANGE_BG, ORANGE_COL),
        (str(n_ok),   'AM Ổn Định',  GREEN_BG, GREEN_COL),
        (str(n_total),'Tổng Issues', 'EBF5FB', GHN_BLUE),
    ]
    for i, (num, lbl, bg, col) in enumerate(stats):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, bg)
        set_cell_borders(cell)
        cell_margins(cell, top=120, bottom=120)
        cell.paragraphs[0].clear()
        p_num = cell.paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_num.paragraph_format.space_after = Pt(2)
        r = p_num.add_run(num); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = col; r.font.name = 'Arial'
        p_lbl = cell.add_paragraph()
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lbl.paragraph_format.space_before = Pt(0); p_lbl.paragraph_format.space_after = Pt(0)
        rl = p_lbl.add_run(lbl); rl.font.size = Pt(9); rl.font.color.rgb = col; rl.font.name = 'Arial'

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Ranking table
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4); p_title.paragraph_format.space_after = Pt(4)
    rt = p_title.add_run('BẢNG XẾP HẠNG HIỆU SUẤT AM')
    rt.bold = True; rt.font.size = Pt(12); rt.font.color.rgb = GHN_BLUE; rt.font.name = 'Arial'

    rank_tbl = doc.add_table(rows=1, cols=4)
    rank_tbl.style = 'Table Grid'
    rank_tbl.columns[0].width = Cm(1.2); rank_tbl.columns[1].width = Cm(6)
    rank_tbl.columns[2].width = Cm(2.5); rank_tbl.columns[3].width = Cm(4.5)
    headers = ['#', 'AM', 'Điểm', 'Trạng thái']
    for i, h in enumerate(headers):
        c = rank_tbl.rows[0].cells[i]
        set_cell_bg(c, '0072BC')
        set_cell_borders(c, color='0072BC')
        cell_margins(c, top=60, bottom=60)
        c.paragraphs[0].clear()
        cp = c.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(0)
        cr = cp.add_run(h); cr.bold = True; cr.font.size = Pt(10); cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); cr.font.name = 'Arial'

    rank_marks = ['🔴', '🟠', '🟡']
    for i, (am, sc) in enumerate(sorted_ams):
        grade, col, bg = score_grade(sc['score'])
        row = rank_tbl.add_row()
        data = [
            (rank_marks[i] if i < 3 else str(i+1), WD_ALIGN_PARAGRAPH.CENTER, bg, col),
            (am, WD_ALIGN_PARAGRAPH.LEFT, WHITE, BLACK),
            (f"{sc['score']}/100", WD_ALIGN_PARAGRAPH.CENTER, bg, col),
            (f"{grade}  |  {sc['n_total']} issues", WD_ALIGN_PARAGRAPH.CENTER, bg, col),
        ]
        for j, (txt, align, bg_, col_) in enumerate(data):
            c = row.cells[j]
            set_cell_bg(c, bg_)
            set_cell_borders(c)
            cell_margins(c, top=50, bottom=50)
            c.paragraphs[0].clear()
            cp = c.paragraphs[0]
            cp.alignment = align
            cp.paragraph_format.space_after = Pt(0)
            cr = cp.add_run(txt)
            cr.font.size = Pt(10); cr.font.name = 'Arial'
            if col_: cr.font.color.rgb = col_
            if j == 2: cr.bold = True

    page_break(doc)


def build_am_page(doc, am, issues, score_info, n1_label, today_str, is_last):
    sc = score_info['score']
    grade, score_col, score_bg = score_grade(sc)

    # Header table: name | score
    hdr = doc.add_table(rows=1, cols=2)
    hdr.style = 'Table Grid'
    hdr.columns[0].width = Cm(11); hdr.columns[1].width = Cm(3.5)

    c0 = hdr.rows[0].cells[0]
    set_cell_bg(c0, '0072BC')
    set_cell_borders(c0, color='0072BC')
    cell_margins(c0, top=120, bottom=120, left=160)
    c0.paragraphs[0].clear()
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run(am); r0.bold = True; r0.font.size = Pt(16); r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r0.font.name = 'Arial'
    p0b = c0.add_paragraph()
    p0b.paragraph_format.space_before = Pt(0); p0b.paragraph_format.space_after = Pt(0)
    r0b = p0b.add_run(f"{score_info['n_total']} vấn đề cần giải trình  |  N-1: {n1_label}")
    r0b.font.size = Pt(9); r0b.font.color.rgb = RGBColor(0xCC, 0xE5, 0xFF); r0b.font.name = 'Arial'

    c1 = hdr.rows[0].cells[1]
    set_cell_bg(c1, score_bg)
    set_cell_borders(c1)
    cell_margins(c1, top=80, bottom=80, left=80, right=80)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c1.paragraphs[0].clear()
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(f"{sc}/100"); r1.bold = True; r1.font.size = Pt(18); r1.font.color.rgb = score_col; r1.font.name = 'Arial'
    p1b = c1.add_paragraph()
    p1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1b.paragraph_format.space_before = Pt(0); p1b.paragraph_format.space_after = Pt(0)
    r1b = p1b.add_run(grade); r1b.font.size = Pt(9); r1b.font.color.rgb = score_col; r1b.font.name = 'Arial'

    # Issue sections
    sections = [
        ('critical', '[!!] NGHIÊM TRỌNG', RED),
        ('high',     '[!]  CAN CHU Y',    ORANGE_COL),
        ('medium',   '[.] THEO DOI THEM', GRAY),
    ]
    for sev, sec_title, sec_col in sections:
        grp = [i for i in issues if i['severity'] == sev]
        if not grp:
            continue

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(8); sp.paragraph_format.space_after = Pt(3)
        sr = sp.add_run(f"{sec_title} — {len(grp)} van de")
        sr.bold = True; sr.font.size = Pt(11); sr.font.color.rgb = sec_col; sr.font.name = 'Arial'

        for issue in grp:
            meta = SEV_META[issue['severity']]
            q = make_question(issue, am)

            itbl = doc.add_table(rows=1, cols=2)
            itbl.style = 'Table Grid'
            itbl.columns[0].width = Cm(2); itbl.columns[1].width = Cm(12.5)

            # Left stripe
            lc = itbl.rows[0].cells[0]
            set_cell_bg(lc, meta['bg'])
            set_cell_borders(lc)
            cell_margins(lc, top=80, bottom=80, left=80, right=60)
            lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            lc.paragraphs[0].clear()
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lp.paragraph_format.space_after = Pt(2)
            lr = lp.add_run(TYPE_ICON.get(issue['type'], '[?]'))
            lr.bold = True; lr.font.size = Pt(9); lr.font.color.rgb = meta['color']; lr.font.name = 'Arial'
            lp2 = lc.add_paragraph()
            lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lp2.paragraph_format.space_before = Pt(0); lp2.paragraph_format.space_after = Pt(0)
            lr2 = lp2.add_run(meta['label']); lr2.font.size = Pt(7); lr2.font.color.rgb = meta['color']; lr2.font.name = 'Arial'

            # Right content
            rc = itbl.rows[0].cells[1]
            set_cell_bg(rc, WHITE)
            set_cell_borders(rc)
            cell_margins(rc, top=80, bottom=80, left=140, right=100)

            rc.paragraphs[0].clear()
            rp1 = rc.paragraphs[0]
            rp1.paragraph_format.space_after = Pt(2)
            r_type = rp1.add_run(f"[{issue['type']}]  ")
            r_type.bold = True; r_type.font.size = Pt(10); r_type.font.color.rgb = meta['color']; r_type.font.name = 'Arial'
            r_bc = rp1.add_run(issue['bc'])
            r_bc.bold = True; r_bc.font.size = Pt(10); r_bc.font.color.rgb = BLACK; r_bc.font.name = 'Arial'

            rp2 = rc.add_paragraph()
            rp2.paragraph_format.space_before = Pt(0); rp2.paragraph_format.space_after = Pt(4)
            r_det = rp2.add_run(issue['detail'])
            r_det.italic = True; r_det.font.size = Pt(9); r_det.font.color.rgb = GRAY; r_det.font.name = 'Arial'

            rp3 = rc.add_paragraph()
            rp3.paragraph_format.space_before = Pt(0); rp3.paragraph_format.space_after = Pt(0)
            r_qlbl = rp3.add_run('Cau hoi: ')
            r_qlbl.bold = True; r_qlbl.font.size = Pt(10); r_qlbl.font.color.rgb = GHN_BLUE; r_qlbl.font.name = 'Arial'
            r_q = rp3.add_run(q)
            r_q.font.size = Pt(10); r_q.font.color.rgb = BLACK; r_q.font.name = 'Arial'

            sp_after = doc.add_paragraph()
            sp_after.paragraph_format.space_after = Pt(3)

    # Response box
    sp3 = doc.add_paragraph()
    sp3.paragraph_format.space_before = Pt(6); sp3.paragraph_format.space_after = Pt(3)
    rb_tbl = doc.add_table(rows=1, cols=1)
    rb_tbl.style = 'Table Grid'
    rb_tbl.columns[0].width = Cm(14.5)
    rb_c = rb_tbl.rows[0].cells[0]
    set_cell_bg(rb_c, 'F0F3F4')
    set_cell_borders(rb_c)
    cell_margins(rb_c, top=100, bottom=100, left=160, right=120)
    rb_c.paragraphs[0].clear()
    rb_p = rb_c.paragraphs[0]
    rb_p.paragraph_format.space_after = Pt(6)
    rb_r = rb_p.add_run('Phan hoi cua AM:')
    rb_r.bold = True; rb_r.font.size = Pt(10); rb_r.font.color.rgb = GHN_BLUE; rb_r.font.name = 'Arial'
    for _ in range(3):
        line_p = rb_c.add_paragraph()
        line_p.paragraph_format.space_before = Pt(0)
        line_p.paragraph_format.space_after = Pt(10)
        pPr = line_p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4'); bot.set(qn('w:color'), BORDER_COL)
        pBdr.append(bot); pPr.append(pBdr)
        line_p.add_run(' ' * 80)

    if not is_last:
        page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
def main():
    today = datetime.now()
    today_str = today.strftime('%d/%m/%Y')
    file_date = today.strftime('%d%m%Y')

    print("=" * 55)
    print("  NTB BRIEFING SANG — Auto Generator")
    print("=" * 55)

    # 1. Load data
    print("\n[1/3] Dang load data...")
    xl = load_data()
    print(f"  Sheets tim thay: {list(xl.keys())}")

    # 2. Analyze
    print("\n[2/3] Dang phan tich chi so...")
    results, scores, n1_label = analyze(xl)
    sorted_ams = sorted(scores.items(), key=lambda x: x[1]['score'])
    print(f"  Phan tich xong {len(results)} AM")
    print(f"  Ngay N-1: {n1_label}")
    crit_ams = [(a, s) for a, s in sorted_ams if s['score'] < 30]
    if crit_ams:
        print(f"  AM bao dong: {', '.join(a for a, _ in crit_ams)}")

    # 3. Build Word
    print("\n[3/3] Dang tao file Word...")
    doc = Document()

    # Page setup: A4
    from docx.oxml import OxmlElement
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Default font
    from docx.oxml.ns import qn
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    # Cover
    build_cover(doc, results, scores, n1_label, today_str)
    print("  Trang bia xong")

    # AM pages
    for i, (am, sc) in enumerate(sorted_ams):
        issues = results.get(am, [])
        is_last = (i == len(sorted_ams) - 1)
        build_am_page(doc, am, issues, sc, n1_label, today_str, is_last)
        print(f"  {am}: {sc['n_total']} issues, diem {sc['score']}")

    # Save
    out_dir = Path(CONFIG["output_dir"])
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"NTB_Briefing_Sang_{file_date}.docx"

    # Fallback: save next to script if output_dir doesn't exist
    if not out_dir.exists():
        out_path = Path(f"NTB_Briefing_Sang_{file_date}.docx")

    doc.save(str(out_path))
    print(f"\nDone! File da luu tai:\n  {out_path}")
    print("=" * 55)

    # Auto-open on Windows
    if sys.platform == "win32":
        os.startfile(str(out_path))


if __name__ == "__main__":
    main()
