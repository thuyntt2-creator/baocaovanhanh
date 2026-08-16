import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_2.docx"

if not os.path.exists(doc_path):
    print("File không tồn tại")
    sys.exit(1)

doc = docx.Document(doc_path)

print("=== SO SÁNH KIỂM TRA SỐ HỌC FILE WORD V8_2 ===")

# 1. Kiểm tra Bảng 11 (Chi tiết bưu cục)
table_11 = doc.tables[10]
errors = 0

print("\n1. Kiểm tra Bảng 11 (Chi tiết bưu cục) vs tổng thực tế:")
for c_idx in range(1, 7):
    nt = float(table_11.rows[1].cells[c_idx].text.replace(',', ''))
    dl = float(table_11.rows[2].cells[c_idx].text.replace(',', ''))
    dd = float(table_11.rows[3].cells[c_idx].text.replace(',', ''))
    dul = float(table_11.rows[4].cells[c_idx].text.replace(',', ''))
    
    total_in_doc = float(table_11.rows[5].cells[c_idx].text.replace(',', ''))
    actual_sum = nt + dl + dd + dul
    
    if abs(total_in_doc - actual_sum) > 0.1:
        print(f"  [LỖI] Tháng {c_idx+6}: Tổng dòng 5 = {total_in_doc} | Tổng cộng thực tế = {actual_sum:.1f} (LỆCH!)")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Tổng dòng 5 = {total_in_doc} | Tổng cộng thực tế = {actual_sum:.1f} (Khớp)")

# 2. Kiểm tra Bảng 10 (Hạng mục chi phí)
table_10 = doc.tables[9]
print("\n2. Kiểm tra Bảng 10 (Hạng mục chi phí) vs tổng thực tế:")
for c_idx in range(1, 7):
    xe = float(table_10.rows[1].cells[c_idx].text.replace(',', ''))
    giao = float(table_10.rows[2].cells[c_idx].text.replace(',', ''))
    kho = float(table_10.rows[3].cells[c_idx].text.replace(',', ''))
    mb = float(table_10.rows[4].cells[c_idx].text.replace(',', ''))
    
    total_in_doc = float(table_10.rows[5].cells[c_idx].text.replace(',', ''))
    actual_sum = xe + giao + kho + mb
    
    if abs(total_in_doc - actual_sum) > 0.1:
        print(f"  [LỖI] Tháng {c_idx+6}: Tổng dòng 5 = {total_in_doc} | Tổng cộng thực tế = {actual_sum:.1f} (LỆCH!)")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Tổng dòng 5 = {total_in_doc} | Tổng cộng thực tế = {actual_sum:.1f} (Khớp)")

# 3. So sánh Bảng 10 vs Bảng 11
print("\n3. Kiểm tra đối chiếu Bảng 10 vs Bảng 11:")
for c_idx in range(1, 7):
    t10_sum = float(table_10.rows[5].cells[c_idx].text.replace(',', ''))
    t11_sum = float(table_11.rows[5].cells[c_idx].text.replace(',', ''))
    
    if abs(t10_sum - t11_sum) > 0.1:
        print(f"  [LỖI] Tháng {c_idx+6}: Bảng 10 Tổng = {t10_sum} | Bảng 11 Tổng = {t11_sum} (LỆCH!)")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Bảng 10 Tổng = {t10_sum} | Bảng 11 Tổng = {t11_sum} (Khớp)")

# 4. Kiểm tra Bảng 7 (Định biên nhân sự)
table_7 = doc.tables[6]
print("\n4. Kiểm tra Bảng 7 (Định biên nhân sự):")
for c_idx in range(1, 7):
    # Trích xuất phần số trước dấu mở ngoặc
    nt_g = int(table_7.rows[2].cells[c_idx].text.split(' ')[0])
    dl_g = int(table_7.rows[3].cells[c_idx].text.split(' ')[0])
    dd_g = int(table_7.rows[4].cells[c_idx].text.split(' ')[0])
    dul_g = int(table_7.rows[5].cells[c_idx].text.split(' ')[0])
    
    total_giao = int(table_7.rows[6].cells[c_idx].text)
    actual_sum_giao = nt_g + dl_g + dd_g + dul_g
    
    if total_giao != actual_sum_giao:
        print(f"  [LỖI] Tháng {c_idx+6}: Dòng tổng NV giao = {total_giao} | Tổng chi tiết = {actual_sum_giao} (LỆCH!)")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Dòng tổng NV giao = {total_giao} | Tổng chi tiết = {actual_sum_giao} (Khớp)")
        
    # Kiểm tra Tổng nhân sự
    total_staff = int(table_7.rows[9].cells[c_idx].text)
    kho_ql = int(table_7.rows[8].cells[c_idx].text)
    actual_sum_staff = total_giao + kho_ql
    
    if total_staff != actual_sum_staff:
        print(f"  [LỖI] Tháng {c_idx+6}: Dòng tổng nhân sự = {total_staff} | Tổng giao + kho + QL = {actual_sum_staff} (LỆCH!)")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Dòng tổng nhân sự = {total_staff} | Tổng giao + kho + QL = {actual_sum_staff} (Khớp)")

if errors == 0:
    print("\n>>> TẤT CẢ KIỂM TRA ĐỀU THÀNH CÔNG! KHÔNG PHÁT HIỆN SAI LỆCH SỐ HỌC <<<")
else:
    print(f"\n>>> CÓ {errors} LỖI LỆCH SỐ LIỆU CẦN KHẮC PHỤC <<<")
