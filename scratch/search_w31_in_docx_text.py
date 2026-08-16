import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

docx_path = r'C:\Users\lap4all\Downloads\NTB - Báo Cáo Tuần (HRBP - ARD) - Tuần 32.docx'
doc = docx.Document(docx_path)

print("=== SEARCHING PARAGRAPHS FOR W31 / Tuần 31 / Bản đồ ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if '31' in txt or 'bản đồ' in txt.lower() or 'map' in txt.lower():
        print(f"P{i:3d}: {txt}")

print("\n=== SEARCHING TABLES FOR W31 IN CAPTIONS OR CELL TEXT ===")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell_txt = cell.text.strip().replace('\n', ' ')
            if '31' in cell_txt or 'bản đồ' in cell_txt.lower() or 'map' in cell_txt.lower():
                print(f"Table {t_idx:2d} ({r_idx},{c_idx}): {cell_txt[:120]}")
