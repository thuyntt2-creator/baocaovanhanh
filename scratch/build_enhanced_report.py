import sys, docx, pandas as pd, json, os
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

sys.stdout.reconfigure(encoding='utf-8')

# File paths
excel_path = r'C:\Users\lap4all\Downloads\NTB_Phan_Tuyen_Hanh_Chinh_Quy_Hoach_Moi.xlsx'
df = pd.read_excel(excel_path, sheet_name='Sheet1')

doc2 = docx.Document(r'C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi.docx')
maps_dir = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\d95894ad-9248-47c3-872b-72d7de66bc20\maps'

am_reasons = {
    24611: '''Lịch sử vận hành: Khu vực Gia Nghĩa từng bể tuyến triền miên nhiều tháng do không tuyển được nhân sự. Phương án tách 3 kho (Bắc Gia Nghĩa, Nam Gia Nghĩa 2, Nhân Cơ) vừa mới ổn định được vận hành và thu hút lực lượng lao động tại chỗ.
Xã Đắk Ha: Nhân sự hiện tại là người xã Quảng Sơn, ưu điểm gần kho Quảng Sơn. Nếu chuyển về kho Gia Nghĩa quản lý sẽ phát sinh di chuyển 25-30km/ngày, rủi ro nhân sự nghỉ việc 100%.
Đặc thù dân cư: Khu vực có tỷ lệ đồng bào dân tộc thiểu số cao. Việc tuyển nhân sự tại chỗ đã khó, nếu đổi BC mới càng không thể tuyển dụng.
Rủi ro nhân sự: Nếu sáp nhập BC theo xã mới, 100% nhân sự hiện tại các điểm xã ven đều xác nhận sẽ KHÔNG di chuyển theo bưu cục mới.
Đề xuất AM: TẠM THỜI GIỮ NGUYÊN 100% quản lý theo các BC cũ trong 3-6 tháng để tránh vỡ vận hành.''',

    24615: '''Lịch sử vận hành: Giữ nguyên phân vùng giao của BC Nam Gia Nghĩa 2 và Bắc Gia Nghĩa để tránh nguy cơ gãy tuyến và biến động nhân sự. Nhân sự hiện tại tại các xã cũ quen thuộc địa bàn và kho hiện tại.
Đề xuất AM: TẠM THỜI GIỮ NGUYÊN 100% Bưu cục và phạm vi quản lý theo ranh giới Xã/Phường cũ trong 3-6 tháng.''',

    24781: '''Bưu cục (LDO) Xuân Hương - Đà Lạt hiện phụ trách sản lượng lớn (2,292 đơn/ngày). Địa hình Đà Lạt bị chia cắt bởi Hồ Xuân Hương, thời tiết cuối năm mưa và lạnh, thiếu nhân sự. Việc dùng 1 bưu cục sẽ gây quá tải trầm trọng.
Đề xuất AM: Tách thêm 01 bưu cục mới là Bưu cục (LDO) Xuân Hương - Đà Lạt 2 (đặt tại Phường 10, cover Phường 3 & Phường 10, sản lượng 1,050 đơn giao, 150 đơn lấy, 8 NVPTTT + 1 NVXL) và giữ BC Xuân Hương - Đà Lạt cũ (cover Phường 1 & Phường 2, sản lượng 1,800 đơn giao, 400 đơn lấy, 13 NVPTTT + 1 NVXL).''',

    24784: '''Phường Lâm Viên (Đà Lạt mới) có đặc thù địa hình bị chia cắt bởi Hồ Xuân Hương. Giữ nguyên 02 bưu cục (LDO) Lâm Viên - Đà Lạt 1 và (LDO) Lâm Viên - Đà Lạt 2 nằm ngay trong khu vực Phường 8 giúp nhân viên chỉ tập trung xử lý một phường, giảm thời gian di chuyển, đảm bảo hiệu quả giao hàng.''',

    24958: '''Trong địa bàn Đức Trọng có 2 bưu cục (LDO Đức Trọng 1 & LDO Đức Trọng 2). Nếu gộp lại 1 bưu cục, tuyến giao của nhân viên xa và dẫn đến chồng chéo tuyến. Khoảng cách giữa 2 BC trên 15km.
Lý do giữ 2 BC: Địa bàn rộng, khó tuyển dụng (đặc biệt mùa cà phê và mùa mưa cuối năm), khoảng cách di chuyển tới địa điểm giao của nhân viên xa, BC cũ không đủ không gian m² để chứa hàng, vị trí nhà tới BC xa gây ảnh hưởng xuất phát giao trễ.
Phương án: Đề xuất gộp tuyến giao theo xã mới về (LDO) Đức Trọng 2 chiếm 79.7% sản lượng (901 đơn/ngày), (LDO) Đức Trọng 1 giữ nhiệm vụ phụ trách các xã lân cận.''',

    22972: '''Vị trí đặt 2 Bưu cục tại 2 thị trấn có sản lượng hàng nhiều nhất. Khoảng cách giữa 2 thị trấn tầm 25km và địa bàn rộng nên không thể gộp lại thành 1 Bưu cục (Phan Rí Cửa và khu vực lân cận). Giữ phân vùng vận hành hiện tại.''',

    23143: '''BTH - Tân Hải phụ trách Tân Thuận, Tân Thành cũ. Đề xuất giữ nguyên phần Thuận Quý cũ cho BC BTH - Hàm Thuận Nam vì địa hình khu vực đặc thù, khoảng cách từ BC Hàm Thuận Nam đến địa điểm gần hơn, dân cư tập trung đông đúc hơn so với khoảng cách từ BC Tân Hải.''',

    23131: '''BTH - Phước Hội (mới) cover Phường Lagi, Phước Hội, Sơn Mỹ (1,200-1,500 đơn giao, 280 đơn lấy, 15 NVPTTT). BTH - Tân Hải (mới) cover Xã Tân Hải, Xã Tân Thành (400-500 đơn giao, 120 đơn lấy, 7 NVPTTT). Phân định rõ ranh giới Lagi và Tân Hải để tối ưu bán kính di chuyển.''',

    22528: '''Huyện Ninh Hòa (cũ) có diện tích lớn nhất cả nước (20 xã/phường). Bưu cục Ninh Hòa 1 được tách ra để cover các tuyến xã xa trung tâm, đặt tại các tuyến xã thường xuyên thiếu hụt nhân sự để lôi kéo nguồn lực nhân sự tại chỗ, tiện giao lấy. Chia nhỏ cụm xã nóng vừa tận dụng được nguồn lực mới vừa đảm bảo thiếu hụt nhân sự không ảnh hưởng vận hành. Diện tích BC Ninh Hòa 2 cũng không đảm bảo m² để gom gộp.''',

    25000: '''BC (LDO) Hòa Ninh hiện cover tuyến Liên Đầm của BC (LDO) Di Linh. Sau sáp nhập Liên Đầm gộp về xã Di Linh, tuy nhiên đặc thù khoảng cách xa với BC hiện tại, nhân sự chưa ổn định. Đề xuất tách Bưu cục Hàng Nhỏ / Hàng Vừa (Xã Đinh Trang Thượng) để cover Đinh Trang Thượng, Di Linh, Phúc Thọ, Liên Đầm, giảm tải cho BC Di Linh (sản lượng 1,300-1,600 đơn/ngày, bán kính xa 22-45km).''',

    22366: '''Sản lượng Phường Nha Trang rất lớn (1,592 đơn/ngày). Đề xuất giữ nguyên/gộp phân vùng về BC KHO Nha Trang chính, đồng thời giữ các bưu cục Nam Nha Trang 1 & 2 phụ trách phân đoạn phụ để tránh quá tải kho bãi và vỡ tuyến.''',

    22402: '''Sản lượng Phường Nam Nha Trang cực lớn (2,418 đơn/ngày). Giữ Bưu cục Nam Nha Trang 1 mới (gộp Nam Nha Trang 3 và phần Nam Nha Trang 1 cũ) và giữ nguyên Bưu cục Nam Nha Trang 5 (phụ trách Phước Đồng). Đóng/bỏ bưu cục Nam Nha Trang 2 & Nam Nha Trang 3 để tối ưu điểm tập kết.''',

    24748: '''Địa hình Xã Quảng Tân (gồm Xã Đắk Ngo & Xã Quảng Tân cũ) có diện tích rất rộng, bán kính di chuyển giao hàng xa và bị chia cắt mạnh bởi địa hình đồi núi phức tạp.
- Bưu cục (DNO) Kiến Đức phụ trách cụm Đắk Ngo cũ (109.0 đơn/ngày, chiếm 54.2% sản lượng).
- Bưu cục (DNO) Quảng Tín phụ trách cụm Quảng Tân cũ (92.0 đơn/ngày, chiếm 45.8% sản lượng).
Nếu gộp về 1 bưu cục duy nhất, thời gian và bán kính giao hàng của shipper sẽ tăng gấp đôi, làm trễ SLA và nguy cơ bỏ tuyến xã xa. Vì vậy, AM ĐỀ XUẤT GIỮ NGUYÊN 02 BƯU CỤC ((DNO) Quảng Tín & (DNO) Kiến Đức) phụ trách song song theo phân vùng cũ.'''
}

ward_maps = {
    24781: ('image6.png', 'Bản đồ địa bàn quy hoạch mới TP. Đà Lạt 2026'),
    24784: ('image6.png', 'Bản đồ địa bàn quy hoạch mới TP. Đà Lạt 2026'),
    25000: ('image14.png', 'Bản đồ địa bàn quy hoạch khu vực Di Linh 2026'),
    24958: ('image13.png', 'Bản đồ địa bàn quy hoạch khu vực Đức Trọng 2026'),
    24823: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    24820: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    24817: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    25084: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    24853: ('image2.png', 'Bản đồ quy hoạch khu vực Lâm Hà - Đam Rông 2026'),
    24883: ('image2.png', 'Bản đồ quy hoạch khu vực Lâm Hà - Đam Rông 2026'),
    22366: ('image5.png', 'Bản đồ địa bàn quy hoạch mới TP. Nha Trang 2026'),
    22402: ('image5.png', 'Bản đồ địa bàn quy hoạch mới TP. Nha Trang 2026'),
    22390: ('image5.png', 'Bản đồ địa bàn quy hoạch mới TP. Nha Trang 2026'),
    22528: ('image8.jpg', 'Bản đồ địa bàn quy hoạch mới khu vực Ninh Hòa 2026'),
    22516: ('image7.jpg', 'Bản đồ địa bàn quy hoạch khu vực Vạn Ninh - Tu Bông 2026'),
    22933: ('image10.jpg', 'Bản đồ địa bàn quy hoạch mới TP. Phan Thiết & Hàm Thắng 2026'),
    22945: ('image10.jpg', 'Bản đồ địa bàn quy hoạch mới TP. Phan Thiết 2026'),
    22960: ('image11.jpg', 'Bản đồ phân vùng quy hoạch mới Phường Bình Thuận 2026'),
    23131: ('image9.png', 'Bản đồ quy hoạch mới khu vực Phường La Gi 2026'),
    23143: ('image9.png', 'Bản đồ quy hoạch mới khu vực Xã Tân Thành 2026'),
    22972: ('image4.png', 'Bản đồ quy hoạch mới khu vực Phan Rí Cửa & Tuy Phong 2026'),
    22759: ('image12.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Phan Rang 2026'),
    22834: ('image12.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Ninh Chử 2026'),
}

ward_list = []
for code, group in df.groupby('Mã Xã mới'):
    prov = group['Tỉnh, thành phố mới'].iloc[0]
    name = group['Tên Xã mới'].iloc[0]
    n_bcs = group['Số BC'].iloc[0]
    old_communes = group['Tên Xã cũ'].tolist()
    total_giao = group['Sản lượng giao/ngày (đơn)'].sum()
    total_lay = group['Sản lượng lấy/ngày (đơn)'].sum()
    total_don = group['TỔNG ĐƠN/NGÀY (Phường mới)'].iloc[0]
    total_kg = group['TỔNG KG/NGÀY (Phường mới)'].iloc[0]
    prop = group['Đánh giá & Phương án đề xuất'].iloc[0]
    reason_excel = group['Lý do & Bố trí nhân sự'].iloc[0]

    bc_details = []
    for bc_name, bc_group in group.groupby('Tên Bưu cục giao'):
        bc_giao = bc_group['Sản lượng giao/ngày (đơn)'].sum()
        bc_lay = bc_group['Sản lượng lấy/ngày (đơn)'].sum()
        bc_total = bc_giao + bc_lay
        pct = round((bc_total / total_don * 100), 1) if total_don > 0 else 0.0
        bc_id = bc_group['ID Bưu cục giao'].iloc[0]
        am = bc_group['Quản lý khu vực (AM)'].iloc[0]
        addr = bc_group['Địa chỉ Bưu cục'].iloc[0]
        gmaps = bc_group['Google Maps Bưu cục'].iloc[0]
        bc_details.append({
            'name': bc_name, 'id': bc_id, 'am': am, 'addr': addr, 'gmaps': gmaps,
            'giao': bc_giao, 'lay': bc_lay, 'total': bc_total, 'pct': pct
        })
    bc_details.sort(key=lambda x: x['total'], reverse=True)

    if code == 24748:
        prop = "GIỮ NGUYÊN 02 BƯU CỤC ((DNO) Quảng Tín & (DNO) Kiến Đức)"

    if code in am_reasons:
        final_reason = am_reasons[code]
    else:
        final_reason = reason_excel

    img_info = ward_maps.get(code, (None, None))

    ward_list.append({
        'code': code, 'name': name, 'prov': prov, 'n_bcs': n_bcs,
        'old_communes': old_communes, 'total_giao': total_giao, 'total_lay': total_lay,
        'total_don': total_don, 'total_kg': total_kg, 'prop': prop,
        'bc_details': bc_details, 'reason': final_reason,
        'img_file': img_info[0], 'img_caption': img_info[1]
    })

ward_list.sort(key=lambda x: x['name'])

# Build Markdown content
md = []
md.append('# BÁO CÁO TOÀN DIỆN: RÀ SOÁT VÀ QUY HOẠCH BƯU CỤC THEO ĐƠN VỊ HÀNH CHÍNH MỚI (VÙNG NTB)\n')

md.append('## I. TỔNG QUAN HIỆN TRẠNG MẠNG LƯỚI BƯU CỤC VÙNG NTB')
md.append('- **Tổng số Bưu cục Express đang vận hành**: 83 Bưu cục (thuộc các tỉnh Khánh Hòa, Ninh Thuận, Bình Thuận, Lâm Đồng, Đắc Nông cũ).')
md.append('- **Tổng số Xã/Phường hành chính mới rà soát**: 36 Xã/Phường mới (gồm 114 xã/phường cũ sáp nhập).')
md.append('- **Số Xã/Phường mới đã quy hoạch chuẩn (01 BC phụ trách)**: 4 Xã/Phường (chiếm 11.1%).')
md.append('- **Số Xã/Phường mới bị CHIA CẮT (2-3 BC phụ trách)**: 32 Xã/Phường (chiếm 88.9%).\n')

md.append('## II. RÀ SOÁT THEO ĐƠN VỊ HÀNH CHÍNH MỚI')
md.append('### 1. Đánh giá độ phủ Bưu cục và ranh giới hành chính mới')
md.append('Sau khi sáp nhập các xã cũ thành xã/phường mới, ranh giới quản lý của các Bưu cục đang xuất hiện hiện tượng chia cắt mảnh, dẫn đến:')
md.append('- **01 Phường mới có tới 3 Bưu cục cùng giao hàng**: Gây chồng chéo tuyến đường, làm shipper di chuyển cắt ngang địa bàn của nhau.')
md.append('- **Tuyến đi chéo xa ranh giới**: Một số xã vùng ven bị giao từ Bưu cục ở huyện/tỉnh khác cách xa 30 - 40 km, trong khi Bưu cục lân cận chỉ cách 7 - 15 km.\n')

md.append('### 2. Danh sách 18 Tuyến giao chéo xa ranh giới cần Reassign ngay')
md.append('Bảng dưới đây liệt kê 18 tuyến xã cũ bị đi chéo ranh giới.')
md.append('> ⚠️ **ĐÁNH GIÁ VẬN HÀNH & BẤT HỢP LÝ CỦA CÁC TUYẾN CHÉO KHÁC TỈNH (ĐẶC BIỆT XÃ ĐA MI)**:')
md.append('> 1. **Trường hợp Xã Đa Mi (Bình Thuận)**: Thuật toán quét khoảng cách gợi ý chuyển về Bưu cục `(LDO) Bảo Lâm 3` (tỉnh Lâm Đồng) vì khoảng cách 15.1km. Tuy nhiên, phương án này **HOÀN TOÀN SAI VỀ MẶT HÀNH CHÍNH VÀ VẬN HÀNH**:')
md.append('>    - **Vi phạm ranh giới tỉnh**: Xã Đa Mi thuộc tỉnh Bình Thuận, việc chuyển cho BC Bảo Lâm 3 (tỉnh Lâm Đồng) phụ trách sẽ gây sai lệch luồng chia chọn liên tỉnh, sai lệch đối soát COD và báo cáo hành chính tỉnh.')
md.append('>    - **Khác AM quản lý**: (BTH) Hàm Thuận thuộc AM Nguyễn Ngọc Khánh, trong khi (LDO) Bảo Lâm 3 thuộc AM Hồng Bích Nga.')
md.append('>    - **Kế hoạch chính thức của AM**: AM Nguyễn Ngọc Khánh đã quy hoạch **(BTH) Hàm Thuận (Mới)** giữ nguyên phụ trách Xã Đa Mi cùng cụm Đông Giang, La Dạ, Thuận Hòa, Hồng Sơn thuộc Bình Thuận.')
md.append('> 2. **Các tuyến chéo Đắk Nông & Liên tỉnh khác**: AM quản lý khu vực đề xuất **TẠM THỜI GIỮ NGUYÊN 100% PHẠM VI QUẢN LÝ THEO XÃ CŨ (trong 3-6 tháng)** để tránh rủi ro 100% nhân sự tại chỗ nghỉ việc và vỡ tuyến giao.\n')

md.append('| STT | Tên Xã/Phường cũ | Tỉnh | BC hiện tại (`from`) | Khoảng cách cũ | BC tối ưu gần nhất (`to`) | Quản lý AM tiếp nhận | Khoảng cách mới | Khoảng cách tiết kiệm | Đề xuất AM |')
md.append('|---|---|---|---|---|---|---|---|---|---|')

for r in doc2.tables[1].rows[1:]:
    cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
    xa = cells[1]
    prov = cells[2]
    bc_from = cells[3]
    bc_to = cells[5]
    
    # check if cross province
    p_from = bc_from.split(')')[0].replace('(', '').strip()[:2]
    p_to = bc_to.split(')')[0].replace('(', '').strip()[:2]
    
    if xa == "Xã Đa Mi":
        am_note = f"GIỮ NGUYÊN {bc_from} (Sai ranh giới tỉnh BTH->LDO)"
    elif p_from != p_to:
        am_note = f"GIỮ NGUYÊN {bc_from} (Tuyến khác tỉnh {p_from}->{p_to})"
    elif "Đắk Nông" in prov or "Lâm Đồng" in prov:
        am_note = f"GIỮ NGUYÊN {bc_from} (Tránh vỡ tuyến DNO/LDO)"
    else:
        am_note = "Reassign tối ưu"
        
    md.append(f'| {cells[0]} | {cells[1]} | {cells[2]} | {cells[3]} | {cells[4]} | {cells[5]} | {cells[6]} | {cells[7]} | {cells[8]} | {am_note} |')

md.append('\n## III. ĐÁNH GIÁ CHI TIẾT SẢN LƯỢNG VÀ ĐỀ XUẤT PHƯƠNG ÁN GỘP / GIỮ BƯU CỤC')
md.append(f'### Danh sách {len(ward_list)} Phường/Xã mới rà soát\n')

for i, w in enumerate(ward_list, 1):
    md.append(f'### {i}. {w["name"]} ({w["prov"]})')
    md.append(f'- **Mã Xã mới**: `{w["code"]}`')
    md.append(f'- **TỔNG SẢN LƯỢNG DỰ KIẾN**: {w["total_don"]:.1f} đơn/ngày ({w["total_kg"]:.1f} kg/ngày) | *Giao: {w["total_giao"]:.1f} đơn · Lấy: {w["total_lay"]:.1f} đơn*')
    
    bcs_str = ', '.join([f'"{b["name"]}"' for b in w['bc_details']])
    md.append(f'- **Các BC hiện phụ trách ({w["n_bcs"]} BC)**: {bcs_str}')
    
    old_str = ', '.join(w['old_communes'])
    md.append(f'- **Các xã cũ sáp nhập**: {old_str}')
    md.append('- **Tỷ lệ phân chia sản lượng thực tế giữa các Bưu cục**:')
    
    for b in w['bc_details']:
        gmaps_str = f' [Google Maps]({b["gmaps"]})' if b['gmaps'] and str(b['gmaps']) != 'nan' else ''
        md.append(f'  - `{b["name"]}` (ID: {b["id"]}) [AM: {b["am"]}]: {b["total"]:.1f} đơn/ngày ({b["pct"]}%) (Giao: {b["giao"]:.1f}, Lấy: {b["lay"]:.1f}) (Địa chỉ: {b["addr"]}){gmaps_str}')
    
    md.append(f'- **ĐỀ XUẤT PHƯƠNG ÁN**: **{w["prop"]}**')
    md.append(f'- **LÝ DO VÀ BỐ TRÍ NHÂN SỰ**:\n{w["reason"]}')
    
    if w['img_file']:
        img_abs = os.path.join(maps_dir, w['img_file'])
        md.append(f'\n![{w["img_caption"]}]({img_abs})')
        md.append(f'*Hình {i}: {w["img_caption"]}*\n')

md.append('## IV. YÊU CẦU 3: ĐÁNH GIÁ NHU CẦU TÁCH BƯU CỤC VÀ TỐI ƯU CÔNG SUẤT KHO BÃI')
md.append('### 1. Đánh giá Bưu cục quá tải (Cần tách bớt Phường hoặc Mở rộng diện tích)')
md.append('| STT | ID Bưu cục | Tên Bưu cục | Tỉnh | Quản lý AM | Địa chỉ Bưu cục | Áp lực m² (`em2`) | Đề xuất giải pháp |')
md.append('|---|---|---|---|---|---|---|---|')

for r in doc2.tables[0].rows[1:]:
    cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
    md.append(f'| {cells[0]} | {cells[1]} | {cells[2]} | {cells[3]} | {cells[4]} | {cells[5]} | {cells[6]} | {cells[7]} |')

md.append('\n### 2. Kế hoạch Mở mới / Tách Bưu cục & Tối ưu Mạng lưới')
md.append('1. **Mở mới Bưu cục (LDO) Xuân Hương - Đà Lạt 2**: Đặt tại khu vực Phường 10 (TP. Đà Lạt), chia tải cho Bưu cục Xuân Hương cũ (phụ trách Phường 3 & Phường 10, Vol giao: 1,050 đơn/ngày, Vol lấy: 150 đơn/ngày, 8 NVPTTT + 1 NVXL).')
md.append('2. **Tách Bưu cục Hàng Nhỏ / Hàng Vừa Di Linh (Xã Đinh Trang Thượng)**: Phụ trách Đinh Trang Thượng, Di Linh, Phúc Thọ Lâm Hà, Liên Đầm để giảm bán kính di chuyển 22-45km cho BC Di Linh.')
md.append('3. **Mở mới Bưu cục Đông Hải (Tỉnh Ninh Thuận)**: Cover khu vực ven biển Đông Hải, tối ưu điểm tập kết sản lượng lấy.')
md.append('4. **Mở mới Bưu cục (LDO) B\'Lao Mới (Bảo Lộc)**: Tối ưu mạng lưới khu vực Phường 1 & Phường 2 Bảo Lộc, đóng cửa BC (LDO) 1 Bảo Lộc cũ.\n')

md.append('## V. TỔNG HỢP BIẾN ĐỘNG MẠNG LƯỚI BƯU CỤC NTB 2026')
md.append('- **Bưu cục Mở mới (04 BC)**: BC Xuân Hương - Đà Lạt 2, BC Di Linh Hàng Nhỏ, BC Đông Hải, BC B\'Lao Mới.')
md.append('- **Bưu cục Đóng cửa (02 BC)**: BC 1 Bảo Lộc, BC Nam Nha Trang 3.')
md.append('- **Bưu cục Gộp tuyến/Điều chỉnh phân vùng (19 BC)**: Gom các tuyến xã lẻ về BC trung tâm theo ĐVHC mới.')
md.append('- **Bưu cục Giữ nguyên vận hành (13 BC)**: Duy trì tại các khu vực đặc thù Gia Nghĩa, Đắc Nông, Đà Lạt, Đức Trọng, Phan Rí Cửa, Ninh Hòa, Xã Đa Mi (Bình Thuận)...')

full_md_text = '\n'.join(md)

# Write Markdown report to Artifact path
artifact_md_path = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\d95894ad-9248-47c3-872b-72d7de66bc20\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi.md'
with open(artifact_md_path, 'w', encoding='utf-8') as f:
    f.write(full_md_text)

print(f'Saved Markdown report to: {artifact_md_path}')

# Write Word DOCX report
new_doc = docx.Document()

# Style configuration
style_normal = new_doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(10.5)

def add_styled_heading(text, level):
    h = new_doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    run = h.runs[0]
    run.font.name = 'Arial'
    if level == 0:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 51, 51)
    return h

def set_table_styling(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(r'<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        if r_idx == 0:
            trPr.append(parse_xml(r'<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="003366"/>'))
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(9)
        else:
            bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{bg_color}"/>'))
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8.5)

def insert_table_18_routes_inline(doc_table_source):
    t_doc = new_doc.add_table(rows=len(doc_table_source.rows), cols=10)
    for r_idx, r in enumerate(doc_table_source.rows):
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        if r_idx == 0:
            row_data = cells + ['Đề xuất AM']
        else:
            xa = cells[1]
            prov = cells[2]
            bc_from = cells[3]
            bc_to = cells[5]
            
            p_from = bc_from.split(')')[0].replace('(', '').strip()[:2]
            p_to = bc_to.split(')')[0].replace('(', '').strip()[:2]
            
            if xa == "Xã Đa Mi":
                am_note = f"GIỮ NGUYÊN {bc_from} (Sai ranh giới tỉnh BTH->LDO)"
            elif p_from != p_to:
                am_note = f"GIỮ NGUYÊN {bc_from} (Tuyến khác tỉnh {p_from}->{p_to})"
            elif "Đắk Nông" in prov or "Lâm Đồng" in prov:
                am_note = f"GIỮ NGUYÊN {bc_from} (Tránh vỡ tuyến DNO/LDO)"
            else:
                am_note = "Reassign tối ưu"
            row_data = cells + [am_note]
        for c_idx, val in enumerate(row_data):
            t_doc.cell(r_idx, c_idx).text = val
    set_table_styling(t_doc)

def insert_table_inline(doc_table_source):
    t_doc = new_doc.add_table(rows=len(doc_table_source.rows), cols=len(doc_table_source.columns))
    for r_idx, r in enumerate(doc_table_source.rows):
        for c_idx, c in enumerate(r.cells):
            t_doc.cell(r_idx, c_idx).text = c.text.strip().replace('\n', ' ')
    set_table_styling(t_doc)

# Add Title
add_styled_heading('BÁO CÁO TOÀN DIỆN: RÀ SOÁT VÀ QUY HOẠCH BƯU CỤC THEO ĐƠN VỊ HÀNH CHÍNH MỚI (VÙNG NTB)', level=0)

i = 0
lines = md[1:]
while i < len(lines):
    line = lines[i].strip()
    if line.startswith('## '):
        add_styled_heading(line.replace('## ', ''), level=1)
    elif line.startswith('### '):
        heading_text = line.replace('### ', '')
        add_styled_heading(heading_text, level=2)
        
        # Check if this heading needs inline table insertion
        if '2. Danh sách 18 Tuyến giao chéo' in heading_text:
            insert_table_18_routes_inline(doc2.tables[1])
        elif '1. Đánh giá Bưu cục quá tải' in heading_text:
            insert_table_inline(doc2.tables[0])

    elif '![' in line:
        import re
        m = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if m:
            cap_text, img_p = m.group(1), m.group(2)
            if os.path.exists(img_p):
                p_img = new_doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_p, width=Inches(5.5))
                
                p_cap = new_doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_cap = p_cap.add_run(f'Hình: {cap_text}')
                r_cap.font.italic = True
                r_cap.font.size = Pt(9.5)
                r_cap.font.color.rgb = RGBColor(100, 100, 100)
    elif line.startswith('*Hình '):
        pass
    elif line.startswith('|') and '---' in line:
        pass
    elif line.startswith('|'):
        pass # markdown tables handled inline above
    elif line.startswith('- '):
        p_b = new_doc.add_paragraph(style='List Bullet')
        txt = line.replace('- ', '')
        p_b.add_run(txt)
    elif line.startswith('  - '):
        p_b2 = new_doc.add_paragraph(style='List Bullet 2')
        txt = line.replace('  - ', '')
        p_b2.add_run(txt)
    elif line.strip():
        p_txt = new_doc.add_paragraph(line.strip())
        p_txt.paragraph_format.space_after = Pt(4)
    i += 1

docx_out_path = r'C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi_v6.docx'
new_doc.save(docx_out_path)
print(f'Saved enhanced DOCX report v6 with Đa Mi fix to: {docx_out_path}')
