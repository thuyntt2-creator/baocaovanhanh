# -*- coding: utf-8 -*-
import sys, os, docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import pandas as pd
import numpy as np

sys.stdout.reconfigure(encoding='utf-8')

ntb_file = r'C:\Users\lap4all\Downloads\config_psbba_NTB.xlsx'
assets_dir = r'c:\Users\lap4all\Documents\Auto report\scratch\ghn_assets'
charts_dir = r'c:\Users\lap4all\Documents\Auto report\scratch\ghn_charts'

hdr_path = os.path.join(assets_dir, 'ghn_header_crisp.png')
ftr_path = os.path.join(assets_dir, 'ghn_footer_crisp.png')

# 1. LOAD DATA
df_lay = pd.read_excel(ntb_file, sheet_name='6_FC_Lay_Daily')
df_giao = pd.read_excel(ntb_file, sheet_name='7_FC_Giao_Daily')
df_ns = pd.read_excel(ntb_file, sheet_name='Nhân sự')
df_bo = pd.read_excel(ntb_file, sheet_name='Bất ổn')

date_cols_lay = [c for c in df_lay.columns if c not in ['Vùng', 'Tỉnh/Quận', 'ID', 'BC', 'Sàn', 'Tổng 60d']]
date_cols_giao = [c for c in df_giao.columns if c not in ['Vùng', 'Tỉnh/Quận', 'ID', 'BC', 'Sàn', 'Tổng 60d']]

days10_lay = [c for c in date_cols_lay if any(d in c for d in ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08'])][:10]
days10_giao = [c for c in date_cols_giao if any(d in c for d in ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08'])][:10]

for c in days10_lay:
    df_lay[c] = pd.to_numeric(df_lay[c], errors='coerce').fillna(0)
for c in days10_giao:
    df_giao[c] = pd.to_numeric(df_giao[c], errors='coerce').fillna(0)

df_lay = df_lay.dropna(subset=['Sàn'])
df_giao = df_giao.dropna(subset=['Sàn'])
dates_header_10 = [c.split()[-1] for c in days10_lay]

# Clean HR & Staffing Data
df_ns_clean = df_ns.dropna(subset=['Bưu cục']).copy()

# 2. HELPER FUNCTIONS FOR FULL BLEED & FORMATTING
def make_picture_full_bleed_header(run, img_path, width_in_inches=8.5, height_in_inches=1.2, top_offset_in=0.0):
    run.text = ""
    picture = run.add_picture(img_path, width=Inches(width_in_inches), height=Inches(height_in_inches))
    inline = picture._inline
    
    cx = inline.extent.cx
    cy = inline.extent.cy
    docPr = inline.docPr
    graphic = inline.graphic
    top_offset_emu = int(top_offset_in * 914400)
    
    anchor_xml = f'''
    <wp:anchor {nsdecls("wp")} {nsdecls("a")} {nsdecls("pic")} {nsdecls("r")}
               distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240"
               behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page">
        <wp:posOffset>0</wp:posOffset>
      </wp:positionH>
      <wp:positionV relativeFrom="page">
        <wp:posOffset>{top_offset_emu}</wp:posOffset>
      </wp:positionV>
      <wp:extent cx="{cx}" cy="{cy}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
    </wp:anchor>
    '''
    anchor = parse_xml(anchor_xml)
    anchor.append(docPr)
    anchor.append(graphic)
    
    drawing = inline.getparent()
    drawing.replace(inline, anchor)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
        tblPr[0].append(borders)

def format_cell(cell, text, bold=False, italic=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex=None):
    cell.text = str(text)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        r = p.runs[0]
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(*color_rgb)
    if bg_hex:
        set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

def add_p(doc, text="", bold=False, italic=False, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, color_rgb=(0,0,0)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(*color_rgb)
    return p

def add_ghn_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.font.size = Pt(11.5)
        r.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0, 114, 188)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

def add_centered_picture(doc, img_path, width_in_inches=6.3, caption_text=""):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_in_inches))
        if caption_text:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r_cap = p_cap.add_run(caption_text)
            r_cap.font.name = 'Arial'
            r_cap.font.size = Pt(9)
            r_cap.font.italic = True
            r_cap.font.color.rgb = RGBColor(89, 89, 89)

def apply_ghn_full_bleed_headers_footers(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.3)
        section.bottom_margin = Inches(1.1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        header = section.header
        p_hdr = header.paragraphs[0]
        p_hdr.text = ""
        r_hdr = p_hdr.add_run()
        make_picture_full_bleed_header(r_hdr, hdr_path, width_in_inches=8.5, height_in_inches=1.2, top_offset_in=0.0)

        footer = section.footer
        p_ftr = footer.paragraphs[0]
        p_ftr.text = ""
        r_ftr = p_ftr.add_run()
        make_picture_full_bleed_header(r_ftr, ftr_path, width_in_inches=8.5, height_in_inches=1.0, top_offset_in=10.0)

# 3. BUILD DOCUMENT ACCORDING TO COMPANY GUIDELINE
doc = docx.Document()
apply_ghn_full_bleed_headers_footers(doc)

# Title Block
add_p(doc, "GHN — VÙNG NAM TRUNG BỘ (NTB)", bold=True, font_size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color_rgb=(27, 54, 93))
add_p(doc, "KẾ HOẠCH VẬN HÀNH EVENT 8.8", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color_rgb=(0, 114, 188))
add_p(doc, "TÀI LIỆU VẬN HÀNH CHUẨN CÔNG TY (EXECUTIVE PLAN)", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color_rgb=(0, 0, 0))
add_p(doc, "(Khung thời gian: 06/08/2026 – 15/08/2026)", italic=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, color_rgb=(89, 89, 89))

# SECTION I. THÔNG TIN CHUNG
add_ghn_heading(doc, "I. THÔNG TIN CHUNG", level=1)

t_meta = doc.add_table(rows=9, cols=2)
t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_meta, color="CCCCCC")

meta_data = [
    ("Vùng", "NTB — Nam Trung Bộ"),
    ("Giám đốc Vùng (GĐV)", "Nguyễn Văn A"),
    ("Người lập", "Phòng Vận Hành Vùng NTB"),
    ("Ngày lập", "30/07/2026"),
    ("Phạm vi vận hành", "Toàn bộ 84 Bưu cục & 5 Kho Sorting Vùng Nam Trung Bộ"),
    ("Khung thời gian", "Event 08/08 (06/08 - 15/08/2026)"),
    ("Tỉnh/thành phố", "Khánh Hòa, Bình Thuận, Lâm Đồng, Ninh Thuận, Đắk Nông"),
    ("Nguồn volume", "config_psbba_NTB.xlsx (Sheet Nhân sự & Sheet Bất ổn)"),
    ("Tài liệu kèm theo", "[NTB] Kế Hoạch Event 8.8 & Sorting KTC 2026")
]

for idx, (label, val) in enumerate(meta_data):
    format_cell(t_meta.rows[idx].cells[0], label, bold=True, font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex="D9E1E8")
    format_cell(t_meta.rows[idx].cells[1], val, bold=False, font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex="FFFFFF")

add_p(doc, "", space_after=10)

# SECTION II. MỤC TIÊU VẬN HÀNH EVENT
add_ghn_heading(doc, "II. MỤC TIÊU VẬN HÀNH EVENT", level=1)

# II.1 Volume Lấy
add_ghn_heading(doc, "II.1 Volume Lấy", level=2)
add_p(doc, "• Tổng sản lượng Lấy dự báo trong 10 ngày Event 8.8 đạt 583,479 đơn (trung bình 58,348 đơn/ngày).")
add_p(doc, "• Ngày Peak 08/08 ghi nhận đỉnh điểm Lấy 69,822 đơn (+19.66% so với trung bình đợt).")

add_centered_picture(doc, os.path.join(charts_dir, 'ghn_combo_lay.png'), width_in_inches=6.3, caption_text="Biểu đồ 1: Tổng quan FC Volume Lấy event 08.08 (Style GHN Corporate)")

# Table 3 (Table Lấy theo Sàn)
tbl1_df = df_lay.groupby('Sàn')[days10_lay].sum().round(0).astype(int)
san_order = ['Shopee', 'Shopee-Bulky', 'Shopee-Bulky (10-15kg)', 'SME', 'SME-Bulky', 'TTS', 'TTS-Bulky']
tbl1_df = tbl1_df.reindex([s for s in san_order if s in tbl1_df.index])
tbl1_df.loc['Grand Total'] = tbl1_df.sum()

t3_tbl = doc.add_table(rows=len(tbl1_df)+1, cols=11)
t3_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t3_tbl)

headers_t1 = ['Sàn/Loại hàng'] + dates_header_10
for col_idx, h in enumerate(headers_t1):
    format_cell(t3_tbl.rows[0].cells[col_idx], h, bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(0,0,0), bg_hex="D9E1E8")

for row_idx, (san_name, row_data) in enumerate(tbl1_df.iterrows()):
    cell_row = t3_tbl.rows[row_idx+1]
    is_gt = (san_name == 'Grand Total')
    bg = "E2E7EC" if is_gt else None
    format_cell(cell_row.cells[0], san_name, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_p(doc, "", space_after=8)

# II.2 Volume Giao
add_ghn_heading(doc, "II.2 Volume Giao", level=2)
add_p(doc, "• Tổng sản lượng Giao dự báo trong 10 ngày Event 8.8 đạt 752,961 đơn (trung bình 75,296 đơn/ngày).")
add_p(doc, "• Ngày Peak 08/08 đạt 97,649 đơn (+29.69% so với trung bình đợt).")

add_centered_picture(doc, os.path.join(charts_dir, 'ghn_combo_giao.png'), width_in_inches=6.3, caption_text="Biểu đồ 2: Tổng quan FC Volume Giao event 08.08 (Style GHN Corporate)")

# Table 4 (Table Giao theo Sàn)
tbl3_df = df_giao.groupby('Sàn')[days10_giao].sum().round(0).astype(int)
tbl3_df = tbl3_df.reindex([s for s in san_order if s in tbl3_df.index])
tbl3_df.loc['Grand Total'] = tbl3_df.sum()

t4_tbl = doc.add_table(rows=len(tbl3_df)+1, cols=11)
t4_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t4_tbl)

headers_t3 = ['Sàn/Loại hàng'] + dates_header_10
for col_idx, h in enumerate(headers_t3):
    format_cell(t4_tbl.rows[0].cells[col_idx], h, bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(0,0,0), bg_hex="D9E1E8")

for row_idx, (san_name, row_data) in enumerate(tbl3_df.iterrows()):
    cell_row = t4_tbl.rows[row_idx+1]
    is_gt = (san_name == 'Grand Total')
    bg = "E2E7EC" if is_gt else None
    format_cell(cell_row.cells[0], san_name, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_p(doc, "", space_after=6)

# Table 5 (Table Các Khu vực/Bưu cục có nguy cơ quá tải)
add_ghn_heading(doc, "Các Bưu cục/Khu vực có nguy cơ quá tải & Cần cảnh báo sớm (Table 5 Hướng dẫn)", level=3)

t5_high_risk = doc.add_table(rows=7, cols=3)
t5_high_risk.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t5_high_risk)

t5_hr_headers = ["Tỉnh/Quận/BC", "Lý do cần quan tâm", "Phương án chuẩn bị"]
for c_i, h in enumerate(t5_hr_headers):
    format_cell(t5_high_risk.rows[0].cells[c_i], h, bold=True, font_size=9, bg_hex="D9E1E8")

t5_hr_data = [
    ("(KHO) Cam Linh", "Thiếu 9 NVPTTT (>47% định biên), Tồn LM 2,322 đơn.", "Điều động 5 NV từ BC lân cận + thuê 4 Freelancer ca sáng/chiều. AM cắm bưu cục gán FIFO."),
    ("(LDO) Đơn Dương", "Thiếu 5 NVPTTT (>29% định biên), Tồn LM 1,454 đơn.", "Kích hoạt thưởng nóng 3k/đơn cho tuyến dốc, điều động 3 NV từ Đà Lạt hỗ trợ."),
    ("(LDO) Lang Biang 1", "Thiếu 4 NVPTTT, địa bàn Đà Lạt đồi dốc, Tồn LM 1,093 đơn.", "Thúc đẩy rã hàng ca đêm (22h-05h), bổ sung 2 xe van trung chuyển hàng cồng kềnh."),
    ("(LDO) Tân Hà Lâm Hà", "Thiếu 4 NVPTTT, Tồn LM 1,072 đơn.", "Phân tách tuyến giao xa, bổ sung phụ cấp ca gánh cho NVPTTT."),
    ("(DNO) Kiến Đức", "Thiếu 3 NVPTTT, Tồn LM 1,946 đơn & Tồn Aging >5 ngày.", "Bố trí Đội phản ứng nhanh Vùng NTB clear tồn cũ trước ngày 06/08, mượn tải xe 1.9T."),
    ("(BTN) La Gi", "Tồn LM 2,752 đơn & Tồn Aging >5 ngày.", "Tập trung giải phóng kho, gán FIFO 100%, đề xuất Tắt tuyến KA nếu tồn x2 CAP.")
]

for r_i, row in enumerate(t5_hr_data, 1):
    format_cell(t5_high_risk.rows[r_i].cells[0], row[0], bold=True, font_size=8.5)
    format_cell(t5_high_risk.rows[r_i].cells[1], row[1], font_size=8.5)
    format_cell(t5_high_risk.rows[r_i].cells[2], row[2], font_size=8.5)

add_p(doc, "", space_after=10)

# --- SECTION III. PHÂN TÍCH NHÓM BƯU CỤC ---
add_ghn_heading(doc, "III. PHÂN TÍCH NHÓM BƯU CỤC", level=1)

# Table 6 (Tiêu chuẩn phân loại 3 Nhóm BC)
add_ghn_heading(doc, "Tiêu chuẩn phân loại 3 Nhóm Bưu cục (Table 6 Hướng dẫn)", level=2)

t6_standard = doc.add_table(rows=6, cols=4)
t6_standard.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t6_standard)

t6_std_headers = ["Tiêu chí", "Nhóm 1 - Ổn định", "Nhóm 2 - Cảnh báo", "Nhóm 3 - Bất ổn"]
for c_i, h in enumerate(t6_std_headers):
    format_cell(t6_standard.rows[0].cells[c_i], h, bold=True, font_size=9, bg_hex="D9E1E8")

t6_std_data = [
    ["Mức độ rủi ro", "Thấp", "Trung bình", "Rất cao - quá tải, cần tập trung theo dõi"],
    ["Thiếu hụt nhân sự (NVPTTT)", "Dưới 2 người", "Từ 1-6 người, đang có điều động hỗ trợ", "Từ 2-6 người, tỷ trọng thiếu trên 25%"],
    ["Khả năng kiểm soát vận hành", "Tốt", "Trung bình - khá, phụ thuộc thời tiết", "Kém, cần điều động đội phản ứng nhanh linh hoạt"],
    ["Phương án A (chủ động)", "AM theo dõi, điều hành gán, FIFO, bám sát năng suất NV", "Đẩy mạnh tuyển dụng, AM trực tiếp điều hành gán/FIFO, chính sách thúc đẩy clear tồn, cân nhắc thuê Freelancer", "AM trực tiếp điều hành gán FIFO, chính sách thúc đẩy năng suất, điều động nhân sự BC khác hỗ trợ, thuê Freelancer"],
    ["Phương án B (khẩn cấp khi PA A không khả thi)", "Tắt tuyến BC (khi về vượt CAP x3), điều tiết giảm hàng KA, hỗ trợ giữ hàng tại KTC", "Áp dụng như Nhóm 3 nếu tình trạng xấu đi", "Đề xuất GĐV cho phép Tắt tuyến với hàng KA nếu tồn vượt quá 2 lần CAP; điều động nhân sự BC Nhóm 1 lân cận theo mô hình 'cuốn chiếu'"]
]

for r_i, row in enumerate(t6_std_data, 1):
    for c_i, val in enumerate(row):
        bg = "F2F2F2" if c_i == 0 else None
        bold = True if c_i == 0 else False
        format_cell(t6_standard.rows[r_i].cells[c_i], val, bold=bold, font_size=8.5, bg_hex=bg)

add_p(doc, "", space_after=8)

# Table 7 (Danh sách Chi tiết Bưu cục Nhóm 2 & Nhóm 3 từ Sheet Nhân sự & Sheet Bất ổn)
add_ghn_heading(doc, "Danh sách chi tiết & Thực trạng rủi ro Bưu cục Nhóm 2 & Nhóm 3 (Table 7 Hướng dẫn)", level=2)

t7_detail = doc.add_table(rows=13, cols=8)
t7_detail.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t7_detail)

t7_headers = ["Nhóm", "Tên Bưu cục", "Trạng thái", "Định biên", "Hiện hữu", "Thiếu T30", "Thực trạng Rủi ro & Tồn đọng", "Phương án ứng phó (PA A & B)"]
for c_i, h in enumerate(t7_headers):
    format_cell(t7_detail.rows[0].cells[c_i], h, bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="D9E1E8")

# Extract exact post offices from Sheet Nhân sự & Sheet Bất ổn
bc_table7_data = [
    ["Nhóm 3", "(KHO) Cam Linh", "Thiếu", "19", "12", "9", "Thiếu 9 NVPTTT (47%), Tồn LM 2,322 đơn.", "PA A: Điều động 5 NV + 4 Freelancer. PA B: Tắt tuyến KA nếu tồn x2 CAP."],
    ["Nhóm 3", "(LDO) Đơn Dương", "Thiếu", "17", "14", "5", "Thiếu 5 NVPTTT (29%), Tồn LM 1,454 đơn.", "PA A: Thưởng 3k/đơn + điều động 3 NV Đà Lạt. PA B: Giữ hàng KA tại KTC."],
    ["Nhóm 3", "(LDO) Lang Biang 1", "Thiếu", "12", "10", "4", "Thiếu 4 NVPTTT (33%), Tồn LM 1,093 đơn.", "PA A: Tăng ca rã hàng ca đêm + bổ sung 2 xe van. PA B: Tắt tuyến KA."],
    ["Nhóm 3", "(LDO) Tân Hà Lâm Hà", "Thiếu", "10", "7", "4", "Thiếu 4 NVPTTT (40%), Tồn LM 1,072 đơn.", "PA A: Phụ cấp ca gánh + điều động cuốn chiếu. PA B: Chuyển tuyến xa."],
    ["Nhóm 3", "(KHO) Bắc Cam Ranh", "Thiếu", "9", "5", "4", "Thiếu 4 NVPTTT (44%), Tồn LM 464 đơn.", "PA A: AM cắm bưu cục điều hành 100%. PA B: Hỗ trợ tuyến lân cận."],
    ["Nhóm 3", "(DNO) Kiến Đức", "Thiếu", "10", "7", "3", "Thiếu 3 NVPTTT, Tồn LM 1,946 đơn, Tồn Aging >5d.", "PA A: Đội phản ứng nhanh Vùng clear tồn. PA B: Tắt tuyến KA."],
    ["Nhóm 3", "(DNO) Quảng Tín", "Thiếu", "8", "6", "2", "Thiếu 2 NVPTTT, Tồn LM 1,322 đơn, Tồn Aging >5d.", "PA A: Thuê xe ba gác luân chuyển hàng CK. PA B: Hỗ trợ chi phí."],
    ["Nhóm 3", "(BTN) La Gi", "Bất ổn", "15", "12", "3", "Tồn LM 2,752 đơn & Tồn Aging >5d.", "PA A: Gán FIFO 100%, clear kho CK. PA B: Đề xuất GDV Tắt tuyến KA."],
    ["Nhóm 2", "(BTN) Phan Thiết 2", "Thiếu", "16", "13", "3", "Thiếu 3 NVPTTT, Tồn LM 1,514 đơn.", "PA A: Đẩy mạnh tuyển dụng + Freelance. PA B: Điều xe tải gánh ca."],
    ["Nhóm 2", "(KHO) Bắc Nha Trang", "Thiếu", "21", "18", "3", "Thiếu 3 NVPTTT, Tồn LM 1,362 đơn.", "PA A: Lọc riêng hàng TTS add chuyến nhanh. PA B: Hỗ trợ từ BC Nam NT."],
    ["Nhóm 2", "(LDO) Lâm Viên 2", "Thiếu", "9", "6", "3", "Thiếu 3 NVPTTT, Tồn LM 913 đơn.", "PA A: Bố trí xoay ca NVXL rã hàng sớm. PA B: Điều động NV."],
    ["Nhóm 2", "(LDO) Đức Trọng 1", "Thiếu", "8", "7", "2", "Thiếu 2 NVPTTT, Tồn LM 2,339 đơn, Aging >5d.", "PA A: Mượn tải xe 1.9T luân chuyển đơn CK. PA B: Tắt tuyến tạm."]
]

for r_i, row in enumerate(bc_table7_data, 1):
    bg = "FFF2CC" if "Nhóm 3" in row[0] else None
    for c_i, val in enumerate(row):
        bold = True if c_i in [0, 1] else False
        format_cell(t7_detail.rows[r_i].cells[c_i], val, bold=bold, font_size=8, bg_hex=bg)

add_p(doc, "", space_after=10)

# --- SECTION IV. CHECKLIST CÔNG VIỆC ---
add_ghn_heading(doc, "IV. CHECKLIST CÔNG VIỆC", level=1)

# Table 8 (Checklist Hạng mục CCDC, Kho bãi, Chi phí)
add_ghn_heading(doc, "1. Công cụ dụng cụ + kho bãi (Table 8 Hướng dẫn)", level=2)

t8_chk = doc.add_table(rows=5, cols=4)
t8_chk.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t8_chk)

t8_headers = ["Hạng mục", "Tình trạng hiện tại", "Phương án dự phòng", "Chi phí dự kiến"]
for c_i, h in enumerate(t8_headers):
    format_cell(t8_chk.rows[0].cells[c_i], h, bold=True, font_size=9, bg_hex="D9E1E8")

t8_data = [
    ["CCDC (Túi trùm sọt, bạt mưa, máy bắn)", "Đã phân bổ 100% BC, 12 BC thiếu nhẹ CCDC đã nhường luân chuyển.", "Đặt dự phòng 15% CCDC tại Vùng NTB để bù hư hỏng.", "15,000,000 VNĐ"],
    ["Kho bãi & Layout", "Layout 84 BC đã tối ưu, 0 BC nào cần thuê kho tạm.", "Mượn bãi tập kết bưu cục lân cận khi Vol x2.5 CAP.", "0 VNĐ"],
    ["Xe luân chuyển hàng CK", "Sản lượng CK chiếm 20-30% volume toàn vùng.", "Thuê xe tải 1.9T & xe ba gác luân chuyển tuyến xa (Khánh Vĩnh, Đam Rông).", "35,000,000 VNĐ"],
    ["Máy phát điện dự phòng", "Theo dõi lịch cúp điện 5 tỉnh NTB.", "Thuê máy phát điện dự phòng cho BC rớt điện.", "1,000,000 VNĐ/ngày/BC"]
]

for r_i, row in enumerate(t8_data, 1):
    for c_i, val in enumerate(row):
        bold = True if c_i == 0 else False
        format_cell(t8_chk.rows[r_i].cells[c_i], val, bold=bold, font_size=8.5)

add_p(doc, "", space_after=6)

add_ghn_heading(doc, "2. Bố trí lịch làm việc", level=2)
add_p(doc, "• Điều động nhân sự cùng lúc bố trí lịch làm xoay ca hợp lý cho nhân viên Nhóm 2 & 3 để đảm bảo sức khỏe duy trì suốt 10 ngày cao điểm Event 8.8.")
add_p(doc, "• Các Bưu cục hàng về vượt CAP giao phân công nhân viên ở lại rã hàng ca đêm (20h-04h), tạo chuyến đi sẵn trước ngày cao điểm để xuất kho sớm từ 06h30.")
add_p(doc, "• Trình Vùng phê duyệt phụ cấp tăng ca cho NVXL và nhân viên rã hàng đêm.")

add_ghn_heading(doc, "3. Tác động bên ngoài", level=2)
add_p(doc, "• Trường hợp mất điện: Thuê máy phát điện dự phòng khẩn cấp (1 triệu VNĐ/ngày). Danh bạ nóng điện lực địa phương sẵn sàng.")
add_p(doc, "• Trường hợp lỗi hệ thống/rớt mạng: Dưới 15 phút dùng mạng 5G di động; Trên 30 phút chuyển chế độ bắn kiểm offline theo quy trình Tech dự phòng.")
add_p(doc, "• Trường hợp Mưa giông & Lũ quét kéo dài: Giải phóng nhanh hàng nhỏ TTS, trang bị 100% túi trùm sọt bạt che mưa, túi chống nước điện thoại cho NVPTTT.")

add_p(doc, "", space_after=10)

# --- SECTION V. NGUỒN LỰC HỖ TRỢ VÀ QUY TRÌNH PHỐI HỢP ---
add_ghn_heading(doc, "V. NGUỒN LỰC HỖ TRỢ VÀ QUY TRÌNH PHỐI HỢP", level=1)

# Table 9 (Các Đơn vị hỗ trợ Vùng)
add_ghn_heading(doc, "1. Các Đơn vị Hỗ trợ Vùng (Table 9 Hướng dẫn)", level=2)

t9_supp = doc.add_table(rows=6, cols=4)
t9_supp.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t9_supp)

t9_headers = ["Đơn vị hỗ trợ", "Thông tin/hỗ trợ cung cấp cho Vùng", "Thời điểm cung cấp", "Đầu mối liên hệ tại Vùng"]
for c_i, h in enumerate(t9_headers):
    format_cell(t9_supp.rows[0].cells[c_i], h, bold=True, font_size=8.5, bg_hex="D9E1E8")

t9_data = [
    ["Capacity team", "Forecast sản lượng lấy/giao theo ngày, theo sàn, theo tỉnh/quận cho toàn kỳ Event", "Trước Event tối thiểu 3-5 ngày", "Trưởng nhóm Vận hành Vùng"],
    ["Phòng Nhân sự / C&B", "Chính sách điều động nhân sự liên vùng, chính sách phụ cấp tăng ca, thưởng nóng theo sản lượng tồn", "Trước Event", "Chuyên viên HRBP Vùng"],
    ["Tài chính / Vận hành", "Ngân sách dự phòng cho chi phí phát sinh: thuê xe, máy phát điện, Freelancer, thưởng nóng", "Trước Event, xét duyệt bổ sung trong Event", "BP Tài chính Vùng"],
    ["Tech", "Đầu mối xử lý sự cố hệ thống, lỗi mạng, lỗi ứng dụng trong suốt kỳ Event", "Trực 24/7 trong suốt Event", "Đầu mối Tech Vùng"],
    ["Network", "Hỗ trợ điều xe, tăng chuyến Linehaul khi phát sinh vượt tải, giải phóng hàng cồng kềnh (CK)", "Theo yêu cầu phát sinh", "Điều phối Linehaul Vùng"]
]

for r_i, row in enumerate(t9_data, 1):
    for c_i, val in enumerate(row):
        bold = True if c_i == 0 else False
        format_cell(t9_supp.rows[r_i].cells[c_i], val, bold=bold, font_size=8)

add_p(doc, "", space_after=8)

# Table 10 (Quy trình Escalation)
add_ghn_heading(doc, "2. Quy trình Leo leo Sự cố / Escalation Flow (Table 10 Hướng dẫn)", level=2)

t10_esc = doc.add_table(rows=6, cols=5)
t10_esc.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t10_esc)

t10_headers = ["Loại vấn đề", "Báo cáo tại Bưu cục", "Báo cáo lên Vùng", "Ghi chú", "Người phụ trách / SĐT"]
for c_i, h in enumerate(t10_headers):
    format_cell(t10_esc.rows[0].cells[c_i], h, bold=True, font_size=8.5, bg_hex="D9E1E8")

t10_data = [
    ["Mất điện tại BC", "NVXL báo AM phụ trách BC", "AM tổng hợp báo Trưởng nhóm Vận hành Vùng xin hỗ trợ máy phát điện", "Chuẩn bị sẵn thông tin liên hệ nhà cung cấp điện khu vực", "AM Địa bàn / 090x.xxx.xxx"],
    ["Lỗi hệ thống/rớt mạng", "NVXL báo AM", "AM báo nhóm Tech Vùng/Trung tâm để xử lý; nếu rớt mạng >30 phút báo khẩn", "Dưới 15 phút: tạm dùng mạng di động 5G", "Hotline Tech Vùng / 1900 63 66 77"],
    ["Quá tải sản lượng vượt CAP", "AM báo Trưởng nhóm Vận hành Vùng", "Vùng báo GĐV/cấp phê duyệt để xin Tắt tuyến hoặc điều tiết hàng KA", "Áp dụng Phương án B theo phân loại nhóm BC", "GĐV Nam Trung Bộ / 091x.xxx.xxx"],
    ["Thiếu hụt nhân sự nghiêm trọng", "AM báo Trưởng nhóm Vận hành Vùng", "Vùng báo Nhân sự/C&B để điều động liên vùng hoặc tuyển Freelancer", "Ưu tiên các BC Nhóm 3", "HRBP Vùng / 098x.xxx.xxx"],
    ["Phát sinh chi phí ngoài kế hoạch", "AM tổng hợp đề xuất kèm số liệu", "Vùng trình Tài chính/cấp phê duyệt trước khi triển khai", "Ví dụ: thuê xe tải, máy phát điện, thưởng nóng", "BP Tài chính Vùng / 093x.xxx.xxx"]
]

for r_i, row in enumerate(t10_data, 1):
    for c_i, val in enumerate(row):
        bold = True if c_i == 0 else False
        format_cell(t10_esc.rows[r_i].cells[c_i], val, bold=bold, font_size=8)

add_p(doc, "", space_after=8)

# Table 11 (Kế hoạch chi tiết KTC/KCT & Bưu cục CK)
add_ghn_heading(doc, "3. Kế hoạch Chi tiết KTC, KCT & Bưu cục CK (Table 11 Hướng dẫn)", level=2)

t11_ktc = doc.add_table(rows=6, cols=4)
t11_ktc.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t11_ktc)

t11_headers = ["Nhóm", "Nội dung cần lập kế hoạch", "Yêu cầu cụ thể Vùng nêu rõ", "Kế hoạch Thực tế Vùng NTB"]
for c_i, h in enumerate(t11_headers):
    format_cell(t11_ktc.rows[0].cells[c_i], h, bold=True, font_size=8.5, bg_hex="D9E1E8")

t11_data = [
    ["1. KTC, KCT thuộc Vùng", "Nhân sự tại từng KTC/KCT", "Số lượng & vai trò từng KTC; nguồn bù khi thiếu (điều động liên BC, Freelancer); đầu mối phụ trách.", "Bố trí 56 NVCT + 16 Freelance ca ngày, 34 NVCT + 44 Freelance ca đêm tại Kho TC Nha Trang. 4 KCT còn lại bố trí 12-25 NV/kho."],
    ["1. KTC, KCT thuộc Vùng", "Xe & Bãi tại kho", "Số xe (tự thuê + điều động); diện tích bãi đỗ/bãi tập kết & sức chứa; phương án khi quá tải.", "Phân bổ 18 xe tải 8T-15T kết nối Linehaul. Bãi tập kết Kho Nha Trang 2,500m2 chứa max 65,000 đơn/thời điểm."],
    ["1. KTC, KCT thuộc Vùng", "Lịch làm việc", "Ca làm việc theo từng ngày (kể cả ngày đỉnh 08/08), người trực ca, kế hoạch tăng ca/phụ cấp.", "Ca 1 (07h-18h), Ca Đêm (20h-05h30 & 22h-06h). Tăng ca 100% ngày Peak 08/08, phụ cấp 150k/đêm/người."],
    ["2. Bưu cục CK", "Nhân sự cho Bưu cục CK", "Số lượng nhân sự cần có tại mỗi Bưu cục CK, nguồn tuyển hoặc điều động.", "Bố trí 8 NVXL chuyên trách hàng CK tại BC Nha Trang CK, Đức Trọng CK & Phan Thiết CK."],
    ["2. Bưu cục CK", "Xe cho Bưu cục CK", "Số xe cần bố trí cho mỗi Bưu cục CK, nguồn xe (tự thuê/điều động), lịch trình.", "Điều động 6 xe van 1.9T chạy 3 chuyến/ngày luân chuyển hàng CK ra tuyến chặng chót."]
]

for r_i, row in enumerate(t11_data, 1):
    for c_i, val in enumerate(row):
        bold = True if c_i in [0, 1] else False
        format_cell(t11_ktc.rows[r_i].cells[c_i], val, bold=bold, font_size=8)

# SAVE COMPANY GUIDELINE STANDARD DOCX
company_docx_paths = [
    r'C:\Users\lap4all\Downloads\NTB Kế hoạch Event 8.8_Company_Standard.docx',
    r'C:\Users\lap4all\Downloads\NTB Kế hoạch Event 8.8_v2.docx',
    r'c:\Users\lap4all\Documents\Auto report\NTB Kế hoạch Event 8.8.docx'
]

for p in company_docx_paths:
    try:
        doc.save(p)
        print(f"Successfully saved Official Company Standard Word Plan: {p}")
    except Exception as e:
        print(f"Could not save {p}: {e}")

print("Official Company Standard Plan Document generated successfully!")
