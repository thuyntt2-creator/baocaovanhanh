import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import re
import os

md_path = r"C:\Users\lap4all\Downloads\AOP_BCCK_Plan_NTB_new.md"
docx_path = r"C:\Users\lap4all\Downloads\AOP_BCCK_Plan_NTB_new.docx"

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="888888"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def build_docx():
    doc = docx.Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style fonts
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if not os.path.exists(md_path):
        print(f"MD file not found: {md_path}")
        return

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()
        
        # Check if we are inside a table
        if stripped.startswith('|'):
            in_table = True
            table_rows.append(stripped)
            continue
        elif in_table:
            # We just finished a table
            create_docx_table(doc, table_rows)
            in_table = False
            table_rows = []

        if not stripped:
            continue

        if stripped.startswith('# '):
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run(stripped[2:])
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(12)
        elif stripped.startswith('## '):
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run(stripped[3:])
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif stripped.startswith('### '):
            h = doc.add_paragraph()
            run = h.add_run(stripped[4:])
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x55, 0x88)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif stripped.startswith('#### '):
            h = doc.add_paragraph()
            run = h.add_run(stripped[5:])
            run.font.size = Pt(11)
            run.bold = True
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(4)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            # Parse simple bold markdown inside list item
            text = stripped[2:]
            parse_styled_text(p, text)
            p.paragraph_format.space_after = Pt(3)
        elif stripped.startswith('1. ') or stripped.startswith('2. ') or stripped.startswith('3. ') or stripped.startswith('4. ') or stripped.startswith('5. '):
            p = doc.add_paragraph(style='List Number')
            text = stripped[3:]
            parse_styled_text(p, text)
            p.paragraph_format.space_after = Pt(3)
        elif stripped.startswith('---'):
            doc.add_page_break()
        else:
            # Normal paragraph
            p = doc.add_paragraph()
            parse_styled_text(p, stripped)
            p.paragraph_format.space_after = Pt(6)

    # In case the file ended with a table
    if in_table and table_rows:
        create_docx_table(doc, table_rows)

    doc.save(docx_path)
    print(f"Docx saved to: {docx_path}")

def parse_styled_text(p, text):
    # Splits text by '**' to alternate bold and regular fonts
    parts = text.split('**')
    is_bold = False
    for part in parts:
        if part:
            run = p.add_run(part)
            if is_bold:
                run.bold = True
        is_bold = not is_bold

def create_docx_table(doc, table_lines):
    # Process lines
    matrix = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        matrix.append(cells)
    
    if not matrix:
        return
        
    # Check if row 1 is separator (contains dashes)
    if len(matrix) > 1 and all(re.match(r'^[-:\s]+$', cell) for cell in matrix[1]):
        header = matrix[0]
        data = matrix[2:]
    else:
        header = matrix[0]
        data = matrix[1:]

    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Format header
    hdr_cells = table.rows[0].cells
    for i in range(cols):
        hdr_cells[i].text = header[i]
        set_cell_background(hdr_cells[i], "003366")
        set_cell_margins(hdr_cells[i], top=120, bottom=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)
            
    # Add data rows
    for r_idx, row_data in enumerate(data):
        row = table.add_row()
        cells = row.cells
        
        # Strip row styling if row is a total row (contains 'tổng' or 'total')
        is_total = any('tổng' in str(cell).lower() or 'total' in str(cell).lower() for cell in row_data)
        bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
        if is_total:
            bg_color = "E6EEF4"
            
        for i in range(min(cols, len(row_data))):
            cells[i].text = row_data[i]
            set_cell_background(cells[i], bg_color)
            set_cell_margins(cells[i], top=90, bottom=90)
            p = cells[i].paragraphs[0]
            # If value is numeric, right align
            val = row_data[i].replace('.', '').replace(',', '').replace('%', '').strip()
            if val.isdigit() or (val.startswith('-') and val[1:].isdigit()):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif is_total:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            for run in p.runs:
                run.font.size = Pt(9.5)
                if is_total:
                    run.font.bold = True

    # Spacer paragraph after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

if __name__ == "__main__":
    build_docx()

