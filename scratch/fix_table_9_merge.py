import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_3.docx"

if not os.path.exists(doc_path):
    print("File không tồn tại")
    sys.exit(1)

doc = docx.Document(doc_path)
table = doc.tables[8]  # index 8 (Bảng 9)

print("=== XỬ LÝ LỖI MERGE CELL DÒNG TỔNG BẢNG 9 ===")

# Kiểm tra xem cell 3 và cell 4 có phải là một không
print("-> Bắt đầu xóa dòng TỔNG cũ và tái tạo dòng mới không bị gộp...")
ref_run = None
for cell in table.rows[4].cells:
    if cell.paragraphs and cell.paragraphs[0].runs:
        ref_run = cell.paragraphs[0].runs[0]
        break
        
# Xóa hàng cuối cùng bị gộp
tbl = table._tbl
tr = table.rows[5]._tr
tbl.remove(tr)
print("-> Đã xóa dòng TỔNG cũ")

# Thêm dòng mới
new_row = table.add_row()
print("-> Đã thêm dòng mới")

# Điền giá trị cho dòng mới
vals = ["", "TỔNG", "2.073", "630", "75.600.000 đ/th", "", "630", ""]
for idx, val in enumerate(vals):
    cell = new_row.cells[idx]
    cell.text = val
    if val:
        p = cell.paragraphs[0]
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT if idx in [2, 3, 4, 6] else docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.bold = True
            if ref_run:
                run.font.name = ref_run.font.name
                run.font.size = ref_run.font.size
                
print("-> Đã điền xong dữ liệu và định dạng dòng TỔNG mới")

doc.save(doc_path)
print(f"=== ĐÃ LƯU FILE THÀNH CÔNG: {doc_path} ===")
