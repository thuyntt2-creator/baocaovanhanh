import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_2.docx"

if not os.path.exists(doc_path):
    print("File không tồn tại")
    sys.exit(1)

doc = docx.Document(doc_path)

print(f"=== ĐỌC TẤT CẢ CÁC BẢNG TRONG FILE: {os.path.basename(doc_path)} ===")
print(f"Số lượng bảng: {len(doc.tables)}")

for t_idx, table in enumerate(doc.tables):
    print(f"\n--- BẢNG {t_idx+1} ---")
    for r_idx, row in enumerate(table.rows):
        row_text = []
        for cell in row.cells:
            row_text.append(cell.text.strip().replace('\n', ' '))
        cleaned_text = []
        for text in row_text:
            if not cleaned_text or cleaned_text[-1] != text:
                cleaned_text.append(text)
        print(f"  Row {r_idx+1:02d}: {cleaned_text}")
