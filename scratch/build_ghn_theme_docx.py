import sys, docx, pandas as pd, json, os, re
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

# Source Template Path (contains GHN header & footer banners, page setup)
tmpl_path = r'C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi.docx'
doc = docx.Document(tmpl_path)

web_data_dir = r'c:\Users\lap4all\Documents\Auto report\scratch\web_data'
rezone_file = os.path.join(web_data_dir, 'rezone.json')
hubs_file = os.path.join(web_data_dir, 'hubs.json')
wards_file = os.path.join(web_data_dir, 'wards.json')
excel_path = r'C:\Users\lap4all\Downloads\NTB_Phan_Tuyen_Hanh_Chinh_Quy_Hoach_Moi.xlsx'

with open(rezone_file, 'r', encoding='utf-8') as f:
    rezone_data = json.load(f)

with open(hubs_file, 'r', encoding='utf-8') as f:
    hubs_data = json.load(f)

with open(wards_file, 'r', encoding='utf-8') as f:
    wards_data = json.load(f)

ward_stats_dict = {}
for w in wards_data:
    code = str(w.get('ward_code', ''))
    ward_stats_dict[code] = {
        'pv': w.get('pv', 0.0),
        'pw': w.get('pw', 0.0),
        'dv': w.get('dv', 0.0),
        'dw': w.get('dw', 0.0),
    }

df_excel = pd.read_excel(excel_path, sheet_name='Sheet1')
excel_wards = df_excel.groupby('Mã Xã mới')

maps_dir = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\d95894ad-9248-47c3-872b-72d7de66bc20\maps'

# Save original tables
tables_saved = [doc.tables[0], doc.tables[1]]

# Clear all paragraphs in document body
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)

# Clear all tables from body
for t in list(doc.tables):
    t._element.getparent().remove(t._element)

web_wards = rezone_data['new_wards']
web_ward_dict = {}
for w in web_wards:
    web_ward_dict[w['name'].lower().strip()] = w

am_proposals_and_reasons = {
    24823: {
        'prop': "Gộp toàn bộ Phường 1 và Phường B'Lao về Bưu cục (LDO) B'Lao Mới. ĐÓNG CỬA bưu cục (LDO) 1 Bảo Lộc cũ.",
        'reason': '''Theo quy hoạch AM Bảo Lộc:
- Bưu cục (LDO) B'Lao Mới được thành lập làm bưu cục trung tâm phụ trách toàn bộ Phường 1 (750 đơn giao full) và Phường B'Lao (600 đơn giao full), tổng giao 1,500 đơn/ngày, lấy 1,000 đơn/ngày, định biên 15 NVPTTT + 2 NVXL.
- Bưu cục (LDO) 1 Bảo Lộc cũ không đủ m² chứa hàng và trùng tuyến nên đề xuất ĐÓNG CỬA, điều chuyển toàn bộ lực lượng lao động về BC B'Lao Mới.'''
    },
    24820: {
        'prop': "Gộp toàn bộ Phường 2 và Phường 3 về Bưu cục (LDO) 3 Bảo Lộc.",
        'reason': '''Theo quy hoạch AM Bảo Lộc: Bưu cục (LDO) 3 Bảo Lộc phụ trách Phường 2 (800 đơn giao full) và Phường 3 (600 đơn giao full), tổng giao 1,400 đơn/ngày, lấy 500 đơn/ngày, định biên 14 NVPTTT + 2 NVXL. Tối ưu tuyến đường giao nhận khu vực Bảo Lộc.'''
    },
    24817: {
        'prop': "Gộp Phường B'Lao và Phường 1 về Bưu cục (LDO) B'Lao Mới.",
        'reason': '''Theo quy hoạch AM Bảo Lộc: BC (LDO) B'Lao Mới phụ trách Phường B'Lao (600 đơn/ngày) và Phường 1 (750 đơn/ngày), tổng sản lượng 2,500 đơn giao lấy/ngày với 15 NVPTTT + 2 NVXL.'''
    },
    25084: {
        'prop': "Gộp tuyến về Bưu cục (LDO) Bảo Lâm 3.",
        'reason': '''AM Bảo Lâm đề xuất: Bưu cục (LDO) Bảo Lâm 3 phụ trách Bảo Lâm 2 (350 đơn full) và Bảo Lâm 3 (300 đơn full), tổng giao 700 đơn/ngày, lấy 30 đơn/ngày, định biên 8 NVPTTT + 1 NVXL. Bưu cục (LDO) Bảo Lâm 1 phụ trách Bảo Lâm 1 (450 đơn) và Bảo Lâm 4 (100 đơn).'''
    },
    24781: {
        'prop': "TÁCH MỚI Bưu cục (LDO) Xuân Hương - Đà Lạt 2 (tại Phường 10) & GIỮ NGUYÊN Bưu cục Xuân Hương cũ (phụ trách Phường 1 & Phường 2).",
        'reason': '''Sản lượng Phường Xuân Hương cực lớn (hơn 2,850 đơn/ngày). Địa hình bị chia cắt bởi Hồ Xuân Hương, thời tiết mưa lạnh cuối năm.
- BC Xuân Hương cũ: Cover Phường 1 (400 đơn), Phường 2 (400 đơn), Phường 4 (600 đơn), Lấy 400 đơn, 13 NVPTTT + 1 NVXL.
- BC Xuân Hương 2 (Mới tại Phường 10): Cover Phường 10 (400 đơn), Phường 3 (500 đơn), Lấy 150 đơn, 8 NVPTTT + 1 NVXL.'''
    },
    24778: {
        'prop': "GIỮ NGUYÊN 02 BƯU CỤC (Lâm Viên - Đà Lạt 1 & Lâm Viên - Đà Lạt 2).",
        'reason': '''Phường Lâm Viên có địa hình bị chia cắt bởi Hồ Xuân Hương. Duy trì 02 Bưu cục (Lâm Viên 1 & Lâm Viên 2) nằm ngay khu vực Phường 8 giúp nhân viên giao nhận tập trung xử lý theo phân khu, giảm quãng đường chạy rỗng.'''
    },
    24958: {
        'prop': "GIỮ NGUYÊN 02 BƯU CỤC ((LDO) Đức Trọng 1 & (LDO) Đức Trọng 2).",
        'reason': '''AM Đức Trọng nêu rõ nguyên nhân giữ 2 BC: Khoảng cách giữa 2 BC trên 15km, địa bàn rất rộng, nhân sự cực kỳ khó tuyển dụng (đặc biệt mùa cà phê và mùa mưa cuối năm), kho bãi cũ diện tích nhỏ hẹp không đủ m² chứa hàng. Nếu gộp lại 1 bưu cục sẽ dẫn đến di chuyển quá xa, trễ checkin và vỡ tuyến giao.'''
    },
    25000: {
        'prop': "TÁCH BƯU CỤC HÀNG NHỎ / HÀNG VỪA (Xã Đinh Trang Thượng) để chia tải cho BC Di Linh.",
        'reason': '''AM Di Linh nêu nguyên nhân: BC Di Linh phụ trách bán kính xa 22-45km (Gia Hiệp, Đinh Trang Thượng, Sơn Điền, Liên Đầm). Đề xuất tách BC Hàng Nhỏ (300-350 giao, 10-20 lấy, 4 NVPTTT cover Đinh Trang Thượng, Di Linh) và BC Hàng Vừa (4 NVPTTT cover Đinh Trang Thượng, Phúc Thọ Lâm Hà, Di Linh, Liên Đầm).'''
    },
    23235: {
        'prop': "Gộp về Bưu cục (BTH) Phước Hội (Mới).",
        'reason': '''AM La Gi quy hoạch: (BTH) Phước Hội (Mới) cover Phường Lagi (600-700 đơn), Phước Hội (250-300 đơn), Sơn Mỹ (200-250 đơn), Lấy 280 đơn, định biên 15 NVPTTT + 1 NVXL. Tối ưu tuyến đường giao nhận nội thị La Gi.'''
    },
    23143: {
        'prop': "Gộp về Bưu cục (BTH) Tân Hải (Mới). Giữ nguyên phần Thuận Quý cũ cho BC Hàm Thuận Nam.",
        'reason': '''AM La Gi & Hàm Thuận Nam giải thích nguyên nhân: BC Tân Hải (Mới) cover Xã Tân Hải và Tân Thành (Tân Thuận, Tân Thành cũ), Vol giao 400-500 đơn, lấy 120 đơn, 7 NVPTTT + 1 NVXL. Giữ nguyên phần Thuận Quý cũ cho BC (BTH) Hàm Thuận Nam vì địa hình đặc thù, khoảng cách từ Hàm Thuận Nam gần hơn và dân cư tập trung đông đúc hơn.'''
    },
    22972: {
        'prop': "GIỮ NGUYÊN 02 BƯU CỤC (Phan Rí Cửa & Liên Hương).",
        'reason': '''AM Tuy Phong giải thích nguyên nhân: Vị trí 2 Bưu cục đặt tại 2 thị trấn có sản lượng hàng nhiều nhất (Phan Rí Cửa 460 đơn, Liên Hương 600 đơn). Khoảng cách giữa 2 thị trấn tầm 25km và địa bàn rộng nên KHÔNG THỂ gộp lại thành 1 Bưu cục.'''
    },
    22528: {
        'prop': "GIỮ NGUYÊN 02 BƯU CỤC (Ninh Hòa 1 & Ninh Hòa 2).",
        'reason': '''AM Ninh Hòa nêu nguyên nhân: Huyện Ninh Hòa cũ có diện tích lớn nhất cả nước (20 xã/phường). Bưu cục Ninh Hòa 1 được tách ra để cover các xã xa trung tâm, kéo nguồn lực nhân sự tại chỗ. Diện tích kho Ninh Hòa 2 nhỏ hẹp không đảm bảo m² để gom gộp.'''
    },
    22759: {
        'prop': "Gộp về Bưu cục chính (NTH) Phan Rang & TÁCH MỚI Bưu cục Đông Hải.",
        'reason': '''AM Ninh Thuận quy hoạch: BC Phan Rang cover Phường Phan Rang và Huyện Thuận Bắc. Đề xuất TÁCH MỚI Bưu cục Đông Hải (cover Đông Hải, Mỹ Bình, Mỹ Đông, Mỹ Hải, Giao 600, Lấy 250, 7 NVPTTT + 1 NVXL) vì Phường Đông Hải địa bàn khó tuyển dụng nhân sự, thường xuyên thiếu người.'''
    },
    24748: {
        'prop': "GIỮ NGUYÊN 02 BƯU CỤC ((DNO) Quảng Tín & (DNO) Kiến Đức).",
        'reason': '''AM Đắk Nông chỉ đạo nguyên nhân: Xã Quảng Tân mới (gồm Xã Đắk Ngo & Quảng Tân cũ) địa hình rất rộng, bán kính xa và bị chia cắt mạnh bởi đồi núi giữa Đắk Ngo và Quảng Tân. Kiến Đức cover Đắk Ngo (109 đơn), Quảng Tín cover Quảng Tân (92 đơn). Nếu gộp 1 BC shipper di chuyển gấp đôi, trễ SLA và nguy cơ bỏ tuyến xã xa. Đề xuất GIỮ NGUYÊN 2 BƯU CỤC.'''
    },
    24611: {
        'prop': "TẠM THỜI GIỮ NGUYÊN 100% PHẠM VI QUẢN LÝ THEO BƯU CỤC CŨ (trong 3-6 tháng).",
        'reason': '''AM Đắk Nông cảnh báo nguy cơ vỡ tuyến: Lịch sử Gia Nghĩa từng bể tuyến triền miên. Phương án tách 3 kho (Bắc, Nam Gia Nghĩa, Nhân Cơ) giúp giữ ổn định vận hành. 100% nhân sự tại chỗ các xã ven (như Xã Đắk Ha) xác nhận sẽ NGHỈ VIỆC nếu chuyển kho mới xa 25-30km. Khu vực có tỷ lệ đồng bào DTTS cao cực khó tuyển mới. AM đề xuất TẠM GIỮ NGUYÊN 100% các kho cũ trong 3-6 tháng.'''
    },
    24615: {
        'prop': "TẠM THỜI GIỮ NGUYÊN 100% PHẠM VI QUẢN LÝ THEO BƯU CỤC CŨ (trong 3-6 tháng).",
        'reason': '''AM Đắk Nông đề xuất: Giữ nguyên phân vùng giao của BC Nam Gia Nghĩa 2 và Bắc Gia Nghĩa để tránh nguy cơ gãy tuyến và biến động nhân sự tại chỗ.'''
    },
    22366: {
        'prop': "Gộp phân vùng về Bưu cục (KHO) Nha Trang chính, giữ Nam Nha Trang 1 & 2 phụ trách phân đoạn phụ.",
        'reason': '''AM Nha Trang quy hoạch: Sản lượng Phường Nha Trang rất lớn (1,592 đơn/ngày). Gộp phân vùng chính về BC KHO Nha Trang, đồng thời duy trì các bưu cục Nam Nha Trang 1 & 2 phụ trách phân đoạn phụ để tránh quá tải kho bãi và vỡ tuyến.'''
    },
    22402: {
        'prop': "Gộp về Bưu cục (KHO) Nam Nha Trang 1 Mới & Giữ Bưu cục Nam Nha Trang 5. ĐÓNG CỬA Nam Nha Trang 2 & 3.",
        'reason': '''AM Nha Trang quy hoạch: Sản lượng Phường Nam Nha Trang cực lớn (2,418 đơn/ngày). Giữ Bưu cục Nam Nha Trang 1 mới (gộp Nam Nha Trang 3 và phần Nam Nha Trang 1 cũ) và giữ nguyên Bưu cục Nam Nha Trang 5 (phụ trách Phước Đồng). Đóng/bỏ bưu cục Nam Nha Trang 2 & Nam Nha Trang 3 để tối ưu điểm tập kết.'''
    }
}

ward_maps = {
    24781: ('image6.png', 'Bản đồ địa bàn quy hoạch mới TP. Đà Lạt 2026'),
    24784: ('image6.png', 'Bản đồ địa bàn quy hoạch mới TP. Đà Lạt 2026'),
    25000: ('image14.png', 'Bản đồ địa bàn quy hoạch khu vực Di Linh 2026'),
    24958: ('image13.png', 'Bản đồ địa bàn quy hoạch khu vực Đức Trọng 2026'),
    24823: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    24820: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    24817: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    25084: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    22366: ('image5.png', 'Bản đồ địa bàn quy hoạch mới TP. Nha Trang 2026'),
    22402: ('image5.png', 'Bản đồ địa bàn quy hoạch mới TP. Nha Trang 2026'),
    22390: ('image5.png', 'Bản đồ địa bàn quy hoạch mới TP. Nha Trang 2026'),
    22528: ('image8.jpg', 'Bản đồ địa bàn quy hoạch mới khu vực Ninh Hòa 2026'),
    22516: ('image7.jpg', 'Bản đồ địa bàn quy hoạch khu vực Vạn Ninh - Tu Bông 2026'),
    22933: ('image10.jpg', 'Bản đồ địa bàn quy hoạch mới TP. Phan Thiết & Hàm Thắng 2026'),
    22945: ('image10.jpg', 'Bản đồ địa bàn quy hoạch mới TP. Phan Thiết 2026'),
    22960: ('image11.jpg', 'Bản đồ phân vùng quy hoạch mới Phường Bình Thuận 2026'),
    23131: ('image9.png', 'Bản đồ quy hoạch mới khu vực Phường La Gi 2026'),
    23143: ('image9.png', 'Bản đồ quy hoạch mới khu vực Xã Tân Thành 2026'),
    22972: ('image4.png', 'Bản đồ quy hoạch mới khu vực Phan Rí Cửa & Tuy Phong 2026'),
    22759: ('image12.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Phan Rang 2026'),
    22834: ('image12.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Ninh Chử 2026'),
}

report_wards = []

for code, group in excel_wards:
    prov = group['Tỉnh, thành phố mới'].iloc[0]
    name = group['Tên Xã mới'].iloc[0]
    n_bcs = group['Số BC'].iloc[0]
    old_communes = group['Tên Xã cũ'].tolist()
    
    web_item = web_ward_dict.get(name.lower().strip(), None)
    if not web_item:
        for k, v in web_ward_dict.items():
            if name.lower().strip() in k or k in name.lower().strip():
                web_item = v
                break

    if web_item:
        sys_dem = web_item.get('dem', group['TỔNG ĐƠN/NGÀY (Phường mới)'].iloc[0])
        sys_kg = web_item.get('dem_kg', group['TỔNG KG/NGÀY (Phường mới)'].iloc[0])
        sys_status = web_item.get('status', 'split')
        assigned_bc = web_item.get('assigned_bc_name', group['Đánh giá & Phương án đề xuất'].iloc[0])
        cands = web_item.get('candidates', [])
        olds_web = web_item.get('olds', [])
    else:
        sys_dem = group['TỔNG ĐƠN/NGÀY (Phường mới)'].iloc[0]
        sys_kg = group['TỔNG KG/NGÀY (Phường mới)'].iloc[0]
        sys_status = 'split' if n_bcs > 1 else 'clean'
        assigned_bc = group['Đánh giá & Phương án đề xuất'].iloc[0]
        cands = []
        olds_web = []

    total_pv = 0.0
    total_pw = 0.0
    olds_enhanced = []
    
    if olds_web:
        for o in olds_web:
            o_code = str(o.get('ward', ''))
            st = ward_stats_dict.get(o_code, {'pv': 0.0, 'pw': 0.0, 'dv': o.get('dem', 0.0), 'dw': o.get('dem_kg', 0.0)})
            pv_item = st['pv']
            pw_item = st['pw']
            dv_item = o.get('dem', st['dv'])
            dw_item = o.get('dem_kg', st['dw'])
            total_pv += pv_item
            total_pw += pw_item
            olds_enhanced.append({
                'name': o['name'], 'bc_name': o.get('bc_name', 'Chưa rõ'),
                'dv': dv_item, 'dw': dw_item, 'pv': pv_item, 'pw': pw_item
            })

    bc_details = []
    if cands:
        for c in cands:
            bc_details.append({
                'name': c['bc_name'], 'id': c.get('bc', ''), 'dem': c['dem'], 'pct': c['share']
            })
    else:
        for bc_name, bc_group in group.groupby('Tên Bưu cục giao'):
            bc_total = bc_group['Sản lượng giao/ngày (đơn)'].sum() + bc_group['Sản lượng lấy/ngày (đơn)'].sum()
            pct = round((bc_total / sys_dem * 100), 1) if sys_dem > 0 else 0.0
            bc_id = bc_group['ID Bưu cục giao'].iloc[0]
            bc_details.append({
                'name': bc_name, 'id': bc_id, 'dem': bc_total, 'pct': pct
            })
    bc_details.sort(key=lambda x: x['dem'], reverse=True)

    if code in am_proposals_and_reasons:
        prop = am_proposals_and_reasons[code]['prop']
        final_reason = am_proposals_and_reasons[code]['reason']
    else:
        prop = f"Quy hoạch Bưu cục phụ trách chính: {assigned_bc}"
        final_reason = group['Lý do & Bố trí nhân sự'].iloc[0]

    img_info = ward_maps.get(code, (None, None))

    report_wards.append({
        'code': code, 'name': name, 'prov': prov, 'n_bcs': n_bcs,
        'old_communes': old_communes, 'olds_web': olds_enhanced,
        'sys_dem': sys_dem, 'sys_kg': sys_kg, 'sys_pv': total_pv, 'sys_pw': total_pw,
        'sys_status': sys_status, 'assigned_bc': assigned_bc, 'prop': prop,
        'bc_details': bc_details, 'reason': final_reason,
        'img_file': img_info[0], 'img_caption': img_info[1]
    })

report_wards.sort(key=lambda x: x['name'])

# Helper functions to build Word elements WITH BUILT-IN STYLES FOR GOOGLE DOCS OUTLINE
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(230, 81, 0) # GHN Orange/Gold

def add_h1(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102) # GHN Dark Navy Blue

def add_h2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    run.font.name = 'Arial'
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153) # GHN Blue

def add_bullet_item(doc, label_bold, value_text, indent_level=1, italic_val=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 * indent_level)
    
    r_bullet = p.add_run("❖ ")
    r_bullet.font.name = 'Arial'
    r_bullet.font.size = Pt(9.5)
    r_bullet.font.color.rgb = RGBColor(0, 102, 153)

    if label_bold:
        r_lbl = p.add_run(label_bold + " ")
        r_lbl.font.name = 'Arial'
        r_lbl.font.size = Pt(10)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(34, 34, 34)

    if value_text:
        r_val = p.add_run(value_text)
        r_val.font.name = 'Arial'
        r_val.font.size = Pt(10)
        r_val.font.italic = italic_val
        r_val.font.color.rgb = RGBColor(51, 51, 51)

def add_sub_bullet_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.5)
    
    r_dash = p.add_run("- ")
    r_dash.font.name = 'Arial'
    r_dash.font.size = Pt(9.5)
    r_dash.font.color.rgb = RGBColor(100, 100, 100)

    r_txt = p.add_run(text)
    r_txt.font.name = 'Arial'
    r_txt.font.size = Pt(9.5)
    r_txt.font.color.rgb = RGBColor(51, 51, 51)

def add_paragraph_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 51, 51)

def set_table_styling(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(r'<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        if r_idx == 0:
            trPr.append(parse_xml(r'<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="003366"/>'))
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = 'Arial'
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(9)
        else:
            bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{bg_color}"/>'))
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Arial'
                        r.font.size = Pt(8.5)

# NOW BUILD THE GHN THEME DOCUMENT CONTENT WITH HEADING STYLES FOR GOOGLE DOCS OUTLINE
add_title(doc, 'RÀ SOÁT VÀ QUY HOẠCH BƯU CỤC THEO ĐƠN VỊ HÀNH CHÍNH MỚI (VÙNG NTB)')

# Section I
add_h1(doc, 'I. TỔNG QUAN HIỆN TRẠNG MẠNG LƯỚI BƯU CỤC VÙNG NTB')
add_bullet_item(doc, 'Tổng số Bưu cục Express đang vận hành:', '83 Bưu cục (thuộc các tỉnh Khánh Hòa, Ninh Thuận, Bình Thuận, Lâm Đồng, Đắc Nông).')
add_bullet_item(doc, 'Tổng số Xã/Phường hành chính mới rà soát:', '36 Xã/Phường mới (gồm 114 xã/phường cũ sáp nhập).')
add_bullet_item(doc, 'Số Xã/Phường mới đã quy hoạch chuẩn (01 BC phụ trách):', f'{sum(1 for w in report_wards if w["sys_status"]=="clean")} Xã/Phường ({sum(1 for w in report_wards if w["sys_status"]=="clean")/len(report_wards)*100:.1f}%).')
add_bullet_item(doc, 'Số Xã/Phường mới bị CHIA CẮT (2-3 BC phụ trách):', f'{sum(1 for w in report_wards if w["sys_status"]=="split")} Xã/Phường ({sum(1 for w in report_wards if w["sys_status"]=="split")/len(report_wards)*100:.1f}%).')
add_bullet_item(doc, 'Tổng sản lượng giao & lấy toàn vùng NTB (Số liệu Web):', f'Giao: {sum(w["sys_dem"] for w in report_wards):,.1f} đơn/ngày ({sum(w["sys_kg"] for w in report_wards):,.1f} kg/ngày) | Lấy: {sum(w["sys_pv"] for w in report_wards):,.1f} đơn/ngày ({sum(w["sys_pw"] for w in report_wards):,.1f} kg/ngày).')

# Section II
add_h1(doc, 'II. RÀ SOÁT THEO ĐƠN VỊ HÀNH CHÍNH MỚI')
add_h2(doc, '1. Đánh giá độ phủ Bưu cục và ranh giới hành chính mới')
add_paragraph_text(doc, 'Sau khi sáp nhập các xã cũ thành xã/phường mới, ranh giới quản lý của các Bưu cục đang xuất hiện hiện tượng chia cắt địa bàn, dẫn đến:')
add_bullet_item(doc, '01 Phường mới có tới 3 Bưu cục cùng giao hàng:', 'Gây chồng chéo tuyến đường, làm shipper di chuyển cắt ngang địa bàn của nhau.')
add_bullet_item(doc, 'Tuyến đi chéo xa ranh giới:', 'Một số xã vùng ven bị giao từ Bưu cục ở huyện/tỉnh khác cách xa 30 - 40 km, trong khi Bưu cục lân cận chỉ cách 7 - 15 km.')

add_h2(doc, '2. Danh sách 18 Tuyến giao chéo xa ranh giới cần Reassign ngay')

# Insert Table 1 (18 cross-boundary routes)
t1_src = tables_saved[1]
t1_doc = doc.add_table(rows=len(t1_src.rows), cols=10)
for r_idx, r in enumerate(t1_src.rows):
    cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
    if r_idx == 0:
        row_data = cells + ['Đề xuất AM']
    else:
        xa = cells[1]
        prov = cells[2]
        bc_from = cells[3]
        bc_to = cells[5]
        
        p_from = bc_from.split(')')[0].replace('(', '').strip()[:2]
        p_to = bc_to.split(')')[0].replace('(', '').strip()[:2]
        
        if xa == "Xã Đa Mi":
            am_note = f"GIỮ NGUYÊN {bc_from} (Sai ranh giới tỉnh BTH->LDO)"
        elif p_from != p_to:
            am_note = f"GIỮ NGUYÊN {bc_from} (Tuyến khác tỉnh {p_from}->{p_to})"
        elif "Đắk Nông" in prov or "Lâm Đồng" in prov:
            am_note = f"GIỮ NGUYÊN {bc_from} (Tránh vỡ tuyến DNO/LDO)"
        else:
            am_note = "Reassign tối ưu"
        row_data = cells + [am_note]
    for c_idx, val in enumerate(row_data):
        t1_doc.cell(r_idx, c_idx).text = val
set_table_styling(t1_doc)

# Section III
add_h1(doc, 'III. ĐÁNH GIÁ CHI TIẾT SẢN LƯỢNG VÀ ĐỀ XUẤT PHƯƠNG ÁN GỘP / GIỮ BƯU CỤC')
add_h2(doc, 'Danh sách 36 Phường/Xã mới rà soát')

for i, w in enumerate(report_wards, 1):
    # Use built-in Heading 2 style so Google Docs outline detects every single ward!
    add_h2(doc, f'{i}. {w["name"]} ({w["prov"]})')
    add_bullet_item(doc, 'Mã Xã mới:', str(w['code']))
    
    vol_str = f'Giao: {w["sys_dem"]:.1f} đơn/ngày ({w["sys_kg"]:.1f} kg/ngày)'
    if w['sys_pv'] > 0:
        vol_str += f' · Lấy: {w["sys_pv"]:.1f} đơn/ngày ({w["sys_pw"]:.1f} kg/ngày)'
    vol_str += f' | Trạng thái Web: {w["sys_status"].upper()}'
    
    add_bullet_item(doc, 'TỔNG SẢN LƯỢNG GIAO & LẤY (SỐ LIỆU WEB):', vol_str)
    
    bcs_str = ', '.join([f'"{b["name"]}"' for b in w['bc_details']])
    add_bullet_item(doc, f'Các BC hiện phụ trách ({len(w["bc_details"])} BC):', bcs_str)
    
    if w['olds_web']:
        old_strs = []
        for o in w['olds_web']:
            s_o = f"{o['name']} (Giao: {o['dv']:.0f} đơn"
            if o['pv'] > 0:
                s_o += f", Lấy: {o['pv']:.0f} đơn"
            s_o += f" - BC: {o['bc_name']})"
            old_strs.append(s_o)
        add_bullet_item(doc, 'Các xã cũ sáp nhập & Sản lượng Giao/Lấy từng xã (Web):', '; '.join(old_strs))
    else:
        old_str = ', '.join(w['old_communes'])
        add_bullet_item(doc, 'Các xã cũ sáp nhập:', old_str)

    add_bullet_item(doc, 'Tỷ lệ phân chia sản lượng thực tế (Web):', '')
    for b in w['bc_details']:
        add_sub_bullet_item(doc, f'{b["name"]} (ID: {b["id"]}): {b["dem"]:.1f} đơn/ngày ({b["pct"]}%)')

    add_bullet_item(doc, 'ĐỀ XUẤT PHƯƠNG ÁN CỦA AM:', w['prop'])
    add_bullet_item(doc, 'LÝ DO VÀ GIẢI THÍCH NGUYÊN NHÂN CHI TIẾT TỪ AM:', '')
    
    reason_lines = w['reason'].strip().split('\n')
    for rl in reason_lines:
        if rl.strip():
            p_r = doc.add_paragraph()
            p_r.paragraph_format.space_before = Pt(1)
            p_r.paragraph_format.space_after = Pt(2)
            p_r.paragraph_format.left_indent = Inches(0.4)
            r_r = p_r.add_run(rl.strip())
            r_r.font.name = 'Arial'
            r_r.font.size = Pt(9.5)
            r_r.font.color.rgb = RGBColor(68, 68, 68)

    if w['img_file']:
        img_abs = os.path.join(maps_dir, w['img_file'])
        if os.path.exists(img_abs):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after = Pt(2)
            run_img = p_img.add_run()
            run_img.add_picture(img_abs, width=Inches(5.5))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(8)
            r_cap = p_cap.add_run(f'Hình {i}: {w["img_caption"]}')
            r_cap.font.name = 'Arial'
            r_cap.font.italic = True
            r_cap.font.size = Pt(9)
            r_cap.font.color.rgb = RGBColor(120, 120, 120)

# Section IV
add_h1(doc, 'IV. YÊU CẦU 3: ĐÁNH GIÁ NHU CẦU TÁCH BƯU CỤC VÀ TỐI ƯU CÔNG SUẤT KHO BÃI')
add_h2(doc, '1. Đánh giá Bưu cục quá tải (Cần tách bớt Phường hoặc Mở rộng diện tích)')

# Insert Table 0 (Overloaded BCs)
t0_src = tables_saved[0]
t0_doc = doc.add_table(rows=len(t0_src.rows), cols=len(t0_src.columns))
for r_idx, r in enumerate(t0_src.rows):
    for c_idx, c in enumerate(r.cells):
        t0_doc.cell(r_idx, c_idx).text = c.text.strip().replace('\n', ' ')
set_table_styling(t0_doc)

add_h2(doc, '2. Kế hoạch Mở mới / Tách Bưu cục & Tối ưu Mạng lưới')
add_bullet_item(doc, '1. Mở mới Bưu cục (LDO) Xuân Hương - Đà Lạt 2:', 'Đặt tại khu vực Phường 10 (TP. Đà Lạt), chia tải cho Bưu cục Xuân Hương cũ (phụ trách Phường 3 & Phường 10, Vol giao: 1,050 đơn/ngày, Vol lấy: 150 đơn/ngày, 8 NVPTTT + 1 NVXL).')
add_bullet_item(doc, '2. Tách Bưu cục Hàng Nhỏ / Hàng Vừa Di Linh (Xã Đinh Trang Thượng):', 'Phụ trách Đinh Trang Thượng, Di Linh, Phúc Thọ Lâm Hà, Liên Đầm để giảm bán kính di chuyển 22-45km cho BC Di Linh.')
add_bullet_item(doc, '3. Mở mới Bưu cục Đông Hải (Tỉnh Ninh Thuận):', 'Cover khu vực ven biển Đông Hải, tối ưu điểm tập kết sản lượng lấy (600 đơn giao, 250 đơn lấy, 7 NVPTTT + 1 NVXL).')
add_bullet_item(doc, '4. Mở mới Bưu cục (LDO) B\'Lao Mới (Bảo Lộc):', 'Cover Phường 1 & Phường B\'Lao (1,500 đơn giao, 1,000 đơn lấy, 15 NVPTTT + 2 NVXL), đóng cửa BC (LDO) 1 Bảo Lộc cũ.')

# Section V
add_h1(doc, 'V. TỔNG HỢP BIẾN ĐỘNG MẠNG LƯỚI BƯU CỤC NTB 2026')
add_bullet_item(doc, 'Bưu cục Mở mới (04 BC):', 'BC Xuân Hương - Đà Lạt 2, BC Di Linh Hàng Nhỏ, BC Đông Hải, BC B\'Lao Mới.')
add_bullet_item(doc, 'Bưu cục Đóng cửa (02 BC):', 'BC 1 Bảo Lộc, BC Nam Nha Trang 3.')
add_bullet_item(doc, 'Bưu cục Gộp tuyến/Điều chỉnh phân vùng (19 BC):', 'Gom các tuyến xã lẻ về BC trung tâm theo ĐVHC mới.')
add_bullet_item(doc, 'Bưu cục Giữ nguyên vận hành (13 BC):', 'Duy trì tại các khu vực đặc thù Gia Nghĩa, Đắc Nông, Đà Lạt, Đức Trọng, Phan Rí Cửa, Ninh Hòa, Xã Đa Mi (Bình Thuận)...')

# Save output DOCX
docx_out_path = r'C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Co_Muc_Luc_Google_Docs.docx'
doc.save(docx_out_path)
print(f'Saved DOCX report WITH GOOGLE DOCS OUTLINE STYLES to: {docx_out_path}')
