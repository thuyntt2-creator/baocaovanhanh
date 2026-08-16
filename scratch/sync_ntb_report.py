import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

src_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_1.docx"
dst_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_2.docx"

if not os.path.exists(src_path):
    print(f"File nguồn không tồn tại: {src_path}")
    sys.exit(1)

doc = docx.Document(src_path)

print("=== BẮT ĐẦU CẬP NHẬT CÁC BẢNG TRONG WORD ===")

# --- 1. CẬP NHẬT BẢNG 6 (Index 5) ---
# SL/ngày — Hàng nặng (BQ) ở hàng 5 (index 5)
table_6 = doc.tables[5]
sl_bq_vals = ['948', '1,036', '1,465', '1,534', '1,830', '1,879']
for c_idx in range(1, 7):
    table_6.rows[5].cells[c_idx].text = sl_bq_vals[c_idx - 1]
print("-> Đã cập nhật Bảng 6")

# --- 2. CẬP NHẬT BẢNG 7 (Index 6) ---
table_7 = doc.tables[6]
# Sửa chi tiết NV giao từng bưu cục ở hàng 2->5 cho khớp tổng mới
nt_giao = ['44 (SL: 935)', '44 (SL: 987)', '70 (SL: 1245)', '72 (SL: 1317)', '80 (SL: 1550)', '82 (SL: 1604)']
dl_giao = ['16 (SL: 273)', '16 (SL: 290)', '18 (SL: 313)', '20 (SL: 344)', '24 (SL: 393)', '24 (SL: 421)']
dd_giao = ['8 (SL: 139)', '10 (SL: 152)', '10 (SL: 165)', '10 (SL: 183)', '8 (SL: 209)', '8 (SL: 224)']
dul_giao = ['8 (SL: 111)', '8 (SL: 123)', '10 (SL: 140)', '8 (SL: 158)', '8 (SL: 181)', '8 (SL: 194)']
for c_idx in range(1, 7):
    table_7.rows[2].cells[c_idx].text = nt_giao[c_idx - 1]
    table_7.rows[3].cells[c_idx].text = dl_giao[c_idx - 1]
    table_7.rows[4].cells[c_idx].text = dd_giao[c_idx - 1]
    table_7.rows[5].cells[c_idx].text = dul_giao[c_idx - 1]

# Sửa Tổng NV giao ở hàng 6
tot_giao = ['76', '78', '108', '110', '120', '122']
for c_idx in range(1, 7):
    table_7.rows[6].cells[c_idx].text = tot_giao[c_idx - 1]

# Sửa dòng 8 (NV kho & QL) thành "NV xử lý kho (46) + NV QL (4)"
table_7.rows[8].cells[0].text = "NV xử lý kho (46) + NV QL (4)"
for c_idx in range(1, 7):
    table_7.rows[8].cells[c_idx].text = "50"

# Sửa dòng 9 (Tổng nhân sự toàn vùng)
tot_staff = ['126', '128', '158', '160', '170', '172']
for c_idx in range(1, 7):
    table_7.rows[9].cells[c_idx].text = tot_staff[c_idx - 1]
print("-> Đã cập nhật Bảng 7")

# --- 3. CẬP NHẬT BẢNG 8 (Index 7) ---
table_8 = doc.tables[7]
# Dòng 1: SL hàng nặng/ngày (BQ)
for c_idx in range(1, 7):
    table_8.rows[1].cells[c_idx].text = sl_bq_vals[c_idx - 1]
# Dòng 2: Tổng đầu xe BQ/ngày
xe_bq_vals = ['20.4', '20.6', '20.8', '21.0', '21.2', '21.3']
for c_idx in range(1, 7):
    table_8.rows[2].cells[c_idx].text = xe_bq_vals[c_idx - 1]
# Dòng 3: Tổng đầu xe ngày cao điểm
for c_idx in range(1, 7):
    table_8.rows[3].cells[c_idx].text = '30'
# Dòng 4: Số người giao (đỉnh x 2)
for c_idx in range(1, 7):
    table_8.rows[4].cells[c_idx].text = tot_giao[c_idx - 1]
# Dòng 5: Mặt bằng tổng 4 BCCK cần
m2_vals = ['550', '584', '834', '850', '1,017', '1,033']
for c_idx in range(1, 7):
    table_8.rows[5].cells[c_idx].text = m2_vals[c_idx - 1]
print("-> Đã cập nhật Bảng 8")

# --- 4. CẬP NHẬT BẢNG 10 (Index 9) ---
table_10 = doc.tables[9]
# Sửa tên dòng 2 thành "Chi phí NV giao hàng (76–122 người, triệu đ)"
table_10.rows[2].cells[0].text = "Chi phí NV giao hàng (76–122 người, triệu đ)"
# Sửa tên dòng 3 thành "Chi phí NV kho & quản lý (50 người cố định, triệu đ)"
table_10.rows[3].cells[0].text = "Chi phí NV kho & quản lý (50 người cố định, triệu đ)"

# Cập nhật chi phí chi tiết
xe_cost = ['715.0', '820.1', '1208.6', '1326.9', '1575.1', '1658.1']
giao_cost = ['690.0', '690.0', '900.0', '900.0', '1050.0', '1080.0']
kho_cost = ['120.0', '120.0', '120.0', '120.0', '120.0', '120.0']
mb_cost = ['220.8', '220.8', '220.8', '220.8', '220.8', '220.8']
tot_cost = ['1745.8', '1850.9', '2449.4', '2567.7', '2965.9', '3078.9']

for c_idx in range(1, 7):
    table_10.rows[1].cells[c_idx].text = xe_cost[c_idx - 1]
    table_10.rows[2].cells[c_idx].text = giao_cost[c_idx - 1]
    table_10.rows[3].cells[c_idx].text = kho_cost[c_idx - 1]
    table_10.rows[4].cells[c_idx].text = mb_cost[c_idx - 1]
    table_10.rows[5].cells[c_idx].text = tot_cost[c_idx - 1]
print("-> Đã cập nhật Bảng 10")

# --- 5. CẬP NHẬT BẢNG 11 (Index 10) ---
table_11 = doc.tables[10]
# Khôi phục số liệu gốc v8_1
nt_vals = ['730.0', '787.6', '1244.1', '1316.8', '1550.2', '1603.7']
dl_vals = ['476.7', '486.3', '551.3', '574.1', '673.4', '695.0']
dd_vals = ['296.3', '299.9', '337.8', '348.6', '391.3', '400.9']
dul_vals = ['242.8', '277.1', '316.2', '328.2', '351.0', '379.3']

for c_idx in range(1, 7):
    table_11.rows[1].cells[c_idx].text = nt_vals[c_idx - 1]
    table_11.rows[2].cells[c_idx].text = dl_vals[c_idx - 1]
    table_11.rows[3].cells[c_idx].text = dd_vals[c_idx - 1]
    table_11.rows[4].cells[c_idx].text = dul_vals[c_idx - 1]
    # Dòng 5: Tổng
    table_11.rows[5].cells[c_idx].text = tot_cost[c_idx - 1]
print("-> Đã cập nhật Bảng 11")

# Lưu lại file
doc.save(dst_path)
print(f"=== ĐÃ LƯU FILE THÀNH CÔNG: {dst_path} ===")
