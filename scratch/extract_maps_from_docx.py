import docx, os, sys
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

src_docx = r'C:\Users\lap4all\Downloads\Quy_Hoach MANG LUOI NTB.docx'
out_dir = r'C:\Users\lap4all\Documents\Auto report\scratch\maps'
os.makedirs(out_dir, exist_ok=True)

doc = docx.Document(src_docx)

# Extract images from docx
image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.target_ref:
        img_part = rel.target_part
        img_filename = os.path.basename(rel.target_ref)
        img_bytes = img_part.blob
        
        # Save image
        img_path = os.path.join(out_dir, f"map_{image_count+1}_{img_filename}")
        with open(img_path, "wb") as f:
            f.write(img_bytes)
        image_count += 1

print(f"Extracted {image_count} map images to {out_dir}!")

# Create a dedicated Document for Administrative Maps
map_doc = docx.Document()
for section in map_doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

p_t = map_doc.add_paragraph()
r_t = p_t.add_run("BỘ BẢN ĐỒ HÀNH CHÍNH QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB 2026")
r_t.font.name = "Calibri"
r_t.font.size = Pt(16)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(0, 51, 102)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(12)

# Insert images into the document
images_files = sorted([os.path.join(out_dir, f) for f in os.listdir(out_dir) if f.startswith("map_")])

for idx, img_p in enumerate(images_files, 1):
    p_h = map_doc.add_paragraph()
    r_h = p_h.add_run(f"BẢN ĐỒ ĐỊA BÀN QUY HOẠCH - KHU VỰC SỐ {idx}")
    r_h.font.name = "Calibri"
    r_h.font.size = Pt(12)
    r_h.font.bold = True
    r_h.font.color.rgb = RGBColor(180, 0, 0)
    p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_h.paragraph_format.space_before = Pt(10)
    p_h.paragraph_format.space_after = Pt(4)
    
    p_img = map_doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_after = Pt(14)
    p_img.add_run().add_picture(img_p, width=Inches(7.2))

map_docx_path = r'C:\Users\lap4all\Downloads\Ban_Do_Hanh_Chinh_Quy_Hoach_NTB_2026.docx'
map_pdf_path = r'C:\Users\lap4all\Downloads\Ban_Do_Hanh_Chinh_Quy_Hoach_NTB_2026.pdf'

map_doc.save(map_docx_path)
print(f"Saved map docx: {map_docx_path}")

# Convert map docx to PDF
try:
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_win = word.Documents.Open(map_docx_path)
    doc_win.SaveAs(map_pdf_path, FileFormat=17)
    doc_win.Close()
    word.Quit()
    print(f"Successfully created PDF: {map_pdf_path}")
except Exception as e:
    print(f"Word PDF conversion error: {e}")
