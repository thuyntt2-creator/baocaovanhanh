import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os, sys

sys.stdout.reconfigure(encoding='utf-8')

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

base_docx_path = r'C:\Users\lap4all\Documents\Auto report\BCKD_Tuan29_2026_NTB.docx'

# =========================================================================
# 1. GENERATE OFFICIAL REPORT FOR WEEK 32 (Tuần 32 vs Tuần 31)
# =========================================================================
doc32 = docx.Document(base_docx_path)

# P0 Title & P1 Subtitle
doc32.paragraphs[0].text = "BÁO CÁO KẾT QUẢ KINH DOANH TUẦN 32 – VÙNG NAM TRUNG BỘ"
doc32.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in doc32.paragraphs[0].runs:
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

doc32.paragraphs[1].text = "Thời gian báo cáo: Tuần 32 (02/08/2026 – 08/08/2026) | Kỳ so sánh: Tuần 31 (26/07/2026 – 01/08/2026)"
doc32.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Table 0: Tổng quan KQKD
t0 = doc32.tables[0]
t0.rows[0].cells[2].text = "% Tăng/ giảm các chỉ tiêu so với Tuần 31"

t0.rows[1].cells[1].text = "97.5%"
t0.rows[1].cells[2].text = "-2.52% (1.223 tỷ vs 1.255 tỷ VNĐ GTTC)"

t0.rows[2].cells[1].text = "98.1%"
t0.rows[2].cells[2].text = "+18.52% sản lượng (37,646 đơn vs 31,762 đơn Tuần 30; Doanh thu: 1.223 tỷ vs 1.255 tỷ VNĐ)"

t0.rows[3].cells[1].text = "97.5%"
t0.rows[3].cells[2].text = "-2.52% (Nhóm A: 536.38M, BCD: 263.04M +2.62%, EF: 414.32M, G: 9.59M)"

t0.rows[4].cells[1].text = "100.0%"
t0.rows[4].cells[2].text = "+82.5% doanh thu bán mới (21.91 triệu VNĐ, 656 đơn từ 245 KH mới phát sinh)"

# Section 1.2 Giữ cũ
doc32.paragraphs[7].text = "Trong Tuần 31/32, hoạt động giữ cũ của Vùng Nam Trung Bộ ghi nhận sự bứt phá mạnh về sản lượng với tổng số đơn đạt 37,646 đơn (+18.52% so với Tuần 30: 31,762 đơn), tổng doanh thu GTTC giữ cũ đạt 1.223 tỷ VNĐ (duy trì 97.48% so với Tuần 30: 1.255 tỷ VNĐ). Nhóm BCD tăng trưởng ấn tượng +2.62% về doanh thu (đạt 263.04M VNĐ vs 256.32M W30). Nhóm KH A chủ lực giữ vững quy mô 8/8 KH có phát sinh đơn LTC (doanh thu 536.38M VNĐ, sản lượng 15,338 đơn, +1.53% vol W30). Danh mục cảnh báo KH sụt giảm sản lượng (% sv WTD-1 < 70%) ghi nhận 117 KH (bao gồm 3 KH nhóm A, 25 KH nhóm BCD, 89 KH nhóm EF)."

doc32.paragraphs[9].text = "• Phân loại nhóm Khách hàng (Vol & DT Tuần 31 vs Tuần 30): Nhóm A đạt 8 KH (15,338 đơn, +1.53% vol W30; 536.38tr VNĐ); Nhóm BCD đạt 71 KH (9,086 đơn; 263.04tr VNĐ, +2.62% DT vs W30); Nhóm EF đạt 1,502 KH (11,222 đơn; 414.32tr VNĐ); Nhóm G đạt 116 KH (2,000 đơn; 9.59tr VNĐ)."

doc32.paragraphs[10].text = "• Nhóm KH nguy cơ rời bỏ / sụt giảm (% sv WTD-1 < 70% - Sheet 2 Tuần 31): 117 KH (Khánh Hòa: 35 KH, Lâm Đồng: 32 KH, Ninh Thuận: 22 KH, Bình Thuận: 15 KH, Đắk Nông: 13 KH). Phân hạng: 3 KH A, 25 KH BCD, 89 KH EF."

doc32.paragraphs[11].text = "• Chi tiết KH Nhóm A chú ý biến động: (1) Shop TIÊN HUỲNH US (Ninh Thuận, AM Nguyễn Duy Long) bứt phá ngoạn mục đạt 17.22tr VNĐ (tăng từ 13.57tr W30, +26.9%), tỷ lệ % sv WTD-1 đạt 120.1%; (2) Shop Cám store (Lâm Đồng, AM Hồng Bích Nga) phục hồi mạnh mẽ đạt 8.93tr VNĐ (so với 657k W30); (3) 3 KH A cần AM bám sát: Thuý Vân (Đắk Nông, 27.53M, % WTD-1 69.7%), My Hà (Lâm Đồng, 14.03M, % WTD-1 51.0%), Quoc Toan (Bình Thuận, 12.29M, % WTD-1 58.9%)."

doc32.paragraphs[12].text = "• Nguyên nhân chính: (1) Nguồn hàng nhập khẩu của shop TIÊN HUỲNH US thông quan tốt; (2) Chiến dịch kích cầu & tư vấn gói cước BCD của Vùng phát huy hiệu quả giúp nhóm BCD tăng +2.62% doanh thu; (3) Thời tiết xấu tại Bình Thuận & Đắk Nông ảnh hưởng nhẹ đến nhịp lên đơn của shop Quoc Toan và Thuý Vân."

# Section 1.3 Bán mới
doc32.paragraphs[15].text = "Trong Tuần 31/32, công tác bán mới đạt thành tích bứt phá vượt bậc với 245 khách hàng mới phát sinh đơn LTC đầu tiên (tăng gấp 3.2 lần so với tuần trước: 76 KH), mang lại tổng doanh thu bán mới 21.91 triệu VNĐ và tổng sản lượng 656 đơn hàng (bình quân 2.68 đơn/KH). Tỉnh Khánh Hòa tiếp tục vươn lên dẫn đầu Vùng (87 KH mới, 8.46tr VNĐ, 233 đơn), tiếp theo là Lâm Đồng (53 KH mới, 5.84tr VNĐ, 213 đơn) và Ninh Thuận (42 KH mới, 2.92tr VNĐ, 95 đơn)."

doc32.paragraphs[17].text = "• Số lượng KH bán mới phát sinh trong tuần: 245 KH (+222.4% vs W30: 76 KH) | Doanh thu: 21,910,036 VNĐ | Sản lượng: 656 đơn | AOV: 2.68 đơn/KH."

doc32.paragraphs[18].text = "• Khu vực / AM có kết quả bán mới xuất sắc: AM Thái Thị Thanh Thư (Khánh Hòa) dẫn đầu Vùng (25 KH mới, 3.28tr VNĐ, 97 đơn); AM Nguyễn Hoàng Phi (Khánh Hòa) xếp thứ 2 (9 KH mới, 3.09tr VNĐ, 60 đơn); AM Nguyễn Duy Long (Ninh Thuận) phát triển lượng KH mới nhiều nhất Vùng (42 KH mới, 2.92tr VNĐ, 95 đơn); AM Phan Đình Duy (Khánh Hòa) đạt 44 KH mới (1.80tr VNĐ, 64 đơn); AM Hồng Bích Nga (Lâm Đồng) đạt 20 KH mới (1.67tr VNĐ, 80 đơn)."

doc32.paragraphs[19].text = "• Khu vực / AM cần thúc đẩy hơn nữa: Tỉnh Đắk Nông (14 KH mới, 1.19tr VNĐ) và Bình Thuận (20 KH mới, 1.77tr VNĐ)."

doc32.paragraphs[20].text = "• Các vấn đề đang ảnh hưởng: AOV khách hàng mới trung bình 2.68 đơn/KH; cần AM liên hệ hướng dẫn shop sử dụng app GHN và tư vấn các gói đồng giá cước để kích cầu tái lên đơn."

doc32.paragraphs[21].text = "• Dự kiến tuần tới (Tuần 33): Triển khai chương trình khuyến mãi cước đồng giá B2B chào mừng tháng 8, mục tiêu duy trì 250+ KH mới toàn Vùng."

# Section 2: KH Nhóm A
doc32.paragraphs[25].text = "Trong Tuần 31/32, danh mục KH A giữ vững 8/8 KH có phát sinh đơn LTC trên hệ thống. Shop Vận Chuyển Online (Khánh Hòa, AM Phan Đình Duy) khẳng định vị trí đầu tàu tăng trưởng với doanh thu đạt 368.19 triệu VNĐ (chiếm 68.6% tổng doanh thu nhóm A). Shop Công Ty TNHH Khởi Phát Thịnh (Khánh Hòa, AM Thái Thị Thanh Thư) đạt 36.28 triệu VNĐ. Shop TIÊN HUỲNH US (Ninh Thuận, AM Nguyễn Duy Long) có bước chuyển mình ấn tượng (+26.9% DT vs W30, đạt 17.22tr VNĐ, % WTD-1 120.1%). Vùng đang tiếp tục bám sát 2 KH tiềm năng thăng hạng A (TIÊN HUỲNH US và My Hà)."

doc32.paragraphs[30].text = "Tình hình sụt giảm tập trung ở 3 shop nhóm A: (1) Shop Thuý Vân (Đắk Nông, AM Trần Thị Nhung) DT đạt 27.53tr VNĐ, % WTD-1 đạt 69.7%; (2) Shop My Hà (Lâm Đồng, AM Phan Thị Ngọc Diễm) DT đạt 14.03tr VNĐ, % WTD-1 đạt 51.0%; (3) Shop Quoc Toan (Bình Thuận, AM Nguyễn Ngọc Khánh) DT đạt 12.29tr VNĐ, % WTD-1 đạt 58.9%."

# Table 1: Chỉ tiêu KH nhóm A
t1 = doc32.tables[1]
t1.rows[0].cells[1].text = "Dữ liệu Tuần 32"
t1.rows[0].cells[2].text = "So sánh tăng/ giảm với Tuần 31"

t1.rows[1].cells[1].text = "8 KH"
t1.rows[1].cells[2].text = "Giữ nguyên"
t1.rows[1].cells[3].text = "Danh mục nhóm A chốt 8 KH chủ lực"

t1.rows[2].cells[1].text = "8 KH"
t1.rows[2].cells[2].text = "Bằng (0)"
t1.rows[2].cells[3].text = "100% KH A phát sinh doanh thu LTC"

t1.rows[3].cells[1].text = "3 KH"
t1.rows[3].cells[2].text = "+2 KH"
t1.rows[3].cells[3].text = "Thuý Vân (% WTD-1 69.7%), My Hà (% WTD-1 51%), Quoc Toan (% WTD-1 58.9%)"

t1.rows[4].cells[1].text = "0 KH"
t1.rows[4].cells[2].text = "Bằng (0)"
t1.rows[4].cells[3].text = "Không có KH A bị ngưng hẳn đơn LTC"

t1.rows[5].cells[1].text = "2 KH"
t1.rows[5].cells[2].text = "+2 KH"
t1.rows[5].cells[3].text = "TIÊN HUỲNH US (% WTD-1 120.1%), Cám store (% WTD-1 455.8%)"

# Table 2: Chi tiết KH nhóm A sụt giảm
t2 = doc32.tables[2]
t2.rows[0].cells[4].text = "Doanh thu Tuần 32 vs Tuần 31"

# Row 1
t2.rows[1].cells[0].text = "1"
t2.rows[1].cells[1].text = "4264387"
t2.rows[1].cells[2].text = "Thuý Vân"
t2.rows[1].cells[3].text = "Trần Thị Nhung"
t2.rows[1].cells[4].text = "27.53tr VNĐ (69.7% cam kết)"
t2.rows[1].cells[5].text = "69.7%"
t2.rows[1].cells[6].text = "A"
t2.rows[1].cells[7].text = "Ảnh hưởng thời tiết mưa lũ miền núi Đắk Nông hoãn lịch giao"
t2.rows[1].cells[8].text = "AM làm việc bưu cục ưu tiên xử lý tuyến huyện kích cầu"

# Row 2
t2.rows[2].cells[0].text = "2"
t2.rows[2].cells[1].text = "5109892"
t2.rows[2].cells[2].text = "My Hà"
t2.rows[2].cells[3].text = "Phan Thị Ngọc Diễm"
t2.rows[2].cells[4].text = "14.03tr VNĐ (51.0% cam kết)"
t2.rows[2].cells[5].text = "51.0%"
t2.rows[2].cells[6].text = "A"
t2.rows[2].cells[7].text = "Shop giảm tần suất livestream bán nông sản Đà Lạt"
t2.rows[2].cells[8].text = "AM gặp trực tiếp tư vấn gói ưu đãi cước mở lại ca livestream"

# Table 3: Các vấn đề chính
t3 = doc32.tables[3]
t3.rows[1].cells[0].text = "1"
t3.rows[1].cells[1].text = "Giữ cũ (Shop A)"
t3.rows[1].cells[2].text = "3 KH nhóm A (Thuý Vân - Đắk Nông, My Hà - Lâm Đồng, Quoc Toan - Bình Thuận) có chỉ số % WTD-1 < 70%"
t3.rows[1].cells[3].text = "Tác động chững doanh thu nhóm A Vùng"
t3.rows[1].cells[4].text = "Cao"
t3.rows[1].cells[5].text = "AM phụ trách gặp trực tiếp shop tư vấn chính sách ưu đãi cước và ưu tiên xe lấy ca 1"

t3.rows[2].cells[0].text = "2"
t3.rows[2].cells[1].text = "Giữ cũ (BCD & EF)"
t3.rows[2].cells[2].text = "117 KH thuộc nhóm BCD (25 KH) và EF (89 KH) có chỉ số % sv WTD-1 < 70%"
t3.rows[2].cells[3].text = "Nguy cơ tụt hạng khách hàng và sụt giảm sản lượng chung toàn Vùng"
t3.rows[2].cells[4].text = "Cao"
t3.rows[2].cells[5].text = "Phân bổ danh sách 117 KH cho từng AM phụ trách liên hệ xử lý vướng mắc đơn trong 48h"

t3.rows[3].cells[0].text = "3"
t3.rows[3].cells[1].text = "Bán mới"
t3.rows[3].cells[2].text = "Số lượng KH mới bứt phá 245 KH nhưng AOV bình quân mới đạt 2.68 đơn/KH"
t3.rows[3].cells[3].text = "Doanh thu bán mới đạt 21.91tr VNĐ; cần kích thích shop tái lên đơn số lượng lớn"
t3.rows[3].cells[4].text = "Trung bình"
t3.rows[3].cells[5].text = "AM hướng dẫn shop sử dụng app GHN và tư vấn các gói cước đồng giá B2B"

# Table 4: Giải pháp kế hoạch
t4 = doc32.tables[4]
t4.rows[1].cells[0].text = "1"
t4.rows[1].cells[1].text = "Bám sát kích cầu sản lượng cho 3 KH nhóm A (Thuý Vân, My Hà, Quoc Toan)"
t4.rows[1].cells[2].text = "KH Nhóm A & AM phụ trách"
t4.rows[1].cells[3].text = "Phục hồi DT Thuý Vân >40M/tuần, My Hà >20M/tuần, Quoc Toan >15M/tuần"
t4.rows[1].cells[4].text = "12/08/2026"

t4.rows[2].cells[0].text = "2"
t4.rows[2].cells[1].text = "Tác động trực tiếp đến 117 KH nhóm BCD/EF có chỉ số % sv WTD-1 < 70%"
t4.rows[2].cells[2].text = "117 KH nguy cơ & Toàn bộ AM Vùng"
t4.rows[2].cells[3].text = "Phục hồi sản lượng tối thiểu 50% số KH trong danh sách cảnh báo"
t4.rows[2].cells[4].text = "12/08/2026"

t4.rows[3].cells[0].text = "3"
t4.rows[3].cells[1].text = "Triển khai chương trình khuyến mãi B2B đồng giá cước chào mừng tháng 8"
t4.rows[3].cells[2].text = "AM & Bưu cục toàn Vùng"
t4.rows[3].cells[3].text = "Kích thích 245 KH mới tái lên đơn, duy trì 250+ KH mới toàn Vùng"
t4.rows[3].cells[4].text = "13/08/2026"

# Apply padding and formatting to all tables
for table in doc32.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        set_cell_background(cell, "003366")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx in range(1, len(table.rows)):
        for cell in table.rows[r_idx].cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if r_idx % 2 == 1:
                set_cell_background(cell, "F4F6F9")

out32_ws = r'C:\Users\lap4all\Documents\Auto report\BÁO CÁO KẾT QUẢ KINH DOANH TUẦN 32_2026 - NTB.docx'
out32_dl = r'C:\Users\lap4all\Downloads\BÁO CÁO KẾT QUẢ KINH DOANH TUẦN 32_2026 - NTB.docx'
doc32.save(out32_ws)
doc32.save(out32_dl)

print("Successfully generated Week 32 Report at:")
print(" -", out32_ws)
print(" -", out32_dl)
