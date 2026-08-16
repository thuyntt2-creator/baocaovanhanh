import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

web_dir = r"C:\Users\lap4all\.gemini\antigravity-ide\brain\5935f00b-cd19-4de8-83bb-03170a8c8606\web_maps"
artifact_dir = r"C:\Users\lap4all\.gemini\antigravity-ide\brain\5935f00b-cd19-4de8-83bb-03170a8c8606"

docx_out = r"C:\Users\lap4all\Downloads\Ban_Do_Quy_Hoach_NTB_Official_Web_Maps.docx"
pdf_out = r"C:\Users\lap4all\Downloads\Ban_Do_Quy_Hoach_Mang_Luoi_NTB_2026.pdf"

doc = docx.Document()

# Page Margins
for sec in doc.sections:
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("BẢN ĐỒ QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NAM TRUNG BỘ (NTB)")
r_title.bold = True
r_title.font.size = Pt(16)
r_title.font.color.rgb = RGBColor(0, 51, 102)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run("TRÍCH XUẤT TỪ HỆ THỐNG QUY HOẠCH BƯU CỤC (quyhoachbuucuc.info)")
r_sub.bold = True
r_sub.font.size = Pt(11)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# List of web what-if map snapshots to include
maps_sections = [
    {
        "province": "1. TỈNH KHÁNH HÒA",
        "items": [
            ("1.1. Bản đồ & Bảng Quy hoạch (What-If) Phường Nha Trang", 
             "Bản đồ phân vùng lân thổ và Bảng Quy hoạch What-if so sánh sản lượng Trước - Sau khi gộp 5 phường về (KHO) Nha Trang.",
             os.path.join(web_dir, "rezone_row_03_Khánh_HòaPhường_Nha_Trang__100.png")),
            ("1.2. Bản đồ & Bảng Quy hoạch (What-If) Phường Tây Nha Trang", 
             "Gộp tuyến 6 xã/phường sáp nhập về Bưu cục (KHO) Tây Nha Trang.",
             os.path.join(web_dir, "rezone_row_07_Khánh_HòaPhường_Tây_Nha_Trang_.png")),
            ("1.3. Bản đồ Kế hoạch Tách Bưu cục Cam Linh & Bắc Cam Ranh", 
             "Tách bưu cục Nam Cam Ranh từ BC Cam Linh và di dời BC Bắc Cam Ranh.",
             os.path.join(artifact_dir, "media__1785907585318.png")),
            ("1.4. Bản đồ Quy hoạch Cụm Ninh Hòa (Phường Ninh Hòa mới)", 
             "Phân vùng tuyến Bưu cục Ninh Hòa 2.",
             os.path.join(web_dir, "rezone_row_05_Khánh_HòaPhường_Ninh_Hòa__1005.png"))
        ]
    },
    {
        "province": "2. TỈNH LÂM ĐỒNG",
        "items": [
            ("2.1. Bản đồ Phân vùng Bưu cục Đơn Dương (Tách Mở mới Lạc Xuân)", 
             "Sơ đồ tách bưu cục Lạc Xuân (4 xã) và giữ bưu cục gốc Nghĩa Đức (6 xã).",
             os.path.join(artifact_dir, "media__1785896953199.png")),
            ("2.2. Bản đồ & Bảng Quy hoạch (What-If) TP. Bảo Lộc (B'Lao Mới)", 
             "Mở mới BC B'Lao Mới thay thế BC 1 Bảo Lộc cũ.",
             os.path.join(web_dir, "rezone_row_18_Lâm_ĐồngPhường_B_Lao__1005911_.png")),
            ("2.3. Bản đồ Quy hoạch Phường Lâm Viên - TP. Đà Lạt", 
             "Sơ đồ tuyến cover khu vực Đà Lạt.",
             os.path.join(artifact_dir, "image1.png"))
        ]
    },
    {
        "province": "3. TỈNH NINH THUẬN",
        "items": [
            ("3.1. Bản đồ & Bảng Quy hoạch (What-If) Phường Phan Rang & Ninh Chử", 
             "Mở mới BC Đông Hải, di dời BC Ninh Chử và gộp tuyến BC Phan Rang.",
             os.path.join(web_dir, "rezone_row_06_Khánh_HòaPhường_Phan_Rang__100.png"))
        ]
    },
    {
        "province": "4. TỈNH BÌNH THUẬN",
        "items": [
            ("4.1. Bản đồ & Bảng Quy hoạch (What-If) TP. Phan Thiết (Phường Bình Thuận mới)", 
             "Gộp tuyến Phường Phan Thiết mới về BC Hàm Thắng.",
             os.path.join(web_dir, "rezone_row_20_Lâm_ĐồngPhường_Bình_Thuận__100.png")),
            ("4.2. Bản đồ Mở mới Bưu cục Nam Thành (Tánh Linh / Đức Linh)", 
             "Mở mới BC Nam Thành cover Nam Thành + Nghị Đức.",
             os.path.join(artifact_dir, "image9.jpg"))
        ]
    },
    {
        "province": "5. TỈNH ĐẮK NÔNG",
        "items": [
            ("5.1. Bản đồ & Bảng Quy hoạch (What-If) TP. Gia Nghĩa (Phường Bắc Gia Nghĩa mới)", 
             "Quy hoạch ranh giới quản lý 2 bưu cục Đông Gia Nghĩa và Bắc Gia Nghĩa.",
             os.path.join(web_dir, "rezone_row_19_Lâm_ĐồngPhường_Bắc_Gia_Nghĩa__.png"))
        ]
    }
]

for section in maps_sections:
    p_prov = doc.add_paragraph()
    r_prov = p_prov.add_run(section["province"])
    r_prov.bold = True
    r_prov.font.size = Pt(13)
    r_prov.font.color.rgb = RGBColor(0, 51, 102)
    
    for title, desc, img_path in section["items"]:
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(title)
        r_t.bold = True
        r_t.font.size = Pt(11)
        
        p_d = doc.add_paragraph()
        r_d = p_d.add_run("Ghi chú: " + desc)
        r_d.font.size = Pt(10)
        r_d.font.color.rgb = RGBColor(80, 80, 80)
        
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(img_path, width=Inches(6.8))
            print(f"Added picture: {os.path.basename(img_path)}")
        else:
            print(f"Image NOT found: {img_path}")
            
        doc.add_paragraph()

doc.save(docx_out)
print(f"Successfully created Official Web Maps DOCX at: {docx_out}")

# Convert to PDF via Word COM
import win32com.client
word = win32com.client.Dispatch('Word.Application')
word.Visible = False
doc_com = word.Documents.Open(docx_out)
doc_com.SaveAs(pdf_out, FileFormat=17)
doc_com.Close()
word.Quit()
print(f"Successfully generated OFFICIAL WEB MAPS PDF at: {pdf_out}")
