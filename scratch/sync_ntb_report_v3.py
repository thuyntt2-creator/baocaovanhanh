import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

src_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_3.docx"
dst_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_4.docx"

if not os.path.exists(src_path):
    print(f"File nguồn không tồn tại: {src_path}")
    sys.exit(1)

doc = docx.Document(src_path)

print("=== BẮT ĐẦU CẬP NHẬT ĐỢT 4 ===")

# 1. Cập nhật các đoạn văn bản (Paragraphs)
target_old_text = "Tiết kiệm 210.000.000 đ/tháng quỹ lương so với phương án cũ (60 người) — tương đương 1,26 tỷ đồng trong H2/2026."
found_p = False
for p in doc.paragraphs:
    if target_old_text in p.text:
        p.text = "Phương án điều chỉnh định biên nhân sự tối ưu theo sản lượng thực tế của vùng NTB."
        print("-> Đã sửa đoạn văn bản Tiết kiệm lương")
        found_p = True
        break
if not found_p:
    # Thử tìm kiếm mờ hơn
    for p in doc.paragraphs:
        if "1,26 tỷ đồng" in p.text and "Tiết kiệm" in p.text:
            p.text = "Phương án điều chỉnh định biên nhân sự tối ưu theo sản lượng thực tế của vùng NTB."
            print("-> Đã sửa đoạn văn bản Tiết kiệm lương (tìm mờ)")
            break

# 2. Cập nhật Bảng 2 (index 1)
table_2 = doc.tables[1]
table_2.rows[3].cells[1].text = "84 người (T7) → 130 người (T12) — linh hoạt theo sản lượng thực tế"
table_2.rows[4].cells[1].text = "30 xe (cap cố định)"
table_2.rows[5].cells[1].text = "1.746 → 2.915 triệu đ/tháng"
table_2.rows[6].cells[1].text = "—"
print("-> Đã cập nhật Bảng 2 (Tóm tắt)")

# 3. Cập nhật Bảng 3 (index 2)
table_3 = doc.tables[2]
table_3.rows[12].cells[4].text = "2 xe"
print("-> Đã cập nhật Bảng 3 (Số xe Đức Linh T12)")

# 4. Cập nhật Bảng 7 (index 6)
table_7 = doc.tables[6]
# Giao Nha Trang T7
table_7.rows[2].cells[1].text = "39 (SL: 1815)"
# Giao Đơn Dương T7
table_7.rows[4].cells[1].text = "10 (SL: 139)"
# Giao Đức Linh T7
table_7.rows[5].cells[1].text = "11 (SL: 111)"

# Hàng 8 (NV kho + QL)
table_7.rows[8].cells[0].text = "NV xử lý kho (4) + NV QL (4)"
for c_idx in range(1, 7):
    table_7.rows[8].cells[c_idx].text = "8"

# Hàng 9 (TỔNG NHÂN SỰ)
tot_ns = ["84", "86", "116", "118", "128", "130"]
for c_idx in range(1, 7):
    table_7.rows[9].cells[c_idx].text = tot_ns[c_idx - 1]
print("-> Đã cập nhật Bảng 7 (Định biên nhân sự)")

# 5. Cập nhật Bảng 10 (index 9)
table_10 = doc.tables[9]
table_10.rows[2].cells[0].text = "Chi phí NV giao hàng (46–80 người, triệu đ)"
print("-> Đã cập nhật Bảng 10 (Mô tả NV giao)")

# 6. Cập nhật Bảng 14 (index 13)
table_14 = doc.tables[13]
table_14.rows[4].cells[1].text = "Tuyển dụng đợt 1: 76 NV giao + 8 NV kho/QL cho T7"
print("-> Đã cập nhật Bảng 14 (Hành động tuyển dụng)")

doc.save(dst_path)
print(f"=== ĐÃ LƯU FILE THÀNH CÔNG: {dst_path} ===")
