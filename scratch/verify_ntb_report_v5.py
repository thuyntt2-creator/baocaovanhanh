import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_5.docx"

if not os.path.exists(doc_path):
    print("File không tồn tại")
    sys.exit(1)

doc = docx.Document(doc_path)
errors = 0

print("=== KIỂM TRA SỐ LIỆU TOÀN DIỆN FILE WORD V8_5 ===")

# 1. Kiểm tra Bảng 10 dòng 3 (NV kho & QL)
table_10 = doc.tables[9]
print("\n1. Kiểm tra Bảng 10 (Mô tả dòng 3):")
desc_row_3 = table_10.rows[3].cells[0].text.strip()
if "8 người cố định" not in desc_row_3:
    print(f"  [LỖI] Mô tả dòng 3 = '{desc_row_3}' (Kỳ vọng chứa: 8 người cố định)")
    errors += 1
else:
    print(f"  [OK] Mô tả dòng 3 = '{desc_row_3}' (Khớp)")

# 2. Kiểm tra Bảng 7 dòng 8 (NV xử lý kho + QL)
table_7 = doc.tables[6]
print("\n2. Kiểm tra Bảng 7 (Nhân sự kho & QL):")
for c_idx in range(1, 7):
    val = table_7.rows[8].cells[c_idx].text.strip()
    if val != "8":
        print(f"  [LỖI] Tháng {c_idx+6}: NV kho & QL = {val} (Kỳ vọng: 8)")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: NV kho & QL = {val} (Khớp)")

# 3. Kiểm tra các câu văn thuyết minh
print("\n3. Kiểm tra thuyết minh văn bản:")
found_1_26 = False
for p in doc.paragraphs:
    if "1,26 tỷ đồng quỹ lương" in p.text:
        print(f"  [LỖI] Vẫn còn cụm từ '1,26 tỷ đồng quỹ lương' trong Paragraph: {p.text}")
        found_1_26 = True
        errors += 1
if not found_1_26:
    print("  [OK] Đã sạch thông tin tiết kiệm 1,26 tỷ quỹ lương")

# 4. Kiểm tra đối chiếu Bảng 10 vs Bảng 11
table_11 = doc.tables[10]
print("\n4. Kiểm tra đối chiếu Bảng 10 vs Bảng 11:")
for c_idx in range(1, 7):
    t10_sum = float(table_10.rows[5].cells[c_idx].text.replace(',', ''))
    t11_sum = float(table_11.rows[5].cells[c_idx].text.replace(',', ''))
    if abs(t10_sum - t11_sum) > 0.1:
        print(f"  [LỖI] Tháng {c_idx+6}: Bảng 10 = {t10_sum} | Bảng 11 = {t11_sum}")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Tổng CP = {t10_sum} (Khớp)")

if errors == 0:
    print("\n>>> TẤT CẢ KIỂM TRA BÁO CÁO V8_5 ĐỀU THÀNH CÔNG! SỐ LIỆU HOÀN THIỆN TUYỆT ĐỐI <<<")
else:
    print(f"\n>>> CÓ {errors} SAI LỆCH SỐ LIỆU CẦN XỬ LÝ <<<")
