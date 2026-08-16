import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

artifact_dir = r"C:\Users\lap4all\.gemini\antigravity-ide\brain\5935f00b-cd19-4de8-83bb-03170a8c8606"
docx_out = r"C:\Users\lap4all\Downloads\Ban_Do_Quy_Hoach_NTB_Complete_Maps.docx"
pdf_out = r"C:\Users\lap4all\Downloads\Ban_Do_Quy_Hoach_Mang_Luoi_NTB_2026.pdf"

doc = docx.Document()

# Page Margins
for sec in doc.sections:
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("BẢN ĐỒ QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NAM TRUNG BỘ (NTB)")
r_title.bold = True
r_title.font.size = Pt(16)
r_title.font.color.rgb = RGBColor(0, 51, 102)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run("TỔNG HỢP TOÀN BỘ SƠ ĐỒ VÀ BẢN ĐỒ PHÂN VÙNG THEO ĐƠN VỊ HÀNH CHÍNH MỚI 2026")
r_sub.bold = True
r_sub.font.size = Pt(11)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

maps_data = [
    {
        "province": "1. KHU VỰC TỈNH KHÁNH HÒA",
        "items": [
            ("1.1. Bản đồ Kế hoạch Tách Bưu cục Cam Linh & Bắc Cam Ranh", 
             "Sơ đồ quy hoạch tách bưu cục Nam Cam Ranh từ BC Cam Linh và di dời BC Bắc Cam Ranh sang vị trí mới.", 
             os.path.join(artifact_dir, "media__1785907585318.png")),
            ("1.2. Bản đồ Quy hoạch Phường Tây Nha Trang", 
             "Gộp tuyến sản lượng các phường/xã cũ (Phương Sài, Ngọc Hiệp, Vĩnh Hiệp, Vĩnh Ngọc, Vĩnh Thạnh, Vĩnh Trung) về Bưu cục (KHO) Tây Nha Trang quản lý chính.", 
             os.path.join(artifact_dir, "image4.png")),
            ("1.3. Bản đồ Quy hoạch Cụm Vạn Ninh - Tu Bông", 
             "Sơ đồ điều chỉnh phân vùng tuyến Xã Vạn Bình gộp về Bưu cục (KHO) Tu Bông.", 
             os.path.join(artifact_dir, "image5.jpg"))
        ]
    },
    {
        "province": "2. KHU VỰC TỈNH LÂM ĐỒNG",
        "items": [
            ("2.1. Bản đồ Phân vùng Bưu cục Đơn Dương (Lâm Đồng)", 
             "Sơ đồ đề xuất mở mới Bưu cục Lạc Xuân (phụ trách 4 xã Lạc Lâm, Lạc Xuân, D'Ran, Ka Đô) và giữ Bưu cục gốc Nghĩa Đức (phụ trách 6 xã Đạ Ròn, Thạnh Mỹ, Tu Tra, Ka Đơn, Quảng Lập, Pró).", 
             os.path.join(artifact_dir, "media__1785896953199.png")),
            ("2.2. Bản đồ Quy hoạch Phường Lâm Viên - TP. Đà Lạt", 
             "Sơ đồ phân chia tuyến cover khu vực Phường Lâm Viên và Bưu cục trung tâm TP. Đà Lạt.", 
             os.path.join(artifact_dir, "image1.png")),
            ("2.3. Sơ đồ Tuyến di chuyển Xã Đam Rông 4 - Thôn Đưng Knớ", 
             "Sơ đồ tuyến giao xa >50km thôn Đưng Knớ do Bưu cục Lang Biang 1 phụ trách giữ nguyên.", 
             os.path.join(artifact_dir, "image2.png")),
            ("2.4. Sơ đồ Tuyến Cover Khu vực Lâm Hà - Đam Rông", 
             "Sơ đồ tổng thể tuyến vận chuyển và phân vùng bưu cục khu vực Lâm Hà - Đam Rông.", 
             os.path.join(artifact_dir, "image3.png"))
        ]
    },
    {
        "province": "3. KHU VỰC TỈNH NINH THUẬN",
        "items": [
            ("3.1. Bản đồ Mở mới Bưu cục (NTH) Đông Hải (TP. Phan Rang - Tháp Chàm)", 
             "Sơ đồ tuyến cover ven biển 6 phường (Đông Hải, Mỹ Bình, Mỹ Đông, Mỹ Hải...) khi mở mới Bưu cục Đông Hải.", 
             os.path.join(artifact_dir, "image6.png"))
        ]
    },
    {
        "province": "4. KHU VỰC TỈNH BÌNH THUẬN",
        "items": [
            ("4.1. Bản đồ Quy hoạch Phường Bình Thuận - TP. Phan Thiết", 
             "Sơ đồ quy hoạch phân vùng tuyến giao các phường trung tâm TP. Phan Thiết.", 
             os.path.join(artifact_dir, "image7.jpg")),
            ("4.2. Bản đồ Cụm Bưu cục Hàm Thuận Bắc - Hàm Liêm", 
             "Sơ đồ phân vùng tuyến Bưu cục Hàm Liêm và Hàm Thuận Bắc.", 
             os.path.join(artifact_dir, "image8.jpg")),
            ("4.3. Bản đồ Mở mới Bưu cục (BTH) Nam Thành (Tánh Linh / Đức Linh)", 
             "Sơ đồ mở mới Bưu cục Nam Thành cover khu vực Nam Thành (250 đơn) & Nghị Đức (200 đơn).", 
             os.path.join(artifact_dir, "image9.jpg"))
        ]
    }
]

for section in maps_data:
    p_prov = doc.add_paragraph()
    r_prov = p_prov.add_run(section["province"])
    r_prov.bold = True
    r_prov.font.size = Pt(13)
    r_prov.font.color.rgb = RGBColor(0, 51, 102)
    
    for title, desc, img_path in section["items"]:
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(title)
        r_t.bold = True
        r_t.font.size = Pt(11)
        
        p_d = doc.add_paragraph()
        r_d = p_d.add_run("Mô tả: " + desc)
        r_d.font.size = Pt(10)
        r_d.font.color.rgb = RGBColor(80, 80, 80)
        
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(img_path, width=Inches(6.5))
            print(f"Added picture: {img_path}")
        else:
            print(f"Image NOT found: {img_path}")
            
        doc.add_paragraph()

doc.save(docx_out)
print(f"Successfully created Complete Maps DOCX at: {docx_out}")

# Convert to PDF via Word COM
import win32com.client
word = win32com.client.Dispatch('Word.Application')
word.Visible = False
doc_com = word.Documents.Open(docx_out)
doc_com.SaveAs(pdf_out, FileFormat=17)
doc_com.Close()
word.Quit()
print(f"Successfully generated PERFECT MAPS PDF at: {pdf_out}")
