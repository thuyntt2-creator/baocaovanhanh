import docx, openpyxl, sys
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

# 1. Read latest user docx
doc_src = docx.Document(r'C:\Users\lap4all\Downloads\Quy_Hoach MANG LUOI NTB.docx')

parsed_36 = []
current = None

for p in doc_src.paragraphs:
    txt = p.text.strip()
    if not txt: continue
    
    is_heading = False
    for n in range(1, 37):
        if txt.startswith(f'{n}. '):
            is_heading = True
            break
            
    if is_heading:
        if current: parsed_36.append(current)
        raw_title = txt
        title_name = raw_title.split('(')[0].strip()
        if '. ' in title_name: title_name = title_name.split('. ', 1)[1].strip()
        prov_name = 'TỈNH LÂM ĐỒNG'
        if '(Tỉnh Khánh Hòa)' in raw_title: prov_name = 'TỈNH KHÁNH HÒA'
        elif '(Tỉnh Ninh Thuận)' in raw_title: prov_name = 'TỈNH NINH THUẬN'
        elif '(Tỉnh Bình Thuận)' in raw_title: prov_name = 'TỈNH BÌNH THUẬN'
        elif '(Tỉnh Đắk Nông)' in raw_title: prov_name = 'TỈNH ĐẮK NÔNG'
        
        current = {'raw_heading': raw_title, 'title': title_name, 'province': prov_name, 'communes': [], 'proposal': '', 'reason': ''}
    elif current:
        if 'ĐỀ XUẤT PHƯƠNG ÁN' in txt:
            current['proposal'] = txt.replace('❖ ĐỀ XUẤT PHƯƠNG ÁN CỦA AM:', '').replace('❖ ĐỀ XUẤT PHƯƠNG ÁN:', '').strip()
        elif 'LÝ DO VÀ GIẢI THÍCH' in txt:
            current['reason'] = txt.replace('❖ LÝ DO VÀ GIẢI THÍCH NGUYÊN NHÂN CHI TIẾT TỪ AM:', '').replace('❖ LÝ DO VÀ GIẢI THÍCH NGUYÊN NHÂN CHI TIẾT:', '').strip()
        elif current['proposal'] and not current['reason'] and not txt.startswith('❖'):
            current['proposal'] += ' ' + txt
        elif current['reason'] and not txt.startswith('❖'):
            current['reason'] += ' ' + txt
        elif not current['proposal'] and not txt.startswith('❖'):
            if ('Giao:' in txt or 'Lấy:' in txt or 'BC:' in txt or 'Phường' in txt or 'Xã' in txt or 'Thị trấn' in txt) and not txt.startswith('TỔNG SẢN LƯỢNG') and not txt.startswith('Các BC hiện') and not txt.startswith('Mã Xã') and not txt.startswith('Tỷ lệ'):
                raw_ward = txt.split('(Giao')[0].split('(Giao :')[0].replace('•', '').replace('-', '').strip()
                if raw_ward and raw_ward not in current['communes']:
                    current['communes'].append(raw_ward)

if current: parsed_36.append(current)

# Group by Province
provinces_dict = {}
for item in parsed_36:
    p = item['province']
    if p not in provinces_dict:
        provinces_dict[p] = []
        
    c_str = ' + '.join(item['communes']) if item['communes'] else item['title']
    prop = item['proposal'] if item['proposal'] else item['reason']
    
    buucuc_cover = "BC phụ trách theo AM"
    short_note = prop
    if 'Bưu cục' in prop or 'BC' in prop or 'GIỮ NGUYÊN' in prop or 'Gộp' in prop or 'TÁCH' in prop or 'MỞ' in prop:
        parts = prop.split('.')
        buucuc_cover = parts[0]
        if len(parts) > 1:
            short_note = '. '.join(parts[1:]).strip()
            
    provinces_dict[p].append([item['title'], c_str, buucuc_cover, short_note if short_note else buucuc_cover])

# 2. Build Excel File
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Quy_Hoach_NTB_2026_Cap_Nhat"

font_prov_header = Font(name="Calibri", size=12, bold=True, color="003366")
font_header = Font(name="Calibri", size=11, bold=True, color="000000")
fill_header = PatternFill(start_color="F7A059", end_color="F7A059", fill_type="solid")

font_dvhc = Font(name="Calibri", size=10, bold=True, color="000000")
font_normal = Font(name="Calibri", size=10, color="000000")
font_buucuc = Font(name="Calibri", size=10, bold=True, color="003399")

align_header = Alignment(horizontal="center", vertical="center", wrap_text=True)
align_cell = Alignment(horizontal="left", vertical="center", wrap_text=True)

thin_border = Border(
    left=Side(style='thin', color='000000'), right=Side(style='thin', color='000000'),
    top=Side(style='thin', color='000000'), bottom=Side(style='thin', color='000000')
)

headers = ["ĐVHC Mới", "Các Xã / Phường Cũ sáp nhập", "Bưu cục Cover", "Phương án & Ghi chú"]

current_row = 1

for p_name, p_rows in provinces_dict.items():
    cell_prov = ws.cell(row=current_row, column=1, value=f"❖ QUY HOẠCH {p_name}")
    cell_prov.font = font_prov_header
    ws.row_dimensions[current_row].height = 24
    current_row += 1
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=current_row, column=col_idx, value=h_text)
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_header
        cell.border = thin_border
    
    ws.row_dimensions[current_row].height = 26
    current_row += 1
    
    for r_data in p_rows:
        for col_idx, val in enumerate(r_data, start=1):
            cell = ws.cell(row=current_row, column=col_idx, value=val)
            cell.border = thin_border
            cell.alignment = align_cell
            if col_idx == 1:
                cell.font = font_dvhc
            elif col_idx == 3:
                cell.font = font_buucuc
            else:
                cell.font = font_normal
        ws.row_dimensions[current_row].height = 24
        current_row += 1
        
    current_row += 1

ws.column_dimensions['A'].width = 25
ws.column_dimensions['B'].width = 48
ws.column_dimensions['C'].width = 38
ws.column_dimensions['D'].width = 44

excel_out = r'C:\Users\lap4all\Downloads\Quy_Hoach_NTB_2026_Moi_Nhat.xlsx'
wb.save(excel_out)
print(f"Generated latest Excel file: {excel_out}")

# 3. Build Word File
doc_out = docx.Document()

for section in doc_out.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_province_heading(text):
    p = doc_out.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

p_t = doc_out.add_paragraph()
r_t = p_t.add_run('TỜ TRÌNH QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB NĂM 2026\n(BẢNG TỔNG HỢP CẬP NHẬT MỚI NHẤT DỮ LIỆU TỔNG HỢP 06 BC MỞ MỚI & BC CK DI LINH)')
r_t.font.name = 'Calibri'
r_t.font.size = Pt(15)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(180, 0, 0)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(12)

widths = [Inches(1.8), Inches(2.6), Inches(2.2), Inches(2.2)]

for p_name, p_rows in provinces_dict.items():
    add_province_heading(f"❖ QUY HOẠCH {p_name}")
    
    table = doc_out.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], 'F7A059')
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
            
    for r_data in p_rows:
        row_cells = table.add_row().cells
        bg_color = 'FFFFFF'
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            row_cells[c_idx].width = widths[c_idx]
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=80, right=80)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                if c_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                elif c_idx == 2:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 51, 153)

word_out = r'C:\Users\lap4all\Downloads\Thu_trinh_Quy_Hoach_Mang_Luoi_NTB_2026_Moi_Nhat.docx'
doc_out.save(word_out)
print(f"Generated latest Word file: {word_out}")
