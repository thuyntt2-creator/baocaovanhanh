import docx, os, sys, hashlib
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

src_docx = r'C:\Users\lap4all\Downloads\Quy_Hoach MANG LUOI NTB.docx'
out_dir = r'C:\Users\lap4all\Documents\Auto report\scratch\maps_dedup'
os.makedirs(out_dir, exist_ok=True)

doc = docx.Document(src_docx)

# Extract images, deduplicate using MD5 hash, match with context
unique_maps = []
seen_hashes = set()

current_title = "VÙNG NTB"
current_province = "TỈNH LÂM ĐỒNG"

for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt: continue
    
    # Track title & province
    for n in range(1, 37):
        if txt.startswith(f'{n}. '):
            current_title = txt
            if '(Tỉnh Khánh Hòa)' in txt: current_province = "TỈNH KHÁNH HÒA"
            elif '(Tỉnh Ninh Thuận)' in txt: current_province = "TỈNH NINH THUẬN"
            elif '(Tỉnh Bình Thuận)' in txt: current_province = "TỈNH BÌNH THUẬN"
            elif '(Tỉnh Đắk Nông)' in txt or '(Tỉnh Đắc Nông)' in txt: current_province = "TỈNH ĐẮK NÔNG"
            elif '(Tỉnh Lâm Đồng)' in txt: current_province = "TỈNH LÂM ĐỒNG"
            break

    # Extract images
    for run in p.runs:
        drawing_elements = run._r.xpath('.//w:drawing')
        for dr in drawing_elements:
            blip_elements = dr.xpath('.//a:blip')
            for blip in blip_elements:
                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id in doc.part.rels:
                    target_part = doc.part.rels[embed_id].target_part
                    img_bytes = target_part.blob
                    
                    # Compute MD5 hash to deduplicate
                    h_val = hashlib.md5(img_bytes).hexdigest()
                    if h_val not in seen_hashes:
                        seen_hashes.add(h_val)
                        img_filename = f"uniq_map_{len(unique_maps)+1}.png"
                        img_path = os.path.join(out_dir, img_filename)
                        with open(img_path, "wb") as f:
                            f.write(img_bytes)
                        
                        unique_maps.append({
                            'province': current_province,
                            'title': current_title,
                            'caption': txt if 'Hình' in txt else f"Bản đồ quy hoạch {current_title}",
                            'img_path': img_path,
                            'hash': h_val
                        })

print(f"Deduplicated images down to {len(unique_maps)} UNIQUE map screenshots!")

# Categorize unique maps
cat_open_new = []
cat_relocate = []
cat_regional = []

for m in unique_maps:
    title_u = m['title'].upper()
    cap_u = m['caption'].upper()
    
    if any(k in title_u or k in cap_u for k in ['XUÂN HƯƠNG', 'DI LINH', 'LẠC XUÂN', 'ĐƠN DƯƠNG', 'NAM CAM RANH', 'ĐÔNG HẢI', 'NAM THÀNH', 'NAM NHA TRANG']):
        cat_open_new.append(m)
    elif any(k in title_u or k in cap_u for k in ['TÂY NHA TRANG', 'BẮC CAM RANH', 'NINH CHỬ']):
        cat_relocate.append(m)
    else:
        cat_regional.append(m)

# Build Categorized Document
map_doc = docx.Document()
for section in map_doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Main Title
p_t = map_doc.add_paragraph()
r_t = p_t.add_run("BỘ BẢN ĐỒ QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB 2026\n(TỔNG HỢP THEO PHƯƠNG ÁN MỞ MỚI / DI DỜI - ĐÃ LỌC TRÙNG 100%)")
r_t.font.name = "Calibri"
r_t.font.size = Pt(15)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(180, 0, 0)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(14)

def add_category_header(text):
    p = map_doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)

# 1. MỞ MỚI / TÁCH BƯU CỤC
add_category_header("📌 PHẦN 1: BẢN ĐỒ CÁC BƯU CỤC MỞ MỚI / TÁCH BƯU CỤC (07 BC)")
for item in cat_open_new:
    p_item = map_doc.add_paragraph()
    r_item = p_item.add_run(f"✨ [{item['province']}] - {item['title']}")
    r_item.font.name = "Calibri"
    r_item.font.size = Pt(11)
    r_item.font.bold = True
    r_item.font.color.rgb = RGBColor(0, 102, 204)
    p_item.paragraph_format.space_before = Pt(6)
    p_item.paragraph_format.space_after = Pt(4)
    
    p_img = map_doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_after = Pt(12)
    p_img.add_run().add_picture(item['img_path'], width=Inches(7.0))

# 2. DI DỜI / MỞ RỘNG MẶT BẰNG KHO
add_category_header("🚚 PHẦN 2: BẢN ĐỒ CÁC BƯU CỤC DI DỜI / MỞ RỘNG MẶT BẰNG KHO")
for item in cat_relocate:
    p_item = map_doc.add_paragraph()
    r_item = p_item.add_run(f"🚚 [{item['province']}] - {item['title']}")
    r_item.font.name = "Calibri"
    r_item.font.size = Pt(11)
    r_item.font.bold = True
    r_item.font.color.rgb = RGBColor(180, 100, 0)
    p_item.paragraph_format.space_before = Pt(6)
    p_item.paragraph_format.space_after = Pt(4)
    
    p_img = map_doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_after = Pt(12)
    p_img.add_run().add_picture(item['img_path'], width=Inches(7.0))

# 3. QUY HOẠCH SÁP NHẬP ĐỊA BÀN THEO TỈNH
if cat_regional:
    add_category_header("🗺️ PHẦN 3: BẢN ĐỒ CỤM ĐỊA BÀN SÁP NHẬP THEO TỈNH (KHÔNG TRÙNG LẶP)")
    for item in cat_regional:
        p_item = map_doc.add_paragraph()
        r_item = p_item.add_run(f"📍 [{item['province']}] - {item['title']}")
        r_item.font.name = "Calibri"
        r_item.font.size = Pt(11)
        r_item.font.bold = True
        r_item.font.color.rgb = RGBColor(0, 128, 0)
        p_item.paragraph_format.space_before = Pt(6)
        p_item.paragraph_format.space_after = Pt(4)
        
        p_img = map_doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(12)
        p_img.add_run().add_picture(item['img_path'], width=Inches(7.0))

map_docx_path = r'C:\Users\lap4all\Downloads\Ban_Do_Quy_Hoach_NTB_2026_Theo_Phuong_An.docx'
map_pdf_path = r'C:\Users\lap4all\Downloads\Ban_Do_Quy_Hoach_NTB_2026_Theo_Phuong_An.pdf'

map_doc.save(map_docx_path)
print(f"Saved smart map docx: {map_docx_path}")

# Convert to PDF
try:
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_win = word.Documents.Open(map_docx_path)
    doc_win.SaveAs(map_pdf_path, FileFormat=17)
    doc_win.Close()
    word.Quit()
    print(f"Successfully generated PDF by method: {map_pdf_path}")
except Exception as e:
    print(f"Error converting map PDF: {e}")
