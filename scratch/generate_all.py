import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import pandas as pd
import os, sys

sys.stdout.reconfigure(encoding='utf-8')

artifact_dir = r"C:\Users\lap4all\.gemini\antigravity-ide\brain\2488f1f5-238f-4151-b3dd-1b59691a739c"
img_dir = os.path.join(artifact_dir, "images")

# Load Master Excel
df_excel = pd.read_excel(r'C:\Users\lap4all\Downloads\NTB_Phan_Tuyen_Hanh_Chinh_Quy_Hoach_Moi.xlsx', sheet_name='Sheet1')

doc = docx.Document()

# Page Setup - Margins 0.75 inch
for s in doc.sections:
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.75)
    s.right_margin = Inches(0.75)

PRIMARY_COLOR = RGBColor(31, 73, 125)    # Navy #1F497D
SECONDARY_COLOR = RGBColor(46, 117, 182) # Slate Blue #2E75B6
DARK_TEXT = RGBColor(38, 38, 38)         # #262626

def set_cell_bg(cell, hex_color):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = SECONDARY_COLOR
    r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(18)

def add_h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True

def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True

def add_h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True

def add_p(text, bold_prefix='', italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.bold = True
        r_pre.font.size = Pt(11)
        r_pre.font.name = 'Calibri'
        r_pre.font.color.rgb = DARK_TEXT
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.italic = italic
    r.font.name = 'Calibri'
    r.font.color.rgb = DARK_TEXT
    return p

def add_bullet(text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.bold = True
        r_pre.font.size = Pt(11)
        r_pre.font.name = 'Calibri'
        r_pre.font.color.rgb = DARK_TEXT
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    r.font.color.rgb = DARK_TEXT
    return p

def add_callout(text, title=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="12" w:color="1F497D"/></w:pBdr>')
    pPr.append(pBdr)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F4F7"/>')
    pPr.append(shd)
    
    if title:
        rt = p.add_run(f'📌 {title}\n')
        rt.font.bold = True
        rt.font.size = Pt(11)
        rt.font.color.rgb = PRIMARY_COLOR
        rt.font.name = 'Calibri'
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.italic = True
    r.font.name = 'Calibri'

def add_img(img_filename, caption='', width=Inches(5.8)):
    imgPath = os.path.join(img_dir, img_filename)
    if os.path.exists(imgPath):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        run.add_picture(imgPath, width=width)
        
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after = Pt(10)
            crun = cp.add_run(f'Sơ đồ/Bản đồ: {caption}')
            crun.font.size = Pt(9.5)
            crun.font.italic = True
            crun.font.color.rgb = RGBColor(100, 100, 100)
            crun.font.name = 'Calibri'

# Document Content Generation
add_title('BÁO CÁO TOÀN DIỆN & PHƯƠNG ÁN QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NAM TRUNG BỘ (NTB) THEO ĐƠN VỊ HÀNH CHÍNH MỚI')
add_subtitle('Đồng bộ Số liệu Sản lượng AM & Phân tích Chi tiết Lý do Giữ nguyên 2-3 Bưu cục Không gộp (Kèm Bản đồ Quy hoạch 2026)')

add_h1('I. TỔNG QUAN HIỆN TRẠNG MẠNG LƯỚI BƯU CỤC VÙNG NTB')
add_p('Thực hiện chủ trương tinh gọn và sáp nhập đơn vị hành chính cấp xã/phường trên phạm vi toàn quốc năm 2026, mạng lưới giao nhận bưu cục vùng Nam Trung Bộ (bao gồm 5 tỉnh: Khánh Hòa, Ninh Thuận, Bình Thuận, Lâm Đồng, Đắc Nông) đứng trước yêu cầu rà soát và tái quy hoạch toàn bộ phạm vi quản lý tuyến, diện tích kho bãi và bố trí nhân sự.')

add_h2('1. Quy mô Mạng lưới & Mức độ Chia cắt Đơn vị Hành chính Mới')
add_bullet('83 Bưu cục Express đang vận hành trên 5 tỉnh thành thuộc khu vực NTB.', '• Tổng số Bưu cục: ')
add_bullet('36 Xã/Phường mới (gồm 114 xã/phường cũ sáp nhập thành), chuẩn hóa theo danh mục Master Dataset NTB_Phan_Tuyen_Hanh_Chinh_Quy_Hoach_Moi.xlsx.', '• Phạm vi rà soát: ')
add_bullet('4 Xã/Phường (chiếm 11.1%) - Ranh giới quản lý thống nhất, do 01 Bưu cục duy nhất phụ trách (Xã Cam Hiệp, Xã Cam Lâm, Xã Phước Dinh, Xã Phan Rí Cửa).', '• Số Xã/Phường đã quy hoạch chuẩn: ')
add_bullet('32 Xã/Phường (chiếm 88.9%) - Xuất hiện tình trạng chia cắt mảnh, có từ 2 đến 3 Bưu cục cùng phụ trách trong 1 xã/phường mới.', '• Số Xã/Phường bị chia cắt: ')

add_h2('2. Bất cập & Rủi ro Vận hành Chính')
add_bullet('Trong cùng 01 phường mới, có tới 3 Bưu cục cùng cử shipper vào giao hàng. Điều này làm tuyến đường bị dẫm chéo, shipper di chuyển cắt ngang địa bàn của nhau, làm gia tăng thời gian và chi phí vận chuyển.', '• Chồng chéo địa bàn & dẫm tuyến: ')
add_bullet('Một số xã vùng ven đang bị phân tuyến giao từ Bưu cục ở huyện/tỉnh khác cách xa 30 - 40 km, trong khi Bưu cục lân cận chỉ cách 7 - 15 km.', '• Tuyến đi chéo xa ranh giới: ')
add_bullet('Nếu cưỡng ép đóng cửa hoặc di dời BC sáp nhập ngay lập tức theo địa giới mới tại các khu vực miền núi/nông thôn (như Đắc Nông), 100% nhân sự điểm xã ven sẽ xin nghỉ việc do không chịu đi xa, dẫn đến nguy cơ bể tuyến cục bộ.', '• Rủi ro đứt gãy nhân sự: ')

add_callout('Mục tiêu của Kế hoạch Quy hoạch 2026: Tái cấu trúc ranh giới giao nhận theo đơn vị hành chính mới, tối ưu hóa năng suất shipper, rút ngắn bán kính di chuyển, mở mới các bưu cục trọng điểm và đóng bưu cục hoạt động chồng chéo để tối ưu chi phí vận hành.', 'ĐỊNH HƯỚNG QUY HOẠCH')

add_h1('II. DANH MỤC 36 XÃ/PHƯỜNG MỚI THEO CHUẨN MASTER DATASET')
add_p('Toàn bộ tên Phường/Xã mới, Mã xã mới (GHN Code) và Tỉnh/thành mới dưới đây được đồng bộ chính xác 100% theo file danh mục chuẩn NTB_Phan_Tuyen_Hanh_Chinh_Quy_Hoach_Moi.xlsx:')

# Master Table 36 Wards
tbl_master = doc.add_table(rows=1, cols=6)
tbl_master.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = tbl_master.rows[0].cells
headers_m = ['STT', 'Mã Xã Mới', 'Tên Xã/Phường Mới (Chuẩn)', 'Tỉnh/Thành Mới', 'Sản Lượng AM (Đơn/ngày)', 'AM Phụ Trách & Phương Án']
for i, h in enumerate(headers_m):
    hdr_cells[i].text = h
    set_cell_bg(hdr_cells[i], '1F497D')
    for p in hdr_cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.name = 'Calibri'
            r.font.size = Pt(9.5)

grouped_master = df_excel.groupby(['Mã Xã mới', 'Tên Xã mới', 'Tỉnh, thành phố mới'], sort=False)
stt = 1
for (code, name, prov), grp in grouped_master:
    r_cells = tbl_master.add_row().cells
    bg = 'F9FAFB' if stt % 2 == 0 else 'FFFFFF'
    vol_sum = grp['Sản lượng giao/ngày (đơn)'].sum() + grp['Sản lượng lấy/ngày (đơn)'].sum()
    am = grp['Quản lý khu vực (AM)'].iloc[0]
    proposal = str(grp['Đánh giá & Phương án đề xuất'].iloc[0])
    
    vals = [str(stt), str(code), str(name), str(prov), f"{vol_sum:.0f}", f"AM {am}: {proposal}"]
    for i, val in enumerate(vals):
        r_cells[i].text = val
        set_cell_bg(r_cells[i], bg)
        for p in r_cells[i].paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9.0)
                r.font.name = 'Calibri'
                if i == 2:
                    r.font.bold = True
    stt += 1

add_h1('III. PHƯƠNG ÁN QUY HOẠCH CHI TIẾT & LÝ DO GIỮ NGUYÊN KHÔNG GỘP CỦA AM')

# 1. LÂM ĐỒNG
add_h2('1. Tỉnh Lâm Đồng')

add_h3('1.1. Khu vực Thành phố Đà Lạt')
add_p('Theo báo cáo đề xuất của AM Lê Văn Trường tại TP. Đà Lạt:')

add_bullet('Tổng sản lượng AM trình bày: 2.292.0 đơn/ngày (Giao: 1.497.0 đơn · Lấy: 795.0 đơn). Phụ trách bởi 3 BC: (LDO) Lang Biang - Đà Lạt 2, (LDO) Xuân Hương - Đà Lạt, (LDO) Lâm Viên - Đà Lạt 1. Sáp nhập từ Phường 1, Phường 2, Phường 3, Phường 4, Phường 10.', '• Phường Xuân Hương - Đà Lạt (Mã: 24781): ')
add_callout('LÝ DO GIỮ NGUYÊN 2 BƯU CỤC & MỞ MỚI 1 BC (AM Lê Văn Trường):\nSản lượng Phường Xuân Hương quá lớn (2.292 đơn/ngày). Một bưu cục đơn lẻ không thể đáp ứng tải kho bãi và nhân sự xử lý. AM đề xuất mở mới 01 Bưu cục (LDO) Xuân Hương - Đà Lạt 2 tại Phường 10. Phân ranh giới: BC Xuân Hương - Đà Lạt phụ trách Phường 1 + 2 cũ; BC Xuân Hương - Đà Lạt 2 phụ trách Phường 3 + 10 cũ.', 'PHƯƠNG ÁN PHƯỜNG XUÂN HƯƠNG - ĐÀ LẠT')

add_bullet('Tổng sản lượng AM trình bày: 1.365.0 đơn/ngày (Giao: 1.066.0 đơn · Lấy: 299.0 đơn). Phụ trách bởi 2 BC: (LDO) Lâm Viên - Đà Lạt 1 và (LDO) Lâm Viên - Đà Lạt 2. Sáp nhập từ Phường 8, Phường 9, Phường 12.', '• Phường Lâm Viên - Đà Lạt (Mã: 24778): ')
add_callout('LÝ DO GIỮ NGUYÊN 02 BƯU CỤC KHÔNG GỘP (AM Lê Văn Trường):\nPhường Lâm Viên mới có địa hình tự nhiên bị chia cắt làm 2 bờ riêng biệt bởi Hồ Xuân Hương. Nếu gộp về 1 Bưu cục, shipper sẽ phải di chuyển vòng qua hồ rất xa để sang phía bờ đối diện, làm gia tăng thời gian di chuyển và tăng chi phí vận hành. Do đó, AM quyết định GIỮ NGUYÊN 02 Bưu cục (Lâm Viên 1 & Lâm Viên 2) hoạt động ở 2 phía bờ hồ để đảm bảo hiệu quả giao hàng.', 'LÝ DO GIỮ NGUYÊN PHƯỜNG LÂM VIÊN - ĐÀ LẠT')

add_img('da_lat_ban_do_quy_hoach.png', 'Bản đồ Quy hoạch Mạng lưới Bưu cục Thành phố Đà Lạt 2026')

add_h3('1.2. Khu vực Thành phố Bảo Lộc')
add_p('Theo báo cáo trình bày của AM Hồng Bích Nga & AM Nguyễn Lê Nguyên Vũ:')
add_bullet('Tổng sản lượng AM trình bày: 1.217.0 đơn/ngày (3.634.0 kg/ngày) | Giao: 618.0 đơn · Lấy: 599.0 đơn. Hiện do 2 BC phụ trách: (LDO) B\'Lao (65.5% - 797 đơn) và (LDO) 1 Bảo Lộc (34.5% - 420 đơn). AM đề xuất GỘP VỀ 01 BC CHÍNH (LDO) B\'Lao, đóng cửa BC 1 Bảo Lộc.', '• Phường 1 Bảo Lộc (Mã: 24823): ')
add_bullet('Tổng sản lượng AM trình bày: 789.0 đơn/ngày (1.438.0 kg/ngày) | Giao: 454.0 đơn · Lấy: 335.0 đơn. Hiện do 2 BC phụ trách: (LDO) B\'Lao (56.3% - 444 đơn) và (LDO) 3 Bảo Lộc (43.7% - 345 đơn). GỘP VỀ 01 BC CHÍNH (LDO) B\'Lao.', '• Phường 2 Bảo Lộc (Mã: 24820): ')
add_bullet('Tổng sản lượng AM trình bày: 678.0 đơn/ngày (1.341.0 kg/ngày) | Giao: 531.0 đơn · Lấy: 147.0 đơn. Hiện do 3 BC phụ trách: (LDO) B\'Lao (49.9% - 338 đơn), (LDO) 3 Bảo Lộc (28.9% - 196 đơn), (LDO) 1 Bảo Lộc (21.2% - 144 đơn). GỘP VỀ 01 BC CHÍNH (LDO) 3 Bảo Lộc.', '• Phường B\'Lao (Mã: 24829): ')
add_bullet('Tổng sản lượng AM trình bày: 324.0 đơn/ngày (770.0 kg/ngày) | Giao: 274.0 đơn · Lấy: 50.0 đơn. Điều chuyển toàn bộ tuyến giao về BC Bảo Lâm 3.', '• Xã Bảo Lâm 2 (Mã: 25084): ')

add_h3('1.3. Khu vực Lâm Hà - Đam Rông')
add_p('Theo báo cáo trình bày của AM Huỳnh Thị Kim Chi:')
add_bullet('Tổng sản lượng AM trình bày: 104.0 đơn/ngày | Giao: 101.0 đơn · Lấy: 3.0 đơn. Hiện do 2 BC phụ trách: (LDO) Đinh Văn Lâm Hà (~700 đơn/ngày/12 NV) và (LDO) Nam Ban Lâm Hà. AM đề xuất chuyển phần xã cũ Phi Tô [60.0 đơn/ngày] từ BC Đinh Văn sang BC Nam Ban. Sau quy hoạch: BC Đinh Văn còn ~640 đơn (11 NV), BC Nam Ban nâng lên ~510 đơn (9 NV). AM tìm mặt bằng mới rộng hơn MB 96m2 hiện tại.', '• Xã Nam Hà Lâm Hà (Mã: 24883): ')
add_bullet('Tổng sản lượng AM trình bày: 59.0 đơn/ngày | Giao: 56.0 đơn · Lấy: 3.0 đơn. Hiện do 2 BC phụ trách: (LDO) Lang Biang - Đà Lạt 1 và (LDO) Đam Rông 3.', '• Xã Đam Rông 4 (Mã: 24853): ')
add_callout('LÝ DO GIỮ NGUYÊN TUYẾN THÔN ĐƯNG KNỚ THUỘC BC LANG BIANG (AM Huỳnh Thị Kim Chi):\nThôn Đưng Knớ (sau sáp nhập thuộc Xã Đam Rông 4) hiện do BC Lang Biang - Đà Lạt 1 phụ trách với sản lượng ~10 đơn/ngày. Khoảng cách từ BC Đam Rông 3 đến Thôn Đưng Knớ là >50 km (phạm vi địa bàn rộng >30 km), đường núi rất xa và ít đơn. Do đó, AM đề xuất GIỮ NGUYÊN tuyến cover Thôn Đưng Knớ thuộc BC Lang Biang - Đà Lạt 1 quản lý như hiện tại.', 'LÝ DO GIỮ NGUYÊN TUYẾN XÃ ĐAM RÔNG 4')

add_img('lam_ha_tuyen_cover.png', 'Sơ đồ Tuyến Cover hiện tại Khu vực Lâm Hà - Đam Rông')
add_img('lam_ha_ban_do_quy_hoach.png', 'Bản đồ Quy hoạch Khu vực Lâm Hà - Đam Rông theo ĐVHC Mới 2026')

# 2. KHÁNH HÒA
add_h2('2. Tỉnh Khánh Hòa')
add_h3('2.1. Thành phố Nha Trang')
add_p('Theo báo cáo trình bày của AM Thái Thị Thanh Thư & AM Phan Đình Duy:')

add_bullet('Tổng sản lượng AM trình bày: 2.418.0 đơn/ngày (4.664.0 kg/ngày) | Giao: 1.340.0 đơn · Lấy: 1.078.0 đơn. Hiện do 3 BC phụ trách: (KHO) Nam Nha Trang 3 (45.9% - 1.109 đơn), (KHO) Nam Nha Trang 1 (34.8% - 842 đơn), (KHO) Nam Nha Trang 5 (19.3% - 467 đơn). Sáp nhập từ Phường Phước Hải, Phước Long, Vĩnh Trường, Vĩnh Thái, Phước Đồng.', '• Phường Nam Nha Trang (Mã: 22402): ')
add_callout('LÝ DO GIỮ NGUYÊN 2 BƯU CỤC PHỤ TRÁCH (AM Thái Thị Thanh Thư):\n1. Sản lượng Phường Nam Nha Trang quá lớn (2.418 đơn/ngày), một bưu cục không thể tải nổi kho bãi và nhân sự.\n2. Bưu cục (KHO) Nam Nha Trang 5 nằm tại Phước Đồng là khu vực xã xa, đường đi hiểm trở, địa bàn nằm biệt lập phía ngoài rìa thành phố giáp ranh đèo và Cam Lâm. Nếu gộp chung về cụm Nam Nha Trang trung tâm sẽ làm khoảng cách di chuyển quá xa (>20km đường đèo).\n=> Phương án AM: Đóng cửa BC Nam Nha Trang 2 & Nam Nha Trang 3. Di dời BC Nam Nha Trang 1 ra vị trí trung tâm phường mới. GIỮ NGUYÊN BC Nam Nha Trang 5 phụ trách cụm xã biệt lập Phước Đồng.', 'LÝ DO GIỮ NGUYÊN BC NAM NHA TRANG 5')

add_bullet('Tổng sản lượng AM trình bày: 1.592.0 đơn/ngày (4.948.0 kg/ngày) | Giao: 1.205.0 đơn · Lấy: 387.0 đơn. Hiện do 3 BC phụ trách: (KHO) Nha Trang (49.7% - 792 đơn), (KHO) Nam Nha Trang 2 (36.9% - 588 đơn), (KHO) Nam Nha Trang 1 (13.3% - 212 đơn). Di dời BC (KHO) Nha Trang ra mặt bằng rộng ở trung tâm phường mới và gộp tuyến các BC cũ về BC Nha Trang duy nhất quản lý.', '• Phường Nha Trang (Mã: 22366): ')
add_bullet('Tổng sản lượng AM trình bày: 747.0 đơn/ngày (Giao: 610.0, Lấy: 137.0). Gộp về 01 BC chính (KHO) Tây Nha Trang (AM Phan Đình Duy).', '• Phường Tây Nha Trang (Mã: 22390): ')

add_img('nha_trang_ban_do_quy_hoach.png', 'Bản đồ Quy hoạch Mạng lưới Bưu cục Thành phố Nha Trang 2026')

add_h3('2.2. Huyện Diên Khánh & Huyện Vạn Ninh')
add_bullet('Tổng sản lượng AM trình bày: 425.0 đơn/ngày. Gộp về BC Diên Khánh 2. BC Diên Khánh 1 phụ trách Vĩnh Thạnh (170 đơn), Vĩnh Trung (110 đơn), Diên An (120 đơn).', '• Xã Diên Khánh (Mã: 22651): ')
add_bullet('Tổng sản lượng AM trình bày: 211.0 đơn/ngày (475.0 kg/ngày) | Giao: 205.0 đơn · Lấy: 6.0 đơn. Hiện do 2 BC phụ trách: (KHO) Tu Bông (72.5% - 153 đơn) và (KHO) Vạn Ninh (27.5% - 58 đơn). AM đề xuất GỘP VỀ 01 BC (KHO) Tu Bông.', '• Xã Vạn Thắng (Mã: 22516): ')
add_callout('LÝ DO GỘP VỀ BC TU BÔNG (AM Phạm Bá Thành Công):\nBưu cục Tu Bông nằm đúng trục địa giới hành chính mới, vị trí kho bãi phù hợp với phạm vi giao nhận của 2 xã cũ (Vạn Thắng + Vạn Bình) và có diện tích kho đủ lớn cho vận hành dài hạn.', 'LÝ DO QUY HOẠCH VẠN NINH')

add_img('van_ninh_ban_do_quy_hoach.jpg', 'Bản đồ Quy hoạch Bưu cục Khu vực Vạn Ninh 2026')

# 3. NINH THUẬN
add_h2('3. Tỉnh Ninh Thuận')
add_p('Theo báo cáo trình bày của AM Nguyễn Duy Long:')
add_bullet('Tổng sản lượng AM trình bày: 802.0 đơn/ngày | Giao: 509.0 đơn · Lấy: 293.0 đơn. Gộp về BC chính (NTH) Phan Rang.', '• Phường Phan Rang (Mã: 22759): ')
add_bullet('Tổng sản lượng AM trình bày: 466.0 đơn/ngày (1.079.0 kg/ngày) | Giao: 299.0 đơn · Lấy: 167.0 đơn. Hiện do 2 BC phụ trách: (NTH) Phan Rang (62.9% - 293 đơn) và (NTH) Ninh Chử (37.1% - 173 đơn). Gộp về BC chính (NTH) Phan Rang.', '• Phường Ninh Chử (Mã: 22834): ')
add_bullet('Tổng sản lượng AM trình bày: 185.0 đơn/ngày (376.0 kg/ngày) | Giao: 178.0 đơn · Lấy: 7.0 đơn. Hiện do 2 BC phụ trách: (NTH) Ninh Chử (75.7% - 140 đơn) và (NTH) Phan Rang (24.3% - 45 đơn). Gộp về BC (NTH) Ninh Chử.', '• Xã Ninh Hải (Mã: 22852): ')
add_bullet('Mở mới 01 Bưu cục: BC (NTH) Đông Hải phụ trách 6 phường ven biển (Đông Hải, Mỹ Bình, Mỹ Đông, Mỹ Hải, Đạo Long, Kinh Dinh). Sản lượng giao: 600 đơn/ngày, Lấy: 250 đơn/ngày. Nhân sự: 1 NVXL - 7 NVPTTT. Đề xuất này giúp giải quyết bài toán khó tuyển dụng nhân sự địa bàn ven biển Đông Hải.', '• Bưu cục Mở mới: ')
add_img('ninh_thuan_ban_do_quy_hoach.png', 'Bản đồ Quy hoạch Mạng lưới Bưu cục Tỉnh Ninh Thuận 2026')

# 4. BÌNH THUẬN
add_h2('4. Tỉnh Bình Thuận')
add_h3('4.1. Cụm Hàm Thuận Bắc & Hàm Liêm')
add_p('Theo báo cáo trình bày của AM Nguyễn Ngọc Khánh:')
add_bullet('Tổng sản lượng AM trình bày: 618.0 đơn/ngày (1.939.0 kg/ngày) | Giao: 545.0 đơn · Lấy: 73.0 đơn. Hiện do 2 BC phụ trách: (BTH) Hàm Liêm (62.1% - 384 đơn) và (BTH) Phú Thủy (37.9% - 234 đơn). GỘP VỀ 01 BC CHÍNH (BTH) Phú Thủy.', '• Phường Hàm Thắng (Mã: 22933): ')
add_bullet('Tổng sản lượng AM trình bày: 905.0 đơn/ngày (2.112.0 kg/ngày) | Giao: 777.0 đơn · Lấy: 128.0 đơn. Hiện do 2 BC phụ trách: (BTH) Hàm Thắng (68.6% - 621 đơn) và (BTH) Phú Thủy (31.4% - 284 đơn). GỘP VỀ 01 BC CHÍNH (BTH) Hàm Thắng.', '• Phường Phan Thiết (Mã: 22945): ')
add_bullet('Tổng sản lượng AM trình bày: 680.0 đơn/ngày (1.577.0 kg/ngày) | Giao: 611.0 đơn · Lấy: 69.0 đơn. Hiện do 2 BC phụ trách: (BTH) Hàm Thắng (82.8% - 563 đơn) và (BTH) Hàm Liêm (17.2% - 117 đơn). GỘP VỀ 01 BC (BTH) Hàm Thắng.', '• Phường Bình Thuận (Mã: 22960): ')
add_img('ham_thuan_bac_ban_do.jpg', 'Bản đồ Quy hoạch Cụm Bưu cục Hàm Thuận Bắc - Hàm Liêm 2026')
add_img('phu_thuy_phan_thiet_ban_do.jpg', 'Bản đồ Quy hoạch Cụm Bưu cục Phú Thủy - Hàm Thắng - Phan Thiết 2026')

add_h3('4.2. Cụm Tánh Linh - Đức Linh')
add_bullet('Mở mới 01 Bưu cục: BC (BTH) Nam Thành (Cover Nam Thành 250 đơn, Nghị Đức 200 đơn; Giao 450-500, Lấy 50-60; Nhân sự: 7 NV, 1 NVXL).', '• Bưu cục Mở mới: ')
add_img('tanh_linh_duc_linh_ban_do.jpg', 'Bản đồ Quy hoạch Bưu cục Tánh Linh - Đức Linh 2026')

# 5. ĐẮK NÔNG
add_h2('5. Tỉnh Đắc Nông')
add_p('Theo báo cáo đánh giá khảo sát thực địa của AM Trần Thị Nhung & AM Trần Văn Phước tại các xã mới: Phường Bắc Gia Nghĩa (Mã: 24611 - 734 đơn/ngày), Phường Nam Gia Nghĩa (Mã: 24615 - 499 đơn/ngày), Xã Tà Đùng (Mã: 24637 - 187 đơn/ngày), Xã Đắc Sắk (Mã: 24678 - 194 đơn/ngày), Xã Đức An (Mã: 24717 - 327 đơn/ngày), Xã Quảng Tân (Mã: 24748 - 201 đơn/ngày):')

add_callout('LÝ DO ĐỀ XUẤT TẠM THỜI GIỮ NGUYÊN 100% BƯU CỤC TẠI ĐẮK NÔNG (AM Trần Thị Nhung):\n1. Rủi ro đứt gãy nhân sự 100%: 100% nhân sự hiện tại tại các điểm Bưu cục xã ven khẳng định KHÔNG THEO BC MỚI nếu sáp nhập cưỡng ép do địa hình miền núi chia cắt, khoảng cách di chuyển quá xa. Việc gộp đường đột sẽ dẫn đến nghỉ việc hàng loạt và sập toàn bộ tuyến giao.\n2. Địa lý chia cắt: Như Xã Đắk Sắk (76.3% do BC Đức Lập đảm nhận, phần Nam Xuân do BC Krông Nô), Xã Đức An (phần Đắk N\'Drung gần BC Trường Xuân hơn BC Đức An).\n3. Tham khảo đối thủ: 100% đối thủ trong ngành (GHTK, Viettel Post, Shopee Express, J&T) đều TẠM THỜI GIỮ NGUYÊN Bưu cục theo xã cũ, chưa đơn vị nào gộp tuyến theo xã mới.\n=> Đề xuất: Kính đề nghị Ban Lãnh đạo cho phép TẠM THỜI GIỮ NGUYÊN 100% Bưu cục và phạm vi quản lý theo địa giới cũ tại Đắk Nông trong 6 tháng tới.', 'ĐỀ XUẤT GIỮ NGUYÊN ĐẮK NÔNG')

add_h1('IV. BẢNG TỔNG HỢP BIẾN ĐỘNG MẠNG LƯỚI BƯU CỤC TOÀN VÙNG NTB')
add_p('Dưới đây là tổng hợp toàn bộ các thay đổi về quy mô bưu cục trên toàn vùng NTB theo Kế hoạch Quy hoạch 2026:')

# Add Table for Network Changes
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
headers = ['Loại Biến Động', 'Tên Bưu Cục', 'Tỉnh / Khu Vực', 'Chi Tiết Phương Án & Nhân Sự']
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    set_cell_bg(hdr_cells[i], '1F497D')
    for p in hdr_cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.name = 'Calibri'

changes_data = [
    ('MỞ MỚI', 'BC (LDO) Xuân Hương - Đà Lạt 2', 'Lâm Đồng (Đà Lạt)', 'Mở mới tại Phường 10 phụ trách Phường 3 + Phường 10 cũ.'),
    ('MỞ MỚI', 'BC (NTH) Đông Hải', 'Ninh Thuận (Phan Rang)', 'Mở mới phụ trách 6 phường coastal, 1 NVXL - 7 NVPTTT.'),
    ('MỞ MỚI', 'BC (BTH) Nam Thành', 'Bình Thuận (Tánh Linh)', 'Mở mới phụ trách Nam Thành & Nghị Đức, 1 NVXL - 7 NVPTTT.'),
    ('ĐÓNG CỬA', 'BC (LDO) 1 Bảo Lộc', 'Lâm Đồng (Bảo Lộc)', 'Đóng cửa, gộp toàn bộ tuyến và nhân sự về BC B\'Lao.'),
    ('ĐÓNG CỬA', 'BC (KHO) Nam Nha Trang 2', 'Khánh Hòa (Nha Trang)', 'Đóng cửa, gộp tuyến về BC Nam Nha Trang 1 & Nam Nha Trang 5.'),
    ('ĐÓNG CỬA', 'BC (KHO) Nam Nha Trang 3', 'Khánh Hòa (Nha Trang)', 'Đóng cửa, gộp tuyến về BC Nam Nha Trang 1 & Nam Nha Trang 5.'),
    ('DI DỜI / MỞ RỘNG', 'BC (KHO) Nha Trang', 'Khánh Hòa (Nha Trang)', 'Di dời ra mặt bằng rộng hơn ở trung tâm Phường Nha Trang mới.'),
    ('DI DỜI / MỞ RỘNG', 'BC (KHO) Nam Nha Trang 1', 'Khánh Hòa (Nha Trang)', 'Di dời ra vị trí trung tâm Phường Nam Nha Trang mới.'),
    ('DI DỜI / MỞ RỘNG', 'BC (LDO) Nam Ban Lâm Hà', 'Lâm Đồng (Lâm Hà)', 'Tìm mặt bằng rộng hơn MB 96m2 hiện tại để gộp tuyến xã Phi Tô.'),
    ('DI DỜI / MỞ RỘNG', 'BC (BTH) Hàm Thuận & Hàm Liêm', 'Bình Thuận (Hàm Thuận B.)', 'Di dời vị trí kho bãi phù hợp với địa giới hành chính mới.')
]

for row_idx, data in enumerate(changes_data):
    row_cells = table.add_row().cells
    bg = 'F9FAFB' if row_idx % 2 == 0 else 'FFFFFF'
    for i, val in enumerate(data):
        row_cells[i].text = val
        set_cell_bg(row_cells[i], bg)
        for p in row_cells[i].paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9.5)
                r.font.name = 'Calibri'
                if i == 0:
                    r.font.bold = True
                    if val == 'MỞ MỚI':
                        r.font.color.rgb = RGBColor(38, 128, 0)
                    elif val == 'ĐÓNG CỬA':
                        r.font.color.rgb = RGBColor(192, 0, 0)
                    else:
                        r.font.color.rgb = SECONDARY_COLOR

add_h1('V. LỘ TRÌNH TRIỂN KHAI & ĐỀ XUẤT BAN LÃNH ĐẠO')
add_bullet('Điều chuyển lập tức 18 tuyến giao chéo xa ranh giới (30-40km) về các Bưu cục lân cận gần hơn để giảm chi phí di chuyển.', '• Giai đoạn 1 (Tháng 1 - 2/2026): ')
add_bullet('Tiến hành khảo sát và ký hợp đồng mặt bằng di dời cho các BC Nha Trang, Nam Nha Trang 1, Nam Ban, Hàm Thuận, Hàm Liêm; Đàm phán thanh lý HĐ mặt bằng các BC đóng cửa (Nam Nha Trang 2, 3, 1 Bảo Lộc).', '• Giai đoạn 2 (Tháng 3 - 4/2026): ')
add_bullet('Khai trương 03 Bưu cục mở mới (Xuân Hương 2, Đông Hải, Nam Thành); Hoàn tất việc điều chuyển nhân sự và chính thức áp dụng ranh giới ĐVHC mới trên hệ thống.', '• Giai đoạn 3 (Tháng 5 - 6/2026): ')

add_callout('Kính đề nghị Ban Lãnh đạo xem xét phê duyệt Kế hoạch Quy hoạch Mạng lưới Bưu cục Vùng NTB 2026 để Khối Vận hành và các AM chủ động triển khai theo lộ trình.', 'KẾT LUẬN & KIẾN NGHỊ')

# Save Docx - Save to V3
out_alt_v3 = r"C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Theo_DVHC_Moi_V3.docx"
out_workspace = r"Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Theo_DVHC_Moi_Co_Hinh_AM.docx"

doc.save(out_alt_v3)
doc.save(out_workspace)
print(f"Successfully saved updated DOCX to:\n  - {out_alt_v3}\n  - {out_workspace}")
