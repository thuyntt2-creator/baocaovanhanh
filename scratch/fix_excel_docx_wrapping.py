import docx, openpyxl, sys
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

# 1. Read source data from docx
doc_src = docx.Document(r'C:\Users\lap4all\Downloads\Quy_Hoach MANG LUOI NTB.docx')

parsed_36 = []
current = None

for p in doc_src.paragraphs:
    txt = p.text.strip()
    if not txt: continue
    
    is_heading = False
    for n in range(1, 37):
        if txt.startswith(f'{n}. '):
            is_heading = True
            break
            
    if is_heading:
        if current: parsed_36.append(current)
        current = {'title': txt, 'communes': '', 'proposal': '', 'reason': ''}
    elif current:
        if 'Các xã cũ sáp nhập' in txt:
            current['communes'] = txt.replace('❖ Các xã cũ sáp nhập & Sản lượng Giao/Lấy từng xã:', '').strip()
        elif 'ĐỀ XUẤT PHƯƠNG ÁN' in txt:
            current['proposal'] = txt.replace('❖ ĐỀ XUẤT PHƯƠNG ÁN CỦA AM:', '').replace('❖ ĐỀ XUẤT PHƯƠNG ÁN:', '').strip()
        elif 'LÝ DO VÀ GIẢI THÍCH' in txt:
            current['reason'] = txt.replace('❖ LÝ DO VÀ GIẢI THÍCH NGUYÊN NHÂN CHI TIẾT TỪ AM:', '').replace('❖ LÝ DO VÀ GIẢI THÍCH NGUYÊN NHÂN CHI TIẾT:', '').strip()
        elif current['communes'] and not current['proposal'] and not txt.startswith('❖'):
            current['communes'] += ' ' + txt
        elif current['proposal'] and not current['reason'] and not txt.startswith('❖'):
            current['proposal'] += ' ' + txt
        elif current['reason'] and not txt.startswith('❖'):
            current['reason'] += ' ' + txt

if current: parsed_36.append(current)

# ==============================================================================
# EXCEL FILE GENERATION (WITH FULL AUTO-WRAP TEXT & DYNAMIC ROW HEIGHTS)
# ==============================================================================
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Quy_Hoach_NTB_36_DVHC"

font_header = Font(name="Calibri", size=11, bold=True, color="000000")
fill_header = PatternFill(start_color="F7A059", end_color="F7A059", fill_type="solid")

font_dvhc = Font(name="Calibri", size=10, bold=True, color="000000")
font_normal = Font(name="Calibri", size=9.5, color="000000")
font_buucuc = Font(name="Calibri", size=10, bold=True, color="003399")

align_header = Alignment(horizontal="center", vertical="center", wrap_text=True)
align_cell = Alignment(horizontal="left", vertical="top", wrap_text=True) # wrap_text=True IS CRITICAL!

thin_border = Border(
    left=Side(style='thin', color='000000'), right=Side(style='thin', color='000000'),
    top=Side(style='thin', color='000000'), bottom=Side(style='thin', color='000000')
)

headers = ["STT & Tên ĐVHC Mới", "Các Xã / Phường Cũ sáp nhập & Sản lượng Giao/Lấy", "Bưu cục Cover (Đề xuất AM)", "Phương án Quy hoạch & Lý do chi tiết từ AM"]

for col_idx, h_text in enumerate(headers, start=1):
    cell = ws.cell(row=1, column=col_idx, value=h_text)
    cell.font = font_header
    cell.fill = fill_header
    cell.alignment = align_header
    cell.border = thin_border

ws.row_dimensions[1].height = 28

ws.column_dimensions['A'].width = 25
ws.column_dimensions['B'].width = 45
ws.column_dimensions['C'].width = 35
ws.column_dimensions['D'].width = 50

for r_idx, item in enumerate(parsed_36, start=2):
    c1 = ws.cell(row=r_idx, column=1, value=item['title'])
    c2 = ws.cell(row=r_idx, column=2, value=item['communes'])
    c3 = ws.cell(row=r_idx, column=3, value=item['proposal'])
    c4 = ws.cell(row=r_idx, column=4, value=item['reason'] if item['reason'] else item['proposal'])
    
    c1.font = font_dvhc
    c2.font = font_normal
    c3.font = font_buucuc
    c4.font = font_normal
    
    # Calculate required row height dynamically so no text is hidden!
    max_len = max(len(item['title'])/20, len(item['communes'])/40, len(item['proposal'])/30, len(str(item['reason']))/45)
    line_count = int(max_len) + 1
    calc_height = max(35, line_count * 18)
    ws.row_dimensions[r_idx].height = calc_height
    
    for col_idx in range(1, 5):
        cell = ws.cell(row=r_idx, column=col_idx)
        cell.border = thin_border
        cell.alignment = align_cell

excel_out = r'C:\Users\lap4all\Downloads\Bao_Cao_Full_36_DVHC_Hien_Thi_100_Phan_Tram.xlsx'
wb.save(excel_out)
print(f"Generated perfect Excel file: {excel_out}")

# ==============================================================================
# WORD FILE GENERATION (WITH FULL WRAPPING & MARGINS & LANDSCAPE A4)
# ==============================================================================
doc_out = docx.Document()

# Set Landscape A4 for wide table display
for section in doc_out.sections:
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Title
p_t = doc_out.add_paragraph()
r_t = p_t.add_run('TỜ TRÌNH QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB NĂM 2026\n(BẢNG TỔNG HỢP FULL 36 ĐƠN VỊ HÀNH CHÍNH MỚI - HIỂN THỊ TRỌN VẸN 100% NỘI DUNG)')
r_t.font.name = 'Calibri'
r_t.font.size = Pt(16)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(180, 0, 0)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(12)

table = doc_out.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

widths = [Inches(2.2), Inches(3.2), Inches(2.3), Inches(2.8)]

hdr_cells = table.rows[0].cells
for i, h_text in enumerate(headers):
    hdr_cells[i].text = h_text
    hdr_cells[i].width = widths[i]
    set_cell_background(hdr_cells[i], 'F7A059')
    set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0, 0, 0)

for r_idx, item in enumerate(parsed_36):
    row_cells = table.add_row().cells
    bg_color = 'FFFFFF'
    
    row_cells[0].text = item['title']
    row_cells[0].width = widths[0]
    
    row_cells[1].text = item['communes']
    row_cells[1].width = widths[1]
    
    row_cells[2].text = item['proposal']
    row_cells[2].width = widths[2]
    
    row_cells[3].text = item['reason'] if item['reason'] else item['proposal']
    row_cells[3].width = widths[3]
    
    for c_idx in range(4):
        set_cell_background(row_cells[c_idx], bg_color)
        set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=80, right=80)
        p = row_cells[c_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            if c_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif c_idx == 2:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 153)

word_out = r'C:\Users\lap4all\Downloads\Bao_Cao_Full_36_DVHC_Hien_Thi_100_Phan_Tram.docx'
doc_out.save(word_out)
print(f"Generated perfect Word file: {word_out}")
