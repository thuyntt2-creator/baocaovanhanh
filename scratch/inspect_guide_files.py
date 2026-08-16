import os
import glob
import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

downloads_dir = r"C:\Users\lap4all\Downloads"

docx_path = os.path.join(downloads_dir, "Huong_dan_AOP_Hang_Nang_cho_GDV.docx")
if os.path.exists(docx_path):
    print("=== paragraphs in docx ===")
    doc = docx.Document(docx_path)
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    for i in range(min(20, len(doc.paragraphs))):
        print(f"P{i}: {doc.paragraphs[i].text}")
else:
    print("Docx not found!")

html_path = os.path.join(downloads_dir, "Huong_dan_chi_tiet_lap_AOP_Hang_Nang.html")
if os.path.exists(html_path):
    print("\n=== lines in html ===")
    with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
        print(f"Length of HTML: {len(content)}")
        lines = content.split('\n')
        for i in range(min(30, len(lines))):
            print(f"L{i}: {lines[i].strip()[:100]}")
else:
    print("HTML not found!")
