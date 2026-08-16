import docx, sys

sys.stdout.reconfigure(encoding='utf-8')

src_path = r'C:\Users\lap4all\Downloads\Quy_Hoach_MANG_LUOI_NTB_Co_Nha_Trang_Final.docx'
doc = docx.Document(src_path)

sec5_p_idx = -1
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if 'V. TỔNG HỢP BIẾN ĐỘNG' in txt or 'V.TỔNG HỢP BIẾN ĐỘNG' in txt:
        sec5_p_idx = i

for p in doc.paragraphs[sec5_p_idx:]:
    txt = p.text.strip()
    if 'Bưu cục Di dời / Mở rộng' in txt:
        p.text = "❖ Bưu cục Di dời / Mở rộng mặt bằng kho (02 BC): BC Tây Nha Trang (Khánh Hòa - di dời MB mới), BC Bắc Cam Ranh (Khánh Hòa - mở rộng từ 100m²)."

doc.save(src_path)
print(f"Updated relocation list in {src_path}!")
