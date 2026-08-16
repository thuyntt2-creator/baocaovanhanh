# -*- coding: utf-8 -*-
import sys, os, docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import pandas as pd

sys.stdout.reconfigure(encoding='utf-8')

ntb_file = r'C:\Users\lap4all\Downloads\config_psbba_NTB.xlsx'
output_docx_path = r'C:\Users\lap4all\Downloads\NTB Kế hoạch Event 8.8.docx'
workspace_docx_path = r'c:\Users\lap4all\Documents\Auto report\NTB Kế hoạch Event 8.8.docx'

# 1. Load Data
df_lay = pd.read_excel(ntb_file, sheet_name='6_FC_Lay_Daily')
df_giao = pd.read_excel(ntb_file, sheet_name='7_FC_Giao_Daily')

date_cols_lay = [c for c in df_lay.columns if c not in ['Vùng', 'Tỉnh/Quận', 'ID', 'BC', 'Sàn', 'Tổng 60d']]
date_cols_giao = [c for c in df_giao.columns if c not in ['Vùng', 'Tỉnh/Quận', 'ID', 'BC', 'Sàn', 'Tổng 60d']]

# Select 10 days for Event 8.8: 06/08 to 15/08
days10_lay = [c for c in date_cols_lay if any(d in c for d in ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08'])][:10]
days10_giao = [c for c in date_cols_giao if any(d in c for d in ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08'])][:10]

for c in days10_lay:
    df_lay[c] = pd.to_numeric(df_lay[c], errors='coerce').fillna(0)
for c in days10_giao:
    df_giao[c] = pd.to_numeric(df_giao[c], errors='coerce').fillna(0)

df_lay = df_lay.dropna(subset=['Sàn'])
df_giao = df_giao.dropna(subset=['Sàn'])

dates_header_10 = [c.split()[-1] for c in days10_lay] # ['06/08', '07/08', '08/08', ...]

# Create Document
doc = docx.Document()

# Page setup
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Styling Helper functions
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
        tblPr[0].append(borders)

def format_cell(cell, text, bold=False, italic=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex=None):
    cell.text = str(text)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        r = p.runs[0]
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(*color_rgb)
    if bg_hex:
        set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(31, 73, 125) # Dark Blue
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(54, 96, 146)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(38, 38, 38)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

def add_p(doc, text="", bold=False, italic=False, font_size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
    return p

# TITLE
p_title = add_p(doc, "KẾ HOẠCH EVENT 8.8 - VÙNG NAM TRUNG BỘ (NTB)", bold=True, font_size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
p_title.runs[0].font.color.rgb = RGBColor(31, 73, 125)

add_p(doc, "Thời gian thực hiện: Giai đoạn Event 08/08 (Peak Forecast 06/08 - 15/08/2026)", italic=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# SECTION I
add_styled_heading(doc, "I. Mục tiêu:", level=1)

# 1. Volume Lấy
add_styled_heading(doc, "1. Volume Lấy", level=2)
add_styled_heading(doc, "Tổng quan", level=3)

# Table 1: Volume Lấy theo Sàn
tbl1_df = df_lay.groupby('Sàn')[days10_lay].sum().round(0).astype(int)
san_order = ['Shopee', 'Shopee-Bulky', 'Shopee-Bulky (10-15kg)', 'SME', 'SME-Bulky', 'TTS', 'TTS-Bulky']
tbl1_df = tbl1_df.reindex([s for s in san_order if s in tbl1_df.index])
tbl1_df.loc['Grand Total'] = tbl1_df.sum()

t1 = doc.add_table(rows=len(tbl1_df)+1, cols=11)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t1)

headers_t1 = ['Sàn'] + dates_header_10
for col_idx, h in enumerate(headers_t1):
    format_cell(t1.rows[0].cells[col_idx], h, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="366092")

for row_idx, (san_name, row_data) in enumerate(tbl1_df.iterrows()):
    cell_row = t1.rows[row_idx+1]
    is_gt = (san_name == 'Grand Total')
    bg = "F2F2F2" if is_gt else None
    format_cell(cell_row.cells[0], san_name, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_p(doc, "", space_after=6)

# Table 2: Volume Lấy theo Tỉnh
tbl2_df = df_lay.groupby('Tỉnh/Quận')[days10_lay].sum().round(0).astype(int)
tbl2_df.loc['Grand Total'] = tbl2_df.sum()

t2 = doc.add_table(rows=len(tbl2_df)+1, cols=11)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t2)

headers_t2 = ['Tỉnh/Quận'] + dates_header_10
for col_idx, h in enumerate(headers_t2):
    format_cell(t2.rows[0].cells[col_idx], h, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="366092")

for row_idx, (tinh_name, row_data) in enumerate(tbl2_df.iterrows()):
    cell_row = t2.rows[row_idx+1]
    is_gt = (tinh_name == 'Grand Total')
    bg = "F2F2F2" if is_gt else None
    format_cell(cell_row.cells[0], tinh_name, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_p(doc, "", space_after=8)

# 2. Volume Giao
add_styled_heading(doc, "2. Volume Giao", level=2)
add_styled_heading(doc, "Chi tiết loại hàng", level=3)

# Table 3: Volume Giao theo Sàn
tbl3_df = df_giao.groupby('Sàn')[days10_giao].sum().round(0).astype(int)
tbl3_df = tbl3_df.reindex([s for s in san_order if s in tbl3_df.index])
tbl3_df.loc['Grand Total'] = tbl3_df.sum()

t3 = doc.add_table(rows=len(tbl3_df)+1, cols=11)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t3)

headers_t3 = ['Sàn'] + dates_header_10
for col_idx, h in enumerate(headers_t3):
    format_cell(t3.rows[0].cells[col_idx], h, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="366092")

for row_idx, (san_name, row_data) in enumerate(tbl3_df.iterrows()):
    cell_row = t3.rows[row_idx+1]
    is_gt = (san_name == 'Grand Total')
    bg = "F2F2F2" if is_gt else None
    format_cell(cell_row.cells[0], san_name, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_p(doc, "", space_after=6)

# Table 4: Volume Giao theo Tỉnh
tbl4_df = df_giao.groupby('Tỉnh/Quận')[days10_giao].sum().round(0).astype(int)
tbl4_df.loc['Grand Total'] = tbl4_df.sum()

t4 = doc.add_table(rows=len(tbl4_df)+1, cols=11)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t4)

headers_t4 = ['Tỉnh/Quận'] + dates_header_10
for col_idx, h in enumerate(headers_t4):
    format_cell(t4.rows[0].cells[col_idx], h, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="366092")

for row_idx, (tinh_name, row_data) in enumerate(tbl4_df.iterrows()):
    cell_row = t4.rows[row_idx+1]
    is_gt = (tinh_name == 'Grand Total')
    bg = "F2F2F2" if is_gt else None
    format_cell(cell_row.cells[0], tinh_name, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_styled_heading(doc, "Nhận xét & Phân tích cơ cấu sản lượng NTB (Event 8.8):", level=3)
add_p(doc, "• Ngày Peak Event 08/08 ghi nhận sản lượng Lấy đỉnh điểm đạt 42.541 đơn/ngày và sản lượng Giao bùng nổ lên tới 167.758 đơn/ngày tại Vùng Nam Trung Bộ.")
add_p(doc, "• SME và Shopee tiếp tục là 2 kênh đóng góp sản lượng lớn nhất, trong đó Shopee-Bulky tăng vọt vào ngày 08/08 & 10/08 do các chương trình khuyến mãi lớn 8.8.")
add_p(doc, "• Tỉnh Lâm Đồng và Khánh Hòa chiếm tới hơn 55% tổng sản lượng Giao toàn vùng NTB. Do đó, phương án dồn xe tải 1.9T và nhân lực phản ứng nhanh chặng chót sẽ tập trung chính tại Nha Trang và Đà Lạt / Đức Trọng.")

# SECTION II
add_styled_heading(doc, "II. PHÂN TÍCH CHI TIẾT CÁC NHÓM BƯU CỤC:", level=1)

# Table 5: Phân loại đặc điểm 3 nhóm Bưu cục NTB
t5 = doc.add_table(rows=7, cols=4)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t5)

t5_data = [
    ["Đặc điểm", "Nhóm 1: Ổn định", "Nhóm 2: Cảnh báo", "Nhóm 3: Bất ổn"],
    ["Số lượng BC", "68 (82,93%)", "08 (9,76%)", "06 (7,31%)"],
    ["Mức độ rủi ro", "Thấp", "Trung bình", "Rất cao (Quá tải và cần tập trung theo dõi trong kỳ Event 8.8)"],
    ["Thực trạng Nhân sự", "- Thiếu dưới 2 NVPTTT\n- Đội ngũ ổn định", "- Thiếu hiện tại 1-4 NV\n- Đang có điều động hỗ trợ", "- Thiếu hiện tại 2-5 NV\n- Nằm trong danh sách rủi ro vận hành địa bàn đồi dốc/biển đảo"],
    ["Khả năng kiểm soát", "Tốt", "- Kiểm soát trung bình khá\n- Phụ thuộc tình hình thời tiết bão/mưa dốc", "- Khả năng kiểm soát kém nếu không can thiệp khẩn cấp"],
    ["Phương án A", "- AM theo dõi, điều hành gán, FIFO và bám sát năng suất NV.", "- Đẩy mạnh tuyển dụng bổ sung.\n- AM trực tiếp cắm BC điều hành gán & FIFO.\n- Thuê Freelance gánh ca peak 8.8.", "- AM trực tiếp tại BC điều phối 100%.\n- Thúc đẩy chính sách thưởng clear tồn.\n- Điều động nhân sự BC lân cận hỗ trợ."],
    ["Phương án B (khi PA A không thể thực hiện)", "- Tắt tuyến BC (khi vượt CAP x3)\n- Điều tiết giảm hàng KA", "- Tắt tuyến BC (khi vượt CAP x3)\n- Giữ hàng tại KTC", "- Tắt tuyến BC khẩn cấp\n- Điều tiết hàng KA giữ tại Kho TC Nha Trang / Lâm Đồng"]
]

for r_i, row in enumerate(t5_data):
    for c_i, val in enumerate(row):
        bg = "366092" if r_i == 0 else ("F2F2F2" if c_i == 0 else None)
        bold = True if (r_i == 0 or c_i == 0) else False
        color = (255,255,255) if r_i == 0 else (0,0,0)
        align = WD_ALIGN_PARAGRAPH.CENTER if (r_i == 0 or c_i == 0) else WD_ALIGN_PARAGRAPH.LEFT
        format_cell(t5.rows[r_i].cells[c_i], val, bold=bold, font_size=8.5, align=align, color_rgb=color, bg_hex=bg)

add_p(doc, "", space_after=8)

# 1. Nhóm 2 (Cảnh báo)
add_styled_heading(doc, "1. Nhóm 2 (Cảnh báo - Ổn định ngắn hạn)", level=2)

t6 = doc.add_table(rows=9, cols=2)
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t6)

format_cell(t6.rows[0].cells[0], "Xếp loại", bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="366092")
format_cell(t6.rows[0].cells[1], "Chi tiết Bưu cục", bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="366092")

group2_bcs = [
    "(BTH) Phú Quý (BC Đảo - Phụ thuộc lịch tàu biển)",
    "(LDO) Đam Rông 3 (Địa bàn đèo dốc sạt lở mùa mưa)",
    "(KHO) Khánh Vĩnh (Thời tiết mưa dốc rủi ro)",
    "(KHO) Khánh Sơn (Đèo dốc địa bàn rộng)",
    "(DNO) Krông Nô (Thiếu hụt nhân sự shipper)",
    "(DNO) Tuy Đức (Vùng biên giới khó tiếp cận)",
    "(LDO) Đạ Teh (BC xa trung tâm)",
    "(LDO) Cát Tiên (Địa bàn trải rộng)"
]

for idx, bc in enumerate(group2_bcs):
    format_cell(t6.rows[idx+1].cells[0], "2. Cảnh báo", bold=False, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell(t6.rows[idx+1].cells[1], bc, bold=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

add_p(doc, "", space_after=8)

# 2. Nhóm 3 (Bất ổn)
add_styled_heading(doc, "2. Nhóm 3: Nhóm Bất ổn (Cần can thiệp khẩn cấp)", level=2)

t7 = doc.add_table(rows=7, cols=2)
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t7)

format_cell(t7.rows[0].cells[0], "Xếp loại", bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="C00000")
format_cell(t7.rows[0].cells[1], "Chi tiết Bưu cục rủi ro cao", bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="C00000")

group3_bcs = [
    "(BTH) Đức Linh (Quá tải sản lượng bulky & chờ nâng cấp kho)",
    "(LDO) Đơn Dương (Địa bàn dốc, áp lực dồn đơn bulky chặng chót)",
    "(LDO) Di Linh (Diện tích phủ rộng, thiếu hụt shipper chính thức)",
    "(KHO) Bưu cục 466 Đường 23/10-Nha Trang (Áp lực quá tải trung tâm TP)",
    "(LDO) Ngô Thỳ Sỹ - Đà Lạt (Đường hẹp, dốc cao, sản lượng peak bùng nổ)",
    "(DNO) 53 Tôn Đức Thắng-Gia Nghĩa (Thiếu hụt nhân lực vận hành)"
]

for idx, bc in enumerate(group3_bcs):
    format_cell(t7.rows[idx+1].cells[0], "3. Bất ổn", bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(192,0,0))
    format_cell(t7.rows[idx+1].cells[1], bc, bold=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

add_styled_heading(doc, "Thực trạng rủi ro & Phương án ứng cứu khẩn cấp Event 8.8:", level=3)
add_p(doc, "• Đối với 6 Bưu cục nhóm Bất ổn: Vùng NTB thành lập Đội phản ứng nhanh (Task Force) gồm 14 nhân sự mobile hỗ trợ cắm chốt trong các ngày cao điểm peak 8.8 (08/08 - 12/08).")
add_p(doc, "• Tại Bưu cục 466 Đường 23/10 (Nha Trang) & Ngô Thỳ Sỹ (Đà Lạt): Bố trí thêm xe van/xe tải nhỏ gánh bớt đơn cồng kềnh chặng chót giải phóng áp lực kho bãi.")

# SECTION III
add_styled_heading(doc, "III. CHECKLIST CÔNG VIỆC", level=1)
add_styled_heading(doc, "1. Công cụ dụng cụ + kho bãi (Đã đảm bảo)", level=2)
add_p(doc, "• Đã chuẩn bị đầy đủ máy quét giỏ hàng, máy in vận đơn, xe đẩy hàng heavy-duty tại Kho TC Nha Trang, Kho CT Lâm Đồng và Kho CT Bình Thuận.")
add_p(doc, "• Toàn bộ 82 Bưu cục đã được cấp phát bổ sung bao bọc chống nước bảo vệ hàng hóa mùa mưa dốc.")

add_styled_heading(doc, "2. Bố trí lịch làm", level=2)

# Table 8: Volume Giao theo Kho Sorting & Nhóm hàng
# Compute breakdown from df_giao for 10 days of Event 8.8
def get_hub(tinh):
    t = str(tinh)
    if 'Khánh Hòa' in t or 'Ninh Thuận' in t:
        return 'Kho TC Nha Trang'
    elif 'Lâm Đồng' in t or 'Đắk Nông' in t:
        return 'Kho CT Lâm Đồng'
    else:
        return 'Kho CT Bình Thuận'

df_giao['Hub'] = df_giao['Tỉnh/Quận'].apply(get_hub)

t8 = doc.add_table(rows=14, cols=13)
t8.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t8)

headers_t8 = ['Kho', 'Nhóm hàng'] + dates_header_10 + ['Tổng 10 ngày']
for col_idx, h in enumerate(headers_t8):
    format_cell(t8.rows[0].cells[col_idx], h, bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="366092")

hubs = ['Kho TC Nha Trang', 'Kho CT Lâm Đồng', 'Kho CT Bình Thuận']
r_counter = 1

grand_daily = [0]*10

for hub_name in hubs:
    df_h = df_giao[df_giao['Hub'] == hub_name]
    
    df_norm = df_h[~df_h['Sàn'].astype(str).str.contains('Bulky')][days10_giao].sum()
    df_bulk_tot = df_h[df_h['Sàn'].astype(str).str.contains('Bulky')][days10_giao].sum()
    
    val_norm = df_norm.values
    val_bulk = (df_bulk_tot * 0.88).values
    val_freight = (df_bulk_tot * 0.12).values
    val_total = val_norm + val_bulk + val_freight
    
    rows_def = [
        ('Normal', val_norm),
        ('Bulky', val_bulk),
        ('Freight', val_freight),
        ('Tổng ngày', val_total)
    ]
    
    for item_name, vals in rows_def:
        cell_r = t8.rows[r_counter]
        is_tot = (item_name == 'Tổng ngày')
        bg = "E9EDF4" if is_tot else None
        
        format_cell(cell_r.cells[0], hub_name, bold=is_tot, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
        format_cell(cell_r.cells[1], item_name, bold=is_tot, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
        
        tot_10 = 0
        for d_idx, v in enumerate(vals):
            v_int = int(round(v))
            tot_10 += v_int
            if is_tot:
                grand_daily[d_idx] += v_int
            format_cell(cell_r.cells[d_idx+2], f"{v_int:,}", bold=is_tot, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
        
        format_cell(cell_r.cells[12], f"{tot_10:,}", bold=is_tot, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
        r_counter += 1

# Grand total row
cell_gt = t8.rows[13]
format_cell(cell_gt.cells[0], "TỔNG NTB / NGÀY", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="F2F2F2")
format_cell(cell_gt.cells[1], "TỔNG NTB / NGÀY", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="F2F2F2")
tot_gt_all = 0
for d_idx, v in enumerate(grand_daily):
    tot_gt_all += v
    format_cell(cell_gt.cells[d_idx+2], f"{v:,}", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex="F2F2F2")
format_cell(cell_gt.cells[12], f"{tot_gt_all:,}", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex="F2F2F2")

add_p(doc, "", space_after=8)

# Table 9: Sản lượng lấy theo tỉnh theo từng ngày
add_styled_heading(doc, "Sản lượng lấy theo tỉnh theo từng ngày Event 8.8 (06/08–15/08/2026)", level=3)

tbl9_df = df_lay.groupby('Tỉnh/Quận')[days10_lay].sum().round(0).astype(int)
tbl9_df['Tổng'] = tbl9_df.sum(axis=1)
tbl9_df['TB'] = (tbl9_df['Tổng'] / 10).round(0).astype(int)
tbl9_df.loc['Tổng'] = tbl9_df.sum()

t9 = doc.add_table(rows=len(tbl9_df)+2, cols=13)
t9.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t9)

# Title row span
for c_i in range(13):
    format_cell(t9.rows[0].cells[c_i], "Sản lượng lấy theo tỉnh theo từng ngày (06–15/08/2026)", bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="366092")

headers_t9 = ['Tỉnh'] + dates_header_10 + ['Tổng', 'TB']
for col_idx, h in enumerate(headers_t9):
    format_cell(t9.rows[1].cells[col_idx], h, bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="366092")

for row_idx, (tinh_name, row_data) in enumerate(tbl9_df.iterrows()):
    cell_row = t9.rows[row_idx+2]
    is_gt = (tinh_name == 'Tổng')
    bg = "F2F2F2" if is_gt else None
    format_cell(cell_row.cells[0], tinh_name, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_p(doc, "", space_after=8)

# Table 10: Ca làm việc tại Kho TC Nha Trang
add_styled_heading(doc, "Bố trí ca làm việc Kho TC Nha Trang Event 8.8 (NVCT & Freelance)", level=3)

t10 = doc.add_table(rows=14, cols=12)
t10.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t10)

headers_t10 = ['Ca Làm Việc'] + dates_header_10 + ['TB/ngày']
for col_idx, h in enumerate(headers_t10):
    format_cell(t10.rows[0].cells[col_idx], h, bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="366092")

# Shifts schedule template tailored for NTB Kho TC Nha Trang for Event 8.8
t10_data = [
    ["Ca Làm Việc (NVCT)", "06/08", "07/08", "08/08", "09/08", "10/08", "11/08", "12/08", "13/08", "14/08", "15/08", "TB"],
    ["07:00 - 18:00", "44", "48", "56", "46", "52", "48", "44", "42", "42", "42", "46"],
    ["18:00 - 03:30", "18", "20", "24", "18", "22", "20", "18", "16", "16", "16", "19"],
    ["20:00 - 05:30", "14", "16", "20", "14", "18", "16", "14", "12", "12", "12", "15"],
    ["22:00 - 06:00", "14", "16", "20", "14", "18", "16", "14", "12", "12", "12", "15"],
    ["Tổng NVCT", "90", "100", "120", "92", "110", "100", "90", "82", "82", "82", "95"],
    ["Ca Làm Việc (Freelance)", "06/08", "07/08", "08/08", "09/08", "10/08", "11/08", "12/08", "13/08", "14/08", "15/08", "TB"],
    ["07:00 - 11:30", "8", "12", "16", "10", "14", "12", "8", "6", "6", "6", "10"],
    ["13:00 - 17:30", "8", "10", "14", "8", "12", "10", "8", "6", "6", "6", "9"],
    ["18:00 - 23:30", "14", "18", "24", "14", "20", "16", "12", "10", "10", "10", "15"],
    ["20:00 - 05:30", "30", "36", "44", "30", "38", "34", "28", "22", "22", "22", "31"],
    ["22:00 - 06:00", "24", "28", "36", "24", "30", "26", "20", "16", "16", "16", "24"],
    ["Tổng Freelance", "84", "104", "134", "86", "114", "98", "76", "60", "60", "60", "88"]
]

for r_i, row in enumerate(t10_data):
    for c_i, val in enumerate(row):
        is_hdr = (r_i == 0 or row[0].startswith("Ca Làm Việc"))
        is_tot = row[0].startswith("Tổng")
        bg = "366092" if is_hdr and r_i==0 else ("E9EDF4" if is_tot else ("F2F2F2" if is_hdr else None))
        bold = True if (is_hdr or is_tot) else False
        color = (255,255,255) if is_hdr and r_i==0 else (0,0,0)
        align = WD_ALIGN_PARAGRAPH.LEFT if c_i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        format_cell(t10.rows[r_i+1].cells[c_i], val, bold=bold, font_size=8, align=align, color_rgb=color, bg_hex=bg)

add_p(doc, "", space_after=8)

# Table 11: Ca làm việc tại Hub Lâm Đồng (Đà Lạt)
add_styled_heading(doc, "Bố trí ca làm việc Hub Lâm Đồng Event 8.8 (Đà Lạt / Đức Trọng)", level=3)

t11 = doc.add_table(rows=6, cols=11)
t11.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t11)

t11_data = [
    ["Ca Làm Việc (NVCT)", "06/08", "07/08", "08/08", "09/08", "10/08", "11/08", "12/08", "13/08", "14/08", "15/08"],
    ["03:00 - 11:30 & 17:45 - 21:30", "5", "6", "8", "6", "7", "6", "5", "4", "4", "4"],
    ["03:00 - 11:30", "4", "5", "7", "5", "6", "5", "4", "3", "3", "3"],
    ["03:00 - 06:30 & 17:45 - 21:30", "5", "6", "8", "6", "7", "6", "5", "4", "4", "4"],
    ["03:00 - 06:30", "2", "2", "3", "2", "2", "2", "1", "1", "1", "1"],
    ["Tổng", "16", "19", "26", "19", "22", "19", "15", "12", "12", "12"]
]

for r_i, row in enumerate(t11_data):
    for c_i, val in enumerate(row):
        is_hdr = (r_i == 0)
        is_tot = (r_i == 5)
        bg = "366092" if is_hdr else ("E9EDF4" if is_tot else None)
        bold = True if (is_hdr or is_tot) else False
        color = (255,255,255) if is_hdr else (0,0,0)
        align = WD_ALIGN_PARAGRAPH.LEFT if c_i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        format_cell(t11.rows[r_i].cells[c_i], val, bold=bold, font_size=8.5, align=align, color_rgb=color, bg_hex=bg)

add_p(doc, "", space_after=8)

# Table 12: Ca làm việc tại Hub Bình Thuận (Phan Thiết)
add_styled_heading(doc, "Bố trí ca làm việc Hub Bình Thuận Event 8.8 (Phan Thiết)", level=3)

t12 = doc.add_table(rows=6, cols=11)
t12.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t12)

t12_data = [
    ["Ca Làm Việc (NVCT)", "06/08", "07/08", "08/08", "09/08", "10/08", "11/08", "12/08", "13/08", "14/08", "15/08"],
    ["04:00 - 07:45", "3", "4", "6", "4", "5", "4", "3", "3", "3", "3"],
    ["04:00 - 07:45 & 10:30 - 13:00", "3", "3", "5", "3", "4", "3", "2", "2", "2", "2"],
    ["04:00 - 07:45 & 17:45 - 20:00", "4", "5", "7", "5", "6", "5", "4", "4", "4", "4"],
    ["Ca Full 04:00 - 20:00", "6", "8", "10", "8", "9", "8", "6", "6", "6", "6"],
    ["Tổng", "16", "20", "28", "20", "24", "20", "15", "15", "15", "15"]
]

for r_i, row in enumerate(t12_data):
    for c_i, val in enumerate(row):
        is_hdr = (r_i == 0)
        is_tot = (r_i == 5)
        bg = "366092" if is_hdr else ("E9EDF4" if is_tot else None)
        bold = True if (is_hdr or is_tot) else False
        color = (255,255,255) if is_hdr else (0,0,0)
        align = WD_ALIGN_PARAGRAPH.LEFT if c_i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        format_cell(t12.rows[r_i].cells[c_i], val, bold=bold, font_size=8.5, align=align, color_rgb=color, bg_hex=bg)

add_p(doc, "", space_after=8)

# 3. Tác động bên ngoài
add_styled_heading(doc, "3. Tác động bên ngoài", level=2)

add_styled_heading(doc, "3.1. Trường hợp mất điện:", level=3)
add_p(doc, "• Kho TC Nha Trang và Hub Lâm Đồng chuẩn bị sẵn máy phát điện dự phòng 100% công suất hoạt động liên tục 24/7.")
add_p(doc, "• Tất cả Bưu cục trang bị đèn tích điện và sạc dự phòng cho máy quét 3G/4G của Shipper.")

add_styled_heading(doc, "3.2. Trường hợp lỗi hệ thống:", level=3)
add_p(doc, "• Kích hoạt quy trình xuất/nhập kho offline qua ứng dụng file Excel dự phòng.")
add_p(doc, "• AM & Trưởng ca chụp snapshot mã vận đơn gán ca để duy trì luồng giao nhận không đứt gãy.")

add_styled_heading(doc, "3.3 Mưa bão / giông kéo dài & Sạt lở đèo dốc:", level=3)
add_p(doc, "• Đặc thù địa bàn Lâm Đồng (đèo Bảo Lộc, đèo Prenn) & Đắk Nông có nguy cơ sạt lở mùa mưa bão giữa tháng 8: Chuẩn bị 02 tuyến đường vòng dự phòng kết nối Kho TC Nha Trang - Lâm Đồng.")
add_p(doc, "• Bưu cục Đảo Phú Quý (Bình Thuận): Bố trí gom hàng sớm trước khi biển động, ưu tiên xuất hàng bãi bến tàu đúng khung ca bến cảng.")
add_p(doc, "• Toàn bộ Shipper chặng chót được trang bị túi chống nước chuyên dụng 100% bảo vệ kiện hàng.")

# Save documents
doc.save(output_docx_path)
doc.save(workspace_docx_path)
print(f"Successfully generated Event 8.8 doc: {output_docx_path}")
print(f"Successfully generated Event 8.8 doc: {workspace_docx_path}")
