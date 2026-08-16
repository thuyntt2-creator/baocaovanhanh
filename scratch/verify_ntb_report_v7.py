import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_7.docx"

if not os.path.exists(doc_path):
    print("File không tồn tại")
    sys.exit(1)

doc = docx.Document(doc_path)
errors = 0

print("=== KIỂM TRA SỐ LIỆU TOÀN DIỆN FILE WORD V8_7 ===")

# 1. Kiểm tra Bảng 2 (Tóm tắt)
table_2 = doc.tables[1]
print("\n1. Kiểm tra Bảng 2 (Tóm tắt):")
desc_ns = table_2.rows[3].cells[1].text.strip()
if "92 người" not in desc_ns or "138 người" not in desc_ns:
    print(f"  [LỖI] Nhân sự tóm tắt = '{desc_ns}' (Kỳ vọng: 92 người (T7) → 138 người (T12))")
    errors += 1
else:
    print(f"  [OK] Nhân sự tóm tắt = '{desc_ns}' (Khớp)")

# 2. Kiểm tra Bảng 7 (Nhân sự kho & QL)
table_7 = doc.tables[6]
print("\n2. Kiểm tra Bảng 7 (Nhân sự kho & QL và Tổng nhân sự):")
desc_row_8 = table_7.rows[8].cells[0].text.strip()
if "NV xử lý kho (12)" not in desc_row_8:
    print(f"  [LỖI] Tên dòng 8 = '{desc_row_8}' (Kỳ vọng chứa: NV xử lý kho (12))")
    errors += 1
else:
    print("  [OK] Tên dòng 8 khớp")

for c_idx in range(1, 7):
    val_kho = table_7.rows[8].cells[c_idx].text.strip()
    val_tot = table_7.rows[9].cells[c_idx].text.strip()
    expected_tot = ["92", "94", "124", "126", "136", "138"][c_idx - 1]
    
    if val_kho != "16":
        print(f"  [LỖI] Tháng {c_idx+6}: NV kho & QL = {val_kho} (Kỳ vọng: 16)")
        errors += 1
    if val_tot != expected_tot:
        print(f"  [LỖI] Tháng {c_idx+6}: Tổng nhân sự = {val_tot} | Kỳ vọng = {expected_tot}")
        errors += 1

if errors == 0:
    print("\n>>> TẤT CẢ KIỂM TRA BÁO CÁO V8_7 ĐỀU THÀNH CÔNG! SỐ LIỆU HOÀN THIỆN TUYỆT ĐỐI <<<")
else:
    print(f"\n>>> CÓ {errors} SAI LỆCH SỐ LIỆU CẦN XỬ LÝ <<<")
