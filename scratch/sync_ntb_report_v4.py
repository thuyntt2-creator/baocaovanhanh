import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_4.docx"
dst_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_5.docx"

if not os.path.exists(doc_path):
    print(f"File không tồn tại: {doc_path}")
    sys.exit(1)

doc = docx.Document(doc_path)

print("=== BẮT ĐẦU CẬP NHẬT ĐỢT 5 ===")

# 1. Cập nhật Bảng 10 (index 9) Row 3 Col 0
table_10 = doc.tables[9]
table_10.rows[3].cells[0].text = "Chi phí NV kho & quản lý (8 người cố định, triệu đ)"
print("-> Đã sửa mô tả NV kho & QL ở Bảng 10 thành 8 người cố định")

# 2. Cập nhật Paragraph 127 (Đơn giá Đức Linh 7.8M)
for p in doc.paragraphs:
    if "đơn giá tạm tính 6.500.000 đ/tháng" in p.text:
        p.text = p.text.replace("đơn giá tạm tính 6.500.000 đ/tháng", "đơn giá tạm tính 7.800.000 đ/tháng")
        print("-> Đã sửa mô tả đơn giá Đức Linh thành 7.8M ở Paragraph 127")
        break

# 3. Cập nhật Paragraph 167 (Xóa tiết kiệm 1.26 tỷ quỹ lương ở kết luận)
for p in doc.paragraphs:
    if "tiết kiệm 1,26 tỷ đồng quỹ lương H2/2026" in p.text:
        p.text = p.text.replace("tiết kiệm 1,26 tỷ đồng quỹ lương H2/2026", "điều chỉnh định biên nhân sự tối ưu theo sản lượng thực tế")
        print("-> Đã sửa kết luận (xóa thông tin tiết kiệm lương) ở Paragraph 167")
        break

doc.save(dst_path)
print(f"=== ĐÃ LƯU FILE THÀNH CÔNG: {dst_path} ===")
