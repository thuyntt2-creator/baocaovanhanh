# -*- coding: utf-8 -*-
import sys, os, docx, openpyxl
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import pandas as pd
import numpy as np

sys.stdout.reconfigure(encoding='utf-8')

ntb_file = r'C:\Users\lap4all\Downloads\config_psbba_NTB.xlsx'
fresh_charts_dir = r'c:\Users\lap4all\Documents\Auto report\scratch\ntb_fresh_charts'

# 1. Load FC Sorting 60d
df_sort = pd.read_excel(ntb_file, sheet_name='FC Sorting 60d')

target_cols_sort = list(range(24, 34))
dates_header_10 = ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08']

hubs_list = [
    'Kho Trung Chuyển Khánh Hòa',
    'Kho Chuyển Tiếp Bình Thuận',
    'Kho Chuyển Tiếp Đức Trọng-Lâm Đồng',
    'Kho Chuyển Tiếp Bảo Lộc-Lâm Đồng',
    'Kho Chuyển Tiếp Đắk Nông'
]

hub_data_dict = {}
for h in hubs_list:
    hub_rows = df_sort[df_sort.iloc[:, 0] == h]
    hub_data_dict[h] = {}
    for idx, r in hub_rows.iterrows():
        cat = str(r.iloc[1]).strip()
        vals = [int(round(r.iloc[c])) if pd.notna(r.iloc[c]) and isinstance(r.iloc[c], (int, float)) else 0 for c in target_cols_sort]
        hub_data_dict[h][cat] = np.array(vals)

total_ktc_daily = np.zeros(10)
for h in hubs_list:
    total_ktc_daily += hub_data_dict[h]['Total']

# BUILD KTC SORTING WORD DOCUMENT WITH FRESH OCEAN TEAL THEME
doc_ktc = docx.Document()
for section in doc_ktc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
        tblPr[0].append(borders)

def format_cell(cell, text, bold=False, italic=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex=None):
    cell.text = str(text)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        r = p.runs[0]
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(*color_rgb)
    if bg_hex:
        set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(15, 76, 129)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0, 128, 128)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(38, 38, 38)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

def add_p(doc, text="", bold=False, italic=False, font_size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
    return p

def add_centered_picture(doc, img_path, width_in_inches=6.2, caption_text=""):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_in_inches))
        if caption_text:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r_cap = p_cap.add_run(caption_text)
            r_cap.font.name = 'Arial'
            r_cap.font.size = Pt(9)
            r_cap.font.italic = True
            r_cap.font.color.rgb = RGBColor(89, 89, 89)

# Title
p_title = add_p(doc_ktc, "KẾ HOẠCH SORTING KHO TRUNG CHUYỂN (KTC) - EVENT 8.8", bold=True, font_size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
p_title.runs[0].font.color.rgb = RGBColor(15, 76, 129)
add_p(doc_ktc, "MẠNG LƯỚI 5 KHO TRUNG CHUYỂN & CHUYỂN TIẾP VÙNG NAM TRUNG BỘ (NTB)", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_p(doc_ktc, "Thời gian thực hiện: Giai đoạn Event 08/08 (Forecast 60d: 06/08 - 15/08/2026)", italic=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

# Section I
add_styled_heading(doc_ktc, "I. TỔNG QUAN DỰ BÁO SẢN LƯỢNG SORTING KTC (EVENT 8.8)", level=1)
add_p(doc_ktc, "• Tổng sản lượng Sorting toàn hệ thống KTC Nam Trung Bộ trong 10 ngày Event 8.8 đạt 752.961 đơn, trung bình 75.296 đơn/ngày.")
add_p(doc_ktc, "• Ngày Peak 08/08 ghi nhận đỉnh điểm 97.649 đơn/ngày (+29.69% so với trung bình), trong đó hàng Normal đạt 76,716 đơn, Bulky đạt 12,525 đơn và Freight đạt 8,408 đơn.")

add_centered_picture(doc_ktc, os.path.join(fresh_charts_dir, 'fresh_combo_ktc_sorting.png'), width_in_inches=6.4, caption_text="Biểu đồ 1: Tổng quan Volume KTC Sorting Event 08.08 (Style Mới Nam Trung Bộ)")

# Section II
add_styled_heading(doc_ktc, "II. PHÂN BỔ SẢN LƯỢNG CHI TIẾT 5 KHO SORTING NTB", level=1)
add_centered_picture(doc_ktc, os.path.join(fresh_charts_dir, 'fresh_ktc_hubs.png'), width_in_inches=6.2, caption_text="Biểu đồ 2: So sánh sản lượng Sorting 5 Kho KTC / Chuyển tiếp Vùng NTB")

# Table KTC Sorting Full 10 days
t_ktc = doc_ktc.add_table(rows=22, cols=13)
t_ktc.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_ktc)

headers_tktc = ['Kho Sorting / Hub', 'Nhóm hàng'] + dates_header_10 + ['Tổng 10 ngày']
for col_idx, h in enumerate(headers_tktc):
    format_cell(t_ktc.rows[0].cells[col_idx], h, bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="0F4C81")

r_count = 1
for h in hubs_list:
    h_short = h.replace('Kho Trung Chuyển ', '').replace('Kho Chuyển Tiếp ', 'KCT ')
    for cat in ['Normal', 'Bulky', 'Freight', 'Total']:
        vals = hub_data_dict[h][cat]
        is_tot = (cat == 'Total')
        bg = "E9EDF4" if is_tot else None
        
        format_cell(t_ktc.rows[r_count].cells[0], h_short, bold=is_tot, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
        format_cell(t_ktc.rows[r_count].cells[1], cat, bold=is_tot, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
        
        tot_val = 0
        for d_i, v in enumerate(vals):
            tot_val += v
            format_cell(t_ktc.rows[r_count].cells[d_i+2], f"{v:,}", bold=is_tot, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
        
        format_cell(t_ktc.rows[r_count].cells[12], f"{tot_val:,}", bold=is_tot, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
        r_count += 1

# Total Row
format_cell(t_ktc.rows[21].cells[0], "TỔNG NTB SORTING", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="F2F2F2")
format_cell(t_ktc.rows[21].cells[1], "TỔNG KTC", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="F2F2F2")
tot_ktc_all = 0
for d_i, v in enumerate(total_ktc_daily):
    v_int = int(round(v))
    tot_ktc_all += v_int
    format_cell(t_ktc.rows[21].cells[d_i+2], f"{v_int:,}", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex="F2F2F2")
format_cell(t_ktc.rows[21].cells[12], f"{tot_ktc_all:,}", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex="F2F2F2")

add_p(doc_ktc, "", space_after=8)

# Section III
add_styled_heading(doc_ktc, "III. PHƯƠNG ÁN BỐ TRÍ NGUỒN LỰC VẬN HÀNH KTC", level=1)
add_styled_heading(doc_ktc, "1. Định biên ca kíp Sorting tại Kho TC Nha Trang (Hub chính)", level=2)
add_p(doc_ktc, "• Ca 1 (07:00 - 18:00): 56 NVCT + 16 Freelance gánh ca chia chọn hàng Normal & phân loại Bulky chặng chót.")
add_p(doc_ktc, "• Ca Đêm Peak (20:00 - 05:30 & 22:00 - 06:00): 34 NVCT + 44 Freelance tập trung đóng chuyển linehaul kết nối Lâm Đồng & Bình Thuận.")

add_styled_heading(doc_ktc, "2. Phương án dự phòng rủi ro & Giải phóng dock xuất:", level=2)
add_p(doc_ktc, "• Điều tiết dòng hàng: Hàng cồng kềnh (Bulky & Freight) tại Khánh Hòa được phân luồng xử lý riêng ở dock số 3 và số 4 để tránh nghẽn luồng hàng Normal.")
add_p(doc_ktc, "• Dự phòng đèo dốc & thời tiết: Tuyến kết nối Nha Trang - Đức Trọng (đèo Ngoạn Mục/Sông Bình) được bố trí xe tải 8T dự phòng trung chuyển khi mưa lớn.")

# Save Word Document
word_ktc_paths = [
    r'C:\Users\lap4all\Downloads\Kế hoạch KTC Sorting Event 8.8 - NTB_Fresh.docx',
    r'C:\Users\lap4all\Downloads\Kế hoạch KTC Sorting Event 8.8 - NTB_v2.docx',
    r'c:\Users\lap4all\Documents\Auto report\Kế hoạch KTC Sorting Event 8.8 - NTB_Fresh.docx'
]
for p in word_ktc_paths:
    try:
        doc_ktc.save(p)
        print(f'Successfully saved Word KTC plan: {p}')
    except Exception as e:
        print(f'Could not save {p}: {e}')

print('KTC Sorting Plan updated with fresh modern charts!')
