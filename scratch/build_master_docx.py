import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

doc = docx.Document()

# Margins
for s in doc.sections:
    s.top_margin = Inches(0.7)
    s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.7)
    s.right_margin = Inches(0.7)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_styled_table(doc, data, header_bg='1F4E79', col_widths=None):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    
    for r_idx, row in enumerate(data):
        for c_idx, val in enumerate(row):
            cell = t.cell(r_idx, c_idx)
            cell.text = str(val)
            if r_idx == 0:
                set_cell_background(cell, header_bg)
            elif r_idx == len(data) - 1:
                set_cell_background(cell, 'D9E1F2')
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(7.0)
                    if r_idx == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    elif 'Tổng' in str(val) or 'GRAND' in str(val) or r_idx == len(data) - 1:
                        run.font.bold = True
    return t

# Title
tp = doc.add_paragraph()
tr = tp.add_run('KẾ HOẠCH VẬN HÀNH CHI TIẾT EVENT 8.8\nKHO TRUNG CHUYỂN (KTC) - VÙNG NAM TRUNG BỘ (NTB)')
tr.font.size = Pt(16)
tr.font.bold = True
tr.font.color.rgb = RGBColor(31, 78, 121)
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER

ip = doc.add_paragraph('Hồ sơ Vận hành Tổng thể Chuẩn hóa 100% theo Mô hình TNB Event 7.7.docx, tích hợp đầy đủ 7 Bảng Số liệu Chi tiết, Biểu đồ Trực quan và Phân tích Vận hành Chuyên sâu từ config_psbba_NTB.xlsx.')
ip.runs[0].font.italic = True

# 1. THÔNG TIN CHUNG
doc.add_heading('1. THÔNG TIN CHUNG & MỤC TIÊU CHIẾN LƯỢC', level=1)
p = doc.add_paragraph()
p.add_run('• Thời gian đợt Event: ').bold = True
p.add_run('06/08/2026 đến 15/08/2026 (10 ngày cao điểm).\n')
p.add_run('• Cụm 5 Kho KTC/Hub: ').bold = True
p.add_run('Khánh Hòa (Super-Hub duyên hải - 38.1%), Bình Thuận (Cửa ngõ phía Nam - 22.6%), Đức Trọng (Trung chuyển Tây Nguyên - 19.0%), Bảo Lộc (12.4%), Đắc Nông (7.9%).\n')
p.add_run('• Mục tiêu KPI: ').bold = True
p.add_run('Tổng 771,977 đơn sorting (TB 77,198 đơn/ngày). Peak 08/08 (97,649 đơn) & 10/08 (89,608 đơn). COT Fulfillment 100%, Backlog >24h = 0%.')

# 2. DỰ BÁO SẢN LƯỢNG & NHẬN XẾT
doc.add_heading('2. CHỈ TIÊU & DỰ BÁO SẢN LƯỢNG CHI TIẾT', level=1)

chart1_img = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\497c8661-37a7-4c4e-986e-e8ef6c144fb3\chart1_daily_sorting_volume.png'
if os.path.exists(chart1_img):
    doc.add_paragraph('Biểu đồ 1: Diễn biến Sản lượng Sorting theo Kho KTC Event 8.8')
    cp1 = doc.add_paragraph()
    cp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp1.add_run().add_picture(chart1_img, width=Inches(6.2))

# Table 1: Sorting
doc.add_heading('Bảng 2.1: Sản Lượng Sorting Theo Kho & Nhóm Hàng (06/08 – 15/08/2026)', level=2)
t1_data = [
    ['Kho KTC / Hub', 'Nhóm hàng', '06/08', '07/08', '08/08 (Peak 1)', '09/08', '10/08 (Peak 2)', '11/08', '12/08', '13/08', '14/08', '15/08', 'Tổng 10d', 'TB/ngày', 'Tỷ trọng'],
    ['KTC Khánh Hòa', 'Normal', '21.652', '22.586', '30.057', '18.795', '26.792', '22.516', '22.667', '22.425', '23.306', '23.281', '234.077', '23.408', '79.7%'],
    ['', 'Bulky', '2.755', '2.852', '3.940', '3.975', '3.524', '4.971', '4.882', '3.431', '2.738', '3.940', '37.008', '3.701', '12.6%'],
    ['', 'Freight', '1.659', '1.718', '2.901', '2.438', '2.725', '2.819', '2.274', '1.779', '1.601', '2.732', '22.646', '2.265', '7.7%'],
    ['', 'Cộng KH', '26.067', '27.156', '36.898', '25.208', '33.041', '30.306', '29.823', '27.635', '27.645', '29.953', '293.732', '29.373', '38.1%'],
    ['CT Bình Thuận', 'Normal', '12.437', '12.973', '17.265', '10.796', '15.389', '12.933', '13.021', '12.881', '13.387', '13.372', '134.454', '13.445', '76.9%'],
    ['', 'Bulky', '1.785', '1.848', '2.347', '2.080', '2.515', '2.599', '3.508', '2.804', '1.940', '2.347', '23.773', '2.377', '13.6%'],
    ['', 'Freight', '1.218', '1.260', '2.334', '1.665', '2.398', '1.849', '1.429', '1.285', '1.175', '2.005', '16.618', '1.662', '9.5%'],
    ['', 'Cộng BTH', '15.440', '16.082', '21.946', '14.541', '20.302', '17.381', '17.958', '16.970', '16.502', '17.724', '174.846', '17.485', '22.6%'],
    ['CT Đức Trọng', 'Normal', '10.447', '10.897', '14.502', '9.069', '12.927', '10.864', '10.937', '10.820', '11.245', '11.233', '112.941', '11.294', '76.8%'],
    ['', 'Bulky', '1.436', '1.487', '3.483', '1.664', '2.898', '1.888', '1.543', '1.453', '1.386', '2.365', '19.603', '1.960', '13.3%'],
    ['', 'Freight', '1.059', '1.096', '2.032', '1.439', '1.918', '1.659', '1.324', '1.161', '1.022', '1.744', '14.454', '1.445', '9.8%'],
    ['', 'Cộng DTR', '12.943', '13.481', '20.017', '12.172', '17.743', '14.411', '13.804', '13.434', '13.653', '15.341', '146.999', '14.700', '19.0%'],
    ['CT Bảo Lộc', 'Normal', '6.858', '7.154', '8.774', '6.250', '8.805', '7.255', '7.180', '7.103', '7.382', '7.374', '74.135', '7.414', '77.5%'],
    ['', 'Bulky', '965', '999', '1.368', '1.285', '1.612', '1.531', '1.489', '1.331', '1.008', '1.405', '12.993', '1.299', '13.6%'],
    ['', 'Freight', '625', '646', '1.112', '888', '1.124', '1.007', '773', '718', '603', '1.028', '8.524', '852', '8.9%'],
    ['', 'Cộng BLO', '8.448', '8.799', '11.254', '8.423', '11.541', '9.793', '9.442', '9.152', '8.993', '9.807', '95.652', '9.565', '12.4%'],
    ['CT Đắc Nông', 'Normal', '4.407', '4.597', '6.118', '3.826', '5.453', '4.583', '4.614', '4.564', '4.744', '4.738', '47.644', '4.764', '78.4%'],
    ['', 'Bulky', '642', '665', '887', '902', '954', '843', '1.128', '979', '705', '887', '8.592', '859', '14.1%'],
    ['', 'Freight', '332', '344', '529', '488', '574', '553', '454', '390', '320', '529', '4.513', '451', '7.4%'],
    ['', 'Cộng DNO', '5.381', '5.605', '7.534', '5.216', '6.981', '5.979', '6.196', '5.933', '5.769', '6.154', '60.748', '6.075', '7.9%'],
    ['TỔNG KTC', 'TỔNG NTB', '68.279', '71.123', '97.649', '65.560', '89.608', '77.870', '77.223', '73.124', '72.562', '78.979', '771.977', '77.198', '100.0%']
]
add_styled_table(doc, t1_data)

# Chart 2 & Insights
chart2_img = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\497c8661-37a7-4c4e-986e-e8ef6c144fb3\chart2_product_group_breakdown.png'
if os.path.exists(chart2_img):
    doc.add_paragraph('\nBiểu đồ 2: Cơ cấu Nhóm hàng tại KTC NTB')
    cp2 = doc.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp2.add_run().add_picture(chart2_img, width=Inches(6.2))

doc.add_heading('Nhận Xét Chi Tiết Từng Kho (Chuẩn hóa từ Mô hình TNB):', level=2)
doc.add_paragraph('• KTC Khánh Hòa: Normal 79.7%, Bulky 12.6% (37k đơn), Freight 7.7% (22.6k đơn). Tổng Bulky+Freight ngày Peak vọt 6,841 đơn/ngày. Sàn cồng kềnh & xe 15-20T quy hoạch theo NGÀY ĐỈNH (~36.9K đơn/ngày).\n'
                  '• CT Bình Thuận: Freight 9.5% CAO NHẤT VÙNG (16.6k đơn). Rủi ro dock nhập hàng liên vùng từ HCM. Vận hành luồng xả bãi nhanh ca đêm.\n'
                  '• CT Đức Trọng: Bulky+Freight ngày 08/08 tăng vọt 5,515 đơn (+130%). Bố trí bãi hạ pallet riêng & xe nâng 2 tầng.\n'
                  '• CT Bảo Lộc & Đắc Nông: Quy mô nhỏ nhưng ngày peak vẫn tăng 1.4 - 1.6 lần. Dùng cơ chế ca kíp linh hoạt & đi chung xe tuyến.')

# Chart 3 & Tables 2.2 - 2.6
chart3_img = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\497c8661-37a7-4c4e-986e-e8ef6c144fb3\chart3_pick_delivery_channel.png'
if os.path.exists(chart3_img):
    doc.add_paragraph('\nBiểu đồ 3: Biến động Sản lượng Lấy & Giao theo Kênh Sàn')
    cp3 = doc.add_paragraph()
    cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp3.add_run().add_picture(chart3_img, width=Inches(6.2))

# Table 2.2: Pick Channel
doc.add_heading('Bảng 2.2: Volume LẤY Theo Sàn / Kênh (06/08 – 15/08)', level=2)
t22_data = [
    ['Sàn / Kênh', '06/08', '07/08', '08/08 (Peak)', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08', 'Tổng 10d', 'TB/ngày'],
    ['Shopee', '2.263', '1.830', '4.155', '1.969', '3.282', '2.484', '2.625', '2.212', '2.227', '2.357', '25.402', '2.540'],
    ['Shopee-Bulky', '184', '197', '679', '533', '529', '283', '179', '195', '170', '378', '3.327', '333'],
    ['(Shopee-Bulky 10-15kg)', '(94)', '(103)', '(479)', '(289)', '(272)', '(176)', '(97)', '(99)', '(91)', '(223)', '(1.923)', '(192)'],
    ['SME (Truyền thống)', '10.139', '10.495', '10.911', '8.353', '11.321', '10.417', '10.166', '10.139', '10.495', '10.215', '102.653', '10.265'],
    ['SME-Bulky', '410', '412', '364', '276', '486', '469', '433', '410', '412', '398', '4.074', '407'],
    ['TikTok Shop (TTS)', '4.104', '4.337', '4.546', '3.110', '4.045', '3.403', '3.366', '3.876', '4.133', '4.066', '38.983', '3.898'],
    ['TTS-Bulky', '150', '166', '136', '102', '166', '130', '126', '132', '155', '135', '1.398', '140'],
    ['GRAND TOTAL LẤY', '17.344', '17.541', '21.270', '14.631', '20.101', '17.362', '16.992', '17.063', '17.683', '17.772', '177.759', '17.776']
]
add_styled_table(doc, t22_data)

# Table 2.3: Pick Province
doc.add_heading('Bảng 2.3: Volume LẤY Theo Tỉnh (06/08 – 15/08)', level=2)
t23_data = [
    ['Tỉnh / Thành', '06/08', '07/08', '08/08 (Peak)', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08', 'Tổng 10d', 'TB/ngày'],
    ['Bình Thuận', '3.327', '3.388', '3.361', '2.810', '3.776', '3.183', '3.089', '3.220', '3.373', '3.431', '32.955', '3.295'],
    ['Khánh Hòa', '4.267', '4.286', '5.450', '3.576', '4.995', '4.292', '4.245', '4.210', '4.352', '4.369', '44.043', '4.404'],
    ['Lâm Đồng', '4.705', '4.720', '6.645', '4.027', '5.521', '4.709', '4.604', '4.608', '4.780', '4.851', '49.170', '4.917'],
    ['Ninh Thuận', '1.358', '1.357', '1.729', '1.177', '1.654', '1.410', '1.377', '1.348', '1.380', '1.401', '14.190', '1.419'],
    ['Đắc Nông', '3.687', '3.791', '4.086', '3.042', '4.155', '3.769', '3.677', '3.677', '3.799', '3.720', '37.401', '3.740'],
    ['GRAND TOTAL LẤY', '17.344', '17.541', '21.270', '14.631', '20.101', '17.362', '16.992', '17.063', '17.683', '17.772', '177.759', '17.776']
]
add_styled_table(doc, t23_data)

# Table 2.4: Delivery Channel
doc.add_heading('Bảng 2.4: Volume GIAO Theo Sàn / Kênh (06/08 – 15/08)', level=2)
t24_data = [
    ['Sàn / Kênh', '06/08', '07/08', '08/08 (Peak)', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08', 'Tổng 10d', 'TB/ngày'],
    ['Shopee', '14.964', '12.090', '33.036', '12.642', '22.526', '16.427', '17.352', '14.667', '14.722', '19.582', '178.007', '17.801'],
    ['Shopee-Bulky', '2.379', '2.470', '7.021', '3.998', '6.012', '3.426', '2.388', '2.325', '2.251', '5.147', '37.418', '3.742'],
    ['(Shopee-Bulky 10-15kg)', '(1.304)', '(1.355)', '(5.254)', '(2.163)', '(3.136)', '(1.877)', '(1.305)', '(1.276)', '(1.230)', '(2.903)', '(21.804)', '(2.180)'],
    ['SME (Truyền thống)', '22.391', '23.184', '23.334', '18.439', '24.985', '23.010', '22.460', '22.391', '23.184', '22.553', '225.931', '22.593'],
    ['SME-Bulky', '1.446', '1.447', '1.341', '910', '1.628', '1.567', '1.516', '1.446', '1.447', '1.354', '14.104', '1.410'],
    ['TikTok Shop (TTS)', '12.085', '12.738', '13.492', '9.200', '11.904', '9.998', '9.888', '11.317', '12.157', '11.980', '114.758', '11.476'],
    ['TTS-Bulky', '318', '343', '400', '225', '361', '281', '278', '291', '321', '311', '3.129', '313'],
    ['GRAND TOTAL GIAO', '54.888', '53.627', '83.879', '47.577', '70.552', '56.586', '55.187', '53.715', '55.312', '63.830', '595.152', '59.515']
]
add_styled_table(doc, t24_data)

# Table 2.5: Delivery Province
doc.add_heading('Bảng 2.5: Volume GIAO Theo Tỉnh (06/08 – 15/08)', level=2)
t25_data = [
    ['Tỉnh / Thành', '06/08', '07/08', '08/08 (Peak)', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08', 'Tổng 10d', 'TB/ngày'],
    ['Bình Thuận', '12.846', '12.451', '20.394', '11.203', '16.722', '13.287', '12.979', '12.569', '12.926', '15.090', '140.466', '14.047'],
    ['Khánh Hòa', '14.760', '14.451', '23.538', '12.784', '18.951', '15.220', '14.834', '14.442', '14.894', '17.152', '161.028', '16.103'],
    ['Lâm Đồng', '15.907', '15.529', '23.879', '13.927', '20.685', '16.544', '16.088', '15.580', '16.038', '18.661', '172.836', '17.284'],
    ['Ninh Thuận', '5.222', '5.079', '7.994', '4.478', '6.699', '5.384', '5.272', '5.114', '5.254', '6.073', '56.568', '5.657'],
    ['Đắc Nông', '6.152', '6.118', '8.074', '5.185', '7.494', '6.152', '6.013', '6.010', '6.201', '6.853', '64.253', '6.425'],
    ['GRAND TOTAL GIAO', '54.888', '53.627', '83.879', '47.577', '70.552', '56.586', '55.187', '53.715', '55.312', '63.830', '595.152', '59.515']
]
add_styled_table(doc, t25_data)

# Table 2.6: Combo Chart FC vs Base
doc.add_heading('Bảng 2.6: Bảng Tính Combo Chart FC Volume Lấy & Giao vs Base vs % Tăng/Giảm', level=2)
t26_data = [
    ['Chỉ tiêu', '06/08', '07/08', '08/08 (Peak 1)', '09/08', '10/08 (Peak 2)', '11/08', '12/08', '13/08', '14/08', '15/08', 'TB Toàn đợt'],
    ['FC Volume LẤY', '17.344', '17.541', '21.270', '14.631', '20.101', '17.362', '16.992', '17.063', '17.683', '17.772', '17.776'],
    ['Trung bình Lấy Base', '17.776', '17.776', '17.776', '17.776', '17.776', '17.776', '17.776', '17.776', '17.776', '17.776', '17.776'],
    ['% Tăng/Giảm FC LẤY', '-2.4%', '-1.3%', '+19.7%', '-17.7%', '+13.1%', '-2.3%', '-4.4%', '-4.0%', '-0.5%', '-0.0%', '0.0%'],
    ['FC Volume GIAO', '54.888', '53.627', '83.879', '47.577', '70.552', '56.586', '55.187', '53.715', '55.312', '63.830', '59.515'],
    ['Trung bình Giao Base', '59.515', '59.515', '59.515', '59.515', '59.515', '59.515', '59.515', '59.515', '59.515', '59.515', '59.515'],
    ['% Tăng/Giảm FC GIAO', '-7.8%', '-9.9%', '+40.9%', '-20.1%', '+18.5%', '-4.9%', '-7.3%', '-9.7%', '-7.1%', '+7.2%', '0.0%']
]
add_styled_table(doc, t26_data)

# 3. NHÂN SỰ
doc.add_heading('3. KẾ HOẠCH PHÂN CÔNG NHÂN SỰ & CA LÀM VIỆC TẤT CẢ 5 KHO', level=1)

chart4_img = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\497c8661-37a7-4c4e-986e-e8ef6c144fb3\chart4_staffing_plan.png'
if os.path.exists(chart4_img):
    doc.add_paragraph('Biểu đồ 4: Phân bổ Nhân sự theo Kho KTC')
    cp4 = doc.add_paragraph()
    cp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp4.add_run().add_picture(chart4_img, width=Inches(6.2))

# Table 3.1: Staffing
doc.add_heading('Bảng 3.1: Tổng Hợp Phân Bổ Nhân Sự Chi Tiết Tất Cả 5 Kho KTC (NTB)', level=2)
t31_data = [
    ['Kho KTC / Hub', 'Ca làm việc', 'NVCT/ngày', 'Freelance/ngày', 'Tổng NV ca/ngày', 'Sản lượng TB/ngày', 'Năng suất TB (Đơn/người/ngày)'],
    ['KTC Khánh Hòa', 'Ca Day (7h00 - 18h00)', '12 NV', '5 NV', '17 NV', '29.373', '839 đơn/người'],
    ['', 'Ca Night (20h00 - 5h30)', '13 NV', '5 NV', '18 NV', '', ''],
    ['', 'Cộng Khánh Hòa', '25 NV', '10 NV', '35 NV', '29.373', '839 đơn/người'],
    ['CT Đức Trọng', 'Ca Day (7h00 - 18h00)', '11 NV', '4 NV', '15 NV', '14.700', '565 đơn/người'],
    ['', 'Ca Night (20h00 - 5h30)', '9 NV', '4 NV', '13 NV', '', ''],
    ['', 'Cộng Đức Trọng', '20 NV', '8 NV', '26 NV', '14.700', '565 đơn/người'],
    ['CT Bảo Lộc', 'Ca Day (7h00 - 18h00)', '4 NV', '3 NV', '7 NV', '9.565', '638 đơn/người'],
    ['', 'Ca Night (20h00 - 5h30)', '4 NV', '4 NV', '8 NV', '', ''],
    ['', 'Cộng Bảo Lộc', '8 NV', '7 NV', '15 NV', '9.565', '638 đơn/người'],
    ['CT Đắc Nông', 'Ca Day (7h00 - 18h00)', '3 NV', '2 NV', '5 NV', '6.075', '357 đơn/người'],
    ['', 'Ca Night (20h00 - 5h30)', '4 NV', '4 NV', '8 NV', '', ''],
    ['', 'Cộng Đắc Nông', '7 NV', '6 NV', '13 NV', '6.075', '467 đơn/người'],
    ['CT Bình Thuận', 'Ca Day (7h00 - 18h00)', '6 NV', '0 NV', '6 NV', '17.485', '1.589 đơn/người (🔴 NỔ KHO)'],
    ['', 'Ca Night (18h00 - 7h00)', '5 NV', '0 NV', '5 NV', '', ''],
    ['', 'Cộng Bình Thuận', '11 NV', '0 NV', '11 NV', '17.485', '1.589 đơn/người']
]
add_styled_table(doc, t31_data)

# 4 & 5
doc.add_heading('4. QUY TRÌNH & CHECKLIST VẬN HÀNH CHI TIẾT', level=1)
doc.add_paragraph('• KTC Khánh Hòa: 2 máng nhập, 4 băng nâng hạ, 4 băng phân chọn, 4 xe nâng điện, 14 xe nâng tay. Từ 9:30 dùng 2 băng nâng hạ xuất xe GXT.')
doc.add_paragraph('• CT Bình Thuận: Khẩn cấp bổ sung 3-4 Freelance ca đêm ngày 08/08 & 10/08 giải tỏa áp lực năng suất 1,995 đơn/người.')
doc.add_paragraph('• CT Đức Trọng: Bố trí bãi hạ pallet Bulky riêng, xe nâng 2 tầng hạ tải nhanh.')
doc.add_paragraph('• CNTT: 100% PDA sạc đủ + SIM 4G dung lượng cao dự phòng cúp điện/wifi chập chờn.')

doc.add_heading('5. PHƯƠNG ÁN NGHỆ THUẬT ỨNG PHÓ HÀNG TĂNG 30%', level=1)
doc.add_paragraph('• Huy động 100% quân số NVCT không off ngày Peak. Cho đi làm ngày off (nghỉ bù sau).')
doc.add_paragraph('• Setup mở thêm cổng nhập & băng tải nâng hạ giảm xe tồn chờ xả.')
doc.add_paragraph('• Book xe xuất tăng cường giải phóng layout. Mở rộng khu bạt che ngoài sân kho cho hàng cồng kềnh.')

# Save Word files
w_doc1 = r'c:\Users\lap4all\Documents\Auto report\KTC_Ke_Hoach_Van_Hanh_Event_8_8_NTB_MASTER.docx'
w_doc2 = r'C:\Users\lap4all\Downloads\KTC_Ke_Hoach_Van_Hanh_Event_8_8_NTB_MASTER.docx'

doc.save(w_doc1)
doc.save(w_doc2)

print('Built Master DOCX files successfully!')
print('1.', w_doc1)
print('2.', w_doc2)
