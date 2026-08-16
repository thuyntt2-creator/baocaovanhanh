import docx
import os
import re
import sys

sys.stdout.reconfigure(encoding='utf-8')

docx_path = r"C:\Users\lap4all\Downloads\AOP_BCCK_Plan_new.docx"
output_path = r"C:\Users\lap4all\Downloads\AOP_BCCK_Plan_new_updated.docx"

def update_docx():
    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        return

    doc = docx.Document(docx_path)
    print("Loaded Word plan document.")

    # 1. Update text paragraphs
    text_replacements = [
        ("Di Linh–Lâm Hà–Ninh Hòa", "Di Linh–Lâm Hà–Hòa Ninh"),
        ("Di Linh, Lâm Hà, Ninh Hòa", "Di Linh, Lâm Hà, Hòa Ninh"),
        ("Di Linh – Lâm Hà – Ninh Hòa", "Di Linh – Lâm Hà – Hòa Ninh"),
        ("gom 7 bưu cục (Di Linh, Lâm Hà, Ninh Hòa", "gom 5 bưu cục (Di Linh, Lâm Hà, Hòa Ninh"),
        ("tăng ~32% (từ ~198K lên ~261K đơn/tháng)", "tăng ~73% (từ ~188K lên ~326K đơn/tháng)"),
        ("xe tải 1.9T cần tăng cường lên 9–12 đầu xe/BCCK", "xe tải 1.9T cần tăng cường lên 4–10 đầu xe/BCCK"),
        ("tiết kiệm 210.000.000 đ/tháng quỹ lương, tương đương 1,26 tỷ đồng trong 6 tháng/2026", "tiết kiệm 210.000.000 đ/tháng quỹ lương, tương đương 1,26 tỷ đồng trong 6 tháng/2026"),
        ("giảm từ 60 người phương án cũ", "giảm từ 60 người phương án cũ"),
        ("Đơn Dương, Đức Trọng, Hiệp Thạnh", "Đơn Dương, Hiệp Thạnh"),
        ("Đơn Dương - Đức Trọng - Hiệp Thạnh", "Đơn Dương - Hiệp Thạnh"),
        ("Đơn Dương – Đức Trọng – Hiệp Thạnh", "Đơn Dương – Hiệp Thạnh"),
        ("diện tích cần 112 m²", "diện tích cần 75 m²"),
        ("Diện tích cần 112 m²", "Diện tích cần 75 m²"),
        ("Đơn Dương, Đức Trọng (cụm Hiệp Thạnh)", "Đơn Dương (cụm Hiệp Thạnh)"),
        ("Đơn Dương, Đức Trọng (cụm Hiệp Thạnh", "Đơn Dương (cụm Hiệp Thạnh")
    ]

    p_count = 0
    for p in doc.paragraphs:
        p_text = p.text
        updated = False
        for old_txt, new_txt in text_replacements:
            if old_txt in p_text:
                p_text = p_text.replace(old_txt, new_txt)
                updated = True
        if updated:
            # Re-write paragraph runs
            # We will just write it to the first run and clear others to keep it simple, or keep run styling if possible.
            # Since these are simple replacements, let's replace inside runs to preserve formatting:
            for run in p.runs:
                for old_txt, new_txt in text_replacements:
                    if old_txt in run.text:
                        run.text = run.text.replace(old_txt, new_txt)
            p_count += 1

    print(f"  - Updated {p_count} text paragraphs.")

    # 2. Update tables
    # Let's search tables by headers
    for idx, table in enumerate(doc.tables):
        if len(table.rows) == 0:
            continue
        hdr = [cell.text.strip() for cell in table.rows[0].cells]
        if not hdr:
            continue
            
        print(f"  Table {idx}: cell(0,0)='{hdr[0]}'")

        # A. Table 2.2: BCCK | Chỉ số | T7 (Khởi động) ...
        if hdr[0] == "BCCK" and "T7 (Khởi động)" in hdr:
            print("    -> Updating Table 2.2 (dự báo GAP)...")
            # Row 1: Nha Trang
            table.cell(1, 2).text = "1.043 đơn\n18 chuyến\n9 xe"  # Wait, let's keep old style or write new?
            # Wait, let's use the new calculated values!
            table.cell(1, 2).text = "660 đơn\n12 chuyến\n6 xe"
            table.cell(1, 3).text = "1.015 đơn\n18 chuyến\n9 xe"
            table.cell(1, 4).text = "1.145 đơn\n20 chuyến\n10 xe"
            # Row 2: Di Linh
            table.cell(2, 2).text = "274 đơn\n5 chuyến\n3 xe"
            table.cell(2, 3).text = "408 đơn\n7 chuyến\n4 xe"
            table.cell(2, 4).text = "467 đơn\n8 chuyến\n4 xe"
            # Row 3: Đơn Dương
            table.cell(3, 2).text = "144 đơn\n3 chuyến\n2 xe"
            table.cell(3, 3).text = "218 đơn\n4 chuyến\n2 xe"
            table.cell(3, 4).text = "247 đơn\n5 chuyến\n3 xe"
            # Row 4: Đức Linh
            table.cell(4, 2).text = "125 đơn\n3 chuyến\n2 xe"
            table.cell(4, 3).text = "189 đơn\n4 chuyến\n2 xe"
            table.cell(4, 4).text = "214 đơn\n4 chuyến\n2 xe"

        # B. Table 5.2: Tỉnh | Band cân nặng | T7 | T8 ...
        elif hdr[0] == "Tỉnh" and hdr[1] == "Band cân nặng":
            print("    -> Updating Table 5.2 (Volume giao tháng)...")
            # Row 1: Bình Thuận Hàng vừa
            r1 = [cell.text for cell in table.rows[1].cells]
            table.cell(1, 2).text = "27.739"
            table.cell(1, 3).text = "33.782"
            table.cell(1, 4).text = "36.556"
            table.cell(1, 5).text = "42.162"
            table.cell(1, 6).text = "44.711"
            table.cell(1, 7).text = "47.729"
            # Row 2: Bình Thuận Hàng nặng
            table.cell(2, 2).text = "16.789"
            table.cell(2, 3).text = "20.781"
            table.cell(2, 4).text = "22.746"
            table.cell(2, 5).text = "26.477"
            table.cell(2, 6).text = "27.967"
            table.cell(2, 7).text = "29.661"
            # Row 3: Khánh Hòa Hàng vừa
            table.cell(3, 2).text = "30.398"
            table.cell(3, 3).text = "37.005"
            table.cell(3, 4).text = "40.004"
            table.cell(3, 5).text = "46.098"
            table.cell(3, 6).text = "48.917"
            table.cell(3, 7).text = "52.262"
            # Row 4: Khánh Hòa Hàng nặng
            table.cell(4, 2).text = "19.253"
            table.cell(4, 3).text = "23.722"
            table.cell(4, 4).text = "25.903"
            table.cell(4, 5).text = "30.098"
            table.cell(4, 6).text = "31.844"
            table.cell(4, 7).text = "33.840"
            # Row 5: Lâm Đồng Hàng vừa
            table.cell(5, 2).text = "37.733"
            table.cell(5, 3).text = "45.792"
            table.cell(5, 4).text = "49.489"
            table.cell(5, 5).text = "57.029"
            table.cell(5, 6).text = "60.576"
            table.cell(5, 7).text = "64.771"
            # Row 6: Lâm Đồng Hàng nặng
            table.cell(6, 2).text = "22.704"
            table.cell(6, 3).text = "27.977"
            table.cell(6, 4).text = "30.570"
            table.cell(6, 5).text = "35.541"
            table.cell(6, 6).text = "37.590"
            table.cell(6, 7).text = "39.929"
            # Row 7: Ninh Thuận Hàng vừa
            table.cell(7, 2).text = "10.378"
            table.cell(7, 3).text = "12.596"
            table.cell(7, 4).text = "13.595"
            table.cell(7, 5).text = "15.647"
            table.cell(7, 6).text = "16.621"
            table.cell(7, 7).text = "17.782"
            # Row 8: Ninh Thuận Hàng nặng
            table.cell(8, 2).text = "5.792"
            table.cell(8, 3).text = "7.130"
            table.cell(8, 4).text = "7.782"
            table.cell(8, 5).text = "9.039"
            table.cell(8, 6).text = "9.567"
            table.cell(8, 7).text = "10.171"
            # Row 9: Đắk Nông Hàng vừa
            table.cell(9, 2).text = "11.275"
            table.cell(9, 3).text = "13.468"
            table.cell(9, 4).text = "14.387"
            table.cell(9, 5).text = "16.425"
            table.cell(9, 6).text = "17.574"
            table.cell(9, 7).text = "18.964"
            # Row 10: Đắk Nông Hàng nặng
            table.cell(10, 2).text = "6.535"
            table.cell(10, 3).text = "7.928"
            table.cell(10, 4).text = "8.592"
            table.cell(10, 5).text = "9.929"
            table.cell(10, 6).text = "10.559"
            table.cell(10, 7).text = "11.291"
            # Row 11: TỔNG GIAO BULKY/THÁNG (merged, so cell(1) to cell(6) correspond to months)
            table.cell(11, 1).text = "188.596"
            table.cell(11, 2).text = "230.179"
            table.cell(11, 3).text = "249.625"
            table.cell(11, 4).text = "288.445"
            table.cell(11, 5).text = "305.927"
            table.cell(11, 6).text = "326.400"

        # C. Table 5.3: Kênh | T7 | T8 ...
        elif hdr[0] == "Kênh" and "T7" in hdr:
            print("    -> Updating Table 5.3 (kênh & nhu cầu)...")
            # Row 1: Tổng hàng vừa (10-20kg)
            table.cell(1, 1).text = "117.523"
            table.cell(1, 2).text = "142.641"
            table.cell(1, 3).text = "154.031"
            table.cell(1, 4).text = "177.361"
            table.cell(1, 5).text = "188.400"
            table.cell(1, 6).text = "201.508"
            # Row 2: Tổng hàng nặng (>=20kg)
            table.cell(2, 1).text = "71.073"
            table.cell(2, 2).text = "87.538"
            table.cell(2, 3).text = "95.594"
            table.cell(2, 4).text = "111.084"
            table.cell(2, 5).text = "117.527"
            table.cell(2, 6).text = "124.892"
            # Row 3: B2B
            table.cell(3, 1).text = "61.655"
            table.cell(3, 2).text = "67.282"
            table.cell(3, 3).text = "70.187"
            table.cell(3, 4).text = "78.558"
            table.cell(3, 5).text = "81.296"
            table.cell(3, 6).text = "88.895"
            # Row 4: SL/tháng — Hàng vừa (Bulky Giao) - actually it's total bulky giao
            table.cell(4, 1).text = "188.596"
            table.cell(4, 2).text = "230.179"
            table.cell(4, 3).text = "249.625"
            table.cell(4, 4).text = "288.445"
            table.cell(4, 5).text = "305.927"
            table.cell(4, 6).text = "326.400"
            # Row 5: SL/ngày — Hàng nặng
            table.cell(5, 1).text = "2.369"
            table.cell(5, 2).text = "2.918"
            table.cell(5, 3).text = "3.186"
            table.cell(5, 4).text = "3.703"
            table.cell(5, 5).text = "3.918"
            table.cell(5, 6).text = "4.163"

        # D. Table 6.2: BCCK | SL/ngày | NV Giao | NV Kho | NV QL | Tổng | Ghi chú ...
        elif hdr[0] == "BCCK" and hdr[1] == "SL/ngày":
            print("    -> Updating Table 6.2 (định biên)...")
            # Row 1: BCCK Nha Trang
            table.cell(1, 1).text = "1.146 đơn/ngày"
            # Row 2: BCCK Di Linh
            table.cell(2, 1).text = "466 đơn/ngày"
            table.cell(2, 6).text = "Vận hành hỗn hợp — gom 5 BC Di Linh, Lâm Hà, Hòa Ninh"
            # Row 3: BCCK Đơn Dương
            table.cell(3, 1).text = "247 đơn/ngày"
            table.cell(3, 6).text = "Mở rộng BC — gom Đơn Dương, Hiệp Thạnh"
            # Row 4: BCCK Đức Linh
            table.cell(4, 1).text = "214 đơn/ngày"
            # Row 5: TỔNG CỘNG
            table.cell(5, 1).text = "2.073 đơn/ngày"

        # E. Table 6.3: Chỉ tiêu | T7 | T8 ...
        elif hdr[0] == "Chỉ tiêu" and "T7" in hdr:
            print("    -> Updating Table 6.3 (kế hoạch xe tải 1.9T)...")
            # Row 1: SL hàng nặng/ngày (BQ)
            table.cell(1, 1).text = "2.369"
            table.cell(1, 2).text = "2.918"
            table.cell(1, 3).text = "3.186"
            table.cell(1, 4).text = "3.703"
            table.cell(1, 5).text = "3.918"
            table.cell(1, 6).text = "4.163"
            # Row 2: Tổng đầu xe 1.9T BQ/ngày
            table.cell(2, 1).text = "41"
            table.cell(2, 2).text = "50"
            table.cell(2, 3).text = "55"
            table.cell(2, 4).text = "64"
            table.cell(2, 5).text = "67"
            table.cell(2, 6).text = "71"
            # Row 3: Tổng đầu xe 1.9T ngày cao điểm
            table.cell(3, 1).text = "~58"
            table.cell(3, 2).text = "~70"
            table.cell(3, 3).text = "~77"
            table.cell(3, 4).text = "~90"
            table.cell(3, 5).text = "~94"
            table.cell(3, 6).text = "~100"
            # Row 4: Số người giao
            table.cell(4, 1).text = "~116"
            table.cell(4, 2).text = "~140"
            table.cell(4, 3).text = "~154"
            table.cell(4, 4).text = "~180"
            table.cell(4, 5).text = "~188"
            table.cell(4, 6).text = "~200"
            # Row 5: Mặt bằng tổng 4 BCCK cần
            table.cell(5, 1).text = "~718"
            table.cell(5, 2).text = "~885"
            table.cell(5, 3).text = "~966"
            table.cell(5, 4).text = "~1.123"
            table.cell(5, 5).text = "~1.188"
            table.cell(5, 6).text = "~1.262"

        # F. Table 7.2: Hạng mục | T7 | T8 ...
        elif hdr[0] == "Hạng mục" and "T7" in hdr:
            print("    -> Updating Table 7.2 (chi phí)...")
            # Row 1: Chi phí xe tải 1.9T
            table.cell(1, 1).text = "1.476"
            table.cell(1, 2).text = "1.800"
            table.cell(1, 3).text = "1.980"
            table.cell(1, 4).text = "2.304"
            table.cell(1, 5).text = "2.412"
            table.cell(1, 6).text = "2.556"
            # Row 2: Chi phí NV giao hàng (38 shipper) - stays 570
            table.cell(2, 1).text = "570"
            table.cell(2, 2).text = "570"
            table.cell(2, 3).text = "570"
            table.cell(2, 4).text = "570"
            table.cell(2, 5).text = "570"
            table.cell(2, 6).text = "570"
            # Row 3: Chi phí NV kho & quản lý (8 người) - stays 120 (60 bốc xếp + 60 quản lý)
            table.cell(3, 1).text = "120"
            table.cell(3, 2).text = "120"
            table.cell(3, 3).text = "120"
            table.cell(3, 4).text = "120"
            table.cell(3, 5).text = "120"
            table.cell(3, 6).text = "120"
            # Row 4: Chi phí thuê mặt bằng
            table.cell(4, 1).text = "86"
            table.cell(4, 2).text = "106"
            table.cell(4, 3).text = "116"
            table.cell(4, 4).text = "135"
            table.cell(4, 5).text = "143"
            table.cell(4, 6).text = "151"
            # Row 5: TỔNG CHI PHÍ
            table.cell(5, 1).text = "2.252"
            table.cell(5, 2).text = "2.596"
            table.cell(5, 3).text = "2.786"
            table.cell(5, 4).text = "3.129"
            table.cell(5, 5).text = "3.245"
            table.cell(5, 6).text = "3.397"
            # Row 6: Tiết kiệm - stays 210
            table.cell(6, 1).text = "210"
            table.cell(6, 2).text = "210"
            table.cell(6, 3).text = "210"
            table.cell(6, 4).text = "210"
            table.cell(6, 5).text = "210"
            table.cell(6, 6).text = "210"

    doc.save(output_path)
    print(f"Saved updated docx to: {output_path}")

    # Copy to downloads target
    final_docx_path = r"C:\Users\lap4all\Downloads\AOP_BCCK_Plan_new_calculated.docx"
    doc.save(final_docx_path)
    print(f"Copied to: {final_docx_path}")

if __name__ == "__main__":
    update_docx()

