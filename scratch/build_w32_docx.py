import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

excel_path = r'C:\Users\lap4all\Downloads\BaoCao_Tuan_NTB_W32_2026.xlsx'
docx_path = r'C:\Users\lap4all\Downloads\NTB - Báo Cáo Tuần (HRBP - ARD) - Tuần 32.docx'
out_path = r'C:\Users\lap4all\Downloads\NTB - Báo Cáo Tuần (HRBP - ARD) - Tuần 32 - Final.docx'

wb = openpyxl.load_workbook(excel_path, data_only=True)
doc = docx.Document(docx_path)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def fmt_pct(val, diff=False):
    if val is None or val == "" or val == "-":
        return "-"
    try:
        v = float(val)
        pct = v * 100 if abs(v) <= 1.0 and v != 0 else v
        if diff:
            sign = "+" if pct > 0 else ""
            return f"{sign}{pct:.1f}%"
        return f"{pct:.1f}%"
    except (ValueError, TypeError):
        return str(val)

def fmt_int(val, diff=False):
    if val is None or val == "" or val == "-":
        return "-"
    try:
        v = int(round(float(val)))
        if diff:
            sign = "+" if v > 0 else ""
            return f"{sign}{v:,}"
        return f"{v:,}"
    except (ValueError, TypeError):
        return str(val)

print("Starting Document Update...")

# 1. Update Title and Paragraphs
if len(doc.paragraphs) > 0:
    doc.paragraphs[0].text = "BÁO CÁO TỔNG QUAN VÙNG NTB TUẦN 32"
if len(doc.paragraphs) > 1:
    doc.paragraphs[1].text = "Thời gian báo cáo: Tuần 32 (03/08/2026 – 09/08/2026) | Kỳ so sánh: Tuần 31 (27/07/2026 – 02/08/2026)"

# 2. Table 0 ( Thẻ chỉ số chính )
t0 = doc.tables[0]
ws_tq = wb['00_Tong quan']
# R6: Full, TTS, %GTC
t0.cell(1, 0).text = fmt_int(ws_tq.cell(7, 1).value)
t0.cell(1, 1).text = fmt_int(ws_tq.cell(7, 5).value)
t0.cell(1, 2).text = fmt_pct(ws_tq.cell(7, 9).value)

t0.cell(2, 0).text = str(ws_tq.cell(9, 1).value or "").strip()
t0.cell(2, 1).text = str(ws_tq.cell(9, 5).value or "").strip()
t0.cell(2, 2).text = str(ws_tq.cell(9, 9).value or "").strip()

# R12: ODR, LTC, Rot LC
t0.cell(3, 0).text = fmt_pct(ws_tq.cell(13, 1).value)
t0.cell(3, 1).text = fmt_pct(ws_tq.cell(13, 5).value)
t0.cell(3, 2).text = fmt_pct(ws_tq.cell(13, 9).value)

t0.cell(4, 0).text = str(ws_tq.cell(15, 1).value or "").strip()
t0.cell(4, 1).text = str(ws_tq.cell(15, 5).value or "").strip()
t0.cell(4, 2).text = str(ws_tq.cell(15, 9).value or "").strip()


# 3. Table 1 ( Điểm nổi bật W32 )
t1 = doc.tables[1]
t1.cell(0, 0).text = (
    "ĐIỂM NỔI BẬT TUẦN W32\n"
    "• Sản lượng Full hàng W32 đạt 344,835 đơn (tăng 18,760 so với W31); TTS đạt 92,317 đơn.\n"
    "• AM có %GTC (full hàng) cải thiện mạnh nhất W32 so với tuần trước: Nguyễn Thanh Long (+7.8%).\n"
    "• AM có %GTC (full hàng) giảm mạnh nhất W32 so với tuần trước: Trầm Hữu Tiến (-7.1%) - cần rà soát nguyên nhân.\n"
    "• Tỉnh có %ODR (full hàng) thấp nhất W32: Đắk Nông (91.7%).\n"
    "• AM có %Rớt LC cao nhất W32: Lê Minh Đại (34.7%)."
)

# 4. Table 2 ( Executive Summary 00_Tong quan )
t2 = doc.tables[2]
# Update header
t2.cell(0, 0).text = "Chỉ số"
t2.cell(0, 1).text = "W29"
t2.cell(0, 2).text = "W30"
t2.cell(0, 3).text = "W31"
t2.cell(0, 4).text = "W32"
t2.cell(0, 5).text = "Δ W32/W31"

for i, r in enumerate(range(19, 31)):
    metric = ws_tq.cell(r, 1).value
    w29 = ws_tq.cell(r, 2).value
    w30 = ws_tq.cell(r, 3).value
    w31 = ws_tq.cell(r, 4).value
    w32 = ws_tq.cell(r, 5).value
    diff = ws_tq.cell(r, 6).value

    row_idx = i + 1
    t2.cell(row_idx, 0).text = str(metric) if metric else ""
    is_pct = "%" in str(metric) or "GTC" in str(metric) or "ODR" in str(metric) or "LTC" in str(metric) or "Rớt" in str(metric) or "OPR" in str(metric)
    if not is_pct and "Sản lượng" not in str(metric):
        is_pct = True
        
    if is_pct:
        t2.cell(row_idx, 1).text = fmt_pct(w29)
        t2.cell(row_idx, 2).text = fmt_pct(w30)
        t2.cell(row_idx, 3).text = fmt_pct(w31)
        t2.cell(row_idx, 4).text = fmt_pct(w32)
        t2.cell(row_idx, 5).text = fmt_pct(diff, diff=True)
    else:
        t2.cell(row_idx, 1).text = fmt_int(w29)
        t2.cell(row_idx, 2).text = fmt_int(w30)
        t2.cell(row_idx, 3).text = fmt_int(w31)
        t2.cell(row_idx, 4).text = fmt_int(w32)
        t2.cell(row_idx, 5).text = fmt_int(diff, diff=True)

# Helper function to update standard 4-week comparison summary tables (Table 3, 6, 10, 13, 16, 21)
def update_4week_summary_table(table, ws, start_r=6, is_pct_table=False):
    table.cell(0, 0).text = "Chỉ tiêu"
    table.cell(0, 1).text = "W29"
    table.cell(0, 2).text = "W30"
    table.cell(0, 3).text = "W31"
    table.cell(0, 4).text = "W32"
    table.cell(0, 5).text = "Δ W32/W31"
    
    # Row 1: Full hàng
    table.cell(1, 0).text = "Full hàng"
    if is_pct_table:
        table.cell(1, 1).text = fmt_pct(ws.cell(start_r, 2).value)
        table.cell(1, 2).text = fmt_pct(ws.cell(start_r, 3).value)
        table.cell(1, 3).text = fmt_pct(ws.cell(start_r, 4).value)
        table.cell(1, 4).text = fmt_pct(ws.cell(start_r, 5).value)
        table.cell(1, 5).text = fmt_pct(ws.cell(start_r, 6).value, diff=True)
    else:
        table.cell(1, 1).text = fmt_int(ws.cell(start_r, 2).value)
        table.cell(1, 2).text = fmt_int(ws.cell(start_r, 3).value)
        table.cell(1, 3).text = fmt_int(ws.cell(start_r, 4).value)
        table.cell(1, 4).text = fmt_int(ws.cell(start_r, 5).value)
        table.cell(1, 5).text = fmt_int(ws.cell(start_r, 6).value, diff=True)
        
    # Row 2: TTS
    if len(table.rows) > 2:
        table.cell(2, 0).text = "TTS"
        if is_pct_table:
            table.cell(2, 1).text = fmt_pct(ws.cell(start_r+1, 2).value)
            table.cell(2, 2).text = fmt_pct(ws.cell(start_r+1, 3).value)
            table.cell(2, 3).text = fmt_pct(ws.cell(start_r+1, 4).value)
            table.cell(2, 4).text = fmt_pct(ws.cell(start_r+1, 5).value)
            table.cell(2, 5).text = fmt_pct(ws.cell(start_r+1, 6).value, diff=True)
        else:
            table.cell(2, 1).text = fmt_int(ws.cell(start_r+1, 2).value)
            table.cell(2, 2).text = fmt_int(ws.cell(start_r+1, 3).value)
            table.cell(2, 3).text = fmt_int(ws.cell(start_r+1, 4).value)
            table.cell(2, 4).text = fmt_int(ws.cell(start_r+1, 5).value)
            table.cell(2, 5).text = fmt_int(ws.cell(start_r+1, 6).value, diff=True)

# Update Table 3 (01_San luong)
update_4week_summary_table(doc.tables[3], wb['01_San luong'], start_r=6, is_pct_table=False)

# Update Table 6 (02_GTC tong)
update_4week_summary_table(doc.tables[6], wb['02_GTC tong'], start_r=6, is_pct_table=True)

# Update Table 7 (02_GTC tong Top/Bottom AM cải thiện/giảm)
t7 = doc.tables[7]
ws_gtc = wb['02_GTC tong']
t7.cell(0, 0).text = "Chỉ tiêu"
t7.cell(0, 1).text = "W29"
t7.cell(0, 2).text = "W30"
t7.cell(0, 3).text = "W31"
t7.cell(0, 4).text = "W32"
t7.cell(0, 5).text = "Δ W32/W31"

t7.cell(1, 0).text = "Top 1 AM cải thiện (%GTC Full)"
t7.cell(1, 1).text = fmt_pct(ws_gtc.cell(9, 2).value)
t7.cell(1, 2).text = fmt_pct(ws_gtc.cell(9, 3).value)
t7.cell(1, 3).text = fmt_pct(ws_gtc.cell(9, 4).value)
t7.cell(1, 4).text = fmt_pct(ws_gtc.cell(9, 5).value)
t7.cell(1, 5).text = fmt_pct(ws_gtc.cell(9, 6).value, diff=True)

t7.cell(2, 0).text = "Bottom 1 AM giảm (%GTC Full)"
t7.cell(2, 1).text = fmt_pct(ws_gtc.cell(10, 2).value)
t7.cell(2, 2).text = fmt_pct(ws_gtc.cell(10, 3).value)
t7.cell(2, 3).text = fmt_pct(ws_gtc.cell(10, 4).value)
t7.cell(2, 4).text = fmt_pct(ws_gtc.cell(10, 5).value)
t7.cell(2, 5).text = fmt_pct(ws_gtc.cell(10, 6).value, diff=True)

# Update Table 10 (03_GTC Ca1+Ton)
update_4week_summary_table(doc.tables[10], wb['03_GTC Ca1+Ton'], start_r=6, is_pct_table=True)

# Update Table 13 (04_GTC Ca2)
update_4week_summary_table(doc.tables[13], wb['04_GTC Ca2'], start_r=6, is_pct_table=True)

# Update Table 16 (05_ODR)
update_4week_summary_table(doc.tables[16], wb['05_ODR'], start_r=6, is_pct_table=True)

# Update Table 21 (06_LTC)
update_4week_summary_table(doc.tables[21], wb['06_LTC'], start_r=6, is_pct_table=True)

# Update Table 24 (09_Rot LC Vùng)
t24 = doc.tables[24]
ws_rot = wb['09_Rot LC']
t24.cell(1, 0).text = "Chỉ tiêu"
t24.cell(1, 1).text = "W31"
t24.cell(1, 2).text = "W32"
t24.cell(1, 3).text = "Δ W32/W31"

t24.cell(2, 0).text = "% Rớt LC"
t24.cell(2, 1).text = fmt_pct(ws_rot.cell(6, 2).value)
t24.cell(2, 2).text = fmt_pct(ws_rot.cell(6, 3).value)
t24.cell(2, 3).text = fmt_pct(ws_rot.cell(6, 4).value, diff=True)

# Update Table 25 (09_Rot LC AM)
t25 = doc.tables[25]
t25.cell(0, 0).text = "AM"
t25.cell(0, 1).text = "Vol cần LC"
t25.cell(0, 2).text = "W31"
t25.cell(0, 3).text = "W32"
t25.cell(0, 4).text = "Δ W32/W31"

for r in range(11, 29):
    row_idx = r - 10
    if row_idx < len(t25.rows):
        t25.cell(row_idx, 0).text = str(ws_rot.cell(r, 1).value or "")
        t25.cell(row_idx, 1).text = fmt_int(ws_rot.cell(r, 2).value)
        t25.cell(row_idx, 2).text = fmt_pct(ws_rot.cell(r, 3).value)
        t25.cell(row_idx, 3).text = fmt_pct(ws_rot.cell(r, 4).value)
        t25.cell(row_idx, 4).text = fmt_pct(ws_rot.cell(r, 5).value, diff=True)

# Update Table 26 (09_Rot LC Top 20 Bưu cục)
t26 = doc.tables[26]
t26.cell(0, 0).text = "TOP 20 BƯU CỤC CÓ TỶ LỆ RỚT LC CAO NHẤT (W32)"
t26.cell(1, 0).text = "STT"
t26.cell(1, 1).text = "Bưu cục"
t26.cell(1, 2).text = "Vol cần LC"
t26.cell(1, 3).text = "Vol rớt LC"
t26.cell(1, 4).text = "% rớt LC"
t26.cell(1, 5).text = "AM"

for r in range(42, 62):
    row_idx = r - 40
    if row_idx < len(t26.rows):
        stt = ws_rot.cell(r, 1).value
        bc = ws_rot.cell(r, 2).value
        vol_can = ws_rot.cell(r, 3).value
        vol_rot = ws_rot.cell(r, 4).value
        pct_rot = ws_rot.cell(r, 5).value
        am = ws_rot.cell(r, 6).value if ws_rot.max_column >= 6 else ""

        t26.cell(row_idx, 0).text = str(stt) if stt is not None else ""
        t26.cell(row_idx, 1).text = str(bc) if bc is not None else ""
        t26.cell(row_idx, 2).text = fmt_int(vol_can)
        t26.cell(row_idx, 3).text = fmt_int(vol_rot)
        t26.cell(row_idx, 4).text = fmt_pct(pct_rot)
        if am:
            t26.cell(row_idx, 5).text = str(am)

# Update Table 19 (08_OPR TTS Tỉnh)
t19 = doc.tables[19]
ws_opr = wb['08_OPR TTS']
t19.cell(0, 2).text = "%OPR 9h-19h W31"
t19.cell(0, 3).text = "%OPR 9h-19h W32"
t19.cell(0, 4).text = "Δ 9h-19h"

t19.cell(0, 6).text = "%OPR 19h-9h W31"
t19.cell(0, 7).text = "%OPR 19h-9h W32"
t19.cell(0, 8).text = "Δ 19h-9h"

t19.cell(0, 10).text = "%OPR cả ngày W31"
t19.cell(0, 11).text = "%OPR cả ngày W32"
t19.cell(0, 12).text = "Δ cả ngày"

for r in range(28, 33):
    row_idx = r - 27
    if row_idx < len(t19.rows):
        t19.cell(row_idx, 0).text = str(ws_opr.cell(r, 1).value or "")
        t19.cell(row_idx, 1).text = fmt_int(ws_opr.cell(r, 2).value)
        t19.cell(row_idx, 2).text = fmt_pct(ws_opr.cell(r, 3).value)
        t19.cell(row_idx, 3).text = fmt_pct(ws_opr.cell(r, 4).value)
        t19.cell(row_idx, 4).text = fmt_pct(ws_opr.cell(r, 5).value, diff=True)

        t19.cell(row_idx, 5).text = fmt_int(ws_opr.cell(r, 6).value)
        t19.cell(row_idx, 6).text = fmt_pct(ws_opr.cell(r, 7).value)
        t19.cell(row_idx, 7).text = fmt_pct(ws_opr.cell(r, 8).value)
        t19.cell(row_idx, 8).text = fmt_pct(ws_opr.cell(r, 9).value, diff=True)

        t19.cell(row_idx, 9).text = fmt_int(ws_opr.cell(r, 10).value)
        t19.cell(row_idx, 10).text = fmt_pct(ws_opr.cell(r, 11).value)
        t19.cell(row_idx, 11).text = fmt_pct(ws_opr.cell(r, 12).value)
        t19.cell(row_idx, 12).text = fmt_pct(ws_opr.cell(r, 13).value, diff=True)

# Update Table 20 (08_OPR TTS AM)
t20 = doc.tables[20]
t20.cell(0, 2).text = "%OPR 9h-19h W31"
t20.cell(0, 3).text = "%OPR 9h-19h W32"
t20.cell(0, 4).text = "Δ 9h-19h"

t20.cell(0, 6).text = "%OPR 19h-9h W31"
t20.cell(0, 7).text = "%OPR 19h-9h W32"
t20.cell(0, 8).text = "Δ 19h-9h"

t20.cell(0, 10).text = "%OPR cả ngày W31"
t20.cell(0, 11).text = "%OPR cả ngày W32"
t20.cell(0, 12).text = "Δ cả ngày"

for r in range(9, 27):
    row_idx = r - 8
    if row_idx < len(t20.rows):
        t20.cell(row_idx, 0).text = str(ws_opr.cell(r, 1).value or "")
        t20.cell(row_idx, 1).text = fmt_int(ws_opr.cell(r, 2).value)
        t20.cell(row_idx, 2).text = fmt_pct(ws_opr.cell(r, 3).value)
        t20.cell(row_idx, 3).text = fmt_pct(ws_opr.cell(r, 4).value)
        t20.cell(row_idx, 4).text = fmt_pct(ws_opr.cell(r, 5).value, diff=True)

        t20.cell(row_idx, 5).text = fmt_int(ws_opr.cell(r, 6).value)
        t20.cell(row_idx, 6).text = fmt_pct(ws_opr.cell(r, 7).value)
        t20.cell(row_idx, 7).text = fmt_pct(ws_opr.cell(r, 8).value)
        t20.cell(row_idx, 8).text = fmt_pct(ws_opr.cell(r, 9).value, diff=True)

        t20.cell(row_idx, 9).text = fmt_int(ws_opr.cell(r, 10).value)
        t20.cell(row_idx, 10).text = fmt_pct(ws_opr.cell(r, 11).value)
        t20.cell(row_idx, 11).text = fmt_pct(ws_opr.cell(r, 12).value)
        t20.cell(row_idx, 12).text = fmt_pct(ws_opr.cell(r, 13).value, diff=True)

# Update Table 40 (10_KinhDoanh_TongQuan)
t40 = doc.tables[40]
ws_kd = wb['10_KinhDoanh_TongQuan']
t40.cell(0, 0).text = "TỔNG QUAN DOANH THU & VOLUME – VÙNG NTB"
t40.cell(1, 2).text = "Tuần 31 (27/7–2/8/2026)"
t40.cell(1, 5).text = "Tuần 32 (3/8–9/8/2026)"
t40.cell(1, 8).text = "So sánh (T32 vs T31)"

for r in range(4, 31):
    row_idx = r - 2
    if row_idx < len(t40.rows):
        stt = ws_kd.cell(r, 1).value
        am = ws_kd.cell(r, 2).value
        vol_w31 = ws_kd.cell(r, 3).value
        dt_w31 = ws_kd.cell(r, 4).value
        pct_w31 = ws_kd.cell(r, 5).value

        vol_w32 = ws_kd.cell(r, 6).value
        dt_w32 = ws_kd.cell(r, 7).value
        pct_w32 = ws_kd.cell(r, 8).value

        d_vol = ws_kd.cell(r, 9).value
        d_dt = ws_kd.cell(r, 10).value
        pct_inc = ws_kd.cell(r, 11).value

        t40.cell(row_idx, 0).text = str(stt) if stt is not None else ""
        t40.cell(row_idx, 1).text = str(am) if am is not None else ""
        t40.cell(row_idx, 2).text = fmt_int(vol_w31)
        t40.cell(row_idx, 3).text = fmt_int(dt_w31)
        t40.cell(row_idx, 4).text = fmt_pct(pct_w31)

        t40.cell(row_idx, 5).text = fmt_int(vol_w32)
        t40.cell(row_idx, 6).text = fmt_int(dt_w32)
        t40.cell(row_idx, 7).text = fmt_pct(pct_w32)

        t40.cell(row_idx, 8).text = fmt_int(d_vol, diff=True)
        t40.cell(row_idx, 9).text = fmt_int(d_dt, diff=True)
        t40.cell(row_idx, 10).text = fmt_pct(pct_inc, diff=True)

# Update Table 41 (11_KinhDoanh_F30)
t41 = doc.tables[41]
ws_f30 = wb['11_KinhDoanh_F30']
t41.cell(0, 0).text = "KHÁCH HÀNG MỚI TUẦN – VÙNG NTB"
t41.cell(1, 1).text = "Tuần 31 (27/7–2/8/2026)"
t41.cell(1, 3).text = "Tuần 32 (3/8–9/8/2026)"
t41.cell(1, 5).text = "So sánh (T32 vs T31)"

for r in range(4, 24):
    row_idx = r - 2
    if row_idx < len(t41.rows):
        am = ws_f30.cell(r, 1).value
        kh_w31 = ws_f30.cell(r, 2).value
        dt_w31 = ws_f30.cell(r, 3).value

        kh_w32 = ws_f30.cell(r, 4).value
        dt_w32 = ws_f30.cell(r, 5).value

        d_kh = ws_f30.cell(r, 6).value
        d_dt = ws_f30.cell(r, 7).value
        pct_inc = ws_f30.cell(r, 8).value

        t41.cell(row_idx, 0).text = str(am) if am is not None else ""
        t41.cell(row_idx, 1).text = fmt_int(kh_w31)
        t41.cell(row_idx, 2).text = fmt_int(dt_w31)

        t41.cell(row_idx, 3).text = fmt_int(kh_w32)
        t41.cell(row_idx, 4).text = fmt_int(dt_w32)

        t41.cell(row_idx, 5).text = fmt_int(d_kh, diff=True)
        t41.cell(row_idx, 6).text = fmt_int(d_dt, diff=True)
        t41.cell(row_idx, 7).text = fmt_pct(pct_inc, diff=True)

orig_path = r'C:\Users\lap4all\Downloads\NTB - Báo Cáo Tuần (HRBP - ARD) - Tuần 32.docx'
doc.save(out_path)
print(f"Successfully saved updated document to: {out_path}")

try:
    doc.save(orig_path)
    print(f"Successfully overwritten original document at: {orig_path}")
except Exception as e:
    print(f"Note: Could not overwrite original file directly (likely open in Word): {e}")

