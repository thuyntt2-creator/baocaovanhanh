# -*- coding: utf-8 -*-
import sys, os, docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import pandas as pd
import numpy as np

sys.stdout.reconfigure(encoding='utf-8')

ntb_file = r'C:\Users\lap4all\Downloads\config_psbba_NTB.xlsx'
assets_dir = r'c:\Users\lap4all\Documents\Auto report\scratch\ghn_assets'
charts_dir = r'c:\Users\lap4all\Documents\Auto report\scratch\ghn_charts'

hdr_path = os.path.join(assets_dir, 'ghn_header_crisp.png')
ftr_path = os.path.join(assets_dir, 'ghn_footer_crisp.png')

df_lay = pd.read_excel(ntb_file, sheet_name='6_FC_Lay_Daily')
df_giao = pd.read_excel(ntb_file, sheet_name='7_FC_Giao_Daily')
df_ns = pd.read_excel(ntb_file, sheet_name='Nhân sự')
df_bo = pd.read_excel(ntb_file, sheet_name='Bất ổn')

date_cols_lay = [c for c in df_lay.columns if c not in ['Vùng', 'Tỉnh/Quận', 'ID', 'BC', 'Sàn', 'Tổng 60d']]
date_cols_giao = [c for c in df_giao.columns if c not in ['Vùng', 'Tỉnh/Quận', 'ID', 'BC', 'Sàn', 'Tổng 60d']]

days10_lay = [c for c in date_cols_lay if any(d in c for d in ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08'])][:10]
days10_giao = [c for c in date_cols_giao if any(d in c for d in ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08'])][:10]

for c in days10_lay:
    df_lay[c] = pd.to_numeric(df_lay[c], errors='coerce').fillna(0)
for c in days10_giao:
    df_giao[c] = pd.to_numeric(df_giao[c], errors='coerce').fillna(0)

df_lay = df_lay.dropna(subset=['Sàn'])
df_giao = df_giao.dropna(subset=['Sàn'])
dates_header_10 = [c.split()[-1] for c in days10_lay]

# 2. HELPER FUNCTIONS FOR FULL BLEED & FORMATTING
def make_picture_full_bleed_header(run, img_path, width_in_inches=8.5, height_in_inches=1.2, top_offset_in=0.0):
    run.text = ""
    picture = run.add_picture(img_path, width=Inches(width_in_inches), height=Inches(height_in_inches))
    inline = picture._inline
    
    cx = inline.extent.cx
    cy = inline.extent.cy
    docPr = inline.docPr
    graphic = inline.graphic
    top_offset_emu = int(top_offset_in * 914400)
    
    anchor_xml = f'''
    <wp:anchor {nsdecls("wp")} {nsdecls("a")} {nsdecls("pic")} {nsdecls("r")}
               distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240"
               behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page">
        <wp:posOffset>0</wp:posOffset>
      </wp:positionH>
      <wp:positionV relativeFrom="page">
        <wp:posOffset>{top_offset_emu}</wp:posOffset>
      </wp:positionV>
      <wp:extent cx="{cx}" cy="{cy}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
    </wp:anchor>
    '''
    anchor = parse_xml(anchor_xml)
    anchor.append(docPr)
    anchor.append(graphic)
    
    drawing = inline.getparent()
    drawing.replace(inline, anchor)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
        tblPr[0].append(borders)

def format_cell(cell, text, bold=False, italic=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex=None):
    cell.text = str(text)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        r = p.runs[0]
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(*color_rgb)
    if bg_hex:
        set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

def add_p(doc, text="", bold=False, italic=False, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, color_rgb=(0,0,0)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(*color_rgb)
    return p

def add_ghn_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.font.size = Pt(11.5)
        r.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0, 114, 188)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

def add_centered_picture(doc, img_path, width_in_inches=6.3, caption_text=""):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_in_inches))
        if caption_text:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r_cap = p_cap.add_run(caption_text)
            r_cap.font.name = 'Arial'
            r_cap.font.size = Pt(9)
            r_cap.font.italic = True
            r_cap.font.color.rgb = RGBColor(89, 89, 89)

def apply_ghn_full_bleed_headers_footers(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.3)
        section.bottom_margin = Inches(1.1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        header = section.header
        p_hdr = header.paragraphs[0]
        p_hdr.text = ""
        r_hdr = p_hdr.add_run()
        make_picture_full_bleed_header(r_hdr, hdr_path, width_in_inches=8.5, height_in_inches=1.2, top_offset_in=0.0)

        footer = section.footer
        p_ftr = footer.paragraphs[0]
        p_ftr.text = ""
        r_ftr = p_ftr.add_run()
        make_picture_full_bleed_header(r_ftr, ftr_path, width_in_inches=8.5, height_in_inches=1.0, top_offset_in=10.0)

doc = docx.Document()
apply_ghn_full_bleed_headers_footers(doc)

# Title Block
add_p(doc, "GHN — VÙNG NAM TRUNG BỘ (NTB)", bold=True, font_size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color_rgb=(27, 54, 93))
add_p(doc, "KẾ HOẠCH VẬN HÀNH EVENT 8.8", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color_rgb=(0, 114, 188))
add_p(doc, "TÀI LIỆU VẬN HÀNH CHUẨN VÙNG (EXECUTIVE OPERATIONAL PLAN)", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color_rgb=(0, 0, 0))
add_p(doc, "(Khung thời gian Peak: 06/08/2026 – 15/08/2026)", italic=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, color_rgb=(89, 89, 89))

# SECTION I. MỤC TIÊU
add_ghn_heading(doc, "I. MỤC TIÊU KẾ HOẠCH VẬN HÀNH EVENT 8.8 (VÙNG NTB)", level=1)

add_p(doc, "1. Mục tiêu Vận hành Cốt lõi & Chỉ số SLA Vùng NTB:", bold=True, font_size=10.5, color_rgb=(0, 114, 188))
add_p(doc, "• Tỷ lệ Giao thành công (GTC): Quyết tâm giữ vững chỉ số GTC toàn vùng đạt tối thiểu ≥ 68% xuyên suốt 10 ngày Event, bảo vệ năng suất chặng chót tại các địa bàn đồi dốc và duyên hải.")
add_p(doc, "• Tỷ lệ Lấy thành công (LTC): Cam kết duy trì tỷ lệ LTC toàn Vùng NTB đạt ≥ 96%, đảm bảo 100% đơn hàng tạo mới từ Shopee/SME/TTS được tiếp nhận và luân chuyển đúng SLA.")
add_p(doc, "• Tỷ lệ Nhân sự đi làm (NVPTTT & NVXL): Đảm bảo tỷ lệ quân số đi làm thực tế của NVPTTT duy trì ≥ 94% toàn vùng từ 06/08 đến 15/08/2026, sẵn sàng lực lượng tăng cường gánh ca Peak.")
add_p(doc, "• Quản trị Tồn kho & Giải phóng mặt sàn: Xử lý triệt để 100% đơn tồn > 24h trước 21h00 mỗi ngày. Ưu tiên giải phóng tối đa luồng hàng nhỏ nhẹ (TTS/SME) để dành mặt bằng bưu cục chứa hàng cồng kềnh (CK).")
add_p(doc, "• Kiểm soát Tỷ lệ rớt Luân chuyển (LC): Giới hạn tỷ lệ rớt luân chuyển hàng lấy dưới mốc < 2.5%, không để đứt gãy kết nối Linehaul giữa bưu cục với Kho Sorting Nha Trang & các Kho Chuyển tiếp.")

add_p(doc, "", space_after=6)
add_p(doc, "2. Bảng Tổng quan Thông tin Vận hành Vùng NTB:", bold=True, font_size=10.5, color_rgb=(0, 114, 188))

t_meta = doc.add_table(rows=9, cols=2)
t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_meta, color="CCCCCC")

meta_data = [
    ("Vùng vận hành", "NTB — Nam Trung Bộ (5 tỉnh: Khánh Hòa, Bình Thuận, Lâm Đồng, Ninh Thuận, Đắk Nông)"),
    ("Giám đốc Vùng (GĐV)", "Nguyễn Văn A"),
    ("Đơn vị xây dựng", "Phòng Vận Hành & Khối Khai thác Vùng NTB"),
    ("Ngày ban hành", "30/07/2026"),
    ("Phạm vi triển khai", "84 Bưu cục chặng chót & 5 Kho Trung chuyển / Chuyển tiếp Vùng NTB"),
    ("Khung thời gian Peak", "Event 8.8 (06/08/2026 – 15/08/2026)"),
    ("Nguồn số liệu Forecast", "config_psbba_NTB.xlsx (Sheet 6_FC_Lay, 7_FC_Giao, Nhân sự & Bất ổn)"),
    ("Tệp kế hoạch đi kèm", "[NTB] Bảng Phân bổ Sorting KTC & Phương án Nhân sự Bưu cục Event 8.8"),
    ("Mục tiêu chiến lược", "Bảo đảm SLA Giao/Lấy, không để tồn đọng kho bãi, chủ động ứng phó thời tiết mưa bão")
]

for idx, (label, val) in enumerate(meta_data):
    format_cell(t_meta.rows[idx].cells[0], label, bold=True, font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex="D9E1E8")
    format_cell(t_meta.rows[idx].cells[1], val, bold=False, font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex="FFFFFF")

add_p(doc, "", space_after=10)

# SECTION II. PHÂN TÍCH FORECAST VOLUME & LOẠI HÀNG
add_ghn_heading(doc, "II. PHÂN TÍCH SẢN LƯỢNG FORECAST & ĐẶC ĐIỂM LOẠI HÀNG", level=1)

add_ghn_heading(doc, "1. Phân tích Volume Lấy Vùng NTB", level=2)
add_ghn_heading(doc, "Tổng quan nhịp tăng trưởng Volume Lấy", level=3)
add_p(doc, "Dự báo tổng sản lượng Lấy toàn Vùng Nam Trung Bộ trong 10 ngày Event 8.8 đạt 583.479 đơn (bình quân 58.348 đơn/ngày). Sản lượng gia tăng liên tục từ 53.120 đơn (ngày 06/08) và lập đỉnh tại ngày 08/08 với 69.822 đơn (tăng +19.66% so với trung bình đợt). Nhịp tăng trưởng ổn định giúp Vùng chủ động điều phối tải kết nối Linehaul mà không bị quá tải đột ngột.")

add_centered_picture(doc, os.path.join(charts_dir, 'ghn_combo_lay.png'), width_in_inches=6.3, caption_text="Biểu đồ 1: Tổng quan FC Volume Lấy event 08.08 (Style GHN Corporate)")

add_ghn_heading(doc, "Chi tiết đặc điểm nhóm hàng Lấy & Phương án điều hành", level=3)
add_p(doc, "• Kênh TikTok Shop (TTS) & SME chiếm đến 81.2% tổng sản lượng Lấy. Đây là lực lượng hàng cốt lõi cần ưu tiên bố trí nhân sự lấy hàng sớm và cấp đủ CCDC (bao sọt, máy bắn, biên bản).")
add_p(doc, "• Đỉnh Lấy tập trung vào ngày 08/08: Khối Vận hành Vùng yêu cầu các AM bám sát chỉ số %OPR từng giờ, phối hợp chặt chẽ với các Shop lớn để gom đơn đúng khung giờ cao điểm, chuẩn bị sẵn biên bản giao nhận và video xác nhận để tránh gán lấy trễ.")
add_p(doc, "• Bố trí máy phát điện dự phòng tại bưu cục để phòng ngừa sự cố mất điện làm gián đoạn khâu đóng hàng, tránh rớt luân chuyển và giảm tỷ lệ lấp đầy thùng xe.")

# Table 1: Volume Lấy theo Sàn
tbl1_df = df_lay.groupby('Sàn')[days10_lay].sum().round(0).astype(int)
san_order = ['Shopee', 'Shopee-Bulky', 'Shopee-Bulky (10-15kg)', 'SME', 'SME-Bulky', 'TTS', 'TTS-Bulky']
tbl1_df = tbl1_df.reindex([s for s in san_order if s in tbl1_df.index])
tot_lay_all = tbl1_df.sum().sum()
tbl1_df['Tổng 10 ngày'] = tbl1_df.sum(axis=1)
tbl1_df['Tỷ trọng %'] = (tbl1_df['Tổng 10 ngày'] / tot_lay_all * 100).round(2).astype(str) + '%'

tbl1_df.loc['Grand Total'] = tbl1_df.sum(numeric_only=True)
tbl1_df.loc['Grand Total', 'Tỷ trọng %'] = '100.00%'

t1_tbl = doc.add_table(rows=len(tbl1_df)+1, cols=13)
t1_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t1_tbl)

headers_t1 = ['Sàn/Loại hàng'] + dates_header_10 + ['Tổng 10 ngày', 'Tỷ trọng %']
for col_idx, h in enumerate(headers_t1):
    format_cell(t1_tbl.rows[0].cells[col_idx], h, bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(0,0,0), bg_hex="D9E1E8")

for row_idx, (san_name, row_data) in enumerate(tbl1_df.iterrows()):
    cell_row = t1_tbl.rows[row_idx+1]
    is_gt = (san_name == 'Grand Total')
    bg = "E2E7EC" if is_gt else None
    format_cell(cell_row.cells[0], san_name, bold=is_gt, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx in range(10):
        val = int(row_data.iloc[c_idx])
        format_cell(cell_row.cells[c_idx+1], f"{val:,}", bold=is_gt, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
    format_cell(cell_row.cells[11], f"{int(row_data['Tổng 10 ngày']):,}", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
    format_cell(cell_row.cells[12], str(row_data['Tỷ trọng %']), bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_centered_picture(doc, os.path.join(charts_dir, 'ghn_lay_san.png'), width_in_inches=6.2, caption_text="Biểu đồ 2: Sản lượng Lấy phân rã theo Kênh/Sàn (Event 8.8 NTB)")

add_p(doc, "", space_after=6)

# Table 2: Volume Lấy theo Tỉnh & Tỷ trọng
add_ghn_heading(doc, "Phân bổ Volume Lấy theo Tỉnh & Tỷ trọng trung bình (06/08 - 15/08)", level=3)
add_p(doc, "• Khánh Hòa & Lâm Đồng đóng vai trò trọng tâm Lấy của Vùng NTB (chiếm hơn 70% tổng volume Lấy toàn vùng). Ngày đỉnh 08/08, Khánh Hòa đạt 24.150 đơn, Lâm Đồng đạt 18.650 đơn.")
add_p(doc, "• Sự phân bổ sản lượng cân bằng giữa các tỉnh giúp phương tiện trung chuyển Linehaul hoạt động tối ưu công suất, giảm thiểu rủi ro ùn tắc.")

tbl2_df = df_lay.groupby('Tỉnh/Quận')[days10_lay].sum().round(0).astype(int)
tbl2_df['Tổng 10 ngày'] = tbl2_df.sum(axis=1)
tbl2_df['Tỷ trọng %'] = (tbl2_df['Tổng 10 ngày'] / tot_lay_all * 100).round(2).astype(str) + '%'

tbl2_df.loc['Grand Total'] = tbl2_df.sum(numeric_only=True)
tbl2_df.loc['Grand Total', 'Tỷ trọng %'] = '100.00%'

t2_tbl = doc.add_table(rows=len(tbl2_df)+1, cols=13)
t2_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t2_tbl)

headers_t2 = ['Tỉnh/Quận'] + dates_header_10 + ['Tổng 10 ngày', 'Tỷ trọng %']
for col_idx, h in enumerate(headers_t2):
    format_cell(t2_tbl.rows[0].cells[col_idx], h, bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(0,0,0), bg_hex="D9E1E8")

for row_idx, (tinh_name, row_data) in enumerate(tbl2_df.iterrows()):
    cell_row = t2_tbl.rows[row_idx+1]
    is_gt = (tinh_name == 'Grand Total')
    bg = "E2E7EC" if is_gt else None
    format_cell(cell_row.cells[0], tinh_name, bold=is_gt, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx in range(10):
        val = int(row_data.iloc[c_idx])
        format_cell(cell_row.cells[c_idx+1], f"{val:,}", bold=is_gt, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
    format_cell(cell_row.cells[11], f"{int(row_data['Tổng 10 ngày']):,}", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
    format_cell(cell_row.cells[12], str(row_data['Tỷ trọng %']), bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_centered_picture(doc, os.path.join(charts_dir, 'ghn_lay_tinh.png'), width_in_inches=6.2, caption_text="Biểu đồ 3: Sản lượng Lấy theo Tỉnh/Quận & Tỷ trọng trung bình (Event 8.8 NTB)")

add_p(doc, "", space_after=8)

# 2. Volume Giao
add_ghn_heading(doc, "2. Phân tích Volume Giao Vùng NTB", level=2)

add_ghn_heading(doc, "Tổng quan nhịp sóng Volume Giao", level=3)
add_p(doc, "Dự báo tổng sản lượng Giao toàn Vùng NTB đạt 752.961 đơn trong 10 ngày Event (trung bình 75.296 đơn/ngày). Đỉnh Giao xuất hiện vào ngày 08/08 với 97.649 đơn (+29.69% so với trung bình đợt). Áp lực giao chặng chót duy trì ở mức rất cao trong liên tiếp 5 ngày từ 07/08 đến 11/08.")

add_centered_picture(doc, os.path.join(charts_dir, 'ghn_combo_giao.png'), width_in_inches=6.3, caption_text="Biểu đồ 4: Tổng quan FC Volume Giao event 08.08 (Style GHN Corporate)")

add_ghn_heading(doc, "Phân tích nhóm hàng Giao & Giải pháp điều phối", level=3)
add_p(doc, "• Shopee tăng bùng nổ: Ngày đỉnh 08/08, Shopee đạt 48.250 đơn (gần gấp đôi ngày thường). Các ngày tiếp theo, sản lượng SME và TTS duy trì ổn định quanh mốc 20.000 đơn/ngày giữ nhịp sóng giao cao.")
add_p(doc, "• Shopee-Bulky tăng gấp 4 lần: Đạt 12.850 đơn vào ngày đỉnh (hàng 10-15kg chiếm trên 6.200 đơn). Hàng cồng kềnh gây áp lực chiếm diện tích kho rất lớn tại bưu cục.")
add_p(doc, "• AM cần kiểm tra kỹ các mã đơn revert từ GTX, đồng thời chủ động bố trí xe ba gác, xe van và xe tải nhỏ 1.9T để giải phóng hàng CK vào khung giờ 14h-16h hàng ngày, giữ thông thoáng mặt sàn cho ca rã hàng đêm.")

# Table 3: Volume Giao theo Sàn & Tỷ trọng
tbl3_df = df_giao.groupby('Sàn')[days10_giao].sum().round(0).astype(int)
tbl3_df = tbl3_df.reindex([s for s in san_order if s in tbl3_df.index])
tot_giao_all = tbl3_df.sum().sum()
tbl3_df['Tổng 10 ngày'] = tbl3_df.sum(axis=1)
tbl3_df['Tỷ trọng %'] = (tbl3_df['Tổng 10 ngày'] / tot_giao_all * 100).round(2).astype(str) + '%'

tbl3_df.loc['Grand Total'] = tbl3_df.sum(numeric_only=True)
tbl3_df.loc['Grand Total', 'Tỷ trọng %'] = '100.00%'

t3_tbl = doc.add_table(rows=len(tbl3_df)+1, cols=13)
t3_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t3_tbl)

headers_t3 = ['Sàn/Loại hàng'] + dates_header_10 + ['Tổng 10 ngày', 'Tỷ trọng %']
for col_idx, h in enumerate(headers_t3):
    format_cell(t3_tbl.rows[0].cells[col_idx], h, bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(0,0,0), bg_hex="D9E1E8")

for row_idx, (san_name, row_data) in enumerate(tbl3_df.iterrows()):
    cell_row = t3_tbl.rows[row_idx+1]
    is_gt = (san_name == 'Grand Total')
    bg = "E2E7EC" if is_gt else None
    format_cell(cell_row.cells[0], san_name, bold=is_gt, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx in range(10):
        val = int(row_data.iloc[c_idx])
        format_cell(cell_row.cells[c_idx+1], f"{val:,}", bold=is_gt, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
    format_cell(cell_row.cells[11], f"{int(row_data['Tổng 10 ngày']):,}", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
    format_cell(cell_row.cells[12], str(row_data['Tỷ trọng %']), bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_centered_picture(doc, os.path.join(charts_dir, 'ghn_giao_san.png'), width_in_inches=6.2, caption_text="Biểu đồ 5: Chi tiết Sản lượng Giao theo Sàn (Event 8.8 NTB)")

add_p(doc, "", space_after=6)

# Table 4: Volume Giao theo Tỉnh & Tỷ trọng trung bình
add_ghn_heading(doc, "Phân bổ Volume Giao theo Tỉnh & Tỷ trọng trung bình (06/08 - 15/08)", level=3)
add_p(doc, "• Khánh Hòa & Lâm Đồng là hai trọng điểm Giao gánh áp lực lớn nhất toàn vùng (Khánh Hòa peak 32.150 đơn, Lâm Đồng peak 24.820 đơn ngày 08/08). Cần chuẩn bị nguồn lực NVPTTT và xe tải phụ gánh ca tại các bưu cục thuộc Đà Lạt, Đức Trọng, Cam Ranh, Nha Trang.")
add_p(doc, "• Bình Thuận & Ninh Thuận giữ nhịp giao ổn định (Bình Thuận peak 19.450 đơn, Ninh Thuận peak 11.200 đơn).")
add_p(doc, "• Đắk Nông (peak 10.029 đơn) tập trung 3 Bưu cục bất ổn do địa hình đồi núi đường khó đi (Kiến Đức, Quảng Tín, Đông Gia Nghĩa), cần được ưu tiên điều động Đội phản ứng nhanh Vùng NTB gánh ca.")

tbl4_df = df_giao.groupby('Tỉnh/Quận')[days10_giao].sum().round(0).astype(int)
tbl4_df['Tổng 10 ngày'] = tbl4_df.sum(axis=1)
tbl4_df['Tỷ trọng %'] = (tbl4_df['Tổng 10 ngày'] / tot_giao_all * 100).round(2).astype(str) + '%'

tbl4_df.loc['Grand Total'] = tbl4_df.sum(numeric_only=True)
tbl4_df.loc['Grand Total', 'Tỷ trọng %'] = '100.00%'

t4_tbl = doc.add_table(rows=len(tbl4_df)+1, cols=13)
t4_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t4_tbl)

headers_t4 = ['Tỉnh/Quận'] + dates_header_10 + ['Tổng 10 ngày', 'Tỷ trọng %']
for col_idx, h in enumerate(headers_t4):
    format_cell(t4_tbl.rows[0].cells[col_idx], h, bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(0,0,0), bg_hex="D9E1E8")

for row_idx, (tinh_name, row_data) in enumerate(tbl4_df.iterrows()):
    cell_row = t4_tbl.rows[row_idx+1]
    is_gt = (tinh_name == 'Grand Total')
    bg = "E2E7EC" if is_gt else None
    format_cell(cell_row.cells[0], tinh_name, bold=is_gt, font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx in range(10):
        val = int(row_data.iloc[c_idx])
        format_cell(cell_row.cells[c_idx+1], f"{val:,}", bold=is_gt, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
    format_cell(cell_row.cells[11], f"{int(row_data['Tổng 10 ngày']):,}", bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)
    format_cell(cell_row.cells[12], str(row_data['Tỷ trọng %']), bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_centered_picture(doc, os.path.join(charts_dir, 'ghn_giao_tinh.png'), width_in_inches=6.2, caption_text="Biểu đồ 6: Sản lượng Giao theo Tỉnh/Quận & Tỷ trọng trung bình (Event 8.8 NTB)")

add_p(doc, "", space_after=10)

# --- SECTION III. PHÂN TÍCH NHÓM BƯU CỤC ---
add_ghn_heading(doc, "III. PHÂN TÍCH VẬN HÀNH & PHƯƠNG ÁN ỨNG PHÓ THEO NHÓM BƯU CỤC", level=1)

# Table 6 (Tiêu chuẩn phân loại 3 Nhóm BC)
add_ghn_heading(doc, "Tiêu chuẩn Phân loại 3 Nhóm Bưu cục Vùng NTB (Table 6 Hướng dẫn)", level=2)

t6_standard = doc.add_table(rows=6, cols=4)
t6_standard.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t6_standard)

t6_std_headers = ["Tiêu chí", "Nhóm 1 - Ổn định", "Nhóm 2 - Cảnh báo", "Nhóm 3 - Bất ổn"]
for c_i, h in enumerate(t6_std_headers):
    format_cell(t6_standard.rows[0].cells[c_i], h, bold=True, font_size=9, bg_hex="D9E1E8")

t6_std_data = [
    ["Mức độ rủi ro", "Thấp", "Trung bình", "Rất cao - quá tải, cần tập trung theo dõi"],
    ["Thiếu hụt nhân sự (NVPTTT)", "Dưới 2 người", "Từ 1-6 người, đang có điều động hỗ trợ", "Từ 2-6 người, tỷ trọng thiếu trên 25%"],
    ["Khả năng kiểm soát vận hành", "Tốt", "Trung bình - khá, phụ thuộc thời tiết", "Kém, cần điều động đội phản ứng nhanh linh hoạt"],
    ["Phương án A (chủ động)", "AM theo dõi, điều hành gán, FIFO, bám sát năng suất NV", "Đẩy mạnh tuyển dụng, AM trực tiếp điều hành gán/FIFO, chính sách thúc đẩy clear tồn, cân nhắc thuê Freelancer", "AM trực tiếp điều hành gán FIFO, chính sách thúc đẩy năng suất, điều động nhân sự BC khác hỗ trợ, thuê Freelancer"],
    ["Phương án B (khẩn cấp khi PA A không khả thi)", "Tắt tuyến BC (khi về vượt CAP x3), điều tiết giảm hàng KA, hỗ trợ giữ hàng tại KTC", "Áp dụng như Nhóm 3 nếu tình trạng xấu đi", "Đề xuất GĐV cho phép Tắt tuyến với hàng KA nếu tồn vượt quá 2 lần CAP; điều động nhân sự BC Nhóm 1 lân cận theo mô hình 'cuốn chiếu'"]
]

for r_i, row in enumerate(t6_std_data, 1):
    for c_i, val in enumerate(row):
        bg = "F2F2F2" if c_i == 0 else None
        bold = True if c_i == 0 else False
        format_cell(t6_standard.rows[r_i].cells[c_i], val, bold=bold, font_size=8.5, bg_hex=bg)

add_p(doc, "", space_after=8)

# TNB OPERATIONAL INSIGHTS FOR GROUP 2 & GROUP 3
add_ghn_heading(doc, "1. Nhóm 2 (Cảnh báo - Ổn định ngắn hạn)", level=2)
add_p(doc, "Thực trạng rủi ro:", bold=True, font_size=10)
add_p(doc, "• Nhân sự và Năng suất: Bưu cục vẫn có khả năng giao và có nhân sự điều động hỗ trợ, nhưng đang thiếu hụt nhẹ (từ 1-6 NVPTTT). Tuy nhiên, áp lực quá tải lâu ngày dễ phát sinh tình trạng nhân viên xin nghỉ phép hoặc nghỉ việc trong kỳ Event.")
add_p(doc, "• Vận hành và Thời tiết: Năng lực kiểm soát ở mức Trung bình - Khá và phụ thuộc rất lớn vào thời tiết. Yếu tố thời tiết bất lợi (mưa bão) dự kiến làm năng suất GTC giảm từ 1–5%. Khi gặp rủi ro kép (Thời tiết xấu + Sản lượng giao/Số chuyến tăng đột biến), nhóm này rất dễ bị quá tải nghiêm trọng và tụt xuống Nhóm 3 (Bất ổn).")

add_p(doc, "Phương án ứng phó chủ động:", bold=True, font_size=10)
add_p(doc, "• Tối ưu mặt sàn và Kho bãi: Tiến hành sắp xếp lại layout bưu cục, tập trung clear hàng CK (cồng kềnh) để giải phóng diện tích sử dụng cho luồng hàng nhỏ nhẹ.")
add_p(doc, "• Nhân sự xử lý (NVXL): Chủ động lọc riêng hàng TTS để add chuyến nhanh cho NVPTTT đi giao, nhằm tối ưu và nâng cao tỷ lệ GTC ngay trong ngày.")
add_p(doc, "• Điều phối tuyến giao: Chủ động phân tách và chuyển bớt các tuyến giao/lấy/trả sang các bưu cục lân cận có CAP giao tốt hơn để gánh bớt số đơn. Sắp xếp lộ trình tuyến phù hợp với năng lực thực tế của từng nhân sự; ưu tiên phân bổ nhân viên mới chạy các tuyến dễ, tránh đẩy vào các tuyến khó gây bào mỏng nguồn lực và nản chí.")
add_p(doc, "• Điều động hoặc sử dụng chính sách nếu cần: Kích hoạt chính sách điều động NVPTTT từ các bưu cục Nhóm 1 lân cận sang hỗ trợ gánh ca cao điểm.")

add_p(doc, "", space_after=8)

add_ghn_heading(doc, "2. Nhóm 3: Nhóm Bất ổn (Cần can thiệp khẩn cấp)", level=2)
add_p(doc, "Thực trạng rủi ro:", bold=True, font_size=10)
add_p(doc, "• Nhân sự: Báo động cần follow kỹ nhóm BC có tỷ trọng thiếu nhân sự trên 25%. Tỷ lệ nghỉ việc cao ở khu vực Phú Quý, Khánh Vĩnh, Đam Rông, Đắk Glong; nguồn tuyển đang được cải thiện.")
add_p(doc, "• Vận hành hiện tại: Năng suất NVPTTT đang bị quá tải lâu ngày nên áp lực đi làm Event nếu không kiểm soát tốt.")

add_p(doc, "Phương án ứng phó khẩn cấp:", bold=True, font_size=10)
add_p(doc, "• Giai đoạn Event (06/08 - 15/08): Duy trì Điều động Đội phản ứng nhanh Vùng NTB từ Nhóm 1 sang hỗ trợ clear tồn cũ. Cân đối ngân sách hỗ trợ bất ổn.")
add_p(doc, "• Kích hoạt chính sách 'Thưởng nóng 3k/đơn' cho các tuyến có Vol tồn vượt CAP 50%.")
add_p(doc, "• Xây dựng đội cứu các đơn GTB (Giao Không Thành Công): Gọi điện xác nhận nhu cầu nhận hàng của khách trước khi xuất kho.")
add_p(doc, "• Quản trị: AM trực chiến 24/7 tại BC để gán đơn FIFO, ưu tiên giải phóng hàng CK (hàng cồng kềnh) chiếm diện tích kho trước.")
add_p(doc, "• Cắt giảm áp lực: Trong trường hợp Vol về vượt quá 3 lần công suất thực tế, đề xuất GĐV cho phép Tắt tuyến tạm thời đối với các Shop có Vol lớn tại khu vực bất ổn để bảo vệ chỉ số chung của vùng.")

add_p(doc, "Phương án A (Chủ động):", bold=True, font_size=10)
add_p(doc, "• Ưu tiên giải phóng hàng CK chiếm diện tích kho vào khung giờ thấp điểm (14h-16h). Các Bưu cục tại Khánh Vĩnh, Đức Trọng đã chủ động liên hệ xe van/tải trung chuyển để mượn tải luân chuyển đơn CK; khu vực đảo Phú Quý đã triển khai xe tải 1,9 tấn hoặc ba gác để luân chuyển hàng CK cho các tuyến xa trên 20km.")

add_p(doc, "Phương án B (Trường hợp khẩn cấp):", bold=True, font_size=10)
add_p(doc, "• Đề xuất GĐV cho phép Tắt tuyến đối với hàng KA tại khu vực này nếu tồn vượt quá 2 lần CAP.")
add_p(doc, "• Điều động nhân sự từ các Bưu cục Nhóm 1 lân cận hỗ trợ theo mô hình 'cuốn chiếu'.")

add_p(doc, "", space_after=8)

# Table 7 (Detail Table)
t7_detail = doc.add_table(rows=13, cols=8)
t7_detail.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t7_detail)

t7_headers = ["Nhóm", "Tên Bưu cục", "Trạng thái", "Định biên", "Hiện hữu", "Thiếu T30", "Thực trạng Rủi ro & Tồn đọng", "Phương án ứng phó (PA A & B)"]
for c_i, h in enumerate(t7_headers):
    format_cell(t7_detail.rows[0].cells[c_i], h, bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="D9E1E8")

bc_table7_data = [
    ["Nhóm 3", "(KHO) Cam Linh", "Thiếu", "19", "12", "9", "Thiếu 9 NVPTTT (47%), Tồn LM 2,322 đơn.", "PA A: Điều động 5 NV + 4 Freelancer. PA B: Tắt tuyến KA nếu tồn x2 CAP."],
    ["Nhóm 3", "(LDO) Đơn Dương", "Thiếu", "17", "14", "5", "Thiếu 5 NVPTTT (29%), Tồn LM 1,454 đơn.", "PA A: Thưởng 3k/đơn + điều động 3 NV Đà Lạt. PA B: Giữ hàng KA tại KTC."],
    ["Nhóm 3", "(LDO) Lang Biang 1", "Thiếu", "12", "10", "4", "Thiếu 4 NVPTTT (33%), Tồn LM 1,093 đơn.", "PA A: Tăng ca rã hàng ca đêm + bổ sung 2 xe van. PA B: Tắt tuyến KA."],
    ["Nhóm 3", "(LDO) Tân Hà Lâm Hà", "Thiếu", "10", "7", "4", "Thiếu 4 NVPTTT (40%), Tồn LM 1,072 đơn.", "PA A: Phụ cấp ca gánh + điều động cuốn chiếu. PA B: Chuyển tuyến xa."],
    ["Nhóm 3", "(KHO) Bắc Cam Ranh", "Thiếu", "9", "5", "4", "Thiếu 4 NVPTTT (44%), Tồn LM 464 đơn.", "PA A: AM cắm bưu cục điều hành 100%. PA B: Hỗ trợ tuyến lân cận."],
    ["Nhóm 3", "(DNO) Kiến Đức", "Thiếu", "10", "7", "3", "Thiếu 3 NVPTTT, Tồn LM 1,946 đơn, Tồn Aging >5d.", "PA A: Đội phản ứng nhanh Vùng clear tồn. PA B: Tắt tuyến KA."],
    ["Nhóm 3", "(DNO) Quảng Tín", "Thiếu", "8", "6", "2", "Thiếu 2 NVPTTT, Tồn LM 1,322 đơn, Tồn Aging >5d.", "PA A: Thuê xe ba gác luân chuyển hàng CK. PA B: Hỗ trợ chi phí."],
    ["Nhóm 3", "(BTN) La Gi", "Bất ổn", "15", "12", "3", "Tồn LM 2,752 đơn & Tồn Aging >5d.", "PA A: Gán FIFO 100%, clear kho CK. PA B: Đề xuất GDV Tắt tuyến KA."],
    ["Nhóm 2", "(BTN) Phan Thiết 2", "Thiếu", "16", "13", "3", "Thiếu 3 NVPTTT, Tồn LM 1,514 đơn.", "PA A: Đẩy mạnh tuyển dụng + Freelance. PA B: Điều xe tải gánh ca."],
    ["Nhóm 2", "(KHO) Bắc Nha Trang", "Thiếu", "21", "18", "3", "Thiếu 3 NVPTTT, Tồn LM 1,362 đơn.", "PA A: Lọc riêng hàng TTS add chuyến nhanh. PA B: Hỗ trợ từ BC Nam NT."],
    ["Nhóm 2", "(LDO) Lâm Viên 2", "Thiếu", "9", "6", "3", "Thiếu 3 NVPTTT, Tồn LM 913 đơn.", "PA A: Bố trí xoay ca NVXL rã hàng sớm. PA B: Điều động NV."],
    ["Nhóm 2", "(LDO) Đức Trọng 1", "Thiếu", "8", "7", "2", "Thiếu 2 NVPTTT, Tồn LM 2,339 đơn, Aging >5d.", "PA A: Mượn tải xe 1.9T luân chuyển đơn CK. PA B: Tắt tuyến tạm."]
]

for r_i, row in enumerate(bc_table7_data, 1):
    bg = "FFF2CC" if "Nhóm 3" in row[0] else None
    for c_i, val in enumerate(row):
        bold = True if c_i in [0, 1] else False
        format_cell(t7_detail.rows[r_i].cells[c_i], val, bold=bold, font_size=8, bg_hex=bg)

add_p(doc, "", space_after=10)

# --- SECTION IV. CHECKLIST CÔNG VIỆC ---
add_ghn_heading(doc, "IV. CHECKLIST CÔNG VIỆC CHUẨN BỊ EVENT 8.8", level=1)

# Table 8 (Checklist Hạng mục CCDC, Kho bãi, Chi phí)
add_ghn_heading(doc, "1. Công cụ dụng cụ + kho bãi (Table 8 Hướng dẫn)", level=2)

t8_chk = doc.add_table(rows=5, cols=4)
t8_chk.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t8_chk)

t8_headers = ["Hạng mục", "Tình trạng hiện tại", "Phương án dự phòng", "Chi phí dự kiến"]
for c_i, h in enumerate(t8_headers):
    format_cell(t8_chk.rows[0].cells[c_i], h, bold=True, font_size=9, bg_hex="D9E1E8")

t8_data = [
    ["CCDC (Túi trùm sọt, bạt mưa, máy bắn)", "Đã phân bổ 100% BC, 12 BC thiếu nhẹ CCDC đã nhường luân chuyển.", "Đặt dự phòng 15% CCDC tại Vùng NTB để bù hư hỏng.", "15,000,000 VNĐ"],
    ["Kho bãi & Layout", "Layout 84 BC đã tối ưu, 0 BC nào cần thuê kho tạm.", "Mượn bãi tập kết bưu cục lân cận khi Vol x2.5 CAP.", "0 VNĐ"],
    ["Xe luân chuyển hàng CK", "Sản lượng CK chiếm 20-30% volume toàn vùng.", "Thuê xe tải 1.9T & xe ba gác luân chuyển tuyến xa (Khánh Vĩnh, Đam Rông).", "35,000,000 VNĐ"],
    ["Máy phát điện dự phòng", "Theo dõi lịch cúp điện 5 tỉnh NTB.", "Thuê máy phát điện dự phòng cho BC rớt điện.", "1,000,000 VNĐ/ngày/BC"]
]

for r_i, row in enumerate(t8_data, 1):
    for c_i, val in enumerate(row):
        bold = True if c_i == 0 else False
        format_cell(t8_chk.rows[r_i].cells[c_i], val, bold=bold, font_size=8.5)

add_p(doc, "", space_after=6)

add_ghn_heading(doc, "2. Bố trí lịch làm việc & Ca kíp tăng ca", level=2)
add_p(doc, "• Điều động nhân sự cùng lúc bố trí lịch làm xoay ca hợp lý cho nhân viên Nhóm 2 & 3 để đảm bảo sức khỏe duy trì suốt 10 ngày cao điểm Event 8.8.")
add_p(doc, "• Các Bưu cục hàng về vượt CAP giao phân công nhân viên ở lại rã hàng ca đêm (20h-04h), tạo chuyến đi sẵn trước ngày cao điểm để xuất kho sớm từ 06h30.")
add_p(doc, "• Trình Vùng phê duyệt phụ cấp tăng ca cho NVXL và nhân viên rã hàng đêm.")

add_ghn_heading(doc, "3. Tác động bên ngoài & Quy trình ứng phó rủi ro", level=2)

add_ghn_heading(doc, "3.1. Trường hợp mất điện:", level=3)
add_p(doc, "• BC theo dõi lịch thông báo cắt điện của khu vực. Chuẩn bị sẵn thông tin liên hệ nhà cung cấp điện, bộ phận Tech Vùng.")
add_p(doc, "• Những khu vực hay có thông báo ngắt điện phải chuẩn bị thuê máy phát điện trước. Yêu cầu AM, NVXL phải nắm tình hình các BC xuyên suốt để kịp thời xử lý (chi phí thuê máy phát điện dự kiến 1,000,000 VNĐ/ngày).")

add_ghn_heading(doc, "3.2. Trường hợp lỗi hệ thống:", level=3)
add_p(doc, "• Nếu rớt mạng yêu cầu NVXL thông báo lại cho AM nắm tình hình (xử lý loại check in cho NV nếu cần). AM có nhiệm vụ báo lên nhóm Tech để kịp thời xử lý.")
add_p(doc, "• Nếu rớt mạng trong vòng 15p: Tạm thời sử dụng mạng di động 5G cá nhân để duy trì bắn kiểm.")
add_p(doc, "• Nếu rớt mạng quá 30p ngay thời điểm xuất hàng / bắn kiểm mà không thao tác được: Bưu cục cần chuyển sang chế độ xuất kho ngoại tuyến offline trên app theo quy trình dự phòng Tech và báo khẩn lên Vùng.")

add_ghn_heading(doc, "3.3. Trường hợp Mưa giông kéo dài:", level=3)
add_p(doc, "• Ưu tiên giải phóng hàng nhỏ TTS, gọi khách xác nhận nhu cầu trước khi giao ra xe ngay khi mưa nhỏ. Hàng CK vượt CAP xin thêm chi phí tải kết nối nâng cao năng suất nhân viên.")
add_p(doc, "• Trang bị đủ CCDC: túi trùm sọt bạc che mưa, túi chống nước điện thoại, áo mưa bộ cho NVPTTT.")
add_p(doc, "• Tắt tuyến (khẩn cấp) + điều động đội giải phóng hàng tồn + kích hoạt chính sách kích thích năng suất.")
add_p(doc, "• Họp Bưu cục để báo NVPTTT chủ động chuẩn bị thiết bị chống nước điện thoại, trùm sọt, thuốc cảm, các lưu ý phòng tránh tai nạn mùa mưa bão đèo dốc.")

add_p(doc, "", space_after=10)

# --- SECTION V. NGUỒN LỰC HỖ TRỢ VÀ QUY TRÌNH PHỐI HỢP ---
add_ghn_heading(doc, "V. NGUỒN LỰC HỖ TRỢ VÀ QUY TRÌNH PHỐI HỢP", level=1)

# Table 9 (Các Đơn vị hỗ trợ Vùng)
add_ghn_heading(doc, "1. Các Đơn vị Hỗ trợ Vùng (Table 9 Hướng dẫn)", level=2)

t9_supp = doc.add_table(rows=6, cols=4)
t9_supp.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t9_supp)

t9_headers = ["Đơn vị hỗ trợ", "Thông tin/hỗ trợ cung cấp cho Vùng", "Thời điểm cung cấp", "Đầu mối liên hệ tại Vùng"]
for c_i, h in enumerate(t9_headers):
    format_cell(t9_supp.rows[0].cells[c_i], h, bold=True, font_size=8.5, bg_hex="D9E1E8")

t9_data = [
    ["Capacity team", "Forecast sản lượng lấy/giao theo ngày, theo sàn, theo tỉnh/quận cho toàn kỳ Event", "Trước Event tối thiểu 3-5 ngày", "Trưởng nhóm Vận hành Vùng"],
    ["Phòng Nhân sự / C&B", "Chính sách điều động nhân sự liên vùng, chính sách phụ cấp tăng ca, thưởng nóng theo sản lượng tồn", "Trước Event", "Chuyên viên HRBP Vùng"],
    ["Tài chính / Vận hành", "Ngân sách dự phòng cho chi phí phát sinh: thuê xe, máy phát điện, Freelancer, thưởng nóng", "Trước Event, xét duyệt bổ sung trong Event", "BP Tài chính Vùng"],
    ["Tech", "Đầu mối xử lý sự cố hệ thống, lỗi mạng, lỗi ứng dụng trong suốt kỳ Event", "Trực 24/7 trong suốt Event", "Đầu mối Tech Vùng"],
    ["Network", "Hỗ trợ điều xe, tăng chuyến Linehaul khi phát sinh vượt tải, giải phóng hàng cồng kềnh (CK)", "Theo yêu cầu phát sinh", "Điều phối Linehaul Vùng"]
]

for r_i, row in enumerate(t9_data, 1):
    for c_i, val in enumerate(row):
        bold = True if c_i == 0 else False
        format_cell(t9_supp.rows[r_i].cells[c_i], val, bold=bold, font_size=8)

add_p(doc, "", space_after=8)

# Table 10 (Quy trình Escalation)
add_ghn_heading(doc, "2. Quy trình Leo leo Sự cố / Escalation Flow (Table 10 Hướng dẫn)", level=2)

t10_esc = doc.add_table(rows=6, cols=5)
t10_esc.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t10_esc)

t10_headers = ["Loại vấn đề", "Báo cáo tại Bưu cục", "Báo cáo lên Vùng", "Ghi chú", "Người phụ trách / SĐT"]
for c_i, h in enumerate(t10_headers):
    format_cell(t10_esc.rows[0].cells[c_i], h, bold=True, font_size=8.5, bg_hex="D9E1E8")

t10_data = [
    ["Mất điện tại BC", "NVXL báo AM phụ trách BC", "AM tổng hợp báo Trưởng nhóm Vận hành Vùng xin hỗ trợ máy phát điện", "Chuẩn bị sẵn thông tin liên hệ nhà cung cấp điện khu vực", "AM Địa bàn / 090x.xxx.xxx"],
    ["Lỗi hệ thống/rớt mạng", "NVXL báo AM", "AM báo nhóm Tech Vùng/Trung tâm để xử lý; nếu rớt mạng >30 phút báo khẩn", "Dưới 15 phút: tạm dùng mạng di động 5G", "Hotline Tech Vùng / 1900 63 66 77"],
    ["Quá tải sản lượng vượt CAP", "AM báo Trưởng nhóm Vận hành Vùng", "Vùng báo GĐV/cấp phê duyệt để xin Tắt tuyến hoặc điều tiết hàng KA", "Áp dụng Phương án B theo phân loại nhóm BC", "GĐV Nam Trung Bộ / 091x.xxx.xxx"],
    ["Thiếu hụt nhân sự nghiêm trọng", "AM báo Trưởng nhóm Vận hành Vùng", "Vùng báo Nhân sự/C&B để điều động liên vùng hoặc tuyển Freelancer", "Ưu tiên các BC Nhóm 3", "HRBP Vùng / 098x.xxx.xxx"],
    ["Phát sinh chi phí ngoài kế hoạch", "AM tổng hợp đề xuất kèm số liệu", "Vùng trình Tài chính/cấp phê duyệt trước khi triển khai", "Ví dụ: thuê xe tải, máy phát điện, thưởng nóng", "BP Tài chính Vùng / 093x.xxx.xxx"]
]

for r_i, row in enumerate(t10_data, 1):
    for c_i, val in enumerate(row):
        bold = True if c_i == 0 else False
        format_cell(t10_esc.rows[r_i].cells[c_i], val, bold=bold, font_size=8)

add_p(doc, "", space_after=8)

# Table 11 (Kế hoạch chi tiết KTC/KCT & Bưu cục CK)
add_ghn_heading(doc, "3. Kế hoạch Chi tiết KTC, KCT & Bưu cục CK (Table 11 Hướng dẫn)", level=2)

t11_ktc = doc.add_table(rows=6, cols=4)
t11_ktc.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t11_ktc)

t11_headers = ["Nhóm", "Nội dung cần lập kế hoạch", "Yêu cầu cụ thể Vùng nêu rõ", "Kế hoạch Thực tế Vùng NTB"]
for c_i, h in enumerate(t11_headers):
    format_cell(t11_ktc.rows[0].cells[c_i], h, bold=True, font_size=8.5, bg_hex="D9E1E8")

t11_data = [
    ["1. KTC, KCT thuộc Vùng", "Nhân sự tại từng KTC/KCT", "Số lượng & vai trò từng KTC; nguồn bù khi thiếu (điều động liên BC, Freelancer); đầu mối phụ trách.", "Bố trí 56 NVCT + 16 Freelance ca ngày, 34 NVCT + 44 Freelance ca đêm tại Kho TC Nha Trang. 4 KCT còn lại bố trí 12-25 NV/kho."],
    ["1. KTC, KCT thuộc Vùng", "Xe & Bãi tại kho", "Số xe (tự thuê + điều động); diện tích bãi đỗ/bãi tập kết & sức chứa; phương án khi quá tải.", "Phân bổ 18 xe tải 8T-15T kết nối Linehaul. Bãi tập kết Kho Nha Trang 2,500m2 chứa max 65,000 đơn/thời điểm."],
    ["1. KTC, KCT thuộc Vùng", "Lịch làm việc", "Ca làm việc theo từng ngày (kể cả ngày đỉnh 08/08), người trực ca, kế hoạch tăng ca/phụ cấp.", "Ca 1 (07h-18h), Ca Đêm (20h-05h30 & 22h-06h). Tăng ca 100% ngày Peak 08/08, phụ cấp 150k/đêm/người."],
    ["2. Bưu cục CK", "Nhân sự cho Bưu cục CK", "Số lượng nhân sự cần có tại mỗi Bưu cục CK, nguồn tuyển hoặc điều động.", "Bố trí 8 NVXL chuyên trách hàng CK tại BC Nha Trang CK, Đức Trọng CK & Phan Thiết CK."],
    ["2. Bưu cục CK", "Xe cho Bưu cục CK", "Số xe cần bố trí cho mỗi Bưu cục CK, nguồn xe (tự thuê/điều động), lịch trình.", "Điều động 6 xe van 1.9T chạy 3 chuyến/ngày luân chuyển hàng CK ra tuyến chặng chót."]
]

for r_i, row in enumerate(t11_data, 1):
    for c_i, val in enumerate(row):
        bold = True if c_i in [0, 1] else False
        format_cell(t11_ktc.rows[r_i].cells[c_i], val, bold=bold, font_size=8)

# SAVE FRESH NTB DOCX
fresh_docx_paths = [
    r'C:\Users\lap4all\Downloads\NTB Kế hoạch Event 8.8_FRESH_NTB.docx',
    r'C:\Users\lap4all\Downloads\NTB Kế hoạch Event 8.8_v2.docx',
    r'c:\Users\lap4all\Documents\Auto report\NTB Kế hoạch Event 8.8.docx'
]

for p in fresh_docx_paths:
    try:
        doc.save(p)
        print(f"Successfully saved Fresh NTB Word Plan Document: {p}")
    except Exception as e:
        print(f"Could not save {p}: {e}")

print("Fresh NTB Plan Document generated successfully!")
