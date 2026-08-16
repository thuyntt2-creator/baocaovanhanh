# -*- coding: utf-8 -*-
import sys, os, docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import pandas as pd
import numpy as np

sys.stdout.reconfigure(encoding='utf-8')

ntb_file = r'C:\Users\lap4all\Downloads\config_psbba_NTB.xlsx'
assets_dir = r'c:\Users\lap4all\Documents\Auto report\scratch\ghn_assets'
charts_dir = r'c:\Users\lap4all\Documents\Auto report\scratch\ghn_charts'

hdr_path = os.path.join(assets_dir, 'ghn_header_crisp.png')
ftr_path = os.path.join(assets_dir, 'ghn_footer_crisp.png')

# 1. LOAD DATA FROM EXCEL
df_lay = pd.read_excel(ntb_file, sheet_name='6_FC_Lay_Daily')
df_giao = pd.read_excel(ntb_file, sheet_name='7_FC_Giao_Daily')

date_cols_lay = [c for c in df_lay.columns if c not in ['Vùng', 'Tỉnh/Quận', 'ID', 'BC', 'Sàn', 'Tổng 60d']]
date_cols_giao = [c for c in df_giao.columns if c not in ['Vùng', 'Tỉnh/Quận', 'ID', 'BC', 'Sàn', 'Tổng 60d']]

days10_lay = [c for c in date_cols_lay if any(d in c for d in ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08'])][:10]
days10_giao = [c for c in date_cols_giao if any(d in c for d in ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08'])][:10]

for c in days10_lay:
    df_lay[c] = pd.to_numeric(df_lay[c], errors='coerce').fillna(0)
for c in days10_giao:
    df_giao[c] = pd.to_numeric(df_giao[c], errors='coerce').fillna(0)

df_lay = df_lay.dropna(subset=['Sàn'])
df_giao = df_giao.dropna(subset=['Sàn'])
dates_header_10 = [c.split()[-1] for c in days10_lay]

# 2. HELPER FUNCTIONS FOR FULL BLEED & FORMATTING
def make_picture_full_bleed_header(run, img_path, width_in_inches=8.5, height_in_inches=1.2, top_offset_in=0.0):
    run.text = ""
    picture = run.add_picture(img_path, width=Inches(width_in_inches), height=Inches(height_in_inches))
    inline = picture._inline
    
    cx = inline.extent.cx
    cy = inline.extent.cy
    docPr = inline.docPr
    graphic = inline.graphic
    top_offset_emu = int(top_offset_in * 914400)
    
    anchor_xml = f'''
    <wp:anchor {nsdecls("wp")} {nsdecls("a")} {nsdecls("pic")} {nsdecls("r")}
               distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240"
               behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page">
        <wp:posOffset>0</wp:posOffset>
      </wp:positionH>
      <wp:positionV relativeFrom="page">
        <wp:posOffset>{top_offset_emu}</wp:posOffset>
      </wp:positionV>
      <wp:extent cx="{cx}" cy="{cy}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
    </wp:anchor>
    '''
    anchor = parse_xml(anchor_xml)
    anchor.append(docPr)
    anchor.append(graphic)
    
    drawing = inline.getparent()
    drawing.replace(inline, anchor)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
        tblPr[0].append(borders)

def format_cell(cell, text, bold=False, italic=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex=None):
    cell.text = str(text)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        r = p.runs[0]
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(*color_rgb)
    if bg_hex:
        set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

def add_p(doc, text="", bold=False, italic=False, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, color_rgb=(0,0,0)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(*color_rgb)
    return p

def add_ghn_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.font.size = Pt(11.5)
        r.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0, 114, 188)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

def add_centered_picture(doc, img_path, width_in_inches=6.3, caption_text=""):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_in_inches))
        if caption_text:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r_cap = p_cap.add_run(caption_text)
            r_cap.font.name = 'Arial'
            r_cap.font.size = Pt(9)
            r_cap.font.italic = True
            r_cap.font.color.rgb = RGBColor(89, 89, 89)

def apply_ghn_full_bleed_headers_footers(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.3)
        section.bottom_margin = Inches(1.1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        header = section.header
        p_hdr = header.paragraphs[0]
        p_hdr.text = ""
        r_hdr = p_hdr.add_run()
        make_picture_full_bleed_header(r_hdr, hdr_path, width_in_inches=8.5, height_in_inches=1.2, top_offset_in=0.0)

        footer = section.footer
        p_ftr = footer.paragraphs[0]
        p_ftr.text = ""
        r_ftr = p_ftr.add_run()
        make_picture_full_bleed_header(r_ftr, ftr_path, width_in_inches=8.5, height_in_inches=1.0, top_offset_in=10.0)

# 3. BUILD FULL COMPREHENSIVE DOCX
doc = docx.Document()
apply_ghn_full_bleed_headers_footers(doc)

# Title Block
add_p(doc, "GHN — VÙNG NAM TRUNG BỘ (NTB)", bold=True, font_size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color_rgb=(27, 54, 93))
add_p(doc, "AOP EVENT 8.8", bold=True, font_size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color_rgb=(0, 114, 188))
add_p(doc, "KẾ HOẠCH VẬN HÀNH EVENT 8.8", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color_rgb=(0, 0, 0))
add_p(doc, "(Tháng 8/2026)", italic=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, color_rgb=(89, 89, 89))

# Metadata Table Block
t_meta = doc.add_table(rows=9, cols=2)
t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_meta, color="CCCCCC")

meta_data = [
    ("Vùng", "NTB — Nam Trung Bộ"),
    ("Giám đốc Vùng (GĐV)", "Nguyễn Văn A"),
    ("Người lập", "Phòng Vận Hành Vùng NTB"),
    ("Ngày lập", "30/07/2026"),
    ("Phạm vi vận hành", "Toàn bộ 82 Bưu cục & 5 Kho Sorting Vùng Nam Trung Bộ"),
    ("Khung thời gian", "Event 08/08 (06/08 - 15/08/2026)"),
    ("Tỉnh/thành phố", "Khánh Hòa, Bình Thuận, Lâm Đồng, Ninh Thuận, Đắk Nông"),
    ("Nguồn volume", "config_psbba_NTB.xlsx"),
    ("Tài liệu kèm theo", "[NTB] Kế Hoạch Event 8.8 & Sorting KTC 2026")
]

for idx, (label, val) in enumerate(meta_data):
    format_cell(t_meta.rows[idx].cells[0], label, bold=True, font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex="D9E1E8")
    format_cell(t_meta.rows[idx].cells[1], val, bold=False, font_size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex="FFFFFF")

add_p(doc, "", space_after=10)

# --- SECTION I: MỤC TIÊU ---
add_ghn_heading(doc, "I. MỤC TIÊU:", level=1)

# 1. Volume Lấy
add_ghn_heading(doc, "1. Volume Lấy", level=2)
add_ghn_heading(doc, "Tổng quan", level=3)

add_centered_picture(doc, os.path.join(charts_dir, 'ghn_combo_lay.png'), width_in_inches=6.3, caption_text="Biểu đồ Tổng quan FC Volume Lấy event 08.08 (Vùng NTB)")

# Table 1: Volume Lấy theo Sàn
tbl1_df = df_lay.groupby('Sàn')[days10_lay].sum().round(0).astype(int)
san_order = ['Shopee', 'Shopee-Bulky', 'Shopee-Bulky (10-15kg)', 'SME', 'SME-Bulky', 'TTS', 'TTS-Bulky']
tbl1_df = tbl1_df.reindex([s for s in san_order if s in tbl1_df.index])
tbl1_df.loc['Grand Total'] = tbl1_df.sum()

t1 = doc.add_table(rows=len(tbl1_df)+1, cols=11)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t1)

headers_t1 = ['Sàn'] + dates_header_10
for col_idx, h in enumerate(headers_t1):
    format_cell(t1.rows[0].cells[col_idx], h, bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(0,0,0), bg_hex="D9E1E8")

for row_idx, (san_name, row_data) in enumerate(tbl1_df.iterrows()):
    cell_row = t1.rows[row_idx+1]
    is_gt = (san_name == 'Grand Total')
    bg = "E2E7EC" if is_gt else None
    format_cell(cell_row.cells[0], san_name, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_p(doc, "", space_after=6)

# Table 2: Volume Lấy theo Tỉnh
tbl2_df = df_lay.groupby('Tỉnh/Quận')[days10_lay].sum().round(0).astype(int)
tbl2_df.loc['Grand Total'] = tbl2_df.sum()

t2 = doc.add_table(rows=len(tbl2_df)+1, cols=11)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t2)

headers_t2 = ['Tỉnh/Quận'] + dates_header_10
for col_idx, h in enumerate(headers_t2):
    format_cell(t2.rows[0].cells[col_idx], h, bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(0,0,0), bg_hex="D9E1E8")

for row_idx, (tinh_name, row_data) in enumerate(tbl2_df.iterrows()):
    cell_row = t2.rows[row_idx+1]
    is_gt = (tinh_name == 'Grand Total')
    bg = "E2E7EC" if is_gt else None
    format_cell(cell_row.cells[0], tinh_name, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_p(doc, "", space_after=8)

# 2. Volume Giao
add_ghn_heading(doc, "2. Volume Giao", level=2)
add_ghn_heading(doc, "Tổng quan & Chi tiết loại hàng", level=3)

add_centered_picture(doc, os.path.join(charts_dir, 'ghn_combo_giao.png'), width_in_inches=6.3, caption_text="Biểu đồ Tổng quan FC Volume Giao event 08.08 (Vùng NTB)")

# Table 3: Volume Giao theo Sàn
tbl3_df = df_giao.groupby('Sàn')[days10_giao].sum().round(0).astype(int)
tbl3_df = tbl3_df.reindex([s for s in san_order if s in tbl3_df.index])
tbl3_df.loc['Grand Total'] = tbl3_df.sum()

t3 = doc.add_table(rows=len(tbl3_df)+1, cols=11)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t3)

headers_t3 = ['Sàn'] + dates_header_10
for col_idx, h in enumerate(headers_t3):
    format_cell(t3.rows[0].cells[col_idx], h, bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(0,0,0), bg_hex="D9E1E8")

for row_idx, (san_name, row_data) in enumerate(tbl3_df.iterrows()):
    cell_row = t3.rows[row_idx+1]
    is_gt = (san_name == 'Grand Total')
    bg = "E2E7EC" if is_gt else None
    format_cell(cell_row.cells[0], san_name, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_p(doc, "", space_after=6)

# Table 4: Volume Giao theo Tỉnh
tbl4_df = df_giao.groupby('Tỉnh/Quận')[days10_giao].sum().round(0).astype(int)
tbl4_df.loc['Grand Total'] = tbl4_df.sum()

t4 = doc.add_table(rows=len(tbl4_df)+1, cols=11)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t4)

headers_t4 = ['Tỉnh/Quận'] + dates_header_10
for col_idx, h in enumerate(headers_t4):
    format_cell(t4.rows[0].cells[col_idx], h, bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(0,0,0), bg_hex="D9E1E8")

for row_idx, (tinh_name, row_data) in enumerate(tbl4_df.iterrows()):
    cell_row = t4.rows[row_idx+1]
    is_gt = (tinh_name == 'Grand Total')
    bg = "E2E7EC" if is_gt else None
    format_cell(cell_row.cells[0], tinh_name, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_p(doc, "", space_after=10)

# --- SECTION II: PHÂN TÍCH CHI TIẾT CÁC NHÓM BƯU CỤC ---
add_ghn_heading(doc, "II. PHÂN TÍCH CHI TIẾT CÁC NHÓM BƯU CỤC:", level=1)

# Table 5: Phân loại 3 nhóm bưu cục
t5 = doc.add_table(rows=7, cols=4)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t5)

t5_data = [
    ["Đặc điểm", "Nhóm 1: Ổn định", "Nhóm 2: Cảnh báo", "Nhóm 3: Bất ổn"],
    ["Số lượng BC", "68 (82,93%)", "08 (9,76%)", "06 (7,31%)"],
    ["Mức độ rủi ro", "Thấp", "Trung bình", "Rất cao (Quá tải và cần tập trung theo dõi trong kỳ Event 8.8)"],
    ["Thực trạng Nhân sự", "- Thiếu dưới 2 NVPTTT\n- Đội ngũ ổn định", "- Thiếu hiện tại 1-4 NV\n- Đang có điều động hỗ trợ", "- Thiếu hiện tại 2-5 NV\n- Nằm trong danh sách rủi ro vận hành địa bàn đồi dốc/biển đảo"],
    ["Khả năng kiểm soát", "Tốt", "- Kiểm soát trung bình khá\n- Phụ thuộc tình hình thời tiết bão/mưa dốc", "- Khả năng kiểm soát kém nếu không can thiệp khẩn cấp"],
    ["Phương án A", "- AM theo dõi, điều hành gán, FIFO và bám sát năng suất NV.", "- Đẩy mạnh tuyển dụng bổ sung.\n- AM trực tiếp cắm BC điều hành gán & FIFO.\n- Thuê Freelance gánh ca peak 8.8.", "- AM trực tiếp tại BC điều phối 100%.\n- Thúc đẩy chính sách thưởng clear tồn.\n- Điều động nhân sự BC lân cận hỗ trợ."],
    ["Phương án B (khi PA A không thể thực hiện)", "- Tắt tuyến BC (khi vượt CAP x3)\n- Điều tiết giảm hàng KA", "- Tắt tuyến BC (khi vượt CAP x3)\n- Giữ hàng tại KTC", "- Tắt tuyến BC khẩn cấp\n- Điều tiết hàng KA giữ tại Kho TC Nha Trang / Lâm Đồng"]
]

for r_i, row in enumerate(t5_data):
    for c_i, val in enumerate(row):
        bg = "D9E1E8" if (r_i == 0 or c_i == 0) else None
        bold = True if (r_i == 0 or c_i == 0) else False
        align = WD_ALIGN_PARAGRAPH.CENTER if (r_i == 0 or c_i == 0) else WD_ALIGN_PARAGRAPH.LEFT
        format_cell(t5.rows[r_i].cells[c_i], val, bold=bold, font_size=8.5, align=align, color_rgb=(0,0,0), bg_hex=bg)

add_p(doc, "", space_after=8)

# Table 6: Nhóm 2 Cảnh báo
add_ghn_heading(doc, "1. Nhóm 2 (Cảnh báo - Ổn định ngắn hạn)", level=2)

t6 = doc.add_table(rows=9, cols=2)
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t6)

t6_headers = ["Xếp loại", "Chi tiết Bưu cục Nhóm 2"]
format_cell(t6.rows[0].cells[0], t6_headers[0], bold=True, font_size=9, bg_hex="D9E1E8")
format_cell(t6.rows[0].cells[1], t6_headers[1], bold=True, font_size=9, bg_hex="D9E1E8")

bc_n2 = [
    ("(KHA) BC Cam Ranh", "Thiếu 2 NVPTTT, sản lượng tăng 35% kỳ Peak, địa bàn trải dài."),
    ("(BTN) BC Phan Thiết 2", "Thiếu 3 NVPTTT, áp lực hàng cồng kềnh khu vực thương mại dịch vụ."),
    ("(LDG) BC Đức Trọng 1", "Năng suất cao nhưng áp lực giao đèo dốc và thời tiết sương mù."),
    ("(LDG) BC Bảo Lộc 2", "Thiếu 2 NVPTTT, tuyến Bảo Lâm địa hình đồi dốc dải hẹp."),
    ("(NTH) BC Ninh Hải", "Thiếu 1 NV, tuyến ven biển chịu tác động gió bão mùa mưa."),
    ("(BTN) BC Di Linh", "Thiếu 2 NVPTTT, tuyến nông thôn bán kính giao xa."),
    ("(BTN) BC Bắc Bình", "Sản lượng tăng vọt 40% ngày 08/08, phụ thuộc xe van kết nối."),
    ("(DNG) BC Gia Nghĩa 2", "Thiếu 2 NV, địa hình đồi núi Tây Nguyên mưa lớn dễ sạt lở.")
]

for idx, (bc_name, bc_desc) in enumerate(bc_n2, 1):
    format_cell(t6.rows[idx].cells[0], bc_name, bold=True, font_size=8.5)
    format_cell(t6.rows[idx].cells[1], bc_desc, font_size=8.5)

add_p(doc, "", space_after=6)
add_p(doc, "Thực trạng rủi ro:", bold=True, font_size=10)
add_p(doc, "• Nhân sự và Năng suất: Các bưu cục Nhóm 2 vẫn duy trì khả năng giao và có nhân sự điều động hỗ trợ, nhưng đang thiếu hụt nhẹ (từ 1-4 NVPTTT). Tuy nhiên, áp lực quá tải lâu ngày dễ phát sinh tình trạng nhân viên xin nghỉ phép hoặc nghỉ việc trong kỳ Event 8.8.")
add_p(doc, "• Vận hành và Thời tiết: Năng lực kiểm soát ở mức Trung bình - Khá và phụ thuộc rất lớn vào thời tiết dải ven biển và đồi dốc. Yếu tố thời tiết bất lợi (mưa giông, lũ quét đèo) dự kiến làm năng suất GTC giảm từ 3-8%. Khi gặp rủi ro kép (Thời tiết xấu + Sản lượng giao tăng đột biến), nhóm này rất dễ bị quá tải nghiêm trọng và tụt xuống Nhóm 3 (Bất ổn).")

add_p(doc, "Phương án ứng phó chủ động:", bold=True, font_size=10)
add_p(doc, "• Tối ưu mặt sàn và Kho bãi: Tiến hành sắp xếp lại layout bưu cục, tập trung clear hàng CK (cồng kềnh) để giải phóng diện tích sử dụng cho luồng hàng nhỏ nhẹ.")
add_p(doc, "• Nhân sự xử lý (NVXL) chủ động lọc riêng hàng TTS để add chuyến nhanh cho NVPTTT đi giao, nhằm tối ưu và nâng cao tỷ lệ GTC ngay trong ngày.")
add_p(doc, "• Điều phối tuyến giao: Chủ động phân tách và chuyển bớt các tuyến giao/lấy sang các bưu cục lân cận có CAP giao tốt hơn để gánh bớt số đơn. Sắp xếp lộ trình tuyến phù hợp với năng lực thực tế của từng nhân sự; ưu tiên phân bổ nhân viên mới chạy các tuyến dễ.")

add_p(doc, "", space_after=8)

# Table 7: Nhóm 3 Bất ổn
add_ghn_heading(doc, "2. Nhóm 3: Nhóm Bất ổn (Cần can thiệp khẩn cấp)", level=2)

t7 = doc.add_table(rows=7, cols=2)
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t7)

format_cell(t7.rows[0].cells[0], "Xếp loại", bold=True, font_size=9, bg_hex="D9E1E8")
format_cell(t7.rows[0].cells[1], "Chi tiết Bưu cục Nhóm 3", bold=True, font_size=9, bg_hex="D9E1E8")

bc_n3 = [
    ("(BTN) BC Phú Quý", "Bưu cục đảo xa, phụ thuộc 100% lịch tàu biển; sóng lớn cấm tàu gây dồn tồn tích lũy."),
    ("(KHA) BC Khánh Vĩnh", "Địa hình đèo dốc hiểm trở (đèo Khánh Lê), thiếu 3 NVPTTT, nguy cơ sạt lở mùa mưa."),
    ("(LDG) BC Đam Rông", "Thiếu 4 NVPTTT (>35% định biên), tuyến giao đồi núi trải dài trên 40km."),
    ("(DNG) BC Đắk Glong", "Thiếu 3 NV, đường đất đỏ trơn trượt mùa mưa, năng suất NVPTTT giảm 40%."),
    ("(BTN) BC Tuy Phong", "Tồn dồn đợt Event cũ chưa clear hết, thiếu 3 NVPTTT, CAP giao chạm trần x2.5."),
    ("(KHA) BC Khánh Sơn", "Bưu cục miền núi cao, thiếu 2 NV, sóng mạng chập chờn khi mưa lớn.")
]

for idx, (bc_name, bc_desc) in enumerate(bc_n3, 1):
    format_cell(t7.rows[idx].cells[0], bc_name, bold=True, font_size=8.5)
    format_cell(t7.rows[idx].cells[1], bc_desc, font_size=8.5)

add_p(doc, "", space_after=6)
add_p(doc, "Thực trạng rủi ro:", bold=True, font_size=10)
add_p(doc, "• Nhân sự: Báo động đỏ cần follow kỹ nhóm BC có tỷ trọng thiếu nhân sự trên 25%. Tỷ lệ nghỉ việc cao ở khu vực miền núi và đảo Phú Quý, nguồn tuyển đang được điều động gấp từ Vùng.")
add_p(doc, "• Vận hành hiện tại: Năng suất NVPTTT đang bị quá tải lâu ngày nên áp lực đi làm Event 8.8 rất lớn nếu không kiểm soát tốt.")

add_p(doc, "Phương án ứng phó khẩn cấp:", bold=True, font_size=10)
add_p(doc, "• Giai đoạn 06/08 - 15/08: Duy trì Điều động Đội phản ứng nhanh Vùng NTB từ Nhóm 1 sang hỗ trợ clear tồn cũ. Cân đối ngân sách hỗ trợ bất ổn địa bàn khó.")
add_p(doc, "• Kích hoạt chính sách 'Thưởng nóng 3k/đơn' cho các tuyến có Vol tồn vượt CAP 50%.")
add_p(doc, "• Xây dựng đội cứu các đơn GTB (Giao Không Thành Công), gọi điện xác nhận nhu cầu nhận hàng của khách trước khi xuất kho.")
add_p(doc, "• Quản trị: AM trực chiến 24/7 tại BC để gán đơn FIFO, ưu tiên giải phóng hàng CK (cồng kềnh) chiếm diện tích kho trước.")
add_p(doc, "• Cắt giảm áp lực: Trong trường hợp Vol về vượt quá 3 lần công suất thực tế, đề xuất GĐV cho phép Tắt tuyến tạm thời đối với các Shop có Vol lớn tại khu vực bất ổn để bảo vệ chỉ số chung của toàn vùng.")

add_p(doc, "Phương án A (Chủ động):", bold=True, font_size=10)
add_p(doc, "• Ưu tiên giải phóng hàng CK chiếm diện tích kho vào khung giờ thấp điểm (14h-16h). Các Bưu cục tại Khánh Vĩnh, Đức Trọng đã chủ động liên hệ xe van/tải trung chuyển để mượn tải luân chuyển đơn CK; khu vực đảo Phú Quý hiệp đồng chặt chẽ với các chủ tàu cao tốc để ưu tiên xếp sọt GHN.")

add_p(doc, "Phương án B (Trường hợp khẩn cấp):", bold=True, font_size=10)
add_p(doc, "• Đề xuất GĐV cho phép Tắt tuyến đối với hàng KA tại khu vực này nếu tồn vượt quá 2 lần CAP.")
add_p(doc, "• Điều động nhân sự từ các Bưu cục Nhóm 1 lân cận hỗ trợ theo mô hình 'cuốn chiếu'.")

add_p(doc, "", space_after=10)

# --- SECTION III: CHECKLIST CÔNG VIỆC ---
add_ghn_heading(doc, "III. CHECKLIST CÔNG VIỆC", level=1)

add_ghn_heading(doc, "1. Công cụ dụng cụ + kho bãi (Đã đảm bảo)", level=2)
add_p(doc, "• Các CCDC trên Vùng NTB đã được AM phân bổ chia nhỏ đặt bổ sung theo ngân sách. Bưu cục nào có dư CCDC chủ động nhường cho các bưu cục thiếu thay thế hư hỏng, đảm bảo 100% BC có đủ CCDC phục vụ Event 8.8.")
add_p(doc, "• Hiện tại với Volume như dự báo, chỉ cần sắp xếp lại theo layout chuẩn, chưa có Bưu cục nào cần đề xuất thuê kho tạm.")
add_p(doc, "• Chủ động liên hệ mạng lưới đối tác thuê xe tải luân chuyển hàng CK, giảm tải cho khu vực tuyến xa địa hình đèo dốc (sản lượng hàng CK chiếm 20% - 30% tổng sản lượng toàn vùng).")

add_ghn_heading(doc, "2. Bố trí lịch làm", level=2)
add_p(doc, "• Điều động nhân sự cùng lúc bố trí lịch làm xoay ca hợp lý cho nhân viên Nhóm 2 & 3 để đảm bảo sức khỏe duy trì suốt 10 ngày cao điểm Event 8.8.")
add_p(doc, "• Các Bưu cục hàng về vượt CAP giao cần phân công nhân viên ở lại rã hàng, phân loại, tạo chuyến đi sẵn trước ngày cao điểm để ngày hôm sau xuất kho sớm từ 06h30.")
add_p(doc, "• Cân nhắc chấm phụ cấp thêm cho NVXL tăng ca rã hàng (AM lập đề xuất kịp thời trình Vùng phê duyệt ngân sách).")

add_ghn_heading(doc, "3. Tác động bên ngoài", level=2)

add_ghn_heading(doc, "3.1. Trường hợp mất điện:", level=3)
add_p(doc, "• Bưu cục bám sát lịch thông báo cắt điện của Công ty Điện lực địa phương. Chuẩn bị sẵn danh bạ nóng điện lực khu vực, bộ phận Tech Vùng.")
add_p(doc, "• Những khu vực có thông báo ngắt điện phải chuẩn bị phương án thuê máy phát điện trước (dự toán chi phí thuê máy phát điện 1,000,000 VNĐ/ngày). Yêu cầu AM và NVXL nắm tình hình liên tục để kịp thời xử lý.")

add_ghn_heading(doc, "3.2. Trường hợp lỗi hệ thống Tech / Rớt mạng:", level=3)
add_p(doc, "• Nếu rớt mạng internet bưu cục, NVXL báo ngay cho AM (hỗ trợ check-in thủ công cho nhân viên nếu cần). AM có nhiệm vụ báo gấp lên nhóm Tech Vùng để xử lý.")
add_p(doc, "• Nếu rớt mạng dưới 15 phút: Tạm thời phát wifi di động 5G từ điện thoại cá nhân để duy trì bắn kiểm.")
add_p(doc, "• Nếu rớt mạng quá 30 phút ngay thời điểm xuất hàng/bắn kiểm: Chuyển sang chế độ xuất kho ngoại tuyến offline trên app theo quy trình dự phòng Tech.")

add_ghn_heading(doc, "3.3. Trường hợp Mưa giông & Lũ quét kéo dài:", level=3)
add_p(doc, "• Ưu tiên giải phóng hàng nhỏ TTS, gọi khách xác nhận nhu cầu trước khi giao ra xe ngay khi mưa nhỏ. Hàng CK vượt CAP xin thêm chi phí xe tải kết nối bảo vệ năng suất nhân viên.")
add_p(doc, "• Trang bị 100% CCDC phòng chống thiên tai: túi trùm sọt bạt che mưa, túi chống nước điện thoại, áo mưa bộ cho NVPTTT.")
add_p(doc, "• Phương án khẩn cấp: Tắt tuyến khẩn cấp + điều động Đội phản ứng nhanh giải phóng hàng tồn + kích hoạt chính sách thưởng năng suất.")
add_p(doc, "• Họp bưu cục trước Event: AM dặn dò NVPTTT chủ động chuẩn bị thiết bị chống nước điện thoại, trùm sọt, thuốc cảm, và các lưu ý an toàn giao thông mùa mưa bão đèo dốc.")

# SAVE FULL DOCX
output_docx_paths = [
    r'C:\Users\lap4all\Downloads\NTB Kế hoạch Event 8.8_Full_TNB_Style.docx',
    r'C:\Users\lap4all\Downloads\NTB Kế hoạch Event 8.8_GHN_FullBleed.docx',
    r'C:\Users\lap4all\Downloads\NTB Kế hoạch Event 8.8_v2.docx',
    r'c:\Users\lap4all\Documents\Auto report\NTB Kế hoạch Event 8.8.docx'
]

for p in output_docx_paths:
    try:
        doc.save(p)
        print(f"Successfully saved Comprehensive Word Plan Document: {p}")
    except Exception as e:
        print(f"Could not save {p}: {e}")

print("Full Comprehensive Event 8.8 Plan Document generated successfully!")
