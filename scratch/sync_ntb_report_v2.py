import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_2.docx"
dst_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_3.docx"

if not os.path.exists(doc_path):
    print(f"File không tồn tại: {doc_path}")
    sys.exit(1)

doc = docx.Document(doc_path)

print("=== BẮT ĐẦU CẬP NHẬT ĐỢT 2 ===")

# 1. Cập nhật Bảng 3 (index 2)
# Dòng 3 (index 3): BCCK Nha Trang xe
# Dòng 6 (index 6): BCCK Di Linh xe
# Dòng 9 (index 9): BCCK Đơn Dương xe
# Dòng 12 (index 12): BCCK Đức Linh xe
table_3 = doc.tables[2]

# Nha Trang xe (Row 3): T7 (col 2), T10 (col 3), T12 (col 4)
table_3.rows[3].cells[2].text = "5 xe"
table_3.rows[3].cells[3].text = "9 xe"
table_3.rows[3].cells[4].text = "10 xe"

# Di Linh xe (Row 6): T7 (col 2), T10 (col 3), T12 (col 4)
table_3.rows[6].cells[2].text = "3 xe"
table_3.rows[6].cells[3].text = "4 xe"
table_3.rows[6].cells[4].text = "5 xe"

# Đơn Dương xe (Row 9): T7 (col 2), T10 (col 3), T12 (col 4)
table_3.rows[9].cells[2].text = "2 xe"
table_3.rows[9].cells[3].text = "2 xe"
table_3.rows[9].cells[4].text = "3 xe"

# Đức Linh xe (Row 12): T7 (col 2), T10 (col 3)
table_3.rows[12].cells[2].text = "2 xe"
table_3.rows[12].cells[3].text = "2 xe"
print("-> Đã cập nhật Bảng 3")

# 2. Cập nhật Bảng 7 (index 6)
# Nha Trang giao T7 (Row 2, Col 1)
table_7 = doc.tables[6]
table_7.rows[2].cells[1].text = "38 (SL: 1815)"
print("-> Đã cập nhật Bảng 7")

# 3. Cập nhật Bảng 10 (index 9)
table_10 = doc.tables[9]
xe_cost = ['853.3', '893.3', '1173.5', '1254.6', '1361.35', '1439.4']
giao_cost = ['690.0', '750.0', '990.0', '1020.0', '1170.0', '1200.0']
kho_cost = ['120.0', '120.0', '120.0', '120.0', '120.0', '120.0']
mb_cost = ['82.5', '87.6', '125.1', '127.5', '152.6', '155.1']
tot_cost = ['1745.8', '1850.9', '2408.6', '2522.1', '2803.9', '2914.5']

for c_idx in range(1, 7):
    table_10.rows[1].cells[c_idx].text = xe_cost[c_idx - 1]
    table_10.rows[2].cells[c_idx].text = giao_cost[c_idx - 1]
    table_10.rows[3].cells[c_idx].text = kho_cost[c_idx - 1]
    table_10.rows[4].cells[c_idx].text = mb_cost[c_idx - 1]
    table_10.rows[5].cells[c_idx].text = tot_cost[c_idx - 1]
print("-> Đã cập nhật Bảng 10")

# 4. Cập nhật Bảng 11 (index 10)
table_11 = doc.tables[10]
nt_vals = ['730.0', '787.6', '1203.3', '1271.2', '1388.2', '1439.3']

for c_idx in range(1, 7):
    table_11.rows[1].cells[c_idx].text = nt_vals[c_idx - 1]
    table_11.rows[5].cells[c_idx].text = tot_cost[c_idx - 1]
print("-> Đã cập nhật Bảng 11")

doc.save(dst_path)
print(f"=== ĐÃ LƯU FILE THÀNH CÔNG: {dst_path} ===")
