import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys

sys.stdout.reconfigure(encoding='utf-8')

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

template_path = r'C:\Users\lap4all\Downloads\[TEMPLATE] BÁO CÁO KẾT QUẢ KINH DOANH [VÙNG] - [TUẦN].docx'
output_path = r'C:\Users\lap4all\Documents\Auto report\BCKD_Tuan30_vs_Tuan29_NTB.docx'

doc = docx.Document(template_path)

# Update Title P1
if len(doc.paragraphs) > 1:
    p1 = doc.paragraphs[1]
    p1.text = "BÁO CÁO KẾT QUẢ KINH DOANH TUẦN 30 – VÙNG NAM TRUNG BỘ"
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p1.runs:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

# Update Table 0 (Section 1 overview)
t0 = doc.tables[0]
# R1: Tổng doanh thu LTC
t0.rows[1].cells[1].text = "100.0%"
t0.rows[1].cells[2].text = "-3.46% (1.255 tỷ VNĐ, sản lượng 36,775 đơn; W29: 1.300 tỷ VNĐ, 38,713 đơn)"

# R2: Tổng doanh thu GTTC
t0.rows[2].cells[1].text = "100.0%"
t0.rows[2].cells[2].text = "-3.46% (Giảm 3.46% so với Tuần 29)"

# R3: Doanh thu giữ cũ
t0.rows[3].cells[1].text = "96.5%"
t0.rows[3].cells[2].text = "-3.46% (Nhóm A: 545.19M -0.9%, BCD: 256.32M -8.5%, EF: 431.07M -3.1%, G: 22.42M -9.8%)"

# R4: Doanh thu bán mới
t0.rows[4].cells[1].text = "100.0%"
t0.rows[4].cells[2].text = "+100% (5,719,011 VNĐ, 161 đơn từ 76 KH mới phát sinh trong tuần)"

# Update Paragraphs P6 to P11 (Section 1.2 Giữ cũ)
p6 = doc.paragraphs[6]
p6.text = "Nhận định chung:\nTrong Tuần 30/2026, tình hình giữ cũ của Vùng Nam Trung Bộ đạt tổng doanh thu 1,254,994,315 VNĐ (sản lượng 36,775 đơn), giảm nhẹ 3.46% về doanh thu và 5.01% về sản lượng so với Tuần 29 (1,299,958,379 VNĐ, 38,713 đơn). Nhóm khách hàng chủ lực (Nhóm A) duy trì sự ổn định cao (đạt 545.19M VNĐ, chiếm 43.4% tổng doanh thu Vùng). Sự sụt giảm tập trung chủ yếu ở nhóm BCD (giảm 8.5% từ 280.0M xuống 256.3M VNĐ) và nhóm EF (giảm 3.1% từ 444.9M xuống 431.1M VNĐ)."

p7 = doc.paragraphs[7]
p7.text = "Số liệu & Các điểm nổi bật:"

p8 = doc.paragraphs[8]
p8.text = "• Tốc độ tăng trưởng từng phân nhóm: Nhóm A đạt 99.1% so với W29 (545.19M vs 550.21M); Nhóm BCD đạt 91.5% (256.32M vs 280.00M); Nhóm EF đạt 96.9% (431.07M vs 444.90M); Nhóm G đạt 90.2% (22.42M vs 24.85M)."

p9 = doc.paragraphs[9]
p9.text = "• Nhóm khách hàng có nguy cơ rời bỏ: Toàn Vùng có 124 khách hàng có chỉ số sản lượng/doanh thu giảm sâu (% sv WTD-1 < 70%), bao gồm 101 KH nhóm EF, 22 KH nhóm BCD và 1 KH nhóm A. Đây là nhóm cần tập trung tác động gấp để tránh nguy cơ ngừng lên đơn hoàn toàn."

p10 = doc.paragraphs[10]
p10.text = "• Nhóm khách hàng có dấu hiệu giảm doanh thu: Chú ý đặc biệt KH Cám store (Nhóm A, AM Hồng Bích Nga - Lâm Đồng) giảm từ 14.63M (W29) xuống 12.75M (W30), chỉ đạt 18.4% so với mức cam kết tuần WTD-1. Ngoài ra KH Quoc Toan (Nhóm A, Bình Thuận) giảm 23.1% từ 13.82M xuống 10.62M."

p11 = doc.paragraphs[11]
p11.text = "• Nguyên nhân chính tác động đến kết quả giữ cũ: Biến động sức mua thị trường hè, một số shop kinh doanh thời trang/nông sản điều chỉnh tồn kho. Bên cạnh đó, một số tuyến giao hàng tại bưu cục địa phương có thời gian xử lý khiếu nại chậm khiến shop chia sẻ sản lượng sang đối thủ."

# Update Paragraphs P13 to P20 (Section 1.3 Bán mới)
p13 = doc.paragraphs[13]
p13.text = "Nhận định chung:\nHoạt động bán mới trong Tuần 30 đạt kết quả tích cực với 76 khách hàng mới phát sinh đơn LTC đầu tiên, tổng doanh thu đóng góp 5,719,011 VNĐ và sản lượng 161 đơn hàng (trung bình 2.12 đơn/KH, doanh thu 75.25k VNĐ/KH)."

p15 = doc.paragraphs[15]
p15.text = "• Số lượng khách hàng bán mới phát sinh: 76 KH mới phát sinh đơn trong tuần trên toàn địa bàn NTB."

p16 = doc.paragraphs[16]
p16.text = "• Khu vực / AM có kết quả bán mới tốt:\n  + Tỉnh Khánh Hòa dẫn đầu với 25 KH mới (Doanh thu 2.84M VNĐ, 70 đơn hàng).\n  + Tỉnh Lâm Đồng xếp thứ hai với 19 KH mới (Doanh thu 967k VNĐ, 41 đơn hàng).\n  + Tỉnh Ninh Thuận đóng góp 14 KH mới (Doanh thu 812k VNĐ, 25 đơn hàng).\n  + Top AM phát triển bán mới xuất sắc: Thái Thị Thanh Thư (4 KH - 1.09M VNĐ), Nguyễn Hoàng Phi (3 KH - 1.06M VNĐ), Nguyễn Duy Long (14 KH - 812k VNĐ), Phan Đình Duy (14 KH - 544k VNĐ)."

p18 = doc.paragraphs[18]
p18.text = "• Khu vực / AM chưa đạt tiến độ: Tỉnh Đắk Nông chỉ phát sinh 2 KH mới (73k VNĐ) và Bình Thuận phát sinh 4 KH mới (438k VNĐ). Cần đẩy mạnh hoạt động đi thị trường của AM tại 2 khu vực này."

p19 = doc.paragraphs[19]
p19.text = "• Các vấn đề ảnh hưởng đến tiến độ bán mới: Mật độ Shop mới mở ở các huyện miền núi Đắk Nông thưa thớt; cạnh tranh giá gắt gao từ các chành xe và đơn vị giao vận giá rẻ tại Bình Thuận."

p20 = doc.paragraphs[20]
p20.text = "• Khách hàng / Doanh thu bán mới dự kiến tuần tới: Mục tiêu tuần 31 phát sinh từ 80 - 90 KH mới, doanh thu kỳ vọng từ 8.0 - 10.0 triệu VNĐ."

# Section 2: KH Nhóm A (P22 to P30)
p22 = doc.paragraphs[22]
p22.text = "Nhận định chính:\nVùng Nam Trung Bộ hiện có 8 khách hàng nhóm A chủ lực. Doanh thu nhóm A trong Tuần 30 đạt 545.19 triệu VNĐ (giảm nhẹ 0.91% so với 550.21 triệu VNĐ tuần 29). Hầu hết các KH nhóm A lớn giữ vững phong độ, trong đó Vận Chuyển Online (AM Phan Đình Duy - Khánh Hòa) tiếp tục là KH lớn nhất Vùng với doanh thu 381.86 triệu VNĐ (tăng 3.16M so với W29). KH TIÊN HUỲNH US (AM Nguyễn Duy Long - Ninh Thuận) tăng trưởng ấn tượng với % sv WTD-1 đạt 395.4% (doanh thu 13.57M)."

p24 = doc.paragraphs[24]
p24.text = "• Tình hình nhóm KH A tại Vùng trong tuần: 8/8 KH nhóm A duy trì lên đơn liên tục, không phát sinh KH ngừng đơn."

p25 = doc.paragraphs[25]
p25.text = "• Các vấn đề nhóm KH A đang gặp phải: KH Cám store (Lâm Đồng) có dấu hiệu sụt giảm doanh thu (-12.9% từ 14.63M W29 xuống 12.75M W30), tỷ lệ đạt cam kết tuần % sv WTD-1 chỉ ở mức 18.4%. KH Quoc Toan (Bình Thuận) giảm 23.1% từ 13.82M xuống 10.62M."

p28 = doc.paragraphs[28]
p28.text = "CHI TIẾT KH NHÓM A RỜI BỎ:\nTrong Tuần 30/2026, Vùng Nam Trung Bộ KHÔNG PHÁT SINH khách hàng nhóm A nào rời bỏ (100% 8 KH nhóm A đều lên đơn ổn định)."

# Update Table 1 (Section 2 - KH Nhóm A)
t1 = doc.tables[1]
t1.rows[1].cells[1].text = "8"
t1.rows[1].cells[2].text = "0"
t1.rows[1].cells[3].text = "Duy trì 100% số lượng KH A"

t1.rows[2].cells[1].text = "8"
t1.rows[2].cells[2].text = "0"
t1.rows[2].cells[3].text = "100% KH A phát sinh doanh thu LTC"

t1.rows[3].cells[1].text = "1"
t1.rows[3].cells[2].text = "+1"
t1.rows[3].cells[3].text = "KH Cám store (Lâm Đồng) - % sv WTD-1 chỉ 18.4%"

t1.rows[4].cells[1].text = "0"
t1.rows[4].cells[2].text = "0"
t1.rows[4].cells[3].text = "Không có KH A bị ngưng đơn"

t1.rows[5].cells[1].text = "2"
t1.rows[5].cells[2].text = "+2"
t1.rows[5].cells[3].text = "KH TIÊN HUỲNH US (Ninh Thuận), My Hà (Lâm Đồng)"

# Update Table 2 (Section 2 - Chi tiết KH A rời bỏ)
t2 = doc.tables[2]
# Row 1 note
t2.rows[1].cells[0].text = "1"
t2.rows[1].cells[1].text = "-"
t2.rows[1].cells[2].text = "Không phát sinh KH nhóm A rời bỏ trong Tuần 30"
t2.rows[1].cells[3].text = "-"
t2.rows[1].cells[4].text = "-"
t2.rows[1].cells[5].text = "-"
t2.rows[1].cells[6].text = "-"
t2.rows[1].cells[7].text = "Không có"
t2.rows[1].cells[8].text = "Duy trì bám sát 8 KH A hiện hữu"

# Clear rows 2 & 3
for r_idx in [2, 3]:
    for c_idx in range(9):
        t2.rows[r_idx].cells[c_idx].text = ""

# Update Table 3 (Section 3 - Các vấn đề chính)
t3 = doc.tables[3]
# Row 1
t3.rows[1].cells[0].text = "1"
t3.rows[1].cells[1].text = "Giữ cũ (Nhóm A)"
t3.rows[1].cells[2].text = "2 KH nhóm A (Cám store - Lâm Đồng và Quoc Toan - Bình Thuận) sụt giảm doanh thu từ 13% - 23% so với W29."
t3.rows[1].cells[3].text = "Làm giảm doanh thu nhóm A toàn Vùng khoảng 5.5 triệu VNĐ/tuần."
t3.rows[1].cells[4].text = "Cao"
t3.rows[1].cells[5].text = "AM Hồng Bích Nga và AM Nguyễn Ngọc Khánh đã gặp trực tiếp shop để hỗ trợ vận hành và tư vấn ưu đãi kích cầu."

# Row 2
t3.rows[2].cells[0].text = "2"
t3.rows[2].cells[1].text = "Bán mới"
t3.rows[2].cells[2].text = "Tiến độ bán mới tại Đắk Nông (2 KH mới) và Bình Thuận (4 KH mới) chưa đạt kỳ vọng."
t3.rows[2].cells[3].text = "Hạn chế tốc độ mở rộng doanh thu bán mới Vùng."
t3.rows[2].cells[4].text = "Trung bình"
t3.rows[2].cells[5].text = "AM địa bàn rà soát danh sách Shop tiềm năng, phối hợp GDV bưu cục tiếp cận trực tiếp."

# Row 3
t3.rows[3].cells[0].text = "3"
t3.rows[3].cells[1].text = "Giữ cũ (BCD & EF)"
t3.rows[3].cells[2].text = "124 KH thuộc nhóm BCD (22 KH) và EF (101 KH) có chỉ số % sv WTD-1 < 70% (nguy cơ giảm hạng/rời bỏ)."
t3.rows[3].cells[3].text = "Tác động giảm 3.46% tổng doanh thu Vùng."
t3.rows[3].cells[4].text = "Cao"
t3.rows[3].cells[5].text = "Phân bổ danh sách 124 KH cho từng AM phụ trách để liên hệ xử lý vướng mắc đơn hàng trong 48h."

# Update Table 4 (Section 3 - Giải pháp kế hoạch)
t4 = doc.tables[4]
# Row 1
t4.rows[1].cells[0].text = "1"
t4.rows[1].cells[1].text = "Bám sát kích cầu sản lượng cho 2 KH nhóm A có dấu hiệu giảm (Cám store, Quoc Toan)"
t4.rows[1].cells[2].text = "KH Nhóm A & AM phụ trách"
t4.rows[1].cells[3].text = "Phục hồi DT Cám store >20M/tuần, Quoc Toan >15M/tuần"
t4.rows[1].cells[4].text = "04/08/2026"
t4.rows[1].cells[5].text = "AM Hồng Bích Nga, AM Nguyễn Ngọc Khánh"
t4.rows[1].cells[6].text = "Đề xuất Trưởng vùng hỗ trợ chính sách giá đặc thù nếu có đối thủ cạnh tranh"

# Row 2
t4.rows[2].cells[0].text = "2"
t4.rows[2].cells[1].text = "Triển khai chiến dịch ra soát & phát triển Shop mới tại Đắk Nông và Bình Thuận"
t4.rows[2].cells[2].text = "AM & Bưu cục địa phương"
t4.rows[2].cells[3].text = "Đạt tối thiểu 10 KH mới/tuần tại Bình Thuận, 5 KH mới/tuần tại Đắk Nông"
t4.rows[2].cells[4].text = "05/08/2026"
t4.rows[2].cells[5].text = "AM Trần Thị Nhung, AM Nguyễn Ngọc Khánh"
t4.rows[2].cells[6].text = "Bộ phận Marketing hỗ trợ voucher chào shop mới"

# Row 3
t4.rows[3].cells[0].text = "3"
t4.rows[3].cells[1].text = "Tác động trực tiếp đến 124 KH nhóm BCD/EF có chỉ số % sv WTD-1 < 70%"
t4.rows[3].cells[2].text = "124 KH nguy cơ & Toàn bộ AM Vùng"
t4.rows[3].cells[3].text = "Phục hồi sản lượng tối thiểu 50% số KH trong danh sách cảnh báo"
t4.rows[3].cells[4].text = "04/08 - 08/08"
t4.rows[3].cells[5].text = "Toàn bộ AM Vùng NTB"
t4.rows[3].cells[6].text = "Trưởng bưu cục hỗ trợ ưu tiên xử lý lấy/giao và xử lý đền bù khiếu nại"

# Format Table headers and styling
for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Format header row
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        set_cell_background(cell, "003366") # Navy header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255) # White text

    # Set padding for data cells
    for r_idx in range(1, len(table.rows)):
        for cell in table.rows[r_idx].cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if r_idx % 2 == 1:
                set_cell_background(cell, "F4F6F9") # Zebra light blue background

doc.save(output_path)
print("Successfully generated Word report at:", output_path)
