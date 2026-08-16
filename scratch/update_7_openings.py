import docx, sys

sys.stdout.reconfigure(encoding='utf-8')

# Open document
src_path = r'C:\Users\lap4all\Downloads\Quy_Hoach_MANG_LUOI_NTB_Co_Nha_Trang.docx'
doc = docx.Document(src_path)

sec5_p_idx = -1
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if 'V. TỔNG HỢP BIẾN ĐỘNG' in txt or 'V.TỔNG HỢP BIẾN ĐỘNG' in txt:
        sec5_p_idx = i

# Update Section V
for p in doc.paragraphs[sec5_p_idx:]:
    txt = p.text.strip()
    if 'Bưu cục Mở mới / Tách bưu cục' in txt:
        p.text = "❖ Bưu cục Mở mới / Tách bưu cục (07 BC): BC Xuân Hương - Đà Lạt 2, BC CK Di Linh, BC Lạc Xuân (Lâm Đồng); BC Nam Nha Trang 1 Mới, BC Nam Cam Ranh (Khánh Hòa); BC Đông Hải (Ninh Thuận); BC Nam Thành (Bình Thuận)."

final_doc_path = r'C:\Users\lap4all\Downloads\Quy_Hoach_MANG_LUOI_NTB_Co_Nha_Trang_Final.docx'
doc.save(final_doc_path)
print(f"Updated Section V with 07 openings (2 in Khánh Hòa / Nha Trang) in {final_doc_path}")
