import sys, docx, pandas as pd, json, os
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

sys.stdout.reconfigure(encoding='utf-8')

# Paths
excel_path = r'C:\Users\lap4all\Downloads\NTB_Phan_Tuyen_Hanh_Chinh_Quy_Hoach_Moi.xlsx'
df = pd.read_excel(excel_path, sheet_name='Sheet1')

doc2 = docx.Document(r'C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi.docx')
maps_dir = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\d95894ad-9248-47c3-872b-72d7de66bc20\maps'

# Detailed AM numbers, commune breakdowns, personnel, and proposals for all 36 wards
ward_am_data = {
    24823: { # Phường 1 Bảo Lộc
        'vol_giao': '1,500 đơn/ngày (AM quy hoạch) | Phường 1 cũ: 750 đơn (full)',
        'vol_lay': '1,000 đơn/ngày (AM quy hoạch)',
        'personnel': '15 NVPTTT + 2 NVXL (tại BC LDO B\'Lao Mới)',
        'prop': 'Gộp toàn bộ Phường 1 và Phường B\'Lao về Bưu cục (LDO) B\'Lao Mới. ĐÓNG CỬA bưu cục (LDO) 1 Bảo Lộc cũ.',
        'reason': '''Theo phương án quy hoạch lại các Bưu cục tiêu chuẩn theo từng địa bàn của AM tại TP. Bảo Lộc:
- Bưu cục (LDO) B'Lao Mới sẽ phụ trách toàn bộ Phường 1 (750 đơn/ngày) và Phường B'Lao (600 đơn/ngày), tổng sản lượng giao 1,500 đơn/ngày, sản lượng lấy 1,000 đơn/ngày với 15 NVPTTT + 2 NVXL.
- Bưu cục (LDO) 1 Bảo Lộc cũ diện tích không tối ưu và chồng chéo tuyến nên đề xuất ĐÓNG CỬA, chuyển toàn bộ nhân sự về BC B'Lao Mới.''',
        'map': 'image3.png',
        'map_caption': 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'
    },
    24820: { # Phường 2 Bảo Lộc
        'vol_giao': '1,400 đơn/ngày (AM quy hoạch) | Phường 2 cũ: 800 đơn (full)',
        'vol_lay': '500 đơn/ngày (AM quy hoạch)',
        'personnel': '14 NVPTTT + 2 NVXL (tại BC LDO 3 Bảo Lộc)',
        'prop': 'Gộp toàn bộ Phường 2 và Phường 3 về Bưu cục (LDO) 3 Bảo Lộc.',
        'reason': '''Kế hoạch quy hoạch AM Bảo Lộc: Bưu cục (LDO) 3 Bảo Lộc phụ trách Phường 2 (800 đơn/ngày) và Phường 3 (600 đơn/ngày), tổng sản lượng giao 1,400 đơn/ngày, lấy 500 đơn/ngày với 14 NVPTTT + 2 NVXL. Phương án này tối ưu tuyến đường giao nhận nội đô Bảo Lộc.''',
        'map': 'image3.png',
        'map_caption': 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'
    },
    24817: { # Phường B'Lao
        'vol_giao': '1,500 đơn/ngày (AM quy hoạch) | Phường B\'Lao cũ: 600 đơn (full)',
        'vol_lay': '1,000 đơn/ngày (AM quy hoạch)',
        'personnel': '15 NVPTTT + 2 NVXL (tại BC LDO B\'Lao Mới)',
        'prop': 'Gộp Phường B\'Lao và Phường 1 về Bưu cục (LDO) B\'Lao Mới.',
        'reason': '''Bưu cục (LDO) B'Lao Mới là bưu cục trung tâm phụ trách cụm Phường B'Lao và Phường 1, đáp ứng công suất 1,500 đơn giao và 1,000 đơn lấy.''',
        'map': 'image3.png',
        'map_caption': 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'
    },
    25084: { # Xã Bảo Lâm 2
        'vol_giao': '700 đơn/ngày (AM quy hoạch) | Bảo Lâm 2: 350 đơn, Bảo Lâm 3: 300 đơn',
        'vol_lay': '30 đơn/ngày (AM quy hoạch)',
        'personnel': '8 NVPTTT + 1 NVXL (tại BC LDO Bảo Lâm 3)',
        'prop': 'Gộp các tuyến xã về Bưu cục (LDO) Bảo Lâm 3.',
        'reason': '''Bưu cục (LDO) Bảo Lâm 3 cover Bảo Lâm 2 (350 đơn full) và Bảo Lâm 3 (300 đơn full), tổng giao 700 đơn/ngày, lấy 30 đơn/ngày với 8 NVPTTT + 1 NVXL.''',
        'map': 'image3.png',
        'map_caption': 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'
    },
    24781: { # Phường Xuân Hương - Đà Lạt
        'vol_giao': '1,800 đơn/ngày (BC cũ) + 1,050 đơn/ngày (BC mới tách)',
        'vol_lay': '400 đơn/ngày (BC cũ) + 150 đơn/ngày (BC mới tách)',
        'personnel': 'BC Xuân Hương cũ: 13 NVPTTT + 1 NVXL | BC Xuân Hương 2 mới: 8 NVPTTT + 1 NVXL',
        'prop': 'TÁCH MỚI Bưu cục (LDO) Xuân Hương - Đà Lạt 2 (tại Phường 10) & GIỮ NGUYÊN Bưu cục Xuân Hương cũ (phụ trách Phường 1 & Phường 2).',
        'reason': '''Phường Xuân Hương mới có sản lượng cực lớn (>2,850 đơn/ngày). Địa hình bị chia cắt bởi Hồ Xuân Hương, thời tiết mương lạnh cuối năm.
- BC Xuân Hương cũ: Cover Phường 1 (400 đơn), Phường 2 (400 đơn), Phường 4 (600 đơn), Lấy 400 đơn với 13 NVPTTT + 1 NVXL.
- BC Xuân Hương 2 (Mới tại Phường 10): Cover Phường 10 (400 đơn), Phường 3 (500 đơn), Lấy 150 đơn với 8 NVPTTT + 1 NVXL.''',
        'map': 'image6.png',
        'map_caption': 'Bản đồ địa bàn quy hoạch mới TP. Đà Lạt 2026'
    },
    24778: { # Phường Lâm Viên - Đà Lạt
        'vol_giao': 'BC Lâm Viên 1: 1,000 đơn/ngày | BC Lâm Viên 2: 1,200 đơn/ngày',
        'vol_lay': 'BC Lâm Viên 1: 250 đơn/ngày | BC Lâm Viên 2: 150 đơn/ngày',
        'personnel': 'BC Lâm Viên 1: 8 NVPTTT + 1 NVXL | BC Lâm Viên 2: 7 NVPTTT + 1 NVXL',
        'prop': 'GIỮ NGUYÊN 02 BƯU CỤC (Lâm Viên - Đà Lạt 1 & Lâm Viên - Đà Lạt 2).',
        'reason': '''Đặc thù địa hình Phường Lâm Viên bị chia cắt bởi Hồ Xuân Hương. Việc duy trì 2 Bưu cục nằm ngay trong khu vực Phường 8 giúp nhân viên chỉ tập trung xử lý một phường, giảm thời gian di chuyển, đảm bảo hiệu quả giao hàng.''',
        'map': 'image6.png',
        'map_caption': 'Bản đồ địa bàn quy hoạch mới TP. Đà Lạt 2026'
    },
    25000: { # Xã Di Linh
        'vol_giao': '1,300 - 1,600 đơn/ngày (BC Di Linh) | *Tách mới BC Hàng Nhỏ: 300-350 đơn/ngày*',
        'vol_lay': '90 - 130 đơn/ngày (BC Di Linh) | *Tách mới BC Hàng Nhỏ: 10-20 đơn/ngày*',
        'personnel': 'BC Di Linh: 16/16 NVPTTT | BC Hàng Nhỏ: 4 NVPTTT | BC Hàng Vừa: 4 NVPTTT',
        'prop': 'TÁCH BƯU CỤC HÀNG NHỎ / HÀNG VỪA (Xã Đinh Trang Thượng) để giảm tải cho BC Di Linh.',
        'reason': '''BC Di Linh hiện phụ trách Di Linh, Đinh Trang Thượng, Bảo Thuận, Gia Hiệp, Sơn Điền với bán kính xa 22-45km. BC Hòa Ninh vẫn còn cover tuyến Liên Đầm. Đề xuất mở BC Hàng Nhỏ & BC Hàng Vừa tại Đinh Trang Thượng để phủ tuyến xa Đinh Trang Thượng, Phúc Thọ Lâm Hà, Di Linh, Liên Đầm.''',
        'map': 'image14.png',
        'map_caption': 'Bản đồ địa bàn quy hoạch khu vực Di Linh & Đức Trọng 2026'
    },
    24958: { # Xã Đức Trọng
        'vol_giao': 'BC Đức Trọng 1: 400 - 600 đơn/ngày | BC Đức Trọng 2: 800 - 1,000 đơn/ngày',
        'vol_lay': 'BC Đức Trọng 1: 30 - 50 đơn/ngày | BC Đức Trọng 2: 90 - 150 đơn/ngày',
        'personnel': 'BC Đức Trọng 1: 6/8 NVPTTT | BC Đức Trọng 2: 10/10 NVPTTT',
        'prop': 'GIỮ NGUYÊN 02 BƯU CỤC (Đức Trọng 1 & Đức Trọng 2).',
        'reason': '''Khoảng cách giữa 2 BC trên 15km, địa bàn rộng, khó tuyển dụng nhân sự (mùa cà phê và mưa cuối năm), kho bãi cũ không đủ m² chứa hàng. Nếu gộp lại 1 bưu cục sẽ dẫn đến di chuyển quá xa và vỡ tuyến giao hàng.''',
        'map': 'image13.png',
        'map_caption': 'Bản đồ địa bàn quy hoạch khu vực Đức Trọng 2026'
    },
    23235: { # Phường La Gi
        'vol_giao': '1,200 - 1,500 đơn/ngày (AM quy hoạch) | Phường Lagi: 600-700, Phước Hội: 250-300, Sơn Mỹ: 200-250',
        'vol_lay': '280 đơn/ngày (AM quy hoạch)',
        'personnel': '15 NVPTTT + 1 NVXL (tại BC BTH Phước Hội Mới)',
        'prop': 'Gộp về Bưu cục (BTH) Phước Hội (Mới).',
        'reason': '''BTH Phước Hội (Mới) cover Phường Lagi, Phước Hội, Sơn Mỹ với sản lượng giao 1,200-1,500 đơn/ngày, lấy 280 đơn/ngày, định biên 15 NVPTTT + 1 NVXL. Tối ưu bán kính vận hành nội thị La Gi.''',
        'map': 'image9.png',
        'map_caption': 'Bản đồ quy hoạch mới khu vực Phường La Gi 2026'
    },
    23143: { # Xã Tân Thành
        'vol_giao': '400 - 500 đơn/ngày (AM quy hoạch) | Xã Tân Hải: 200-250, Tân Thành: 250-300',
        'vol_lay': '120 đơn/ngày (AM quy hoạch)',
        'personnel': '7 NVPTTT + 1 NVXL (tại BC BTH Tân Hải Mới)',
        'prop': 'Gộp về Bưu cục (BTH) Tân Hải (Mới). Giữ nguyên Thuận Quý cũ cho BC Hàm Thuận Nam.',
        'reason': '''BTH Tân Hải (Mới) cover Xã Tân Hải và Tân Thành (gồm Tân Thuận, Tân Thành cũ). Giữ nguyên phần Thuận Quý cũ cho BC (BTH) Hàm Thuận Nam vì địa hình đặc thù, khoảng cách từ Hàm Thuận Nam gần hơn và dân cư đông đúc hơn.''',
        'map': 'image9.png',
        'map_caption': 'Bản đồ quy hoạch mới khu vực Xã Tân Thành 2026'
    },
    22972: { # Xã Phan Rí Cửa
        'vol_giao': '460 đơn/ngày (BC Phan Rí Cửa: 200 Phan Rí Cửa, 100 Chí Công, 60 Hòa Phú, 100 Hòa Minh)',
        'vol_lay': '50 đơn/ngày',
        'personnel': '* [Định biên nhân sự: Chờ AM bổ sung cụ thể]*',
        'prop': 'GIỮ NGUYÊN 02 BƯU CỤC (Phan Rí Cửa & Liên Hương).',
        'reason': '''Vị trí 2 Bưu cục đặt tại 2 thị trấn có sản lượng lớn nhất. Khoảng cách giữa 2 thị trấn tầm 25km và địa bàn rộng nên không thể gộp lại thành 1 Bưu cục.''',
        'map': 'image4.png',
        'map_caption': 'Bản đồ quy hoạch mới khu vực Phan Rí Cửa & Tuy Phong 2026'
    },
    22528: { # Phường Ninh Hòa
        'vol_giao': 'BC Ninh Hòa 1: 390 đơn/ngày | BC Ninh Hòa 2: 800 đơn/ngày',
        'vol_lay': '* [Sản lượng lấy chi tiết: Chờ AM bổ sung]*',
        'personnel': '* [Định biên nhân sự chi tiết: Chờ AM bổ sung]*',
        'prop': 'GIỮ NGUYÊN 02 BƯU CỤC (Ninh Hòa 1 & Ninh Hòa 2).',
        'reason': '''Huyện Ninh Hòa cũ có diện tích lớn nhất cả nước (20 xã/phường). BC Ninh Hòa 1 đặt tại các xã xa trung tâm để lôi kéo nguồn lực nhân sự tại chỗ. Diện tích kho BC Ninh Hòa 2 cũng không đảm bảo m² để gom gộp.''',
        'map': 'image8.jpg',
        'map_caption': 'Bản đồ địa bàn quy hoạch mới khu vực Ninh Hòa 2026'
    },
    22759: { # Phường Phan Rang
        'vol_giao': 'BC Phan Rang: 604 đơn/ngày | BC Bảo An: 392 đơn/ngày (Tổng 996 đơn)',
        'vol_lay': 'BC Phan Rang: 153 đơn/ngày | BC Bảo An: 44 đơn/ngày (Tổng 197 đơn)',
        'personnel': 'BC Phan Rang: AM Nguyễn Duy Long | Tách mới BC Đông Hải: 7 NVPTTT + 1 NVXL',
        'prop': 'Gộp về Bưu cục chính (NTH) Phan Rang & TÁCH MỚI Bưu cục Đông Hải.',
        'reason': '''(NTH) Phan Rang cover Phường Phan Rang và Huyện Thuận Bắc. Đề xuất tách mới BC Đông Hải (cover Đông Hải, Mỹ Bình, Mỹ Đông, Mỹ Hải, Giao 600, Lấy 250, 7 NVPTTT + 1 NVXL) để rút ngắn khoảng cách và dễ tuyển dụng.''',
        'map': 'image12.png',
        'map_caption': 'Bản đồ quy hoạch mạng lưới Bưu cục tỉnh Ninh Thuận 2026'
    },
    24748: { # Xã Quảng Tân
        'vol_giao': 'BC Kiến Đức: 108 đơn/ngày (54.2%) | BC Quảng Tín: 91 đơn/ngày (45.8%)',
        'vol_lay': 'BC Kiến Đức: 1 đơn/ngày | BC Quảng Tín: 1 đơn/ngày',
        'personnel': 'AM Trần Văn Phước | * [Định biên NVPTTT/NVXL: Chờ AM bổ sung]*',
        'prop': 'GIỮ NGUYÊN 02 BƯU CỤC ((DNO) Quảng Tín & (DNO) Kiến Đức).',
        'reason': '''Địa hình Xã Quảng Tân (gồm Xã Đắk Ngo & Xã Quảng Tân cũ) có diện tích rất rộng, bán kính di chuyển xa và bị chia cắt mạnh bởi địa hình đồi núi phức tạp.
- (DNO) Kiến Đức phụ trách cụm Đắk Ngo cũ (109 đơn).
- (DNO) Quảng Tín phụ trách cụm Quảng Tân cũ (92 đơn).
Nếu gộp về 1 bưu cục, bán kính giao hàng sẽ tăng gấp đôi, gây trễ SLA. AM đề xuất GIỮ NGUYÊN 02 BC phụ trách song song.''',
        'map': None,
        'map_caption': None
    },
    24611: { # Phường Bắc Gia Nghĩa
        'vol_giao': 'BC Bắc Gia Nghĩa: 559 đơn/ngày | BC Đông Gia Nghĩa: 94 đơn/ngày | BC Quảng Sơn: 81 đơn/ngày',
        'vol_lay': '* [Sản lượng lấy chi tiết: Chờ AM bổ sung]*',
        'personnel': 'AM Trần Văn Phước / AM Trần Thị Nhung | * [Định biên NVPTTT/NVXL: Chờ AM bổ sung]*',
        'prop': 'TẠM THỜI GIỮ NGUYÊN 100% BƯU CỤC VÀ TUYẾN CŨ (trong 3-6 tháng).',
        'reason': '''Lịch sử Gia Nghĩa từng bể tuyến triền miên. Phương án tách 3 BC Bắc - Đông - Nam Gia Nghĩa giúp giữ ổn định SLA. 100% nhân sự tại chỗ ven (như Xã Đắk Ha) xác nhận sẽ NGHỈ VIỆC nếu sáp nhập kho mới xa 25-30km. Tỷ lệ đồng bào dân tộc thiểu số cao, cực khó tuyển mới. AM đề xuất TẠM GIỮ NGUYÊN 100% các kho Đắk Nông.''',
        'map': None,
        'map_caption': None
    }
}

# Build Ward list for all 36 wards
ward_list = []
for code, group in df.groupby('Mã Xã mới'):
    prov = group['Tỉnh, thành phố mới'].iloc[0]
    name = group['Tên Xã mới'].iloc[0]
    n_bcs = group['Số BC'].iloc[0]
    old_communes = group['Tên Xã cũ'].tolist()
    total_giao = group['Sản lượng giao/ngày (đơn)'].sum()
    total_lay = group['Sản lượng lấy/ngày (đơn)'].sum()
    total_don = group['TỔNG ĐƠN/NGÀY (Phường mới)'].iloc[0]
    total_kg = group['TỔNG KG/NGÀY (Phường mới)'].iloc[0]
    prop_excel = group['Đánh giá & Phương án đề xuất'].iloc[0]
    reason_excel = group['Lý do & Bố trí nhân sự'].iloc[0]

    bc_details = []
    for bc_name, bc_group in group.groupby('Tên Bưu cục giao'):
        bc_giao = bc_group['Sản lượng giao/ngày (đơn)'].sum()
        bc_lay = bc_group['Sản lượng lấy/ngày (đơn)'].sum()
        bc_total = bc_giao + bc_lay
        pct = round((bc_total / total_don * 100), 1) if total_don > 0 else 0.0
        bc_id = bc_group['ID Bưu cục giao'].iloc[0]
        am = bc_group['Quản lý khu vực (AM)'].iloc[0]
        addr = bc_group['Địa chỉ Bưu cục'].iloc[0]
        gmaps = bc_group['Google Maps Bưu cục'].iloc[0]
        bc_details.append({
            'name': bc_name, 'id': bc_id, 'am': am, 'addr': addr, 'gmaps': gmaps,
            'giao': bc_giao, 'lay': bc_lay, 'total': bc_total, 'pct': pct
        })
    bc_details.sort(key=lambda x: x['total'], reverse=True)

    # Check if exact AM data exists in ward_am_data
    if code in ward_am_data:
        info = ward_am_data[code]
        vol_giao_str = info['vol_giao']
        vol_lay_str = info['vol_lay']
        personnel_str = info['personnel']
        final_prop = info['prop']
        final_reason = info['reason']
        img_file = info['map']
        img_cap = info['map_caption']
    else:
        vol_giao_str = f"{total_giao:.1f} đơn/ngày"
        vol_lay_str = f"{total_lay:.1f} đơn/ngày"
        personnel_str = "* [Định biên nhân sự NVPTTT/NVXL: Chờ AM bổ sung]*"
        final_prop = prop_excel
        final_reason = reason_excel
        img_file = None
        img_cap = None

    ward_list.append({
        'code': code, 'name': name, 'prov': prov, 'n_bcs': n_bcs,
        'old_communes': old_communes, 'total_giao': total_giao, 'total_lay': total_lay,
        'total_don': total_don, 'total_kg': total_kg,
        'vol_giao_str': vol_giao_str, 'vol_lay_str': vol_lay_str, 'personnel_str': personnel_str,
        'prop': final_prop, 'bc_details': bc_details, 'reason': final_reason,
        'img_file': img_file, 'img_caption': img_cap
    })

ward_list.sort(key=lambda x: x['name'])

# Build Markdown content
md = []
md.append('# BÁO CÁO TOÀN DIỆN: RÀ SOÁT VÀ QUY HOẠCH BƯU CỤC THEO ĐƠN VỊ HÀNH CHÍNH MỚI (VÙNG NTB)\n')

md.append('## I. TỔNG QUAN HIỆN TRẠNG MẠNG LƯỚI BƯU CỤC VÙNG NTB')
md.append('- **Tổng số Bưu cục Express đang vận hành**: 83 Bưu cục (thuộc các tỉnh Khánh Hòa, Ninh Thuận, Bình Thuận, Lâm Đồng, Đắc Nông cũ).')
md.append('- **Tổng số Xã/Phường hành chính mới rà soát**: 36 Xã/Phường mới (gồm 114 xã/phường cũ sáp nhập).')
md.append('- **Số Xã/Phường mới đã quy hoạch chuẩn (01 BC phụ trách)**: 4 Xã/Phường (chiếm 11.1%).')
md.append('- **Số Xã/Phường mới bị CHIA CẮT (2-3 BC phụ trách)**: 32 Xã/Phường (chiếm 88.9%).\n')

md.append('## II. RÀ SOÁT THEO ĐƠN VỊ HÀNH CHÍNH MỚI')
md.append('### 1. Đánh giá độ phủ Bưu cục và ranh giới hành chính mới')
md.append('Sau khi sáp nhập các xã cũ thành xã/phường mới, ranh giới quản lý của các Bưu cục đang xuất hiện hiện tượng chia cắt mảnh, dẫn đến:')
md.append('- **01 Phường mới có tới 3 Bưu cục cùng giao hàng**: Gây chồng chéo tuyến đường, làm shipper di chuyển cắt ngang địa bàn của nhau.')
md.append('- **Tuyến đi chéo xa ranh giới**: Một số xã vùng ven bị giao từ Bưu cục ở huyện/tỉnh khác cách xa 30 - 40 km, trong khi Bưu cục lân cận chỉ cách 7 - 15 km.\n')

md.append('### 2. Danh sách 18 Tuyến giao chéo xa ranh giới cần Reassign ngay')
md.append('Bảng dưới đây liệt kê 18 tuyến xã cũ bị đi chéo ranh giới.')
md.append('> ⚠️ **ĐÁNH GIÁ VẬN HÀNH & BẤT HỢP LÝ CỦA CÁC TUYẾN CHÉO KHÁC TỈNH (ĐẶC BIỆT XÃ ĐA MI)**:')
md.append('> 1. **Trường hợp Xã Đa Mi (Bình Thuận)**: Thuật toán quét khoảng cách gợi ý chuyển về Bưu cục `(LDO) Bảo Lâm 3` (tỉnh Lâm Đồng) vì khoảng cách 15.1km. Tuy nhiên, phương án này **HOÀN TOÀN SAI VỀ MẶT HÀNH CHÍNH VÀ VẬN HÀNH**:')
md.append('>    - **Vi phạm ranh giới tỉnh**: Xã Đa Mi thuộc tỉnh Bình Thuận, việc chuyển cho BC Bảo Lâm 3 (tỉnh Lâm Đồng) phụ trách sẽ gây sai lệch luồng chia chọn kho inter-provincial, sai lệch đối soát COD và báo cáo hành chính tỉnh.')
md.append('>    - **Khác AM quản lý**: (BTH) Hàm Thuận thuộc AM Nguyễn Ngọc Khánh, trong khi (LDO) Bảo Lâm 3 thuộc AM Hồng Bích Nga.')
md.append('>    - **Kế hoạch chính thức của AM**: AM Nguyễn Ngọc Khánh đã quy hoạch **(BTH) Hàm Thuận (Mới)** giữ nguyên phụ trách Xã Đa Mi cùng cụm Đông Giang, La Dạ, Thuận Hòa, Hồng Sơn thuộc Bình Thuận.')
md.append('> 2. **Các tuyến chéo Đắk Nông & Liên tỉnh khác**: AM quản lý khu vực đề xuất **TẠM THỜI GIỮ NGUYÊN 100% PHẠM VI QUẢN LÝ THEO XÃ CŨ (trong 3-6 tháng)** để tránh rủi ro 100% nhân sự tại chỗ nghỉ việc và vỡ tuyến giao.\n')

md.append('| STT | Tên Xã/Phường cũ | Tỉnh | BC hiện tại (`from`) | Khoảng cách cũ | BC tối ưu gần nhất (`to`) | Quản lý AM tiếp nhận | Khoảng cách mới | Khoảng cách tiết kiệm | Đề xuất AM |')
md.append('|---|---|---|---|---|---|---|---|---|---|')

for r in doc2.tables[1].rows[1:]:
    cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
    xa = cells[1]
    prov = cells[2]
    bc_from = cells[3]
    bc_to = cells[5]
    
    p_from = bc_from.split(')')[0].replace('(', '').strip()[:2]
    p_to = bc_to.split(')')[0].replace('(', '').strip()[:2]
    
    if xa == "Xã Đa Mi":
        am_note = f"GIỮ NGUYÊN {bc_from} (Sai ranh giới tỉnh BTH->LDO)"
    elif p_from != p_to:
        am_note = f"GIỮ NGUYÊN {bc_from} (Tuyến khác tỉnh {p_from}->{p_to})"
    elif "Đắk Nông" in prov or "Lâm Đồng" in prov:
        am_note = f"GIỮ NGUYÊN {bc_from} (Tránh vỡ tuyến DNO/LDO)"
    else:
        am_note = "Reassign tối ưu"
        
    md.append(f'| {cells[0]} | {cells[1]} | {cells[2]} | {cells[3]} | {cells[4]} | {cells[5]} | {cells[6]} | {cells[7]} | {cells[8]} | {am_note} |')

md.append('\n## III. ĐÁNH GIÁ CHI TIẾT SẢN LƯỢNG VÀ ĐỀ XUẤT PHƯƠNG ÁN GỘP / GIỮ BƯU CỤC')
md.append(f'### Danh sách {len(ward_list)} Phường/Xã mới rà soát theo Đơn vị Hành chính mới\n')

for i, w in enumerate(ward_list, 1):
    md.append(f'### {i}. {w["name"]} ({w["prov"]})')
    md.append(f'- **Mã Xã mới**: `{w["code"]}`')
    md.append(f'- **SẢN LƯỢNG GIAO (AM ghi / rà soát)**: {w["vol_giao_str"]}')
    md.append(f'- **SẢN LƯỢNG LẤY (AM ghi / rà soát)**: {w["vol_lay_str"]}')
    md.append(f'- **ĐỊNH BIÊN NHÂN SỰ (AM ghi)**: {w["personnel_str"]}')
    
    bcs_str = ', '.join([f'"{b["name"]}"' for b in w['bc_details']])
    md.append(f'- **Các BC hiện phụ trách ({w["n_bcs"]} BC)**: {bcs_str}')
    
    old_str = ', '.join(w['old_communes'])
    md.append(f'- **Các xã cũ sáp nhập**: {old_str}')
    md.append('- **Tỷ lệ phân chia sản lượng thực tế giữa các Bưu cục**:')
    
    for b in w['bc_details']:
        gmaps_str = f' [Google Maps]({b["gmaps"]})' if b['gmaps'] and str(b['gmaps']) != 'nan' else ''
        md.append(f'  - `{b["name"]}` (ID: {b["id"]}) [AM: {b["am"]}]: {b["total"]:.1f} đơn/ngày ({b["pct"]}%) (Giao: {b["giao"]:.1f}, Lấy: {b["lay"]:.1f}) (Địa chỉ: {b["addr"]}){gmaps_str}')
    
    md.append(f'- **ĐỀ XUẤT PHƯƠNG ÁN (AM)**: **{w["prop"]}**')
    md.append(f'- **LÝ DO VÀ BỐ TRÍ NHÂN SỰ CHI TIẾT (AM ghi)**:\n{w["reason"]}')
    
    if w['img_file']:
        img_abs = os.path.join(maps_dir, w['img_file'])
        md.append(f'\n![{w["img_caption"]}]({img_abs})')
        md.append(f'*Hình {i}: {w["img_caption"]}*\n')

md.append('## IV. YÊU CẦU 3: ĐÁNH GIÁ NHU CẦU TÁCH BƯU CỤC VÀ TỐI ƯU CÔNG SUẤT KHO BÃI')
md.append('### 1. Đánh giá Bưu cục quá tải (Cần tách bớt Phường hoặc Mở rộng diện tích)')
md.append('| STT | ID Bưu cục | Tên Bưu cục | Tỉnh | Quản lý AM | Địa chỉ Bưu cục | Áp lực m² (`em2`) | Đề xuất giải pháp |')
md.append('|---|---|---|---|---|---|---|---|')

for r in doc2.tables[0].rows[1:]:
    cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
    md.append(f'| {cells[0]} | {cells[1]} | {cells[2]} | {cells[3]} | {cells[4]} | {cells[5]} | {cells[6]} | {cells[7]} |')

md.append('\n### 2. Kế hoạch Mở mới / Tách Bưu cục & Tối ưu Mạng lưới')
md.append('1. **Mở mới Bưu cục (LDO) Xuân Hương - Đà Lạt 2**: Đặt tại khu vực Phường 10 (TP. Đà Lạt), chia tải cho Bưu cục Xuân Hương cũ (phụ trách Phường 3 & Phường 10, Vol giao: 1,050 đơn/ngày, Vol lấy: 150 đơn/ngày, 8 NVPTTT + 1 NVXL).')
md.append('2. **Tách Bưu cục Hàng Nhỏ / Hàng Vừa Di Linh (Xã Đinh Trang Thượng)**: Phụ trách Đinh Trang Thượng, Di Linh, Phúc Thọ Lâm Hà, Liên Đầm để giảm bán kính di chuyển 22-45km cho BC Di Linh.')
md.append('3. **Mở mới Bưu cục Đông Hải (Tỉnh Ninh Thuận)**: Cover khu vực ven biển Đông Hải, tối ưu điểm tập kết sản lượng lấy (600 đơn giao, 250 đơn lấy, 7 NVPTTT + 1 NVXL).')
md.append('4. **Mở mới Bưu cục (LDO) B\'Lao Mới (Bảo Lộc)**: Cover Phường 1 & Phường B\'Lao (1,500 đơn giao, 1,000 đơn lấy, 15 NVPTTT + 2 NVXL), đóng cửa BC (LDO) 1 Bảo Lộc cũ.\n')

md.append('## V. TỔNG HỢP BIẾN ĐỘNG MẠNG LƯỚI BƯU CỤC NTB 2026')
md.append('- **Bưu cục Mở mới (04 BC)**: BC Xuân Hương - Đà Lạt 2, BC Di Linh Hàng Nhỏ, BC Đông Hải, BC B\'Lao Mới.')
md.append('- **Bưu cục Đóng cửa (02 BC)**: BC 1 Bảo Lộc, BC Nam Nha Trang 3.')
md.append('- **Bưu cục Gộp tuyến/Điều chỉnh phân vùng (19 BC)**: Gom các tuyến xã lẻ về BC trung tâm theo ĐVHC mới.')
md.append('- **Bưu cục Giữ nguyên vận hành (13 BC)**: Duy trì tại các khu vực đặc thù Gia Nghĩa, Đắc Nông, Đà Lạt, Đức Trọng, Phan Rí Cửa, Ninh Hòa, Xã Đa Mi (Bình Thuận)...')

full_md_text = '\n'.join(md)

# Write Markdown report to Artifact path
artifact_md_path = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\d95894ad-9248-47c3-872b-72d7de66bc20\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi.md'
with open(artifact_md_path, 'w', encoding='utf-8') as f:
    f.write(full_md_text)

print(f'Saved Smart Ward Markdown report to: {artifact_md_path}')

# Write Word DOCX report
new_doc = docx.Document()

# Style configuration
style_normal = new_doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(10.5)

def add_styled_heading(text, level):
    h = new_doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    run = h.runs[0]
    run.font.name = 'Arial'
    if level == 0:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 51, 51)
    return h

def set_table_styling(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(r'<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        if r_idx == 0:
            trPr.append(parse_xml(r'<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="003366"/>'))
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(9)
        else:
            bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{bg_color}"/>'))
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8.5)

def insert_table_18_routes_inline(doc_table_source):
    t_doc = new_doc.add_table(rows=len(doc_table_source.rows), cols=10)
    for r_idx, r in enumerate(doc_table_source.rows):
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        if r_idx == 0:
            row_data = cells + ['Đề xuất AM']
        else:
            xa = cells[1]
            prov = cells[2]
            bc_from = cells[3]
            bc_to = cells[5]
            
            p_from = bc_from.split(')')[0].replace('(', '').strip()[:2]
            p_to = bc_to.split(')')[0].replace('(', '').strip()[:2]
            
            if xa == "Xã Đa Mi":
                am_note = f"GIỮ NGUYÊN {bc_from} (Sai ranh giới tỉnh BTH->LDO)"
            elif p_from != p_to:
                am_note = f"GIỮ NGUYÊN {bc_from} (Tuyến khác tỉnh {p_from}->{p_to})"
            elif "Đắk Nông" in prov or "Lâm Đồng" in prov:
                am_note = f"GIỮ NGUYÊN {bc_from} (Tránh vỡ tuyến DNO/LDO)"
            else:
                am_note = "Reassign tối ưu"
            row_data = cells + [am_note]
        for c_idx, val in enumerate(row_data):
            t_doc.cell(r_idx, c_idx).text = val
    set_table_styling(t_doc)

def insert_table_inline(doc_table_source):
    t_doc = new_doc.add_table(rows=len(doc_table_source.rows), cols=len(doc_table_source.columns))
    for r_idx, r in enumerate(doc_table_source.rows):
        for c_idx, c in enumerate(r.cells):
            t_doc.cell(r_idx, c_idx).text = c.text.strip().replace('\n', ' ')
    set_table_styling(t_doc)

# Add Title
add_styled_heading('BÁO CÁO TOÀN DIỆN: RÀ SOÁT VÀ QUY HOẠCH BƯU CỤC THEO ĐƠN VỊ HÀNH CHÍNH MỚI (VÙNG NTB)', level=0)

i = 0
lines = md[1:]
while i < len(lines):
    line = lines[i].strip()
    if line.startswith('## '):
        add_styled_heading(line.replace('## ', ''), level=1)
    elif line.startswith('### '):
        heading_text = line.replace('### ', '')
        add_styled_heading(heading_text, level=2)
        
        if '2. Danh sách 18 Tuyến giao chéo' in heading_text:
            insert_table_18_routes_inline(doc2.tables[1])
        elif '1. Đánh giá Bưu cục quá tải' in heading_text:
            insert_table_inline(doc2.tables[0])

    elif line.startswith('#### '):
        heading_text = line.replace('#### ', '')
        add_styled_heading(heading_text, level=3)

    elif '![' in line:
        import re
        m = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if m:
            cap_text, img_p = m.group(1), m.group(2)
            if os.path.exists(img_p):
                p_img = new_doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_p, width=Inches(5.5))
                
                p_cap = new_doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_cap = p_cap.add_run(f'Hình: {cap_text}')
                r_cap.font.italic = True
                r_cap.font.size = Pt(9.5)
                r_cap.font.color.rgb = RGBColor(100, 100, 100)
    elif line.startswith('*Hình '):
        pass
    elif line.startswith('|') and '---' in line:
        pass
    elif line.startswith('|'):
        pass
    elif line.startswith('- '):
        p_b = new_doc.add_paragraph(style='List Bullet')
        txt = line.replace('- ', '')
        
        # Check if text contains italic missing marker
        if '*[' in txt or 'Chờ AM' in txt:
            run = p_b.add_run(txt)
            run.font.italic = True
            run.font.color.rgb = RGBColor(180, 50, 0)
        else:
            p_b.add_run(txt)

    elif line.startswith('  - '):
        p_b2 = new_doc.add_paragraph(style='List Bullet 2')
        txt = line.replace('  - ', '')
        p_b2.add_run(txt)
    elif line.strip():
        p_txt = new_doc.add_paragraph(line.strip())
        p_txt.paragraph_format.space_after = Pt(4)
    i += 1

docx_out_path = r'C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi_Hoan_Chinh_Theo_Ward.docx'
new_doc.save(docx_out_path)
print(f'Saved SMART WARD DOCX report to: {docx_out_path}')
