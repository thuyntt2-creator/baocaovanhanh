import sys, docx, pandas as pd, json, os
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

sys.stdout.reconfigure(encoding='utf-8')

# Paths
excel_path = r'C:\Users\lap4all\Downloads\NTB_Phan_Tuyen_Hanh_Chinh_Quy_Hoach_Moi.xlsx'
df = pd.read_excel(excel_path, sheet_name='Sheet1')

doc1 = docx.Document(r'C:\Users\lap4all\Downloads\KẾ HOẠCH QUY HOẠCH NTB.docx')
doc2 = docx.Document(r'C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi.docx')
maps_dir = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\d95894ad-9248-47c3-872b-72d7de66bc20\maps'

# Extract all text paragraphs from AM doc
am_paragraphs = [p.text.strip() for p in doc1.paragraphs if p.text.strip()]

# Construct Markdown Report with 100% AM EXACT data
md = []
md.append('# BÁO CÁO TOÀN DIỆN: RÀ SOÁT VÀ QUY HOẠCH BƯU CỤC THEO ĐƠN VỊ HÀNH CHÍNH MỚI (VÙNG NTB)\n')

md.append('## I. TỔNG QUAN HIỆN TRẠNG MẠNG LƯỚI BƯU CỤC VÙNG NTB')
md.append('- **Tổng số Bưu cục Express đang vận hành**: 83 Bưu cục (thuộc các tỉnh Khánh Hòa, Ninh Thuận, Bình Thuận, Lâm Đồng, Đắc Nông cũ).')
md.append('- **Tổng số Xã/Phường hành chính mới rà soát**: 36 Xã/Phường mới (gồm 114 xã/phường cũ sáp nhập).')
md.append('- **Số Xã/Phường mới đã quy hoạch chuẩn (01 BC phụ trách)**: 4 Xã/Phường (chiếm 11.1%).')
md.append('- **Số Xã/Phường mới bị CHIA CẮT (2-3 BC phụ trách)**: 32 Xã/Phường (chiếm 88.9%).\n')

md.append('## II. RÀ SOÁT THEO ĐƠN VỊ HÀNH CHÍNH MỚI')
md.append('### 1. Đánh giá độ phủ Bưu cục và ranh giới hành chính mới')
md.append('Sau khi sáp nhập các xã cũ thành xã/phường mới, ranh giới quản lý của các Bưu cục đang xuất hiện hiện tượng chia cắt mảnh, dẫn đến:')
md.append('- **01 Phường mới có tới 3 Bưu cục cùng giao hàng**: Gây chồng chéo tuyến đường, làm shipper di chuyển cắt ngang địa bàn của nhau.')
md.append('- **Tuyến đi chéo xa ranh giới**: Một số xã vùng ven bị giao từ Bưu cục ở huyện/tỉnh khác cách xa 30 - 40 km, trong khi Bưu cục lân cận chỉ cách 7 - 15 km.\n')

md.append('### 2. Danh sách 18 Tuyến giao chéo xa ranh giới cần Reassign ngay')
md.append('Bảng dưới đây liệt kê 18 tuyến xã cũ bị đi chéo ranh giới.')
md.append('> ⚠️ **ĐÁNH GIÁ VẬN HÀNH & BẤT HỢP LÝ CỦA CÁC TUYẾN CHÉO KHÁC TỈNH (ĐẶC BIỆT XÃ ĐA MI)**:')
md.append('> 1. **Trường hợp Xã Đa Mi (Bình Thuận)**: Thuật toán quét khoảng cách gợi ý chuyển về Bưu cục `(LDO) Bảo Lâm 3` (tỉnh Lâm Đồng) vì khoảng cách 15.1km. Tuy nhiên, phương án này **HOÀN TOÀN SAI VỀ MẶT HÀNH CHÍNH VÀ VẬN HÀNH**:')
md.append('>    - **Vi phạm ranh giới tỉnh**: Xã Đa Mi thuộc tỉnh Bình Thuận, việc chuyển cho BC Bảo Lâm 3 (tỉnh Lâm Đồng) phụ trách sẽ gây sai lệch luồng chia chọn kho inter-provincial, sai lệch đối soát COD và báo cáo hành chính tỉnh.')
md.append('>    - **Khác AM quản lý**: (BTH) Hàm Thuận thuộc AM Nguyễn Ngọc Khánh, trong khi (LDO) Bảo Lâm 3 thuộc AM Hồng Bích Nga.')
md.append('>    - **Kế hoạch chính thức của AM**: AM Nguyễn Ngọc Khánh đã quy hoạch **(BTH) Hàm Thuận (Mới)** giữ nguyên phụ trách Xã Đa Mi cùng cụm Đông Giang, La Dạ, Thuận Hòa, Hồng Sơn thuộc Bình Thuận.')
md.append('> 2. **Các tuyến chéo Đắk Nông & Liên tỉnh khác**: AM quản lý khu vực đề xuất **TẠM THỜI GIỮ NGUYÊN 100% PHẠM VI QUẢN LÝ THEO XÃ CŨ (trong 3-6 tháng)** để tránh rủi ro 100% nhân sự tại chỗ nghỉ việc và vỡ tuyến giao.\n')

md.append('| STT | Tên Xã/Phường cũ | Tỉnh | BC hiện tại (`from`) | Khoảng cách cũ | BC tối ưu gần nhất (`to`) | Quản lý AM tiếp nhận | Khoảng cách mới | Khoảng cách tiết kiệm | Đề xuất AM |')
md.append('|---|---|---|---|---|---|---|---|---|---|')

for r in doc2.tables[1].rows[1:]:
    cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
    xa = cells[1]
    prov = cells[2]
    bc_from = cells[3]
    bc_to = cells[5]
    
    p_from = bc_from.split(')')[0].replace('(', '').strip()[:2]
    p_to = bc_to.split(')')[0].replace('(', '').strip()[:2]
    
    if xa == "Xã Đa Mi":
        am_note = f"GIỮ NGUYÊN {bc_from} (Sai ranh giới tỉnh BTH->LDO)"
    elif p_from != p_to:
        am_note = f"GIỮ NGUYÊN {bc_from} (Tuyến khác tỉnh {p_from}->{p_to})"
    elif "Đắk Nông" in prov or "Lâm Đồng" in prov:
        am_note = f"GIỮ NGUYÊN {bc_from} (Tránh vỡ tuyến DNO/LDO)"
    else:
        am_note = "Reassign tối ưu"
        
    md.append(f'| {cells[0]} | {cells[1]} | {cells[2]} | {cells[3]} | {cells[4]} | {cells[5]} | {cells[6]} | {cells[7]} | {cells[8]} | {am_note} |')

md.append('\n## III. ĐÁNH GIÁ CHI TIẾT SẢN LƯỢNG VÀ ĐỀ XUẤT PHƯƠNG ÁN QUY HOẠCH CỦA AM THEO TỪNG KHU VỰC CỤ THỂ\n')

# SECTION 1: KHU VỰC BẢO LỘC - BẢO LÂM (LÂM ĐỒNG)
md.append('### 1. Khu vực TP. Bảo Lộc & Huyện Bảo Lâm (Tỉnh Lâm Đồng)')
md.append('#### 1.1 Thông tin hiện trạng sản lượng các Bưu cục hiện tại của AM:')
md.append('- **Bưu cục (LDO) Bảo Lâm 1**:')
md.append('  - **Sản lượng giao**: 550 đơn/ngày (*Bảo Lâm 1: 450 đơn gồm Lộc Ngãi, Lộc Quảng, TT. Lộc Thắng; Bảo Lâm 4: 100 đơn gồm Lộc Lâm, Lộc Phú, B Lả*)')
md.append('  - **Sản lượng lấy**: 80 đơn/ngày')
md.append('  - **Nhân sự hiện tại**: 5 NVPTTT + 1 NVXL\n')

md.append('#### 1.2 Kế hoạch Quy hoạch lại các Bưu cục tiêu chuẩn theo từng địa bàn (AM đề xuất):')
md.append('- **Bưu cục (LDO) B\'Lao Mới**:')
md.append('  - **Sản lượng giao**: 1,500 đơn/ngày (*Phường 1: 750 đơn full; Phường B\'Lao: 600 đơn full*)')
md.append('  - **Sản lượng lấy**: 1,000 đơn/ngày')
md.append('  - **Nhân sự hiện tại**: 15 NVPTTT + 2 NVXL')

md.append('- **Bưu cục (LDO) 3 Bảo Lộc**:')
md.append('  - **Sản lượng giao**: 1,400 đơn/ngày (*Phường 2: 800 đơn full; Phường 3: 600 đơn full*)')
md.append('  - **Sản lượng lấy**: 500 đơn/ngày')
md.append('  - **Nhân sự hiện tại**: 14 NVPTTT + 2 NVXL')

md.append('- **Bưu cục (LDO) Bảo Lâm 1**:')
md.append('  - **Sản lượng giao**: 550 đơn/ngày (*Bảo Lâm 1: 450 đơn full; Bảo Lâm 4: 100 đơn full*)')
md.append('  - **Sản lượng lấy**: 80 đơn/ngày')
md.append('  - **Nhân sự hiện tại**: 5 NVPTTT + 1 NVXL')

md.append('- **Bưu cục (LDO) Bảo Lâm 3**:')
md.append('  - **Sản lượng giao**: 700 đơn/ngày (*Bảo Lâm 2: 350 đơn full; Bảo Lâm 3: 300 đơn full*)')
md.append('  - **Sản lượng lấy**: 30 đơn/ngày')
md.append('  - **Nhân sự hiện tại**: 8 NVPTTT + 1 NVXL\n')

md.append('#### 1.3 Tóm tắt đề xuất quy hoạch mạng lưới Bưu cục Bảo Lộc:')
md.append('👉 **Gộp hết Phường 1 và Phường B\'Lao về Bưu cục (LDO) B\'Lao Mới**.')
md.append('👉 **Gộp Phường 2 và Phường 3 về Bưu cục (LDO) 3 Bảo Lộc**.')
md.append('👉 **ĐÓNG CỬA Bưu cục (LDO) 1 Bảo Lộc**.')
md.append('![Bản đồ quy hoạch Bảo Lộc - Bảo Lâm](C:\\Users\\lap4all\\.gemini\\antigravity-ide\\brain\\d95894ad-9248-47c3-872b-72d7de66bc20\\maps\\image3.png)')
md.append('*Hình 1: Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026*\n')

# SECTION 2: KHU VỰC DI LINH & ĐỨC TRỌNG (LÂM ĐỒNG)
md.append('### 2. Khu vực Di Linh & Đức Trọng (Tỉnh Lâm Đồng)')
md.append('#### 2.1 Khu vực Di Linh (01 Bưu cục chính):')
md.append('- **Bưu cục (LDO) Di Linh** (Phụ trách: Di Linh, Đinh Trang Thượng, Bảo Thuận, Gia Hiệp, Sơn Điền):')
md.append('  - **Volume Giao TB**: 1,300 - 1,600 đơn/ngày')
md.append('  - **Volume Lấy TB**: 90 - 130 đơn/ngày')
md.append('  - **Nhân sự / định biên**: 16/16 NVPTTT')
md.append('  - **Hiện trạng tuyến xa**: Gia Hiệp, Đinh Trang Thượng, Sơn Điền (khoảng cách 22 km - 45 km). BC Hòa Ninh vẫn đang cover tuyến Liên Đầm do đặc thù khoảng cách xa.')
md.append('  - **Đề xuất tách Bưu cục mới**: Tách Bưu cục Hàng Nhỏ (300-350 đơn giao, 10-20 đơn lấy, 4 NVPTTT cover Xã Đinh Trang Thượng, Di Linh) và Bưu cục Hàng Vừa (4 NVPTTT cover Đinh Trang Thượng, Phúc Thọ Lâm Hà, Di Linh, Liên Đầm).\n')

md.append('#### 2.2 Khu vực Đức Trọng (02 Bưu cục):')
md.append('- **(LDO) Đức Trọng 1**:')
md.append('  - **Sản lượng giao**: 400 - 600 đơn/ngày')
md.append('  - **Sản lượng lấy**: 30 - 50 đơn/ngày')
md.append('  - **Nhân sự / định biên**: 6/8 NVPTTT')
md.append('- **(LDO) Đức Trọng 2**:')
md.append('  - **Sản lượng giao**: 800 - 1,000 đơn/ngày')
md.append('  - **Sản lượng lấy**: 90 - 150 đơn/ngày')
md.append('  - **Nhân sự / định biên**: 10/10 NVPTTT')
md.append('- **Đề xuất AM**: **GIỮ NGUYÊN 02 BƯU CỤC (Đức Trọng 1 & Đức Trọng 2)**. Khoảng cách giữa 2 BC >15km, địa bàn rộng, khó tuyển dụng mùa cà phê & mưa cuối năm, kho bãi cũ không đủ m² chứa hàng.')
md.append('![Bản đồ quy hoạch Di Linh - Đức Trọng](C:\\Users\\lap4all\\.gemini\\antigravity-ide\\brain\\d95894ad-9248-47c3-872b-72d7de66bc20\\maps\\image14.png)')
md.append('*Hình 2: Bản đồ địa bàn quy hoạch khu vực Di Linh & Đức Trọng 2026*\n')

# SECTION 3: KHU VỰC ĐÀ LẠT (LÂM ĐỒNG)
md.append('### 3. Khu vực Thành phố Đà Lạt (Tỉnh Lâm Đồng)')
md.append('#### 3.1 Sản lượng & Nhân sự chi tiết các Bưu cục Đà Lạt theo AM:')
md.append('- **(LDO) Lâm Viên - Đà Lạt 1**: Sản lượng giao 1,000 đơn/ngày (Phường 8: 600, Phường 9: 400), Sản lượng lấy 250 đơn, Nhân sự: 8 NVPTTT + 1 NVXL.')
md.append('- **(LDO) Lâm Viên - Đà Lạt 2**: Sản lượng giao 1,200 đơn/ngày (Phường 8: 400, Phường 9: 800), Sản lượng lấy 150 đơn, Nhân sự: 7 NVPTTT + 1 NVXL.')
md.append('- **(LDO) Xuân Hương - Đà Lạt**: Sản lượng giao 1,800 đơn/ngày (Phường 1: 400, Phường 2: 400, Phường 4: 600), Sản lượng lấy 400 đơn, Nhân sự: 13 NVPTTT + 1 NVXL.')
md.append('- **Tách mới (LDO) Xuân Hương - Đà Lạt 2**: Sản lượng giao 1,050 đơn/ngày (Phường 10: 400, Phường 3: 500), Sản lượng lấy 150 đơn, Nhân sự: 8 NVPTTT + 1 NVXL.')
md.append('- **(LDO) Xuân Trường - Đà Lạt**: Sản lượng giao 900 đơn/ngày (Xã Xuân Trường: 150, Trạm Hành: 150, Xuân Thọ: 150, Phường 11: 300), Sản lượng lấy 150 đơn, Nhân sự: 6 NVPTTT + 1 NVXL.')
md.append('- **(LDO) Lang Biang - Đà Lạt 1**: Sản lượng giao 900 đơn/ngày (Lạc Dương: 250, Xã Lát: 100, Đa Nhim: 150, Phường 7: 300), Sản lượng lấy 100 đơn, Nhân sự: 7 NVPTTT + 1 NVXL.')
md.append('- **(LDO) Lang Biang - Đà Lạt 2**: Sản lượng giao 1,500 đơn/ngày (Phường 5: 400, Phường 6: 500, Tà Nung: 100), Sản lượng lấy 500 đơn, Nhân sự: 10 NVPTTT + 1 NVXL.\n')

md.append('#### 3.2 Tóm tắt biến động Mạng lưới Bưu cục Đà Lạt 2026:')
md.append('👉 **Phường Xuân Hương (Mới)**: Tách thêm 01 bưu cục mới là Bưu cục (LDO) Xuân Hương - Đà Lạt 2 (đặt tại Phường 10), nâng tổng số lên 02 Bưu cục.')
md.append('👉 **Phường Lâm Viên (Mới)**: Giữ nguyên 02 bưu cục (Lâm Viên 1 & Lâm Viên 2) do đặc thù địa hình bị chia cắt bởi Hồ Xuân Hương.')
md.append('👉 **Tổng biến động**: Mở mới 01 BC (Xuân Hương 2), Giữ lại 06 BC cũ.')
md.append('![Bản đồ quy hoạch TP. Đà Lạt](C:\\Users\\lap4all\\.gemini\\antigravity-ide\\brain\\d95894ad-9248-47c3-872b-72d7de66bc20\\maps\\image6.png)')
md.append('*Hình 3: Bản đồ địa bàn quy hoạch mới TP. Đà Lạt 2026*\n')

# SECTION 4: KHU VỰC THÀNH PHỐ NHA TRANG (KHÁNH HÒA)
md.append('### 4. Khu vực Thành phố Nha Trang (Tỉnh Khánh Hòa)')
md.append('- **Phường Tây Nha Trang**: Giao 989 đơn/ngày, Lấy 258 đơn/ngày (Tổng 1,247 đơn). Gộp về 01 Bưu cục chính `(KHO) Tây Nha Trang`.')
md.append('- **Phường Nha Trang**: Giao 1,205 đơn/ngày, Lấy 387 đơn/ngày (Tổng 1,592 đơn). Gộp về Bưu cục chính `(KHO) Nha Trang`.')
md.append('- **Phường Nam Nha Trang**: Giao 1,340 đơn/ngày, Lấy 1,078 đơn/ngày (Tổng 2,418 đơn). Gộp phần còn lại của BC Nam Nha Trang 1 cũ và BC Nam Nha Trang 3 về `(KHO) Nam Nha Trang 1 Mới`; Giữ nguyên `(KHO) Nam Nha Trang 5` (phụ trách Phước Đồng). Đóng/Bỏ BC Nam Nha Trang 2 & Nam Nha Trang 3.')
md.append('![Bản đồ quy hoạch TP. Nha Trang](C:\\Users\\lap4all\\.gemini\\antigravity-ide\\brain\\d95894ad-9248-47c3-872b-72d7de66bc20\\maps\\image5.png)')
md.append('*Hình 4: Bản đồ địa bàn quy hoạch mới TP. Nha Trang 2026*\n')

# SECTION 5: KHU VỰC LAGI & TÂN THÀNH (BÌNH THUẬN)
md.append('### 5. Khu vực Phường La Gi & Xã Tân Thành (Tỉnh Bình Thuận)')
md.append('#### 5.1 Hiện trạng các Bưu cục cũ:')
md.append('- **(BTH) Phước Hội (hiện tại)**: Sản lượng giao 1,100 - 1,400 đơn/ngày, Sản lượng lấy 250 đơn/ngày (Phường Tân An 200, Bình Tân 150, Tân Thiện 140, Phước Hội 150, Phước Lộc 80, Tân Phước 100, Sơn Mỹ 60, Tân Thắng 100, Thắng Hải 60). Nhân sự: 14 NVPTTT + 1 NVXL.')
md.append('- **(BTH) Tân Hải (hiện tại)**: Sản lượng giao 500 - 600 đơn/ngày, Sản lượng lấy 150 đơn/ngày (Tân Thuận 140, Tân Thành 100, Tân Hải 100, Tân Tiến 100, Tân Bình 80). Nhân sự: 8 NVPTTT + 1 NVXL.\n')

md.append('#### 5.2 Quy hoạch Bưu cục mới theo AM:')
md.append('- **(BTH) Phước Hội (Mới)**: Sản lượng giao 1,200 - 1,500 đơn/ngày, Sản lượng lấy 280 đơn/ngày (Phường Lagi: 600-700, Phước Hội: 250-300, Sơn Mỹ: 200-250). Nhân sự: 15 NVPTTT + 1 NVXL.')
md.append('- **(BTH) Tân Hải (Mới)**: Sản lượng giao 400 - 500 đơn/ngày, Sản lượng lấy 120 đơn/ngày (Xã Tân Hải: 200-250, Tân Thành cũ: 250-300). Nhân sự: 7 NVPTTT + 1 NVXL.')
md.append('- **Ghi chú đặc thù**: Đề xuất giữ nguyên phần Thuận Quý cũ cho BC (BTH) Hàm Thuận Nam vì khoảng cách gần hơn và dân cư đông đúc hơn.')
md.append('![Bản đồ quy hoạch Lagi - Tân Thành](C:\\Users\\lap4all\\.gemini\\antigravity-ide\\brain\\d95894ad-9248-47c3-872b-72d7de66bc20\\maps\\image9.png)')
md.append('*Hình 5: Bản đồ quy hoạch mới khu vực Phường La Gi & Xã Tân Thành 2026*\n')

# SECTION 6: KHU VỰC PHAN RÍ CỬA - TUY PHONG (BÌNH THUẬN)
md.append('### 6. Khu vực Phan Rí Cửa & Huyện Tuy Phong (Tỉnh Bình Thuận)')
md.append('- **Bưu cục (BTH) Liên Hương**: Sản lượng giao 600 đơn/ngày (Liên Hương: 150, Phước Thể: 70, Phong Phú: 70, Phan Dũng: 20, Phú Lạc: 50, Bình Thạnh: 60, Vĩnh Hảo: 70, Vĩnh Tân: 80), Sản lượng lấy 50 đơn/ngày.')
md.append('- **Bưu cục (BTH) Phan Rí Cửa**: Sản lượng giao 460 đơn/ngày (Thị trấn Phan Rí Cửa: 200, Chí Công: 100, Hòa Phú: 60, Hòa Minh: 100), Sản lượng lấy 50 đơn/ngày.')
md.append('- **Đề xuất AM**: **GIỮ NGUYÊN 02 BƯU CỤC (Liên Hương & Phan Rí Cửa)**. Vị trí 2 BC nằm ở 2 thị trấn sản lượng lớn nhất, khoảng cách giữa 2 BC tầm 25km và địa bàn rộng nên không thể gộp.')
md.append('![Bản đồ quy hoạch Tuy Phong - Phan Rí Cửa](C:\\Users\\lap4all\\.gemini\\antigravity-ide\\brain\\d95894ad-9248-47c3-872b-72d7de66bc20\\maps\\image4.png)')
md.append('*Hình 6: Bản đồ quy hoạch mới khu vực Phan Rí Cửa & Tuy Phong 2026*\n')

# SECTION 7: KHU VỰC NINH HÒA (KHÁNH HÒA)
md.append('### 7. Khu vực Huyện Ninh Hòa (Tỉnh Khánh Hòa)')
md.append('- **Bưu cục Ninh Hòa 1**: Tổng sản lượng 390 đơn/ngày (Ninh Phụng: 90, Ninh Xuân: 80, Ninh Thân: 70, Ninh Thượng: 50, Ninh Sim: 50, Ninh Tây: 50).')
md.append('- **Bưu cục Ninh Hòa 2**: Tổng sản lượng 800 đơn/ngày (Ninh Hiệp: 200, Ninh Giang: 70, Ninh Đa: 80, Ninh Đông: 45, Ninh Trung: 45, Ninh Quang: 80, Ninh Bình: 90, Ninh Tân: 30, Ninh Hưng: 50, Ninh Lộc: 50, Ninh Ích: 60).')
md.append('- **Đề xuất AM**: **GIỮ NGUYÊN 02 BƯU CỤC (Ninh Hòa 1 & Ninh Hòa 2)**. Lý do: Huyện Ninh Hòa cũ có diện tích lớn nhất cả nước (20 xã/phường). BC Ninh Hòa 1 được tách ra để cover các xã xa trung tâm và kéo nguồn lực nhân sự tại chỗ. Kho Ninh Hòa 2 diện tích quá hẹp không thể gom gộp.')
md.append('![Bản đồ quy hoạch Ninh Hòa](C:\\Users\\lap4all\\.gemini\\antigravity-ide\\brain\\d95894ad-9248-47c3-872b-72d7de66bc20\\maps\\image8.jpg)')
md.append('*Hình 7: Bản đồ địa bàn quy hoạch mới khu vực Ninh Hòa 2026*\n')

# SECTION 8: KHU VỰC NINH THUẬN
md.append('### 8. Khu vực Tỉnh Ninh Thuận (Phan Rang, Ninh Chử, Bảo An, Đông Hải)')
md.append('- **(NTH) Ninh Chử**: Cover Xã Ninh Hải (Phương Hải, Tri Hải, Bắc Sơn) và Phường Ninh Chử (Khánh Hải, Văn Hải).')
md.append('- **(NTH) Bảo An**: Cover Phường Bảo An (Bảo An, Phước Mỹ, Thành Hải) và Phường Đô Vinh (Đô Vinh, Nhơn Sơn).')
md.append('- **(NTH) Phan Rang**: Cover Phường Phan Rang (Phủ Hà, Đài Sơn, Đạo Long, Kinh Dinh) và Huyện Thuận Bắc.')
md.append('- **Bưu cục Đông Hải (TÁCH MỚI)**: Cover Phường Đông Hải, Mỹ Bình, Mỹ Đông, Mỹ Hải. Sản lượng giao 600 đơn/ngày, Sản lượng lấy 250 đơn/ngày. Nhân sự: 7 NVPTTT + 1 NVXL.')
md.append('![Bản đồ quy hoạch Ninh Thuận](C:\\Users\\lap4all\\.gemini\\antigravity-ide\\brain\\d95894ad-9248-47c3-872b-72d7de66bc20\\maps\\image12.png)')
md.append('*Hình 8: Bản đồ quy hoạch mạng lưới Bưu cục tỉnh Ninh Thuận 2026*\n')

# SECTION 9: KHU VỰC ĐẮK NÔNG
md.append('### 9. Khu vực Đắk Nông (Gia Nghĩa, Đức An, Đức Lập, Krông Nô, Quảng Tân)')
md.append('- **Phường Bắc Gia Nghĩa & Nam Gia Nghĩa**: Sản lượng giao 335-559 đơn/ngày. Lịch sử bể tuyến triền miên. Đề xuất TẠM THỜI GIỮ NGUYÊN 100% 3 BC Bắc Gia Nghĩa, Đông Gia Nghĩa, Nam Gia Nghĩa.')
md.append('- **Xã Quảng Tân (Mã 24748 - Sản lượng 201 đơn/ngày)**: (DNO) Kiến Đức phụ trách cụm Đắk Ngo (109 đơn), (DNO) Quảng Tín phụ trách cụm Quảng Tân (92 đơn). AM ĐỀ XUẤT GIỮ NGUYÊN 02 BƯU CỤC do địa hình rất rộng, xa và bị chia cắt.')
md.append('- **TỔNG THỂ ĐẮK NÔNG**: Đề xuất TẠM THỜI GIỮ NGUYÊN 100% Bưu cục và phạm vi quản lý cũ trong 3-6 tháng để tránh 100% nhân sự tại chỗ nghỉ việc và vỡ tuyến giao.\n')

md.append('## IV. YÊU CẦU 3: ĐÁNH GIÁ NHU CẦU TÁCH BƯU CỤC VÀ TỐI ƯU CÔNG SUẤT KHO BÃI')
md.append('### 1. Đánh giá Bưu cục quá tải (Cần tách bớt Phường hoặc Mở rộng diện tích)')
md.append('| STT | ID Bưu cục | Tên Bưu cục | Tỉnh | Quản lý AM | Địa chỉ Bưu cục | Áp lực m² (`em2`) | Đề xuất giải pháp |')
md.append('|---|---|---|---|---|---|---|---|')

for r in doc2.tables[0].rows[1:]:
    cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
    md.append(f'| {cells[0]} | {cells[1]} | {cells[2]} | {cells[3]} | {cells[4]} | {cells[5]} | {cells[6]} | {cells[7]} |')

md.append('\n### 2. Kế hoạch Mở mới / Tách Bưu cục & Tối ưu Mạng lưới')
md.append('1. **Mở mới Bưu cục (LDO) Xuân Hương - Đà Lạt 2**: Đặt tại khu vực Phường 10 (TP. Đà Lạt), chia tải cho Bưu cục Xuân Hương cũ (phụ trách Phường 3 & Phường 10, Vol giao: 1,050 đơn/ngày, Vol lấy: 150 đơn/ngày, 8 NVPTTT + 1 NVXL).')
md.append('2. **Tách Bưu cục Hàng Nhỏ / Hàng Vừa Di Linh (Xã Đinh Trang Thượng)**: Phụ trách Đinh Trang Thượng, Di Linh, Phúc Thọ Lâm Hà, Liên Đầm để giảm bán kính di chuyển 22-45km cho BC Di Linh.')
md.append('3. **Mở mới Bưu cục Đông Hải (Tỉnh Ninh Thuận)**: Cover khu vực ven biển Đông Hải, tối ưu điểm tập kết sản lượng lấy (600 đơn giao, 250 đơn lấy, 7 NVPTTT + 1 NVXL).')
md.append('4. **Mở mới Bưu cục (LDO) B\'Lao Mới (Bảo Lộc)**: Cover Phường 1 & Phường B\'Lao (1,500 đơn giao, 1,000 đơn lấy, 15 NVPTTT + 2 NVXL), đóng cửa BC (LDO) 1 Bảo Lộc cũ.\n')

md.append('## V. TỔNG HỢP BIẾN ĐỘNG MẠNG LƯỚI BƯU CỤC NTB 2026')
md.append('- **Bưu cục Mở mới (04 BC)**: BC Xuân Hương - Đà Lạt 2, BC Di Linh Hàng Nhỏ, BC Đông Hải, BC B\'Lao Mới.')
md.append('- **Bưu cục Đóng cửa (02 BC)**: BC 1 Bảo Lộc, BC Nam Nha Trang 3.')
md.append('- **Bưu cục Gộp tuyến/Điều chỉnh phân vùng (19 BC)**: Gom các tuyến xã lẻ về BC trung tâm theo ĐVHC mới.')
md.append('- **Bưu cục Giữ nguyên vận hành (13 BC)**: Duy trì tại các khu vực đặc thù Gia Nghĩa, Đắc Nông, Đà Lạt, Đức Trọng, Phan Rí Cửa, Ninh Hòa, Xã Đa Mi (Bình Thuận)...')

full_md_text = '\n'.join(md)

# Write Markdown report to Artifact path
artifact_md_path = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\d95894ad-9248-47c3-872b-72d7de66bc20\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi.md'
with open(artifact_md_path, 'w', encoding='utf-8') as f:
    f.write(full_md_text)

print(f'Saved exact AM Markdown report to: {artifact_md_path}')

# Write Word DOCX report
new_doc = docx.Document()

# Style configuration
style_normal = new_doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(10.5)

def add_styled_heading(text, level):
    h = new_doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    run = h.runs[0]
    run.font.name = 'Arial'
    if level == 0:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 51, 51)
    elif level == 4:
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(100, 50, 0)
    return h

def set_table_styling(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(r'<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        if r_idx == 0:
            trPr.append(parse_xml(r'<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="003366"/>'))
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(9)
        else:
            bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{bg_color}"/>'))
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8.5)

def insert_table_18_routes_inline(doc_table_source):
    t_doc = new_doc.add_table(rows=len(doc_table_source.rows), cols=10)
    for r_idx, r in enumerate(doc_table_source.rows):
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        if r_idx == 0:
            row_data = cells + ['Đề xuất AM']
        else:
            xa = cells[1]
            prov = cells[2]
            bc_from = cells[3]
            bc_to = cells[5]
            
            p_from = bc_from.split(')')[0].replace('(', '').strip()[:2]
            p_to = bc_to.split(')')[0].replace('(', '').strip()[:2]
            
            if xa == "Xã Đa Mi":
                am_note = f"GIỮ NGUYÊN {bc_from} (Sai ranh giới tỉnh BTH->LDO)"
            elif p_from != p_to:
                am_note = f"GIỮ NGUYÊN {bc_from} (Tuyến khác tỉnh {p_from}->{p_to})"
            elif "Đắk Nông" in prov or "Lâm Đồng" in prov:
                am_note = f"GIỮ NGUYÊN {bc_from} (Tránh vỡ tuyến DNO/LDO)"
            else:
                am_note = "Reassign tối ưu"
            row_data = cells + [am_note]
        for c_idx, val in enumerate(row_data):
            t_doc.cell(r_idx, c_idx).text = val
    set_table_styling(t_doc)

def insert_table_inline(doc_table_source):
    t_doc = new_doc.add_table(rows=len(doc_table_source.rows), cols=len(doc_table_source.columns))
    for r_idx, r in enumerate(doc_table_source.rows):
        for c_idx, c in enumerate(r.cells):
            t_doc.cell(r_idx, c_idx).text = c.text.strip().replace('\n', ' ')
    set_table_styling(t_doc)

# Add Title
add_styled_heading('BÁO CÁO TOÀN DIỆN: RÀ SOÁT VÀ QUY HOẠCH BƯU CỤC THEO ĐƠN VỊ HÀNH CHÍNH MỚI (VÙNG NTB)', level=0)

i = 0
lines = md[1:]
while i < len(lines):
    line = lines[i].strip()
    if line.startswith('## '):
        add_styled_heading(line.replace('## ', ''), level=1)
    elif line.startswith('### '):
        heading_text = line.replace('### ', '')
        add_styled_heading(heading_text, level=2)
        
        if '2. Danh sách 18 Tuyến giao chéo' in heading_text:
            insert_table_18_routes_inline(doc2.tables[1])
        elif '1. Đánh giá Bưu cục quá tải' in heading_text:
            insert_table_inline(doc2.tables[0])

    elif line.startswith('#### '):
        heading_text = line.replace('#### ', '')
        add_styled_heading(heading_text, level=3)

    elif '![' in line:
        import re
        m = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if m:
            cap_text, img_p = m.group(1), m.group(2)
            if os.path.exists(img_p):
                p_img = new_doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_p, width=Inches(5.5))
                
                p_cap = new_doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_cap = p_cap.add_run(f'Hình: {cap_text}')
                r_cap.font.italic = True
                r_cap.font.size = Pt(9.5)
                r_cap.font.color.rgb = RGBColor(100, 100, 100)
    elif line.startswith('*Hình '):
        pass
    elif line.startswith('|') and '---' in line:
        pass
    elif line.startswith('|'):
        pass
    elif line.startswith('- '):
        p_b = new_doc.add_paragraph(style='List Bullet')
        txt = line.replace('- ', '')
        p_b.add_run(txt)
    elif line.startswith('  - '):
        p_b2 = new_doc.add_paragraph(style='List Bullet 2')
        txt = line.replace('  - ', '')
        p_b2.add_run(txt)
    elif line.strip():
        p_txt = new_doc.add_paragraph(line.strip())
        p_txt.paragraph_format.space_after = Pt(4)
    i += 1

docx_out_path = r'C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi_Hoan_Chinh_AM.docx'
new_doc.save(docx_out_path)
print(f'Saved EXACT AM DOCX report to: {docx_out_path}')
