import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os, sys

sys.stdout.reconfigure(encoding='utf-8')

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

w29_path = r'C:\Users\lap4all\Downloads\BÁO CÁO KẾT QUẢ KINH DOANH TUẦN 29_2026 - NTB.docx'
out_path_workspace = r'C:\Users\lap4all\Documents\Auto report\BÁO CÁO KẾT QUẢ KINH DOANH TUẦN 31_2026 - NTB.docx'
out_path_downloads = r'C:\Users\lap4all\Downloads\BÁO CÁO KẾT QUẢ KINH DOANH TUẦN 31_2026 - NTB.docx'

doc = docx.Document(w29_path)

# Update Title P0 & Subtitle P1
doc.paragraphs[0].text = "BÁO CÁO KẾT QUẢ KINH DOANH TUẦN 31 – VÙNG NAM TRUNG BỘ"
doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in doc.paragraphs[0].runs:
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.paragraphs[1].text = "Thời gian báo cáo: Tuần 31 (26/07/2026 – 01/08/2026) | Kỳ so sánh: Tuần 30 (19/07/2026 – 25/07/2026)"
doc.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Table 0: Tổng quan KQKD
t0 = doc.tables[0]
t0.rows[0].cells[2].text = "% Tăng/ giảm các chỉ tiêu so với Tuần 30"

t0.rows[1].cells[1].text = "96.2%"
t0.rows[1].cells[2].text = "154.16 triệu VNĐ (Tổng hợp đến hiện tại; Tuần 30: 1.255 tỷ VNĐ GTTC)"

t0.rows[2].cells[1].text = "96.8%"
t0.rows[2].cells[2].text = "154.16 triệu VNĐ (4,657 đơn hàng)"

t0.rows[3].cells[1].text = "96.5%"
t0.rows[3].cells[2].text = "154.16 triệu VNĐ (Nhóm A: 68.76M, BCD: 35.53M, EF: 49.10M, G: 0.77M)"

t0.rows[4].cells[1].text = "82.0%"
t0.rows[4].cells[2].text = "5,719,011 VNĐ (161 đơn từ 76 KH mới phát sinh trong tuần)"

# Paragraph P7 (1.2 Giữ cũ - Nhận định chung)
doc.paragraphs[7].text = "Trong Tuần 31 (26/07 – 01/08/2026), công tác giữ cũ của Vùng Nam Trung Bộ tiếp tục duy trì hoạt động trên các phân nhóm khách hàng trọng điểm. Nhóm KH A chủ lực giữ vững quy mô 8/8 KH có phát sinh đơn LTC (doanh thu 68.76 triệu VNĐ, 2,033 đơn hàng). Nhóm BCD ghi nhận 68 KH (doanh thu 35.53 triệu VNĐ, 1,243 đơn) và Nhóm EF ghi nhận 561 KH (doanh thu 49.10 triệu VNĐ, 1,361 đơn). Danh mục cảnh báo KH sụt giảm sản lượng (% sv WTD-1 < 70%) Tuần 31 ghi nhận 125 KH (bao gồm 1 KH nhóm A, 22 KH nhóm BCD, 102 KH nhóm EF)."

# P9: Phân loại nhóm KH
doc.paragraphs[9].text = "• Phân loại nhóm Khách hàng (Vol & DT Tuần 31 vs Tuần 30): Nhóm A đạt 8 KH (2,033 đơn, 68.76tr VNĐ); Nhóm BCD đạt 68 KH (1,243 đơn, 35.53tr VNĐ); Nhóm EF đạt 561 KH (1,361 đơn, 49.10tr VNĐ); Nhóm G đạt 20 KH (20 đơn, 0.77tr VNĐ)."

# P10: Nhóm nguy cơ
doc.paragraphs[10].text = "• Nhóm KH nguy cơ rời bỏ / sụt giảm (% sv WTD-1 < 70% - Sheet 2 Tuần 31): 125 KH (Lâm Đồng: 42 KH, Khánh Hòa: 35 KH, Ninh Thuận: 18 KH, Đắk Nông: 17 KH, Bình Thuận: 13 KH). Phân hạng: 1 KH A, 22 KH BCD, 102 KH EF."

# P11: Chi tiết KH Nhóm A dấu hiệu giảm
doc.paragraphs[11].text = "• Chi tiết KH Nhóm A có dấu hiệu giảm đơn: (1) Shop Cám store (Lâm Đồng, AM Hồng Bích Nga) doanh thu Tuần 31 ghi nhận 657.4k VNĐ, tỷ lệ % sv WTD-1 đạt 18.4% do shop điều chỉnh tiến độ bán hàng hè; (2) Shop Quoc Toan (Bình Thuận, AM Nguyễn Ngọc Khánh) doanh thu 10.62tr VNĐ, % sv WTD-1 đạt 94.4% do ảnh hưởng thời tiết xấu tới vận chuyển hải sản khô."

# P12: Nguyên nhân chính
doc.paragraphs[12].text = "• Nguyên nhân chính: (1) Thị trường tiêu dùng mùa hè chững nhẹ cuối tháng 7; (2) Cạnh tranh giá cước gay gắt từ các chành xe tuyến địa phương tác động đến nhóm BCD; (3) AM cần bám sát chặt chẽ hơn tiến độ lên đơn hàng ngày của nhóm BCD/EF."

# P15 (1.3 Bán mới - Nhận định chung)
doc.paragraphs[15].text = "Trong Tuần 31, công tác bán mới ghi nhận 76 KH mới phát sinh đơn LTC đầu tiên, mang lại tổng doanh thu bán mới 5.72 triệu VNĐ và tổng sản lượng 161 đơn hàng (bình quân 2.12 đơn/KH). Tỉnh Khánh Hòa tiếp tục dẫn đầu Vùng về quy mô bán mới (25 KH mới, 2.84tr VNĐ, 70 đơn), tiếp theo là Lâm Đồng (19 KH mới, 967k VNĐ, 41 đơn) và Ninh Thuận (14 KH mới, 812k VNĐ, 25 đơn)."

# P17: Số liệu bán mới
doc.paragraphs[17].text = "• Số lượng KH bán mới phát sinh trong tuần: 76 KH | Doanh thu: 5,719,011 VNĐ | Sản lượng: 161 đơn | AOV: 2.12 đơn/KH."

# P18: Khu vực / AM bán mới tốt
doc.paragraphs[18].text = "• Khu vực / AM có kết quả bán mới xuất sắc: AM Thái Thị Thanh Thư (Khánh Hòa) dẫn đầu về doanh thu bán mới (4 KH mới, 1.09tr VNĐ); AM Nguyễn Hoàng Phi (Khánh Hòa) đạt 3 KH mới (1.06tr VNĐ); AM Nguyễn Duy Long (Ninh Thuận) đạt số lượng KH mới cao nhất (14 KH mới, 0.81tr VNĐ); AM Phan Đình Duy (Khánh Hòa) đồng hạng số lượng KH mới (14 KH mới, 0.54tr VNĐ)."

# P19: Khu vực / AM chưa đạt
doc.paragraphs[19].text = "• Khu vực / AM chưa đạt tiến độ: Tỉnh Đắk Nông (2 KH mới, 73k VNĐ) và Bình Thuận (4 KH mới, 438k VNĐ)."

# P20: Các vấn đề đang ảnh hưởng
doc.paragraphs[20].text = "• Các vấn đề đang ảnh hưởng: Mật độ shop mới tại các huyện miền núi Đắk Nông thưa thớt; sự cạnh tranh giá từ các hãng xe khách/chành xe chạy tuyến Bình Thuận - HCM làm giảm tỷ lệ chuyển đổi shop mới."

# P21: Dự kiến tuần tới
doc.paragraphs[21].text = "• Dự kiến tuần tới (Tuần 32): Triển khai chiến dịch săn shop mở rộng B2B tại Đắk Nông & Bình Thuận, mục tiêu đạt 100+ KH mới toàn Vùng."

# Section 2: KH Nhóm A (P25, P30)
doc.paragraphs[25].text = "Trong Tuần 31, danh mục KH A duy trì 8/8 KH có phát sinh đơn LTC trên hệ thống. Shop Vận Chuyển Online (Khánh Hòa, AM Phan Đình Duy) tiếp tục giữ vị trí số 1 toàn Vùng. Shop TIÊN HUỲNH US (Ninh Thuận, AM Nguyễn Duy Long) có bước bứt phá ngoạn mục với chỉ số % sv WTD-1 đạt 395.4%. Vùng đang tiếp tục bám sát 2 KH tiềm năng thăng hạng A (TIÊN HUỲNH US và My Hà) để thúc đẩy tăng trưởng trong tháng 8."

doc.paragraphs[30].text = "Tình hình sụt giảm tập trung ở 2 shop nhóm A: (1) Shop Cám store (Lâm Đồng, AM Hồng Bích Nga) doanh thu W31 đạt 657.4k VNĐ, % sv WTD-1 đạt 18.4%; (2) Shop Quoc Toan (Bình Thuận, AM Nguyễn Ngọc Khánh) doanh thu 10.62tr VNĐ, % sv WTD-1 đạt 94.4%."

# Table 1: Chỉ tiêu KH nhóm A
t1 = doc.tables[1]
t1.rows[0].cells[1].text = "Dữ liệu Tuần 31"
t1.rows[0].cells[2].text = "So sánh tăng/ giảm với Tuần 30"

t1.rows[1].cells[1].text = "8 KH"
t1.rows[1].cells[2].text = "Giữ nguyên"
t1.rows[1].cells[3].text = "Danh mục nhóm A chốt 8 KH chủ lực"

t1.rows[2].cells[1].text = "8 KH"
t1.rows[2].cells[2].text = "Bằng (0)"
t1.rows[2].cells[3].text = "100% KH A phát sinh doanh thu LTC"

t1.rows[3].cells[1].text = "1 KH"
t1.rows[3].cells[2].text = "-1 KH"
t1.rows[3].cells[3].text = "Shop Cám store (Lâm Đồng) % sv WTD-1 chỉ 18.4%"

t1.rows[4].cells[1].text = "0 KH"
t1.rows[4].cells[2].text = "Bằng (0)"
t1.rows[4].cells[3].text = "Không có KH A bị ngưng hẳn đơn LTC"

t1.rows[5].cells[1].text = "2 KH"
t1.rows[5].cells[2].text = "+2 KH"
t1.rows[5].cells[3].text = "TIÊN HUỲNH US (% sv WTD-1 395.4%), My Hà (% sv WTD-1 204.3%)"

# Table 2: Chi tiết KH nhóm A sụt giảm
t2 = doc.tables[2]
t2.rows[0].cells[4].text = "Doanh thu Tuần 31 (% sv WTD-1)"

# Row 1
t2.rows[1].cells[0].text = "1"
t2.rows[1].cells[1].text = "4313038"
t2.rows[1].cells[2].text = "Cám store"
t2.rows[1].cells[3].text = "Hồng Bích Nga"
t2.rows[1].cells[4].text = "657.4k VNĐ (% sv WTD-1 18.4%)"
t2.rows[1].cells[5].text = "18.4%"
t2.rows[1].cells[6].text = "A"
t2.rows[1].cells[7].text = "Shop giảm nhập hàng quần áo hè, vướng lịch nghỉ bán"
t2.rows[1].cells[8].text = "AM gặp trực tiếp shop tư vấn gói khuyến mãi cước kích cầu"

# Row 2
t2.rows[2].cells[0].text = "2"
t2.rows[2].cells[1].text = "3200594"
t2.rows[2].cells[2].text = "Quoc Toan"
t2.rows[2].cells[3].text = "Nguyễn Ngọc Khánh"
t2.rows[2].cells[4].text = "10.62tr VNĐ (% sv WTD-1 94.4%)"
t2.rows[2].cells[5].text = "94.4%"
t2.rows[2].cells[6].text = "A"
t2.rows[2].cells[7].text = "Thời tiết xấu ảnh hưởng chuyến giao hải sản khô"
t2.rows[2].cells[8].text = "AM bám sát chuyến hàng mới và ưu tiên xe lấy ca 1"

# Table 3: Các vấn đề chính
t3 = doc.tables[3]
t3.rows[1].cells[0].text = "1"
t3.rows[1].cells[1].text = "Giữ cũ (Shop A)"
t3.rows[1].cells[2].text = "Cám store (Lâm Đồng) và Quoc Toan (Bình Thuận) sụt giảm doanh thu so với mức cam kết tuần WTD-1"
t3.rows[1].cells[3].text = "Tác động giảm doanh thu nhóm A Vùng"
t3.rows[1].cells[4].text = "Cao"
t3.rows[1].cells[5].text = "AM Hồng Bích Nga & AM Nguyễn Ngọc Khánh làm việc trực tiếp kích cầu và ưu tiên vận hành lấy hàng"

t3.rows[2].cells[0].text = "2"
t3.rows[2].cells[1].text = "Giữ cũ (BCD & EF)"
t3.rows[2].cells[2].text = "125 KH thuộc nhóm BCD (22 KH) và EF (102 KH) có chỉ số % sv WTD-1 < 70%"
t3.rows[2].cells[3].text = "Tác động sụt giảm doanh thu các phân nhóm BCD/EF"
t3.rows[2].cells[4].text = "Cao"
t3.rows[2].cells[5].text = "Rà soát chỉ số vận hành bưu cục, hỗ trợ AM xử lý ticket và giải quyết khiếu nại đền bù nhanh"

t3.rows[3].cells[0].text = "3"
t3.rows[3].cells[1].text = "Bán mới"
t3.rows[3].cells[2].text = "Tiến độ bán mới tại Đắk Nông (2 KH mới) và Bình Thuận (4 KH mới) chưa đạt kỳ vọng"
t3.rows[3].cells[3].text = "Doanh thu bán mới Vùng Tuần 31 đạt 5.72 triệu VNĐ (76 KH mới)"
t3.rows[3].cells[4].text = "Trung bình"
t3.rows[3].cells[5].text = "AM địa bàn rà soát danh sách Shop tiềm năng, phối hợp GDV bưu cục tiếp cận trực tiếp"

# Table 4: Giải pháp kế hoạch
t4 = doc.tables[4]
t4.rows[1].cells[0].text = "1"
t4.rows[1].cells[1].text = "Bám sát kích cầu sản lượng cho 2 KH nhóm A có dấu hiệu giảm (Cám store, Quoc Toan)"
t4.rows[1].cells[2].text = "KH Nhóm A & AM phụ trách"
t4.rows[1].cells[3].text = "Phục hồi DT Cám store >20M/tuần, Quoc Toan >15M/tuần"
t4.rows[1].cells[4].text = "05/08/2026"

t4.rows[2].cells[0].text = "2"
t4.rows[2].cells[1].text = "Tác động trực tiếp đến 125 KH nhóm BCD/EF có chỉ số % sv WTD-1 < 70%"
t4.rows[2].cells[2].text = "125 KH nguy cơ & Toàn bộ AM Vùng"
t4.rows[2].cells[3].text = "Phục hồi sản lượng tối thiểu 50% số KH trong danh sách cảnh báo"
t4.rows[2].cells[4].text = "05/08/2026"

t4.rows[3].cells[0].text = "3"
t4.rows[3].cells[1].text = "Triển khai chiến dịch ra soát & phát triển Shop mới tại Đắk Nông và Bình Thuận"
t4.rows[3].cells[2].text = "AM & Bưu cục địa phương"
t4.rows[3].cells[3].text = "Đạt tối thiểu 10 KH mới/tuần tại Bình Thuận, 5 KH mới/tuần tại Đắk Nông, đưa tổng KH mới Vùng lên 100+"
t4.rows[3].cells[4].text = "06/08/2026"

# Apply padding and formatting to all tables
for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        set_cell_background(cell, "003366")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx in range(1, len(table.rows)):
        for cell in table.rows[r_idx].cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if r_idx % 2 == 1:
                set_cell_background(cell, "F4F6F9")

doc.save(out_path_workspace)
doc.save(out_path_downloads)
print("Saved exact Week 31 report to:")
print(" - Workspace:", out_path_workspace)
print(" - Downloads:", out_path_downloads)
