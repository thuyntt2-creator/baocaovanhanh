import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_3.docx"

if not os.path.exists(doc_path):
    print("File không tồn tại")
    sys.exit(1)

doc = docx.Document(doc_path)

print("=== SO SÁNH KIỂM TRA SỐ HỌC FILE WORD V8_3 ===")
errors = 0

# 1. Kiểm tra Bảng 3 (Năng lực thiết bị)
table_3 = doc.tables[2]
print("\n1. Kiểm tra Bảng 3 (Số đầu xe):")
expected_table_3 = {
    'Nha Trang': ('5 xe', '9 xe', '10 xe'),
    'Di Linh': ('3 xe', '4 xe', '5 xe'),
    'Đơn Dương': ('2 xe', '2 xe', '3 xe'),
    'Đức Linh': ('2 xe', '2 xe')
}

# Nha Trang ở row 3, Di Linh ở row 6, Đơn Dương ở row 9, Đức Linh ở row 12
rows_map = {
    'Nha Trang': 3,
    'Di Linh': 6,
    'Đơn Dương': 9,
    'Đức Linh': 12
}

for name, row_idx in rows_map.items():
    actual_vals = []
    if name == 'Đức Linh':
        # Đức Linh chỉ có T7 và T10
        actual_vals = [table_3.rows[row_idx].cells[2].text.strip(), table_3.rows[row_idx].cells[3].text.strip()]
    else:
        actual_vals = [table_3.rows[row_idx].cells[2].text.strip(), table_3.rows[row_idx].cells[3].text.strip(), table_3.rows[row_idx].cells[4].text.strip()]
    
    expected = expected_table_3[name]
    if tuple(actual_vals) != expected:
        print(f"  [LỖI] {name} xe: Thực tế = {actual_vals} | Kỳ vọng = {expected} (LỆCH!)")
        errors += 1
    else:
        print(f"  [OK] {name} xe: Thực tế = {actual_vals} | Kỳ vọng = {expected} (Khớp)")

# 2. Kiểm tra Bảng 7 (Định biên nhân sự) Nha Trang
table_7 = doc.tables[6]
print("\n2. Kiểm tra Bảng 7 Nha Trang giao T7:")
nt_giao_t7 = table_7.rows[2].cells[1].text.strip()
if nt_giao_t7 != "38 (SL: 1815)":
    print(f"  [LỖI] Nha Trang giao T7 = {nt_giao_t7} (Kỳ vọng: 38 (SL: 1815)) (LỆCH!)")
    errors += 1
else:
    print(f"  [OK] Nha Trang giao T7 = {nt_giao_t7} (Khớp)")

# 3. Kiểm tra Bảng 10 (Hạng mục chi phí)
table_10 = doc.tables[9]
print("\n3. Kiểm tra Bảng 10 (Hạng mục chi phí) vs tổng thực tế:")
for c_idx in range(1, 7):
    xe = float(table_10.rows[1].cells[c_idx].text.replace(',', ''))
    giao = float(table_10.rows[2].cells[c_idx].text.replace(',', ''))
    kho = float(table_10.rows[3].cells[c_idx].text.replace(',', ''))
    mb = float(table_10.rows[4].cells[c_idx].text.replace(',', ''))
    
    total_in_doc = float(table_10.rows[5].cells[c_idx].text.replace(',', ''))
    actual_sum = xe + giao + kho + mb
    
    if abs(total_in_doc - actual_sum) > 0.1:
        print(f"  [LỖI] Tháng {c_idx+6}: Tổng dòng 5 = {total_in_doc} | Tổng cộng thực tế = {actual_sum:.2f} (LỆCH!)")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Tổng dòng 5 = {total_in_doc} | Tổng cộng thực tế = {actual_sum:.2f} (Khớp)")

# 4. Kiểm tra Bảng 11 (Chi tiết bưu cục)
table_11 = doc.tables[10]
print("\n4. Kiểm tra Bảng 11 (Chi tiết bưu cục) vs tổng thực tế:")
for c_idx in range(1, 7):
    nt = float(table_11.rows[1].cells[c_idx].text.replace(',', ''))
    dl = float(table_11.rows[2].cells[c_idx].text.replace(',', ''))
    dd = float(table_11.rows[3].cells[c_idx].text.replace(',', ''))
    dul = float(table_11.rows[4].cells[c_idx].text.replace(',', ''))
    
    total_in_doc = float(table_11.rows[5].cells[c_idx].text.replace(',', ''))
    actual_sum = nt + dl + dd + dul
    
    if abs(total_in_doc - actual_sum) > 0.1:
        print(f"  [LỖI] Tháng {c_idx+6}: Tổng dòng 5 = {total_in_doc} | Tổng cộng thực tế = {actual_sum:.2f} (LỆCH!)")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Tổng dòng 5 = {total_in_doc} | Tổng cộng thực tế = {actual_sum:.2f} (Khớp)")

# 5. So sánh Bảng 10 vs Bảng 11
print("\n5. Kiểm tra đối chiếu Bảng 10 vs Bảng 11:")
for c_idx in range(1, 7):
    t10_sum = float(table_10.rows[5].cells[c_idx].text.replace(',', ''))
    t11_sum = float(table_11.rows[5].cells[c_idx].text.replace(',', ''))
    
    if abs(t10_sum - t11_sum) > 0.1:
        print(f"  [LỖI] Tháng {c_idx+6}: Bảng 10 Tổng = {t10_sum} | Bảng 11 Tổng = {t11_sum} (LỆCH!)")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Bảng 10 Tổng = {t10_sum} | Bảng 11 Tổng = {t11_sum} (Khớp)")

if errors == 0:
    print("\n>>> TẤT CẢ KIỂM TRA ĐỀU THÀNH CÔNG! KHÔNG PHÁT HIỆN SAI LỆCH SỐ HỌC <<<")
else:
    print(f"\n>>> CÓ {errors} LỖI LỆCH SỐ LIỆU CẦN KHẮC PHỤC <<<")
