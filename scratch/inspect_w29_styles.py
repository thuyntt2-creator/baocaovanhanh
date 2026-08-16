import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

w29_path = r'C:\Users\lap4all\Downloads\BÁO CÁO KẾT QUẢ KINH DOANH TUẦN 29_2026 - NTB.docx'
doc = docx.Document(w29_path)

print('=== PARAGRAPH COUNT ===', len(doc.paragraphs))
print('=== TABLE COUNT ===', len(doc.tables))

for i, p in enumerate(doc.paragraphs):
    print(f"P{i:2d} (style={p.style.name}): '{p.text}'")
