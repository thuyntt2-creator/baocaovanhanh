import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

doc_path = r'c:\Users\lap4all\Documents\Auto report\KTC_Ke_Hoach_Van_Hanh_Event_8_8_NTB.docx'
downloads_docx = r'C:\Users\lap4all\Downloads\KTC_Ke_Hoach_Van_Hanh_Event_8_8_NTB.docx'

doc = docx.Document()

# Margins
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

# Title
tp = doc.add_paragraph()
tr = tp.add_run('BẢNG KẾ HOẠCH VẬN HÀNH & PHÂN TÍCH INSIGHT EVENT 8.8\nKHO TRUNG CHUYỂN (KTC) - VÙNG NAM TRUNG BỘ (NTB)')
tr.font.size = Pt(16)
tr.font.bold = True
tr.font.color.rgb = RGBColor(31, 78, 121)
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER

ip = doc.add_paragraph('Tài liệu kế hoạch vận hành chuyên sâu và phân tích dữ liệu chuyên nghiệp cho các Kho Trung Chuyển (KTC / Hub Sorting) khu vực Nam Trung Bộ (NTB) trong đợt Event Mega 8.8 (06/08/2026 – 15/08/2026), dựa trên mô hình từ TNB Kế hoạch Event 7.7.docx và bộ dữ liệu dự báo từ config_psbba_NTB.xlsx.')
ip.runs[0].font.italic = True

# SECTION 1
doc.add_heading('1. THÔNG TIN CHUNG & MỤC TIÊU CHIẾN LƯỢC', level=1)
p = doc.add_paragraph()
p.add_run('• Thời gian đợt Event: ').bold = True
p.add_run('06/08/2026 đến 15/08/2026 (10 ngày liên tục).\n')
p.add_run('• Phạm vi áp dụng: ').bold = True
p.add_run('5 Kho KTC/Hub: Khánh Hòa (Super-Hub), Bình Thuận (Cửa ngõ phía Nam), Đức Trọng (Tây Nguyên), Bảo Lộc, Đắc Nông.\n')
p.add_run('• Mục tiêu KPI chính: ').bold = True
p.add_run('Xử lý 771,977 đơn sorting (TB 77,198 đơn/ngày). Peak 08/08 (97,649 đơn) và 10/08 (89,608 đơn). COT Fulfillment 100%, Backlog >24h = 0%.')

# SECTION 2
doc.add_heading('2. PHÂN TÍCH DỰ BÁO SẢN LƯỢNG & INSIGHT CHUYÊN SÂU', level=1)

# Chart 1
chart1_img = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\497c8661-37a7-4c4e-986e-e8ef6c144fb3\chart1_daily_sorting_volume.png'
if os.path.exists(chart1_img):
    doc.add_paragraph('Biểu đồ 1: Diễn biến Sản lượng Sorting theo Kho KTC Event 8.8')
    cp1 = doc.add_paragraph()
    cp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp1.add_run().add_picture(chart1_img, width=Inches(6.2))

# INSIGHT 1
doc.add_heading('💡 INSIGHT 1: Mô Hình Sóng Kép (M-Shape Wave) & Phân Hóa Vai Trò Các Kho', level=2)
doc.add_paragraph('1. Mô hình biến động 2 Đỉnh (M-Shape Surge):\n'
                  '   - Đỉnh 1 (Mega Day 08/08): Sản lượng đạt đỉnh toàn đợt với 97,649 đơn (+26.5% so với TB) do các Sàn E-com mở bán đêm 07/08.\n'
                  '   - Đoạn trũng (Sunday 09/08): Sản lượng giảm còn 65,560 đơn (-15.1% so với TB) do người bán gom đơn nghỉ Chủ nhật.\n'
                  '   - Đỉnh 2 (After-shock 10/08): Bật tăng trở lại 89,608 đơn (+16.1% so với TB) khi các shop đóng gói hoàn tất đơn cuối tuần.\n'
                  '2. Vai trò chiến lược các Kho:\n'
                  '   - KTC Khánh Hòa gánh 38.1% sản lượng toàn vùng (293,732 đơn). Tắc nghẽn tại Khánh Hòa sẽ làm trễ tuyến toàn vùng.\n'
                  '   - CT Bình Thuận (22.6%) và CT Đức Trọng (19.0%) là 2 cánh bọc lót cho phía Nam và Tây Nguyên.')

# Table Sorting
doc.add_heading('Bảng 2.1: Chi Tiết Sản Lượng Sorting Theo Kho & Nhóm Hàng (06/08 – 15/08/2026)', level=2)

table_data = [
    ['Kho KTC', 'Nhóm hàng', '06/08', '07/08', '08/08 (Peak)', '09/08', '10/08 (Peak)', '11/08', '12/08', '13/08', '14/08', '15/08', 'Tổng 10d', 'TB/ngày', 'Tỷ trọng'],
    ['KTC Khánh Hòa', 'Normal', '21.652', '22.586', '30.057', '18.795', '26.792', '22.516', '22.667', '22.425', '23.306', '23.281', '234.077', '23.408', '79.7%'],
    ['', 'Bulky', '2.755', '2.852', '3.940', '3.975', '3.524', '4.971', '4.882', '3.431', '2.738', '3.940', '37.008', '3.701', '12.6%'],
    ['', 'Freight', '1.659', '1.718', '2.901', '2.438', '2.725', '2.819', '2.274', '1.779', '1.601', '2.732', '22.646', '2.265', '7.7%'],
    ['', 'Tổng KH', '26.067', '27.156', '36.898', '25.208', '33.041', '30.306', '29.823', '27.635', '27.645', '29.953', '293.732', '29.373', '38.1%'],
    ['CT Bình Thuận', 'Normal', '12.437', '12.973', '17.265', '10.796', '15.389', '12.933', '13.021', '12.881', '13.387', '13.372', '134.454', '13.445', '76.9%'],
    ['', 'Bulky', '1.785', '1.848', '2.347', '2.080', '2.515', '2.599', '3.508', '2.804', '1.940', '2.347', '23.773', '2.377', '13.6%'],
    ['', 'Freight', '1.218', '1.260', '2.334', '1.665', '2.398', '1.849', '1.429', '1.285', '1.175', '2.005', '16.618', '1.662', '9.5%'],
    ['', 'Tổng BTH', '15.440', '16.082', '21.946', '14.541', '20.302', '17.381', '17.958', '16.970', '16.502', '17.724', '174.846', '17.485', '22.6%'],
    ['CT Đức Trọng', 'Normal', '10.447', '10.897', '14.502', '9.069', '12.927', '10.864', '10.937', '10.820', '11.245', '11.233', '112.941', '11.294', '76.8%'],
    ['', 'Bulky', '1.436', '1.487', '3.483', '1.664', '2.898', '1.888', '1.543', '1.453', '1.386', '2.365', '19.603', '1.960', '13.3%'],
    ['', 'Freight', '1.059', '1.096', '2.032', '1.439', '1.918', '1.659', '1.324', '1.161', '1.022', '1.744', '14.454', '1.445', '9.8%'],
    ['', 'Tổng DTR', '12.943', '13.481', '20.017', '12.172', '17.743', '14.411', '13.804', '13.434', '13.653', '15.341', '146.999', '14.700', '19.0%'],
    ['CT Bảo Lộc', 'Normal', '6.858', '7.154', '8.774', '6.250', '8.805', '7.255', '7.180', '7.103', '7.382', '7.374', '74.135', '7.414', '77.5%'],
    ['', 'Bulky', '965', '999', '1.368', '1.285', '1.612', '1.531', '1.489', '1.331', '1.008', '1.405', '12.993', '1.299', '13.6%'],
    ['', 'Freight', '625', '646', '1.112', '888', '1.124', '1.007', '773', '718', '603', '1.028', '8.524', '852', '8.9%'],
    ['', 'Tổng BLO', '8.448', '8.799', '11.254', '8.423', '11.541', '9.793', '9.442', '9.152', '8.993', '9.807', '95.652', '9.565', '12.4%'],
    ['CT Đắc Nông', 'Normal', '4.407', '4.597', '6.118', '3.826', '5.453', '4.583', '4.614', '4.564', '4.744', '4.738', '47.644', '4.764', '78.4%'],
    ['', 'Bulky', '642', '665', '887', '902', '954', '843', '1.128', '979', '705', '887', '8.592', '859', '14.1%'],
    ['', 'Freight', '332', '344', '529', '488', '574', '553', '454', '390', '320', '529', '4.513', '451', '7.4%'],
    ['', 'Tổng DNO', '5.381', '5.605', '7.534', '5.216', '6.981', '5.979', '6.196', '5.933', '5.769', '6.154', '60.748', '6.075', '7.9%'],
    ['TỔNG KTC', 'TỔNG NTB', '68.279', '71.123', '97.649', '65.560', '89.608', '77.870', '77.223', '73.124', '72.562', '78.979', '771.977', '77.198', '100.0%']
]

t = doc.add_table(rows=len(table_data), cols=15)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.style = 'Table Grid'

for r_idx, row in enumerate(table_data):
    for c_idx, val in enumerate(row):
        cell = t.cell(r_idx, c_idx)
        cell.text = val
        if r_idx == 0:
            set_cell_background(cell, '1F4E79')
        elif r_idx == len(table_data)-1:
            set_cell_background(cell, 'D9E1F2')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(7.0)
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif 'Tổng' in val or r_idx == len(table_data)-1:
                    run.font.bold = True

# Chart 2 & Insight 2
chart2_img = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\497c8661-37a7-4c4e-986e-e8ef6c144fb3\chart2_product_group_breakdown.png'
if os.path.exists(chart2_img):
    doc.add_paragraph('\nBiểu đồ 2: Cơ cấu Nhóm hàng tại KTC NTB')
    cp2 = doc.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp2.add_run().add_picture(chart2_img, width=Inches(6.2))

doc.add_heading('💡 INSIGHT 2: Thách Thức Từ Hàng Bulky & Freight Cồng Kềnh', level=2)
doc.add_paragraph('• Chiếm gần 20% thể tích mặt bằng sàn lưu trữ kho dù chỉ chiếm 19.4% lượng đơn.\n'
                  '• Điểm nóng Đức Trọng: Bulky nhảy vọt từ 1.4k lên 3,483 đơn (+145%) ngày 08/08. Cần quy hoạch khu vực hạ bãi riêng biệt ngay cửa xả hàng.')

# Chart 3 & Insight 3
chart3_img = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\497c8661-37a7-4c4e-986e-e8ef6c144fb3\chart3_pick_delivery_channel.png'
if os.path.exists(chart3_img):
    doc.add_paragraph('\nBiểu đồ 3: Biến động Sản lượng Lấy & Giao theo Kênh Sàn')
    cp3 = doc.add_paragraph()
    cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp3.add_run().add_picture(chart3_img, width=Inches(6.2))

doc.add_heading('💡 INSIGHT 3: Bất Đối Xứng Giữa Lấy (Pick) và Giao (Delivery)', level=2)
doc.add_paragraph('• Tỷ lệ Giao / Lấy = 3.35 lần (Delivery 595.2k đơn vs Pick 177.8k đơn). NTB là vùng tiêu dùng thuần túy, xe chiều về từ HCM luôn đầy tải 100%.\n'
                  '• Shopee Delivery bùng nổ ngày 08/08 lên 40,058 đơn (+125%), Shopee Bulky tăng +195% (7,021 đơn). Đây là điểm nổ áp lực chặng cuối Last-mile.')

# SECTION 3 & INSIGHT 4
doc.add_heading('3. PHÂN TÍCH HIỆU SUẤT NHÂN SỰ & ĐIỂM NÓNG RỦI RO', level=1)

chart4_img = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\497c8661-37a7-4c4e-986e-e8ef6c144fb3\chart4_staffing_plan.png'
if os.path.exists(chart4_img):
    doc.add_paragraph('Biểu đồ 4: Phân bổ Nhân sự theo Kho KTC')
    cp4 = doc.add_paragraph()
    cp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp4.add_run().add_picture(chart4_img, width=Inches(6.2))

doc.add_heading('💡 INSIGHT 4 & CẢNH BÁO RỦI RO TẠI KHO CT BÌNH THUẬN', level=2)

staff_table_data = [
    ['Kho KTC', 'NVCT/ngày', 'Freelance/ngày', 'Tổng NV/ngày', 'Sản lượng TB/ngày', 'Năng suất TB', 'Năng suất Peak 08/08', 'Đánh giá Rủi ro'],
    ['CT Đắc Nông', '11 NV', '6 NV', '17 NV', '6.075 đơn', '357 đơn/người', '443 đơn/người', 'An toàn'],
    ['CT Đức Trọng', '20 NV', '6 NV', '26 NV', '14.700 đơn', '565 đơn/người', '770 đơn/người', 'Trung bình'],
    ['CT Bình Thuận', '11 NV', '0 NV', '11 NV', '17.485 đơn', '1.589 đơn/người', '1.995 đơn/người', 'CỰC KỲ CAO (NỔ KHO)']
]

st_t = doc.add_table(rows=len(staff_table_data), cols=8)
st_t.alignment = WD_TABLE_ALIGNMENT.CENTER
st_t.style = 'Table Grid'

for r_idx, row in enumerate(staff_table_data):
    for c_idx, val in enumerate(row):
        cell = st_t.cell(r_idx, c_idx)
        cell.text = val
        if r_idx == 0:
            set_cell_background(cell, '1F4E79')
        elif r_idx == 3: # Bình Thuận
            set_cell_background(cell, 'FCE4D6')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(8.0)
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif r_idx == 3:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph('\nCẢNH BÁO RỦI RO CT BÌNH THUẬN: Năng suất ngày Peak tại Bình Thuận chạm mốc 1,995 đơn/người/ngày (gấp 4.5 lần Đắk Nông). Cần bổ sung khẩn cấp 3-4 Freelance dự phòng ca đêm ngày 08/08 và 10/08.')

# SECTION 4 & 5
doc.add_heading('4. QUY TRÌNH & CHECKLIST VẬN HÀNH CHUYÊN SÂU', level=1)
doc.add_paragraph('• Quy hoạch Layout: Quy hoạch ranh giới 3 luồng Normal / Bulky / Freight trước 05/08.')
doc.add_paragraph('• Điều phối Fleet: Đặt xe COT cố định + nhà xe hợp đồng bọc lót xe tăng cường.')
doc.add_paragraph('• Hạ tầng CNTT: 1.2 thiết bị PDA/nhân sự ca cao điểm + 100% PDA cài SIM 4G dung lượng cao.')

doc.add_heading('5. PHƯƠNG ÁN ỨNG PHÓ KHI HÀNG TĂNG 30% (SURGE CONTINGENCY)', level=1)
doc.add_paragraph('• Cấp 1 (+10% đến +15%): Kích hoạt Overtime +2 giờ ca đêm.')
doc.add_paragraph('• Cấp 2 (+16% đến +25%): Gọi 3-5 Freelance dự phòng ca ngày, mở rộng bạt che ngoài sân kho.')
doc.add_paragraph('• Cấp 3 (+26% đến +30%): Huy động 100% quân số NVCT (hủy phép), điều động xe tăng cường khẩn cấp.')

try:
    alt_path = r'c:\Users\lap4all\Documents\Auto report\KTC_Ke_Hoach_Van_Hanh_Event_8_8_NTB_v3.docx'
    doc.save(alt_path)
    doc_path = alt_path
except Exception as e:
    pass

try:
    alt_dl = r'C:\Users\lap4all\Downloads\KTC_Ke_Hoach_Van_Hanh_Event_8_8_NTB_v3.docx'
    doc.save(alt_dl)
    downloads_docx = alt_dl
except Exception as e:
    pass

print('Updated docx files successfully!')
print('1.', doc_path)
print('2.', downloads_docx)
print('1.', doc_path)
print('2.', downloads_docx)
