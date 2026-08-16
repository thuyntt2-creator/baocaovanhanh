import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math
import os
from PIL import Image, ImageDraw, ImageFont
import sys

sys.stdout.reconfigure(encoding='utf-8')

artifact_dir = r"C:\Users\lap4all\.gemini\antigravity-ide\brain\5935f00b-cd19-4de8-83bb-03170a8c8606"
web_dir = os.path.join(artifact_dir, "web_maps")
custom_dir = os.path.join(web_dir, "custom_drawn")
os.makedirs(custom_dir, exist_ok=True)

docx_out = r"C:\Users\lap4all\Downloads\Ban_Do_Quy_Hoach_NTB_Perfect_v8.docx"
pdf_out = r"C:\Users\lap4all\Downloads\Ban_Do_Quy_Hoach_Mang_Luoi_NTB_2026.pdf"

don_duong_web_drawn_img = os.path.join(artifact_dir, "don_duong_official_web_drawn_map.png")
da_lat_infographic_img = os.path.join(artifact_dir, "da_lat_official_map_infographic.png")
di_linh_web_drawn_img = os.path.join(artifact_dir, "di_linh_official_web_drawn_map.png")

def draw_star(draw, cx, cy, radius, fill_color, outline_color=None):
    points = []
    for i in range(10):
        r = radius if i % 2 == 0 else radius * 0.45
        angle = i * math.pi / 5 - math.pi / 2
        x = cx + r * math.cos(angle)
        y = cy + r * math.sin(angle)
        points.append((x, y))
    draw.polygon(points, fill=fill_color, outline=outline_color)

def decorate_map_clean(src_path, dst_name, has_new=True, has_move=True):
    if not os.path.exists(src_path):
        return src_path
    dst_path = os.path.join(custom_dir, dst_name)
    im = Image.open(src_path).convert("RGBA")
    w, h = im.size
    overlay = Image.new("RGBA", (w, h), (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)
    
    try:
        font_leg = ImageFont.truetype("arial.ttf", max(16, int(h * 0.024)))
    except:
        font_leg = ImageFont.load_default()
        
    items = []
    if has_new:
        items.append(((30, 64, 175, 255), "Bưu cục tách mới / mở mới"))
    if has_move:
        items.append(((220, 38, 38, 255), "Bưu cục hiện hữu / di dời"))
        
    box_w = int(w * 0.32)
    box_h = int(h * 0.04) + len(items) * int(h * 0.048)
    box_x1 = w - int(w * 0.03) - box_w
    box_y1 = int(h * 0.03)
    box_x2 = w - int(w * 0.03)
    box_y2 = box_y1 + box_h
    
    draw.rectangle([box_x1, box_y1, box_x2, box_y2], fill=(255, 255, 255, 242), outline=(160, 160, 160, 255), width=2)
    
    y_c = box_y1 + int(h * 0.022)
    star_r = int(h * 0.014)
    
    for color, txt in items:
        star_cx = box_x1 + int(w * 0.028)
        star_cy = y_c + int(h * 0.012)
        draw_star(draw, star_cx, star_cy, star_r, fill_color=color, outline_color=(255, 255, 255, 255))
        draw.text((box_x1 + int(w * 0.055), y_c), txt, fill=(15, 23, 42, 255), font=font_leg)
        y_c += int(h * 0.048)
        
    out_im = Image.alpha_composite(im, overlay)
    out_im.convert("RGB").save(dst_path)
    return dst_path

# Create doc
doc = docx.Document()
for sec in doc.sections:
    sec.top_margin = Inches(0.4)
    sec.bottom_margin = Inches(0.4)
    sec.left_margin = Inches(0.4)
    sec.right_margin = Inches(0.4)

# Document Header
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("BẢN ĐỒ QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NAM TRUNG BỘ (NTB) 2026")
r_title.bold = True
r_title.font.size = Pt(16)
r_title.font.color.rgb = RGBColor(0, 51, 102)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run("BẢN ĐỒ PHÂN VÙNG LÃNH THỔ & ĐỒ HỌA INFOGRAPHIC ĐỒNG BỘ NỀN WEB ANH THIÊN (quyhoachbuucuc.info)")
r_sub.bold = True
r_sub.font.size = Pt(11)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

sections = [
    {
        "province": "I. TỈNH KHÁNH HÒA",
        "items": [
            {
                "title": "1.1. Bản đồ Kế hoạch Tách Bưu cục Cam Linh & Bắc Cam Ranh",
                "img": os.path.join(artifact_dir, "media__1785907585318.png"),
                "desc": "Khu Vực Cam Ranh & Nam Cam Ranh\n- Mở Mới BC Nam Cam Ranh (tách ra từ kho Cam Linh) cover 6 xã/phường mới: Ba Ngòi, Cam Bình, Cam Lập, Cam Phước Đông, Cam Thịnh Đông, Cam Thịnh Tây (Vol: 600 đơn/ngày, 10 NV).\n- Thu gọn BC Cam Linh cover 6 phường trung tâm (Cam Linh, Cam Thuận, Cam Lộc, Cam Phú, Cam Phúc Nam, Cam Lợi).\n- Di dời BC Bắc Cam Ranh kho cũ 100m² quá chật hẹp cover 3 xã/phường Cam Thành Nam, Cam Nghĩa, Cam Phúc Bắc."
            },
            {
                "title": "1.2. Bản đồ Quy hoạch Phân vùng Cụm TP. Nha Trang (BC Nha Trang & BC Tây Nha Trang)",
                "img": decorate_map_clean(os.path.join(web_dir, "map_whatif_nha_trang.png"), "clean_nha_trang.png"),
                "desc": "Khu Vực TP. Nha Trang\n- Di dời Bưu cục (KHO) Nha Trang sang mặt bằng mới đáp ứng Vol lớn 1.592 đơn/ngày, cover 5 phường sáp nhập (Vạn Thạnh, Lộc Thọ, Tân Tiến, Phước Hòa, Vĩnh Nguyên).\n- Đóng cửa BC Nam Nha Trang 2 dồn sản lượng về BC KHO Nha Trang.\n- Giữ lại BC Tây Nha Trang cover 6 xã/phường (Phương Sài, Ngọc Hiệp, Vĩnh Hiệp, Vĩnh Ngọc, Vĩnh Thạnh, Vĩnh Trung)."
            },
            {
                "title": "1.3. Bản đồ Quy hoạch Cụm Ninh Hòa 2 (Phường Ninh Hòa mới)",
                "img": decorate_map_clean(os.path.join(web_dir, "map_whatif_ninh_hoa.png"), "clean_ninh_hoa.png"),
                "desc": "Khu Vực Huyện Ninh Hòa\n- Dồn toàn bộ sản lượng Phường Ninh Hòa mới (Ninh Đa, Ninh Phụng, Ninh Đông, Ninh Hiệp) về kho trung tâm Bưu cục Ninh Hòa 2."
            }
        ]
    },
    {
        "province": "II. TỈNH LÂM ĐỒNG",
        "items": [
            {
                "title": "2.1. Bản đồ Phân vùng Tách Bưu cục Lạc Xuân - Huyện Đơn Dương (Trích xuất Web quyhoachbuucuc.info)",
                "img": don_duong_web_drawn_img,
                "desc": "Khu Vực Huyện Đơn Dương\n- Mở Mới BC Lạc Xuân (tách ra từ Bưu cục gốc Nghĩa Đức) cover 4 xã phía Đông Bắc: Xã Lạc Lâm, Xã Lạc Xuân, TT. D'Ran, Xã Ka Đô (Vol: 400 - 480 đơn/ngày, 7 NV).\n- Thu gọn BC gốc Nghĩa Đức cover 6 xã/thị trấn trung tâm: Đạ Ròn, Thạnh Mỹ, Tu Tra, Ka Đơn, Quảng Lập, Pró (Vol: 600 - 720 đơn/ngày, 9 NV)."
            },
            {
                "title": "2.2. Infographic Quy hoạch Mạng lưới Bưu cục TP. Đà Lạt - Đề xuất Mở mới BC (LDO) Xuân Hương 2",
                "img": da_lat_infographic_img,
                "desc": "Khu Vực TP. Đà Lạt\n- Mở Mới BC (LDO) Xuân Hương 2 (đặt tại Phường 10, TP. Đà Lạt) cover Phường 10 (400 đơn) & Phường 3 (500 đơn) -> Tổng giao: 900 đơn, Lấy: 150 đơn, Định biên: 8 NVPTTT + 1 NVXL.\n- Giữ nguyên BC Xuân Hương cũ (phụ trách Phường 1, Phường 2, Phường 4) -> Tổng giao: 1.400 đơn, Lấy: 400 đơn, Định biên: 13 NVPTTT + 1 NVXL."
            },
            {
                "title": "2.3. Bản đồ Quy hoạch Tách Bưu cục Hàng Nhỏ / Hàng Vừa (Xã Đinh Trang Thượng - Di Linh)",
                "img": di_linh_web_drawn_img,
                "desc": "Khu Vực Huyện Di Linh\n- Tách Bưu cục Hàng Nhỏ / Hàng Vừa (đặt tại Xã Đinh Trang Thượng) để chia tải cho BC (LDO) Di Linh do bán kính di chuyển hiện tại quá xa (22km - 45km).\n- BC Hàng Nhỏ: 300 - 350 đơn giao, 10 - 20 đơn lấy, 4 NVPTTT cover Đinh Trang Thượng & Di Linh.\n- BC Hàng Vừa: 4 NVPTTT cover Đinh Trang Thượng, Phúc Thọ Lâm Hà, Di Linh, Liên Đầm.\n- Trả tuyến Liên Đầm về đúng địa giới bưu cục quản lý."
            },
            {
                "title": "2.4. Bản đồ Quy hoạch Mở mới Bưu cục B'Lao Mới - TP. Bảo Lộc",
                "img": decorate_map_clean(os.path.join(web_dir, "map_whatif_bao_loc.png"), "clean_bao_loc.png"),
                "desc": "Khu Vực TP. Bảo Lộc\n- Đóng cửa BC 1 Bảo Lộc cũ do mặt bằng xuống cấp chật hẹp.\n- Mở Mới BC B'Lao Mới đặt tại vị trí trung tâm hơn cover Phường 1, Phường B'Lao, Phường Lộc Phát (Vol: 2.500 đơn/ngày, 17 NV)."
            }
        ]
    },
    {
        "province": "III. TỈNH NINH THUẬN",
        "items": [
            {
                "title": "3.1. Bản đồ Quy hoạch Mở mới BC Đông Hải & Di dời BC Ninh Chử - TP. Phan Rang",
                "img": decorate_map_clean(os.path.join(web_dir, "map_whatif_phan_rang.png"), "clean_phan_rang.png"),
                "desc": "Khu Vực TP. Phan Rang - Tháp Chàm & Ninh Hải\n- Mở Mới BC Đông Hải đặt tại Phường Đông Hải cover 4 phường ven biển: Đông Hải, Mỹ Bình, Mỹ Đông, Mỹ Hải (Vol: 850 đơn/ngày, 8 NV).\n- Di dời BC Ninh Chử về trung tâm Xã Ninh Hải mới cover 4 xã/thị trấn: Bắc Sơn, Phương Hải, Tri Hải, TT. Khánh Hải."
            }
        ]
    },
    {
        "province": "IV. TỈNH BÌNH THUẬN",
        "items": [
            {
                "title": "4.1. Bản đồ Quy hoạch Mở mới Bưu cục Nam Thành (Tánh Linh / Đức Linh)",
                "img": decorate_map_clean(os.path.join(web_dir, "map_whatif_nam_thanh.png"), "clean_nam_thanh.png"),
                "desc": "Khu Vực Tánh Linh & Đức Linh\n- Mở Mới BC Nam Thành đặt tại Xã Nam Thành cover 2 xã vùng ven: Xã Nam Thành & Xã Nghị Đức để giảm bán kính di chuyển đồi dốc và giảm tải kho Đồng Kho (Vol: 450 đơn/ngày, 8 NV)."
            },
            {
                "title": "4.2. Bản đồ Quy hoạch TP. Phan Thiết - BC Hàm Thắng",
                "img": decorate_map_clean(os.path.join(web_dir, "map_whatif_phan_thiet.png"), "clean_phan_thiet.png"),
                "desc": "Khu Vực TP. Phan Thiết\n- Dồn toàn bộ sản lượng Phường Phan Thiết mới (Bình Hưng, Lạc Đạo, Phú Trinh, Phú Tài, Hàm Hiệp, Phong Nẫm) về BC Hàm Thắng."
            }
        ]
    },
    {
        "province": "V. TỈNH ĐẮK NÔNG",
        "items": [
            {
                "title": "5.1. Bản đồ Quy hoạch Phân ranh giới TP. Gia Nghĩa (Đông Gia Nghĩa & Bắc Gia Nghĩa)",
                "img": decorate_map_clean(os.path.join(web_dir, "map_whatif_gia_nghia.png"), "clean_gia_nghia.png"),
                "desc": "Khu Vực TP. Gia Nghĩa\n- Phân ranh giới quản lý TP. Gia Nghĩa theo 2 bưu cục: BC Đông Gia Nghĩa và BC Bắc Gia Nghĩa cover 2 phường mới (Bắc Gia Nghĩa & Nam Gia Nghĩa)."
            }
        ]
    }
]

for sec in sections:
    p_prov = doc.add_paragraph()
    r_prov = p_prov.add_run(sec["province"])
    r_prov.bold = True
    r_prov.font.size = Pt(13)
    r_prov.font.color.rgb = RGBColor(0, 51, 102)
    
    for item in sec["items"]:
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(item["title"])
        r_t.bold = True
        r_t.font.size = Pt(11)
        
        # Image
        if os.path.exists(item["img"]):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(item["img"], width=Inches(7.0))
            
        # Description matching Telegram text style
        p_d = doc.add_paragraph()
        r_d = p_d.add_run(item["desc"])
        r_d.font.size = Pt(10)
        r_d.font.color.rgb = RGBColor(30, 41, 59)
        
        doc.add_paragraph()

doc.save(docx_out)
print(f"Successfully created Final Perfect DOCX v8 at: {docx_out}")

# Convert to PDF
import win32com.client
word = win32com.client.Dispatch('Word.Application')
word.Visible = False
doc_com = word.Documents.Open(docx_out)
doc_com.SaveAs(pdf_out, FileFormat=17)
doc_com.Close()
word.Quit()
print(f"Successfully generated PERFECT PDF V8 at: {pdf_out}")
