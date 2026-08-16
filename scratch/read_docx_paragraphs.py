import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_3.docx"

if not os.path.exists(doc_path):
    print("File không tồn tại")
    sys.exit(1)

doc = docx.Document(doc_path)

print("=== ĐỌC CÁC ĐOẠN VĂN BẢN TRONG FILE WORD ===")
for p_idx, p in enumerate(doc.paragraphs[:100]):
    text = p.text.strip()
    if text:
        # In ra các đoạn có chứa các từ khóa cần tìm
        keywords = ['nhân sự', 'xe', 'tiết kiệm', 'tỷ', 'chi phí', 'triệu', 'quy hoạch', 'tóm tắt']
        if any(k in text.lower() for k in keywords) or len(text) < 100:
            print(f"Paragraph {p_idx:03d}: {text}")
            
# Kiểm tra xem có bảng nào chứa văn bản này không (nếu nằm trong bảng tóm tắt)
for t_idx, table in enumerate(doc.tables[:3]):
    print(f"\n--- BẢNG {t_idx+1} ---")
    for r_idx, row in enumerate(table.rows):
        row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        print(f"Row {r_idx:02d}: {row_text}")
