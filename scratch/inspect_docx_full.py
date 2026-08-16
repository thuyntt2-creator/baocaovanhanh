import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document(r'C:\Users\lap4all\Downloads\[TEMPLATE] BÁO CÁO KẾT QUẢ KINH DOANH [VÙNG] - [TUẦN].docx')

print('=== PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    print(f'P{i:2d}: "{p.text}"')

print('\n=== TABLES ===')
for t_idx, table in enumerate(doc.tables):
    print(f'=== TABLE {t_idx} ({len(table.rows)} rows x {len(table.columns)} cols) ===')
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            print(f'  Table {t_idx} R{r_idx} C{c_idx}: "{cell.text.strip()}"')
