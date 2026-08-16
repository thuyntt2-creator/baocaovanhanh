import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

web_dir = r"C:\Users\lap4all\.gemini\antigravity-ide\brain\5935f00b-cd19-4de8-83bb-03170a8c8606\web_maps"
artifact_dir = r"C:\Users\lap4all\.gemini\antigravity-ide\brain\5935f00b-cd19-4de8-83bb-03170a8c8606"

docx_out = r"C:\Users\lap4all\Downloads\Ban_Do_Quy_Hoach_NTB_Official_Drawn_Maps.docx"
pdf_out = r"C:\Users\lap4all\Downloads\Ban_Do_Quy_Hoach_Mang_Luoi_NTB_2026.pdf"

doc = docx.Document()

# Page Margins
for sec in doc.sections:
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

# Document Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("BẢN ĐỒ QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NAM TRUNG BỘ (NTB)")
r_title.bold = True
r_title.font.size = Pt(16)
r_title.font.color.rgb = RGBColor(0, 51, 102)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run("BẢN ĐỒ PHÂN VÙNG LÃNH THỔ & THẺ QUY HOẠCH (WHAT-IF) TRÍCH XUẤT TỪ quyhoachbuucuc.info")
r_sub.bold = True
r_sub.font.size = Pt(11)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

maps_data = [
    {
        "province": "1. KHU VỰC TỈNH KHÁNH HÒA",
        "items": [
            ("1.1. Bản đồ & Thẻ Quy hoạch (What-If) Phường Nha Trang", 
             "Bản đồ vẽ lãnh thổ phân vùng và Thẻ Quy hoạch What-If dồn 5 phường sáp nhập về Bưu cục (KHO) Nha Trang làm đầu mối chính.", 
             os.path.join(web_dir, "map_whatif_nha_trang.png")),
            ("1.2. Bản đồ & Thẻ Quy hoạch (What-If) Phường Tây Nha Trang", 
             "Bản đồ vẽ phân vùng gộp 6 xã/phường sáp nhập về Bưu cục (KHO) Tây Nha Trang.", 
             os.path.join(web_dir, "map_whatif_tay_nha_trang.png")),
            ("1.3. Bản đồ & Thẻ Quy hoạch (What-If) Cụm Cam Linh & Nam Cam Ranh", 
             "Sơ đồ tách Bưu cục Nam Cam Ranh từ BC Cam Linh và di dời BC Bắc Cam Ranh.", 
             os.path.join(web_dir, "map_whatif_cam_linh.png")),
            ("1.4. Bản đồ & Thẻ Quy hoạch (What-If) Cụm Ninh Hòa 2", 
             "Sơ đồ gộp tuyến Phường Ninh Hòa mới về Bưu cục Ninh Hòa 2.", 
             os.path.join(web_dir, "map_whatif_ninh_hoa.png"))
        ]
    },
    {
        "province": "2. KHU VỰC TỈNH LÂM ĐỒNG",
        "items": [
            ("2.1. Bản đồ & Thẻ Quy hoạch (What-If) Cụm Đơn Dương (BC Lạc Xuân Mở Mới)", 
             "Sơ đồ đề xuất mở mới Bưu cục Lạc Xuân (phụ trách 4 xã Đông Bắc) và Bưu cục gốc Nghĩa Đức (phụ trách 6 xã trung tâm).", 
             os.path.join(web_dir, "map_whatif_don_duong.png")),
            ("2.2. Bản đồ & Thẻ Quy hoạch (What-If) TP. Đà Lạt (BC Xuân Hương 2 Mở Mới)", 
             "Bản đồ phân vùng chia tải TP. Đà Lạt mở mới BC Xuân Hương 2.", 
             os.path.join(web_dir, "map_whatif_da_lat.png")),
            ("2.3. Bản đồ & Thẻ Quy hoạch (What-If) TP. Bảo Lộc (BC B'Lao Mới)", 
             "Bản đồ mở mới BC B'Lao Mới thay thế BC 1 Bảo Lộc cũ.", 
             os.path.join(web_dir, "map_whatif_bao_loc.png"))
        ]
    },
    {
        "province": "3. KHU VỰC TỈNH NINH THUẬN",
        "items": [
            ("3.1. Bản đồ & Thẻ Quy hoạch (What-If) TP. Phan Rang - BC Đông Hải", 
             "Bản đồ phân vùng mở mới BC Đông Hải ven biển Phan Rang, di dời BC Ninh Chử và gộp tuyến BC Phan Rang.", 
             os.path.join(web_dir, "map_whatif_phan_rang.png"))
        ]
    },
    {
        "province": "4. KHU VỰC TỈNH BÌNH THUẬN",
        "items": [
            ("4.1. Bản đồ & Thẻ Quy hoạch (What-If) TP. Phan Thiết - BC Hàm Thắng", 
             "Bản đồ gộp tuyến Phường Phan Thiết mới về Bưu cục Hàm Thắng.", 
             os.path.join(web_dir, "map_whatif_phan_thiet.png")),
            ("4.2. Bản đồ & Thẻ Quy hoạch (What-If) Cụm Nam Thành (Tánh Linh / Đức Linh)", 
             "Bản đồ mở mới Bưu cục Nam Thành cover khu vực Nam Thành và Nghị Đức.", 
             os.path.join(web_dir, "map_whatif_nam_thanh.png"))
        ]
    },
    {
        "province": "5. KHU VỰC TỈNH ĐẮK NÔNG",
        "items": [
            ("5.1. Bản đồ & Thẻ Quy hoạch (What-If) TP. Gia Nghĩa (Đắk Nông)", 
             "Bản đồ quy hoạch phân chia ranh giới quản lý 2 bưu cục Đông Gia Nghĩa và Bắc Gia Nghĩa.", 
             os.path.join(web_dir, "map_whatif_gia_nghia.png"))
        ]
    }
]

for section in maps_data:
    p_prov = doc.add_paragraph()
    r_prov = p_prov.add_run(section["province"])
    r_prov.bold = True
    r_prov.font.size = Pt(13)
    r_prov.font.color.rgb = RGBColor(0, 51, 102)
    
    for title, desc, img_path in section["items"]:
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(title)
        r_t.bold = True
        r_t.font.size = Pt(11)
        
        p_d = doc.add_paragraph()
        r_d = p_d.add_run("Mô tả quy hoạch: " + desc)
        r_d.font.size = Pt(10)
        r_d.font.color.rgb = RGBColor(80, 80, 80)
        
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(img_path, width=Inches(7.0))
            print(f"Added picture: {os.path.basename(img_path)}")
        else:
            print(f"Image NOT found: {img_path}")
            
        doc.add_paragraph()

doc.save(docx_out)
print(f"Successfully created Official Drawn Maps DOCX at: {docx_out}")

# Convert to PDF via Word COM
import win32com.client
word = win32com.client.Dispatch('Word.Application')
word.Visible = False
doc_com = word.Documents.Open(docx_out)
doc_com.SaveAs(pdf_out, FileFormat=17)
doc_com.Close()
word.Quit()
print(f"Successfully generated PERFECT DRAWN MAPS PDF at: {pdf_out}")
