import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

src_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_5.docx"
dst_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_6.docx"

if not os.path.exists(src_path):
    print(f"File nguồn không tồn tại: {src_path}")
    sys.exit(1)

doc = docx.Document(src_path)

print("=== BẮT ĐẦU CẬP NHẬT ĐỢT 6 ===")

# Sửa Paragraph 133
found_mb = False
for p in doc.paragraphs:
    if "150.000 đ/m²/tháng" in p.text:
        p.text = p.text.replace("150.000 đ/m²/tháng", "120.000 đ/m²/tháng")
        print("-> Đã sửa đơn giá thuê mặt bằng từ 150k thành 120k ở Paragraph 133")
        found_mb = True
        break

if not found_mb:
    for p in doc.paragraphs:
        if "150.000" in p.text and "mặt bằng" in p.text:
            p.text = p.text.replace("150.000", "120.000")
            print("-> Đã sửa đơn giá thuê mặt bằng (tìm mờ) từ 150k thành 120k")
            break

doc.save(dst_path)
print(f"=== ĐÃ LƯU FILE THÀNH CÔNG: {dst_path} ===")
