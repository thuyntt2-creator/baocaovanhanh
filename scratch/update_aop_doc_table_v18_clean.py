import win32com.client
import docx
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

# 1. Fetch exact values from V18 using COM
excel = win32com.client.Dispatch('Excel.Application')
excel.Visible = False
wb = excel.Workbooks.Open(r'C:\Users\lap4all\Downloads\AOP_Hang_Nang_MAU_Vung_T7-T12_2026_final_v18.xlsx')
sheet = wb.Sheets('Nguồn lực & chi phí')

months_cols = {
    'T7': 2, 'T8': 3, 'T9': 4, 'T10': 5, 'T11': 6, 'T12': 7
}

# We need rows:
# Row 13 (Mặt bằng m²? No, row 15: Chi phí thuê mặt bằng)
# Row 16 (Chi phí xe)
# Row 17 (Chi phí NV xử lý kho)
# Row 18 (Chi phí NV giao hàng)
# Row 19 (Chi phí NV quản lý)
# Row 20 (Tổng chi phí)

cost_data = {}
for m, col in months_cols.items():
    rent = sheet.Cells(15, col).Value / 1e6
    xe = sheet.Cells(16, col).Value / 1e6
    kho = sheet.Cells(17, col).Value / 1e6
    giao = sheet.Cells(18, col).Value / 1e6
    ql = sheet.Cells(19, col).Value / 1e6
    tong = sheet.Cells(20, col).Value / 1e6
    cost_data[m] = {
        'rent': rent, 'xe': xe, 'kho': kho, 'giao': giao, 'ql': ql, 'tong': tong
    }

wb.Close(False)
excel.Quit()

print("Excel V18 values loaded:", cost_data)

# 2. Update the Word document
in_path = r'C:\Users\lap4all\Downloads\AOP_BCCK_Plan_new_final_v4.docx'
out_path = r'C:\Users\lap4all\Downloads\AOP_BCCK_Plan_new_final_v5.docx'

doc = docx.Document(in_path)

# Update paragraphs
new_6_1 = """Nhân sự được định biên linh hoạt theo sản lượng dự báo (Forecast) thực tế từng tháng của từng trạm, không cào bằng định mức tĩnh như phương án cũ.
Định biên nhân sự cho 4 BCCK (Tổng T7 là 31 người, tăng dần lên 48 người vào cao điểm T12):
- Nhân sự xử lý kho: 4 người (cố định 1 người/BCCK).
- Nhân sự quản lý & backup: 4 người (cố định 1 người/BCCK).
- Nhân sự giao hàng (shipper): 23 người (T7), tăng lên 34 người (T10) và 40 người (T12) dựa trên sản lượng Peak thực tế từng trạm.
Công thức tính nhân sự: Nhân sự 1 BCCK = 3 NV Xử lý & Quản lý + ROUNDUP(Sản lượng dự báo ngày / Năng suất định mức 1 shipper)."""

new_7_1 = """Để đảm bảo ngân sách Vùng phản ánh chính xác nguồn lực thực tế của 4 BCCK, Cấu trúc Chi phí Topline được thiết kế như sau:
- Chi phí thuê mặt bằng: diện tích m² thực tế × đơn giá 150.000 đ/m²/tháng (riêng Đức Linh tính theo giá thuê thực tế).
- Chi phí xe tải 1.9T: Xe bình quân/ngày của 4 BCCK (T7 là 21 xe) × đơn giá xe 1.200.000 đ/xe/ngày × số ngày làm việc.
- Chi phí nhân sự (được bóc tách rõ làm 3 dòng):
  + Chi phí NV xử lý kho (4 người): 60.000.000 đ/tháng.
  + Chi phí NV giao hàng (23-40 người tùy tháng): 345 - 600 Triệu/tháng.
  + Chi phí NV quản lý & backup (4 người): 60.000.000 đ/tháng."""

new_7_2 = f"""Tổng chi phí thực tế (không tính lặp tiền xe) của 4 BCCK theo tháng:
- T7: {cost_data['T7']['tong']*1e6:,.0f} VNĐ
- T8: {cost_data['T8']['tong']*1e6:,.0f} VNĐ
- T9: {cost_data['T9']['tong']*1e6:,.0f} VNĐ
- T10: {cost_data['T10']['tong']*1e6:,.0f} VNĐ
- T11: {cost_data['T11']['tong']*1e6:,.0f} VNĐ
- T12: {cost_data['T12']['tong']*1e6:,.0f} VNĐ
* Chi tiết tính toán tự động ăn theo công thức được trình bày ở sheet "Nguồn lực & chi phí" trong file Excel AOP_Hang_Nang_MAU_Vung_T7-T12_2026_final_v18.xlsx."""

for para in doc.paragraphs:
    if "Nhân sự được định biên linh hoạt theo sản lượng dự báo" in para.text or "Định biên nhân sự cho 4 BCCK (Tổng T7 là 31 người" in para.text:
        para.text = new_6_1
    if "Để đảm bảo ngân sách Vùng phản ánh chính xác nguồn lực thực tế của 4 BCCK" in para.text:
        para.text = new_7_1
    if "Tổng chi phí thực tế (không tính lặp tiền xe) của 4 BCCK" in para.text or "Tổng chi phí tính toán sát với sản lượng T7-T12/2026:" in para.text:
        para.text = new_7_2

# Update Cost Table (it has row headers like 'Chi phí xe', 'Chi phí xe tải', 'Chi phí NV giao hàng', etc. and 7 columns)
for table in doc.tables:
    first_col_text = [row.cells[0].text.strip() for row in table.rows]
    # Check if this is the cost table
    if 'Chi phí xe tải 1.9T' in first_col_text or 'Chi phí thuê xe tải 1.9T' in first_col_text:
        print("Found Cost Table. Overwriting values...")
        for row in table.rows:
            header = row.cells[0].text.strip()
            # map headers
            if 'Chi phí xe tải' in header or 'Chi phí thuê xe tải' in header:
                for idx, m in enumerate(['T7', 'T8', 'T9', 'T10', 'T11', 'T12']):
                    row.cells[idx+1].text = f"{cost_data[m]['xe']:.1f}"
            elif 'Chi phí NV giao hàng' in header:
                for idx, m in enumerate(['T7', 'T8', 'T9', 'T10', 'T11', 'T12']):
                    row.cells[idx+1].text = f"{cost_data[m]['giao']:.1f}"
            elif 'Chi phí NV kho' in header or 'Chi phí NV xử lý kho' in header:
                for idx, m in enumerate(['T7', 'T8', 'T9', 'T10', 'T11', 'T12']):
                    row.cells[idx+1].text = f"{cost_data[m]['kho']:.1f}"
            elif 'Chi phí NV quản lý' in header or 'Chi phí NV điều phối' in header:
                for idx, m in enumerate(['T7', 'T8', 'T9', 'T10', 'T11', 'T12']):
                    row.cells[idx+1].text = f"{cost_data[m]['ql']:.1f}"
            elif 'Chi phí thuê mặt bằng' in header:
                for idx, m in enumerate(['T7', 'T8', 'T9', 'T10', 'T11', 'T12']):
                    row.cells[idx+1].text = f"{cost_data[m]['rent']:.1f}"
            elif 'TỔNG CHI PHÍ' in header:
                for idx, m in enumerate(['T7', 'T8', 'T9', 'T10', 'T11', 'T12']):
                    row.cells[idx+1].text = f"{cost_data[m]['tong']:.1f}"

doc.save(out_path)
print(f"Word doc successfully updated to V5: {out_path}")
