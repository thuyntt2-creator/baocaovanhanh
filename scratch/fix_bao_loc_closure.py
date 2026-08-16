import docx, sys

sys.stdout.reconfigure(encoding='utf-8')

# Open document
src_path = r'C:\Users\lap4all\Downloads\Quy_Hoach_MANG_LUOI_NTB_Co_Nha_Trang_Final.docx'
doc = docx.Document(src_path)

sec5_p_idx = -1
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if 'V. TỔNG HỢP BIẾN ĐỘNG' in txt or 'V.TỔNG HỢP BIẾN ĐỘNG' in txt:
        sec5_p_idx = i

# Update Section V to remove BC 1 Bảo Lộc from closure list
for p in doc.paragraphs[sec5_p_idx:]:
    txt = p.text.strip()
    if 'Bưu cục Đóng cửa' in txt:
        p.text = "❖ Bưu cục Đóng cửa gộp kho (03 BC): BC Diên Khánh 1 (Khánh Hòa - khi BC Tây Nha Trang có MB mới), BC Nam Nha Trang 2 & BC Nam Nha Trang 3 (Khánh Hòa)."

doc.save(src_path)
print(f"Removed BC 1 Bảo Lộc from closure list in {src_path}!")
