import docx, openpyxl, sys
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

# 5 Province Groups with Clean Ward Names in Col 2
province_data = [
    {
        "province": "TỈNH LÂM ĐỒNG",
        "rows": [
            ["Phường 1 Bảo Lộc", "P.1 + P. Lộc Phát + X. Lộc Thanh", "BC (LDO) B'Lao Mới", "Gộp 100% ➔ Đóng BC 1 Bảo Lộc cũ"],
            ["Phường B'Lao", "P. Lộc Sơn + P. B'Lao + X. Lộc Nga", "BC (LDO) B'Lao Mới", "Gộp sản lượng về BC mới"],
            ["Phường 2 Bảo Lộc", "P.2 + X. Đạm Bri + X. Lộc Tân", "BC (LDO) 3 Bảo Lộc", "Gộp về BC 3 Bảo Lộc"],
            ["Phường Xuân Hương - Đà Lạt", "P.1 + P.2 + P.3 + P.4 + P.10", "BC Xuân Hương 1 & BC Xuân Hương 2 (MỚI)", "TÁCH MỚI BC Xuân Hương 2 tại P.10 (cover P.3 & P.10)"],
            ["Phường Lâm Viên - Đà Lạt", "P.8 + P.12 + P.9", "BC Lâm Viên 1 & BC Lâm Viên 2", "GIỮ NGUYÊN 02 BƯU CỤC chia tải địa bàn rộng"],
            ["Cụm Đơn Dương (Bắc)", "Lạc Lâm + Lạc Xuân + D'Ran + Ka Đô", "BC (LDO) Lạc Xuân (MỚI)", "Mở mới BC (vùng màu vàng)"],
            ["Cụm Đơn Dương (Nam)", "Đạ Ròn + TT. Thạnh Mỹ + Tu Tra + Ka Đơn...", "BC Nghĩa Đức (GỐC)", "Giữ BC gốc Nghĩa Đức cover 6 xã"],
            ["Xã Di Linh", "TT. Di Linh + X. Liên Đầm + X. Gung Ré + Tân Châu", "BC Hàng Vừa Di Linh (MỚI) & BC Di Linh", "TÁCH BC HÀNG VỪA tại Đinh Trang Thượng giảm 22-45km"],
            ["Xã Bảo Lâm 2", "X. Lộc Đức + X. Lộc An + X. Tân Lạc", "BC (LDO) Bảo Lâm 3", "Gộp toàn bộ tuyến địa bàn về BC Bảo Lâm 3"],
            ["Xã Nam Hà Lâm Hà", "X. Nam Hà + X. Phi Tô", "BC (LDO) Nam Ban Lâm Hà", "Gộp về BC Nam Ban Lâm Hà"],
            ["Xã Đam Rông 4", "X. Đạ Tông + X. Đạ Long + X. Đưng KNớ", "BC Lang Biang 1 & BC Đam Rông 3", "GIỮ NGUYÊN 02 BƯU CỤC (khoảng cách Đưng KNớ >50km)"],
            ["Xã Đức Trọng", "TT. Liên Nghĩa + X. Phú Hội", "BC Đức Trọng 1 & BC Đức Trọng 2", "GIỮ NGUYÊN 02 BƯU CỤC (khoảng cách xa >15km)"]
        ]
    },
    {
        "province": "TỈNH KHÁNH HÒA",
        "rows": [
            ["Phường Nam Nha Trang", "P. Phước Hải + P. Phước Long + P. Vĩnh Trường", "BC Nam Nha Trang 1 Mới & BC Nam Nha Trang 5", "Gộp về BC mới (ĐÓNG CỬA BC 2 & BC 3)"],
            ["Phường Nha Trang", "P. Vạn Thạnh + P. Lộc Thọ + P. Tân Tiến + P. Phước Hòa + P. Vĩnh Nguyên", "BC (KHO) Nha Trang", "Gộp phân vùng về BC Nha Trang"],
            ["Phường Tây Nha Trang", "P. Ngọc Hiệp + P. Phương Sài + X. Vĩnh Ngọc + X. Vĩnh Thạnh + X. Vĩnh Trung + X. Vĩnh Hiệp", "BC (KHO) Tây Nha Trang", "Gộp về BC Tây Nha Trang"],
            ["Xã Diên Khánh", "TT. Diên Khánh + X. Diên Toàn + X. Diên An", "BC (KHO) Diên Khánh 2", "Gộp về BC Diên Khánh 2 (ĐÓNG CỬA BC DK1 <40m²)"],
            ["Xã Vạn Thắng", "X. Vạn Thắng + X. Vạn Bình", "BC (KHO) Tu Bông", "Gộp sản lượng Xã Vạn Bình về BC Tu Bông"],
            ["Cụm TP. Cam Ranh", "Ba Ngòi + Cam Bình + Cam Lập + Cam Phước Đông...", "BC Nam Cam Ranh (MỚI) & BC Cam Linh", "TÁCH MỚI BC cover 6 xã phía Nam"],
            ["Xã Cam Hiệp", "X. Sơn Tân + X. Cam Hiệp", "BC (KHO) Cam Lâm 1", "Quy hoạch đồng nhất về BC Cam Lâm 1"],
            ["Xã Cam Lâm", "X. Cam Hải Đông + X. Cam Hải Tây + TT. Cam Đức + X. Cam Thành Bắc", "BC (KHO) Cam Lâm 2", "Quy hoạch đồng nhất về BC Cam Lâm 2"],
            ["Xã Hòa Trí", "X. Ninh Thượng + X. Ninh Trung + X. Ninh Thân", "BC Ninh Hòa 1 & BC Ninh Hòa 2", "GIỮ NGUYÊN 02 BƯU CỤC chia tuyến cover"],
            ["Xã Tân Định", "X. Ninh Quang + X. Ninh Xuân + X. Ninh Bình", "BC Ninh Hòa 1 & BC Ninh Hòa 2", "GIỮ NGUYÊN 02 BƯU CỤC phụ trách chia tuyến"]
        ]
    },
    {
        "province": "TỈNH NINH THUẬN",
        "rows": [
            ["Phường Ninh Chử", "P. Văn Hải + TT. Khánh Hải", "BC (NTH) Ninh Chử", "Di dời kho về trung tâm ĐVHC mới"],
            ["Phường Phan Rang", "P. Phủ Hà + P. Kinh Dinh + P. Đạo Long + P. Đài Sơn", "BC (NTH) Phan Rang", "Gộp phân vùng về BC Phan Rang"],
            ["Xã Ninh Hải", "X. Tri Hải + X. Bắc Sơn + X. Phương Hải", "BC (NTH) Ninh Chử", "BC Ninh Chử cover 75.7% sản lượng toàn xã"],
            ["Xã Phước Dinh", "X. An Hải + X. Phước Dinh + P. Đông Hải", "BC Đông Hải (MỚI) & BC Phước Dinh", "MỜ MỚI BC Đông Hải ven biển (600 giao, 250 lấy/ngày)"]
        ]
    },
    {
        "province": "TỈNH BÌNH THUẬN",
        "rows": [
            ["Phường La Gi", "P. Tân Thiện + P. Tân An + P. Bình Tân + X. Tân Bình", "BC (BTH) Phước Hội & BC Tân Hải", "Phước Hội cover 3 phường, Tân Hải cover X. Tân Bình"],
            ["Phường Bình Thuận", "P. Phú Tài + X. Phong Nẫm + X. Hàm Hiệp", "BC (BTH) Hàm Thắng & BC Hàm Liêm", "Gộp sản lượng sáp nhập về BC Hàm Thắng"],
            ["Phường Phan Thiết", "P. Phú Trinh + P. Lạc Đạo + P. Bình Hưng", "BC (BTH) Hàm Thắng & BC Phú Thủy", "Gộp sản lượng sáp nhập về BC Hàm Thắng"],
            ["Phường Hàm Thắng", "P. Xuân An + TT. Phú Long + X. Hàm Thắng", "BC (BTH) Phú Thủy & BC Hàm Liêm", "Gộp sản lượng về BC Phú Thủy"],
            ["Xã Phan Rí Cửa", "TT. Phan Rí Cửa + X. Hòa Minh + X. Chí Công", "BC Phan Rí Cửa & BC Liên Hương", "GIỮ NGUYÊN 02 BƯU CỤC cover song song"],
            ["Xã Tân Thành", "X. Tân Thuận + X. Tân Thành + X. Thuận Quý", "BC Tân Hải & BC Hàm Thuận Nam", "GIỮ NGUYÊN 02 BƯU CỤC (trục bờ biển dài >20km)"],
            ["Cụm Nam Thành", "X. Nam Thành + X. Nghị Đức", "BC (BTH) Nam Thành (MỚI)", "MỜ MỚI xóa chạy chéo tuyến xa"]
        ]
    },
    {
        "province": "TỈNH ĐẮK NÔNG",
        "rows": [
            ["Phường Bắc Gia Nghĩa", "P. Nghĩa Đức + P. Nghĩa Thành + P. Quảng Thành + X. Đắc Ha", "BC (DNO) Bắc Gia Nghĩa", "Gộp địa bàn về BC Bắc Gia Nghĩa"],
            ["Phường Nam Gia Nghĩa", "P. Nghĩa Phú + P. Nghĩa Tân + X. Đắk R'Moan", "BC Nam Gia Nghĩa", "TẠM THỜI GIỮ NGUYÊN phạm vi quản lý cũ (3-6 tháng)"],
            ["Xã Đắc Sắc", "X. Đắc Sắc + X. Nam Xuân + X. Long Sơn", "BC Đức Lập & BC Krông Nô", "GIỮ NGUYÊN BC Đức Lập & BC Krông Nô (tránh đứt gãy nhân sự)"],
            ["Xã Đức An", "TT. Đức An + X. Nam Bình + X. Đắk N'Drung", "BC Đức An & BC Trường Xuân", "GIỮ NGUYÊN BC Đức An & BC Trường Xuân (gần kho hơn)"],
            ["Xã Tà Đùng", "X. Đắc Som + X. Đắc R'Măng", "BC Quảng Sơn & BC Quảng Khê", "GIỮ NGUYÊN 02 BƯU CỤC (đường đèo dốc >35km)"],
            ["Xã Quảng Tân", "X. Quảng Tân + X. Đắk Ngo", "BC (DNO) Quảng Tín & BC Kiến Đức", "GIỮ NGUYÊN 02 BƯU CỤC cover song song"]
        ]
    }
]

# ==============================================================================
# EXCEL FILE GENERATION (SPLIT BY PROVINCE WITH ORANGE HEADERS)
# ==============================================================================
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Quy_Hoach_NTB_2026_Theo_Tinh"

font_prov_header = Font(name="Calibri", size=12, bold=True, color="003366")
font_header = Font(name="Calibri", size=11, bold=True, color="000000")
fill_header = PatternFill(start_color="F7A059", end_color="F7A059", fill_type="solid")

font_dvhc = Font(name="Calibri", size=10, bold=True, color="000000")
font_normal = Font(name="Calibri", size=10, color="000000")
font_buucuc = Font(name="Calibri", size=10, bold=True, color="003399")

align_header = Alignment(horizontal="center", vertical="center", wrap_text=True)
align_cell = Alignment(horizontal="left", vertical="center", wrap_text=True)

thin_border = Border(
    left=Side(style='thin', color='000000'), right=Side(style='thin', color='000000'),
    top=Side(style='thin', color='000000'), bottom=Side(style='thin', color='000000')
)

headers = ["ĐVHC Mới", "Các Xã / Phường Cũ sáp nhập", "Bưu cục Cover", "Phương án & Ghi chú"]

current_row = 1

for t_data in province_data:
    # Province Title Row
    cell_prov = ws.cell(row=current_row, column=1, value=f"❖ QUY HOẠCH {t_data['province']}")
    cell_prov.font = font_prov_header
    ws.row_dimensions[current_row].height = 24
    current_row += 1
    
    # Table Header Row
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=current_row, column=col_idx, value=h_text)
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_header
        cell.border = thin_border
    
    ws.row_dimensions[current_row].height = 26
    current_row += 1
    
    # Data Rows
    for r_data in t_data["rows"]:
        for col_idx, val in enumerate(r_data, start=1):
            cell = ws.cell(row=current_row, column=col_idx, value=val)
            cell.border = thin_border
            cell.alignment = align_cell
            if col_idx == 1:
                cell.font = font_dvhc
            elif col_idx == 3:
                cell.font = font_buucuc
            else:
                cell.font = font_normal
        ws.row_dimensions[current_row].height = 24
        current_row += 1
        
    current_row += 1 # Space row

ws.column_dimensions['A'].width = 25
ws.column_dimensions['B'].width = 44
ws.column_dimensions['C'].width = 38
ws.column_dimensions['D'].width = 44

excel_out = r'C:\Users\lap4all\Downloads\Bang_Quy_Hoach_NTB_2026_Clean_Phan_Theo_Tinh.xlsx'
wb.save(excel_out)
print(f"Generated Excel file split by province: {excel_out}")

# ==============================================================================
# WORD FILE GENERATION (SPLIT BY PROVINCE WITH ORANGE HEADERS)
# ==============================================================================
doc_out = docx.Document()

for section in doc_out.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_province_heading(text):
    p = doc_out.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

# Title
p_t = doc_out.add_paragraph()
r_t = p_t.add_run('TỜ TRÌNH QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB NĂM 2026\n(BẢNG TỔNG HỢP PHÂN THEO 5 TỈNH - GIAO DIỆN CHUẨN MẪU)')
r_t.font.name = 'Calibri'
r_t.font.size = Pt(15)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(180, 0, 0)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(12)

widths = [Inches(1.8), Inches(2.6), Inches(2.2), Inches(2.2)]

for t_data in province_data:
    add_province_heading(f"❖ QUY HOẠCH {t_data['province']}")
    
    table = doc_out.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], 'F7A059')
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
            
    for r_data in t_data["rows"]:
        row_cells = table.add_row().cells
        bg_color = 'FFFFFF'
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            row_cells[c_idx].width = widths[c_idx]
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=80, right=80)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                if c_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                elif c_idx == 2:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 51, 153)

word_out = r'C:\Users\lap4all\Downloads\Bang_Quy_Hoach_NTB_2026_Clean_Phan_Theo_Tinh.docx'
doc_out.save(word_out)
print(f"Generated Word file split by province: {word_out}")
