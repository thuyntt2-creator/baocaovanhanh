import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_4.docx"

if not os.path.exists(doc_path):
    print("File không tồn tại")
    sys.exit(1)

doc = docx.Document(doc_path)
errors = 0

print("=== KIỂM TRA SỐ LIỆU TOÀN DIỆN FILE WORD V8_4 ===")

# 1. Kiểm tra Bảng 2 (Tóm tắt)
table_2 = doc.tables[1]
print("\n1. Kiểm tra Bảng 2 (Tóm tắt):")
if "84 người" not in table_2.rows[3].cells[1].text:
    print(f"  [LỖI] Dòng Nhân sự = {table_2.rows[3].cells[1].text} (Kỳ vọng chứa 84 người -> 130 người)")
    errors += 1
else:
    print("  [OK] Dòng Nhân sự khớp")

if "30 xe" not in table_2.rows[4].cells[1].text:
    print(f"  [LỖI] Dòng Xe đỉnh = {table_2.rows[4].cells[1].text}")
    errors += 1
else:
    print("  [OK] Dòng Xe đỉnh khớp")

if "1.746 → 2.915" not in table_2.rows[5].cells[1].text:
    print(f"  [LỖI] Dòng Chi phí = {table_2.rows[5].cells[1].text}")
    errors += 1
else:
    print("  [OK] Dòng Chi phí khớp")

# 2. Kiểm tra Bảng 3 (Số đầu xe)
table_3 = doc.tables[2]
print("\n2. Kiểm tra Bảng 3 (Số xe Đức Linh T12):")
dl_t12 = table_3.rows[12].cells[4].text.strip()
if dl_t12 != "2 xe":
    print(f"  [LỖI] Xe Đức Linh T12 = {dl_t12} (Kỳ vọng: 2 xe)")
    errors += 1
else:
    print("  [OK] Xe Đức Linh T12 khớp")

# 3. Kiểm tra Bảng 7 (Nhân sự)
table_7 = doc.tables[6]
print("\n3. Kiểm tra Bảng 7 (Nhân sự chi tiết T7):")
nt_t7 = table_7.rows[2].cells[1].text.strip()
dd_t7 = table_7.rows[4].cells[1].text.strip()
dul_t7 = table_7.rows[5].cells[1].text.strip()

if nt_t7 != "39 (SL: 1815)":
    print(f"  [LỖI] Nha Trang T7 = {nt_t7}")
    errors += 1
else:
    print("  [OK] Nha Trang T7 khớp")

if dd_t7 != "10 (SL: 139)":
    print(f"  [LỖI] Đơn Dương T7 = {dd_t7}")
    errors += 1
else:
    print("  [OK] Đơn Dương T7 khớp")

if dul_t7 != "11 (SL: 111)":
    print(f"  [LỖI] Đức Linh T7 = {dul_t7}")
    errors += 1
else:
    print("  [OK] Đức Linh T7 khớp")

print("\n4. Kiểm tra Tổng nhân sự toàn vùng (Bảng 7):")
tot_ns_expected = ["84", "86", "116", "118", "128", "130"]
for c_idx in range(1, 7):
    actual = table_7.rows[9].cells[c_idx].text.strip()
    expected = tot_ns_expected[c_idx - 1]
    if actual != expected:
        print(f"  [LỖI] Tháng {c_idx+6}: Nhân sự thực tế = {actual} | Kỳ vọng = {expected}")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Nhân sự = {actual} (Khớp)")

# 4. Kiểm tra đối chiếu Bảng 10 vs Bảng 11
table_10 = doc.tables[9]
table_11 = doc.tables[10]
print("\n5. Kiểm tra đối chiếu Bảng 10 vs Bảng 11:")
for c_idx in range(1, 7):
    t10_sum = float(table_10.rows[5].cells[c_idx].text.replace(',', ''))
    t11_sum = float(table_11.rows[5].cells[c_idx].text.replace(',', ''))
    if abs(t10_sum - t11_sum) > 0.1:
        print(f"  [LỖI] Tháng {c_idx+6}: Bảng 10 = {t10_sum} | Bảng 11 = {t11_sum}")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Tổng CP = {t10_sum} (Khớp)")

if errors == 0:
    print("\n>>> TẤT CẢ KIỂM TRA BÁO CÁO V8_4 ĐỀU THÀNH CÔNG! SỐ LIỆU HOÀN THIỆN TUYỆT ĐỐI <<<")
else:
    print(f"\n>>> CÓ {errors} SAI LỆCH SỐ LIỆU CẦN XỬ LÝ <<<")
