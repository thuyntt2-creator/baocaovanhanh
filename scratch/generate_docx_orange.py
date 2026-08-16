import docx, sys
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

doc_out = docx.Document()

for section in doc_out.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_header_p(text):
    p = doc_out.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

# Title
p_t = doc_out.add_paragraph()
r_t = p_t.add_run('TỜ TRÌNH QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB NĂM 2026\n(BẢNG THEO TỪNG TỈNH - GIAO DIỆN CHUẨN FILE ẢNH EXCEL)')
r_t.font.name = 'Calibri'
r_t.font.size = Pt(15)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(180, 0, 0)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(12)

headers = ["ĐVHC Mới", "Các Xã / Phường Cũ sáp nhập", "Bưu cục Cover", "Phương án & Ghi chú"]
widths = [Inches(1.8), Inches(2.6), Inches(2.2), Inches(2.2)]

province_tables = [
    {
        "province": "TỈNH LÂM ĐỒNG",
        "rows": [
            ["Phường 1 Bảo Lộc", "P.1 + P. Lộc Phát + X. Lộc Thanh", "BC (LDO) B'Lao Mới", "Gộp 100% ➔ Đóng BC 1 Bảo Lộc cũ"],
            ["Phường B'Lao", "P. Lộc Sơn + P. B'Lao + X. Lộc Nga", "BC (LDO) B'Lao Mới", "Gộp sản lượng về BC mới"],
            ["Phường 2 Bảo Lộc", "P.2 + X. Đạm Bri + X. Lộc Tân", "BC (LDO) 3 Bảo Lộc", "Gộp về BC 3 Bảo Lộc"],
            ["Cụm Đơn Dương (Bắc)", "Lạc Lâm + Lạc Xuân + D'Ran + Ka Đô", "BC (LDO) Lạc Xuân (MỚI)", "Mở mới BC (vùng màu vàng)"],
            ["Cụm Đơn Dương (Nam)", "Đạ Ròn + TT. Thạnh Mỹ + Tu Tra + Ka Đơn...", "BC Nghĩa Đức (GỐC)", "Giữ BC gốc Nghĩa Đức cover 6 xã"],
            ["Xã Di Linh", "TT. Di Linh + X. Liên Đầm + Tân Châu...", "BC Hàng Vừa Di Linh (MỚI) & BC Di Linh", "Tách BC tại Đinh Trang Thượng"],
            ["Xã Đam Rông 4", "X. Đạ Tông + X. Đạ Long + X. Đưng KNớ", "BC Đam Rông 3 & BC Lang Biang 1", "Giữ 2 BC (địa hình xa >50km)"],
            ["Xã Đức Trọng", "TT. Liên Nghĩa + X. Phú Hội", "BC Đức Trọng 2 & BC Đức Trọng 1", "Giữ 2 BC (địa bàn xa >15km)"]
        ]
    },
    {
        "province": "TỈNH KHÁNH HÒA",
        "rows": [
            ["Phường Nam Nha Trang", "P. Vĩnh Trường + P. Phước Long + X. Vĩnh Hiệp", "BC Nam Nha Trang 1 Mới & BC Nam Nha Trang 5", "Gộp về BC mới (Đóng BC 3)"],
            ["Phường Nha Trang", "P. Vĩnh Nguyên + P. Phước Tiến + P. Phước Tân", "BC (KHO) Nha Trang", "Gộp phân vùng về BC Nha Trang"],
            ["Phường Tây Nha Trang", "P. Vĩnh Hải + X. Vĩnh Ngọc + X. Vĩnh Thạnh", "BC (KHO) Tây Nha Trang", "Gộp về BC Tây Nha Trang"],
            ["Xã Diên Khánh", "TT. Diên Khánh + X. Diên An + X. Diên Toàn", "BC (KHO) Diên Khánh 2", "Gộp về DK2 (Đóng BC DK1 <40m²)"],
            ["Xã Vạn Thắng", "X. Vạn Thắng + X. Vạn Bình", "BC (KHO) Tu Bông", "Gộp dồn sản lượng về BC Tu Bông"],
            ["Cụm TP. Cam Ranh", "Ba Ngòi + Cam Bình + Cam Lập + Cam Phước Đông...", "BC Nam Cam Ranh (MỚI) & BC Cam Linh", "Tách mới BC cover 6 xã phía Nam"]
        ]
    },
    {
        "province": "TỈNH NINH THUẬN",
        "rows": [
            ["Phường Ninh Chử", "TT. Khánh Hải + P. Văn Hải", "BC (NTH) Ninh Chử", "Di dời kho về trung tâm ĐVHC mới"],
            ["Phường Phan Rang", "P. Kinh Dinh + P. Phủ Hà + P. Đạo Long + P. Đài Sơn", "BC (NTH) Phan Rang", "Gộp phân vùng về BC Phan Rang"],
            ["Xã Ninh Hải", "X. Phương Hải + X. Tri Hải + X. Bắc Sơn", "BC (NTH) Ninh Chử", "BC Ninh Chử cover 75.7% sản lượng"],
            ["Xã Phước Dinh", "X. An Hải + X. Phước Dinh + P. Đông Hải", "BC Đông Hải (MỚI) & BC Phước Dinh", "Mở BC Đông Hải gánh đơn hải sản"]
        ]
    },
    {
        "province": "TỈNH BÌNH THUẬN",
        "rows": [
            ["Phường La Gi", "P. Tân Thiện + P. Tân An + P. Bình Tân + X. Tân Bình", "BC (BTH) Phước Hội & BC (BTH) Tân Hải", "Phước Hội cover 3 phường, Tân Hải cover X. Tân Bình"],
            ["Phường Bình Thuận", "P. Phú Tài + X. Phong Nẫm + X. Hàm Hiệp", "BC (BTH) Hàm Thắng & BC (BTH) Hàm Liêm", "Gộp địa bàn về BC Hàm Thắng & Hàm Liêm"],
            ["Phường Phan Thiết", "P. Phú Trinh + P. Lạc Đạo + P. Bình Hưng", "BC (BTH) Hàm Thắng & BC (BTH) Phú Thủy", "Gộp địa bàn về BC Hàm Thắng & Phú Thủy"],
            ["Phường Hàm Thắng", "P. Xuân An + TT. Phú Long + X. Hàm Thắng", "BC (BTH) Phú Thủy & BC (BTH) Hàm Liêm", "Gộp về BC Phú Thủy & BC Hàm Liêm"],
            ["Xã Phan Rí Cửa", "TT. Phan Rí Cửa + X. Chí Công + X. Hòa Minh", "BC Phan Rí Cửa & BC Liên Hương", "Giữ 2 BC cover song song"],
            ["Xã Tân Thành", "X. Tân Thuận + X. Tân Thành + X. Thuận Quý", "BC Tân Hải & BC Hàm Thuận Nam", "Giữ 2 BC (trục bờ biển dài >20km)"],
            ["Cụm Nam Thành", "X. Nam Thành + X. Nghị Đức", "BC (BTH) Nam Thành (MỚI)", "Mở mới xóa chạy chéo tuyến xa"]
        ]
    },
    {
        "province": "TỈNH ĐẮK NÔNG",
        "rows": [
            ["Phường Bắc Gia Nghĩa", "P. Nghĩa Đức + P. Nghĩa Thành + P. Quảng Thành + X. Đắc Ha", "BC (DNO) Bắc Gia Nghĩa", "Gộp địa bàn về BC Bắc Gia Nghĩa"],
            ["Xã Đắc Sắc", "X. Đắc Sắc + X. Nam Xuân + X. Long Sơn", "BC Đức Lập & BC Krông Nô", "Giữ 2 BC (NV Krông Nô không đi)"],
            ["Xã Đức An", "TT. Đức An + X. Đắc N'Drung + X. Nam Bình", "BC Đức An & BC Trường Xuân", "Giữ 2 BC (Đắk N'Drung gần TX hơn)"],
            ["Xã Tà Đùng", "X. Đắc Som + X. Đắc R'Măng", "BC Quảng Khê & BC Quảng Sơn", "Giữ 2 BC (đường đèo dốc >35km)"],
            ["Xã Quảng Tân", "X. Quảng Tân + X. Đắc Ngo", "BC Kiến Đức & BC Quảng Tín", "Giữ 2 BC cover song song"]
        ]
    }
]

for t_data in province_tables:
    add_header_p(f"❖ QUY HOẠCH {t_data['province']}")
    
    table = doc_out.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Orange Header Row matching screenshot
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], 'F7A059') # Orange fill
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0) # Black text
            
    for r_idx, r_data in enumerate(t_data["rows"]):
        row_cells = table.add_row().cells
        bg_color = 'FFFFFF'
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            row_cells[c_idx].width = widths[c_idx]
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=80, right=80)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                if c_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                elif c_idx == 2:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 51, 153) # Dark Blue Bưu cục Cover

out_docx = r'C:\Users\lap4all\Downloads\Thu_trinh_Quy_Hoach_Mang_Luoi_NTB_2026_Theo_Tinh_Bang_Cam.docx'
doc_out.save(out_docx)
print(f'Successfully generated docx with orange headers matching screenshot: {out_docx}')
