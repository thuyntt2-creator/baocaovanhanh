import docx
from docx.shared import Pt, RGBColor
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

source_path = r"C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_ MANG LUOI NTB_FIXED.docx"
target_path = r"C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_ MANG LUOI NTB_FIXED.docx"
orig_path = r"C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_ MANG LUOI NTB.docx"

if not os.path.exists(source_path):
    source_path = orig_path

doc = docx.Document(source_path)

# Find paragraph where IV starts
iv_idx = None
for i, p in enumerate(doc.paragraphs):
    if "IV. YÊU CẦU 3" in p.text or "IV. YÊU CẦU" in p.text:
        iv_idx = i
        break

print(f"Found IV heading at paragraph index {iv_idx}")

# Delete all paragraphs from iv_idx onwards
if iv_idx is not None:
    for p in list(doc.paragraphs[iv_idx:]):
        p._element.getparent().remove(p._element)

# Remove ALL tables in the document
for tbl in list(doc.tables):
    tbl._element.getparent().remove(tbl._element)

# Now rebuild Section IV (ONLY Item 1 & Item 2) & Section V

# Section IV Heading
p_iv = doc.add_paragraph()
r_iv = p_iv.add_run("IV. YÊU CẦU 3: ĐÁNH GIÁ NHU CẦU TÁCH BƯU CỤC VÀ TỐI ƯU MẠNG LƯỚI")
r_iv.bold = True
r_iv.font.size = Pt(13)
r_iv.font.color.rgb = RGBColor(0, 51, 102)

# Item 1: 07 BC Mở mới / Tách bưu cục
p = doc.add_paragraph()
r = p.add_run("1. Kế hoạch Mở mới / Tách Bưu cục & Tối ưu Mạng lưới (07 Bưu cục)")
r.bold = True
r.font.size = Pt(11)

items_iv1 = [
    ("1. Mở mới Bưu cục (LDO) Xuân Hương - Đà Lạt 2: ", "Đặt tại Phường 10 (TP. Đà Lạt), chia tải cho BC Xuân Hương cũ (phụ trách Phường 3 & Phường 10, Vol giao: 1,050 đơn/ngày, Vol lấy: 150 đơn/ngày, định biên 8 NVPTTT + 1 NVXL)."),
    ("2. Tách Bưu cục Hàng Nhỏ / Hàng Vừa Di Linh: ", "Phụ trách Xã Đinh Trang Thượng, Di Linh, Phúc Thọ Lâm Hà, Liên Đầm nhằm giảm bán kính di chuyển 22–45km cho BC Di Linh."),
    ("3. Mở mới Bưu cục Đông Hải (Tỉnh Ninh Thuận): ", "Cover khu vực ven biển Phường Đông Hải, tối ưu điểm tập kết sản lượng lấy (Vol giao: 600 đơn/ngày, Vol lấy: 250 đơn/ngày, định biên 7 NVPTTT + 1 NVXL)."),
    ("4. Mở mới Bưu cục (LDO) B'Lao Mới (Bảo Lộc): ", "Cover Phường 1 & Phường B'Lao (Vol giao: 1,500 đơn/ngày, Vol lấy: 1,000 đơn/ngày, định biên 15 NVPTTT + 2 NVXL), đóng cửa BC (LDO) 1 Bảo Lộc cũ."),
    ("5. Mở mới Bưu cục (BTH) Nam Thành (Tỉnh Bình Thuận): ", "Cover khu vực Nam Thành (250 đơn/ngày) & Nghị Đức (200 đơn/ngày), định biên 7 NVPTTT + 1 NVXL."),
    ("6. Mở mới / Tách Bưu cục Nam Cam Ranh (Tỉnh Khánh Hòa): ", "Tách từ BC (KHO) Cam Linh cũ (ID: 22830000) để chia tải cho 6 xã/phường phía Nam (Ba Ngòi, Cam Bình, Cam Lập, Cam Phước Đông, Cam Thịnh Đông, Cam Thịnh Tây; Vol giao: 500 đơn/ngày, Vol lấy: 100 đơn/ngày, định biên 1 AM + 1 NVXL + 8 NVPTTT)."),
    ("7. Mở mới / Tách Bưu cục (LDO) Lạc Xuân - Đơn Dương (Tỉnh Lâm Đồng): ", "Tách từ Bưu cục gốc Nghĩa Đức (TT. Thạnh Mỹ) để mở bưu cục mới tại Xã Lạc Xuân. BC Lạc Xuân mới cover 4 xã phía Đông Bắc (Lạc Lâm, Lạc Xuân, D'Ran, Ka Đô; Vol giao 400 - 480 đơn/ngày, định biên 7 NV). BC gốc Nghĩa Đức giữ cover 6 xã/thị trấn (Đạ Ròn, TT. Thạnh Mỹ, Tu Tra, Ka Đơn, Quảng Lập, Pró; Vol giao 600 - 720 đơn/ngày, định biên 9 NV). Tổng toàn cụm Đơn Dương 1.000 - 1.200 đơn/ngày, 16 nhân sự.")
]

for title, desc in items_iv1:
    p = doc.add_paragraph()
    r1 = p.add_run("❖ " + title)
    r1.bold = True
    r2 = p.add_run(desc)

# Item 2: Heading 2 -> Table 1 (Overloaded / Relocated Hubs ONLY)
p = doc.add_paragraph()
r = p.add_run("2. Danh sách Bưu cục Quá tải & Đề xuất Di dời / Mở rộng mặt bằng (14 Bưu cục)")
r.bold = True
r.font.size = Pt(11)

table1_data = [
    ["STT", "ID BC", "Tên Bưu cục", "Tỉnh", "Quản lý AM", "Địa chỉ / Quy mô", "Hiện trạng m² / Vol", "Đề xuất giải pháp"],
    ["1", "21479000", "(KHO) Bắc Cam Ranh", "Khánh Hòa", "Thái Thị Thanh Thư", "Phụ trách 3 xã/phường Cam Ranh", "Diện tích 100 m² (Giao 600, Lấy 150)", "Di dời sang mặt bằng mới rộng hơn"],
    ["2", "1896", "(LDO) Đạ Huoai", "Lâm Đồng", "Nguyễn Lê Nguyên Vũ", "Ma Đa Guôi, Đạ Huoai, Lâm Đồng", "Áp lực em2: 59.2", "Tách bớt phường / Mở rộng m²"],
    ["3", "21456000", "(BTH) Đức Linh", "Bình Thuận", "Nguyễn Ngọc Khánh", "38 DT766, xã Nam Chính, Đức Linh", "Áp lực em2: 33.2", "Tách bớt phường / Mở rộng m²"],
    ["4", "21320000", "(BTH) Phú Thủy", "Bình Thuận", "Nguyễn Ngọc Khánh", "188 Trương Hán Siêu, Phan Thiết", "Áp lực em2: 32.0", "Tách bớt phường / Mở rộng m²"],
    ["5", "22452000", "(BTH) Phú Quý", "Bình Thuận", "Nguyễn Ngọc Khánh", "454 Võ Văn Kiệt, Phú Quý", "Áp lực em2: 28.9", "Tách bớt phường / Mở rộng m²"],
    ["6", "20543000", "(KHO) Ninh Hòa 2", "Khánh Hòa", "Phạm Bá Thành Công", "06 đường 2/4 Ninh Hiệp", "Áp lực em2: 28.3", "Tách bớt phường / Mở rộng m²"],
    ["7", "2399", "(KHO) Bắc Nha Trang", "Khánh Hòa", "Phạm Bá Thành Công", "195 đường 2/4, Vĩnh Hòa, Nha Trang", "Áp lực em2: 25.5", "Tách bớt phường / Mở rộng m²"],
    ["8", "20144000", "(BTH) Mũi Né", "Bình Thuận", "Nguyễn Ngọc Khánh", "16 Huỳnh Thúc Kháng, Mũi Né", "Áp lực em2: 24.4", "Tách bớt phường / Mở rộng m²"],
    ["9", "21403000", "(BTH) Hàm Liêm", "Bình Thuận", "Nguyễn Ngọc Khánh", "Thôn 2, xã Hàm Liêm", "Áp lực em2: 24.3", "Tách bớt phường / Mở rộng m²"],
    ["10", "2357", "(BTH) Đồng Kho", "Bình Thuận", "Nguyễn Ngọc Khánh", "158 QL55, Đức Bình, Tánh Linh", "Áp lực em2: 24.2", "Tách bớt phường / Mở rộng m²"],
    ["11", "22861000", "(NTH) Phan Rang", "Ninh Thuận", "Nguyễn Duy Long", "95 Thống Nhất, Phan Rang", "Áp lực em2: 23.8", "Tách bớt phường / Mở rộng m²"],
    ["12", "20648000", "(NTH) Ninh Chử", "Ninh Thuận", "Nguyễn Duy Long", "51 Phạm Ngọc Thạch, Ninh Hải", "Áp lực em2: 22.1", "Tách bớt phường / Mở rộng m²"],
    ["13", "20785000", "(LDO) B'Lao", "Lâm Đồng", "Hồng Bích Nga", "95a Đội Cấn, Lộc Sơn, Bảo Lộc", "Áp lực em2: 22.1", "Tách bớt phường / Mở rộng m²"],
    ["14", "20150000", "(NTH) Bảo An", "Ninh Thuận", "Nguyễn Duy Long", "254A Đường 21/8, Phan Rang", "Áp lực em2: 21.7", "Tách bớt phường / Mở rộng m²"]
]

t1 = doc.add_table(rows=len(table1_data), cols=8)
for r_idx, row in enumerate(table1_data):
    for c_idx, val in enumerate(row):
        cell = t1.cell(r_idx, c_idx)
        cell.text = val
        if r_idx == 0:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

doc.add_paragraph()

# Section V
p_v = doc.add_paragraph()
r_v = p_v.add_run("V. TỔNG HỢP BIẾN ĐỘNG MẠNG LƯỚI BƯU CỤC NTB 2026")
r_v.bold = True
r_v.font.size = Pt(13)
r_v.font.color.rgb = RGBColor(0, 51, 102)

summary_items = [
    ("Bưu cục Mở mới / Tách bưu cục (07 BC): ", "BC Xuân Hương - Đà Lạt 2, BC Di Linh Hàng Vừa, BC Đông Hải, BC B'Lao Mới, BC Nam Thành, BC Nam Cam Ranh, BC Lạc Xuân (Đơn Dương)."),
    ("Bưu cục Di dời / Mở rộng mặt bằng kho (02 BC): ", "BC Bắc Cam Ranh (Khánh Hòa - mở rộng từ 100m²), BC Ninh Chử (Ninh Thuận)."),
    ("Bưu cục Đóng cửa (02 BC): ", "BC 1 Bảo Lộc, BC Nam Nha Trang 3 (gộp tuyến/chuyển đổi địa bàn)."),
    ("Bưu cục Gộp tuyến/Tối ưu phân vùng (21 BC): ", "Quy hoạch dồn sản lượng về 01 Bưu cục phụ trách chính tại các ĐVHC sáp nhập."),
    ("Bưu cục Giữ nguyên vận hành chia tuyến (13 BC): ", "Duy trì 2-3 bưu cục song song theo đề xuất AM do địa hình đặc thù, đồi dốc hoặc chia cắt địa lý (Đam Rông, Đức Trọng, Đắk Sắk, Đức An, Phan Rí Cửa, Ninh Hòa...).")
]

for title, desc in summary_items:
    p = doc.add_paragraph()
    r1 = p.add_run("❖ " + title)
    r1.bold = True
    r2 = p.add_run(desc)

# Try saving
try:
    doc.save(target_path)
    print(f"Successfully removed item 3 and saved DOCX to: {target_path}")
except Exception as e:
    print("Error saving target_path:", e)
    alt_path = r"C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_ MANG LUOI NTB_NO_ITEM3.docx"
    doc.save(alt_path)
    print("Saved to:", alt_path)
