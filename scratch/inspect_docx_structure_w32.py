import docx
import openpyxl
import sys

sys.stdout.reconfigure(encoding='utf-8')

docx_path = r'C:\Users\lap4all\Downloads\NTB - Báo Cáo Tuần (HRBP - ARD) - Tuần 32.docx'
excel_path = r'C:\Users\lap4all\Downloads\BaoCao_Tuan_NTB_W32_2026.xlsx'

doc = docx.Document(docx_path)
wb = openpyxl.load_workbook(excel_path, data_only=True)

print(f"=== DOCX PARAGRAPHS OVERVIEW ({len(doc.paragraphs)}) ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt and (txt.startswith('I.') or txt.startswith('II.') or txt.startswith('III.') or txt.startswith('IV.') or txt.startswith('V.') or txt.startswith('VI.') or txt.startswith('VII.') or txt.startswith('VIII.') or 'TUẦN' in txt.upper() or 'BÁO CÁO' in txt.upper()):
        print(f"P{i:3d}: {txt}")

print(f"\n=== DOCX TABLES OVERVIEW ({len(doc.tables)}) ===")
for i, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.columns)
    r0 = [c.text.strip().replace('\n', ' ') for c in t.rows[0].cells]
    print(f"Table {i:2d} ({rows:2d}x{cols:2d}): Header={r0[:4]}")
