import sys, docx, pandas as pd, json, os
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

sys.stdout.reconfigure(encoding='utf-8')

# Load official web data from Anh Thien's team
web_data_dir = r'c:\Users\lap4all\Documents\Auto report\scratch\web_data'
rezone_file = os.path.join(web_data_dir, 'rezone.json')
excel_path = r'C:\Users\lap4all\Downloads\NTB_Phan_Tuyen_Hanh_Chinh_Quy_Hoach_Moi.xlsx'

with open(rezone_file, 'r', encoding='utf-8') as f:
    rezone_data = json.load(f)

df_excel = pd.read_excel(excel_path, sheet_name='Sheet1')
excel_wards = df_excel.groupby('Mã Xã mới')

doc2 = docx.Document(r'C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi.docx')
maps_dir = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\d95894ad-9248-47c3-872b-72d7de66bc20\maps'

# Build a lookup for rezone items by name or code
web_wards = rezone_data['new_wards']
web_ward_dict = {}
for w in web_wards:
    web_ward_dict[w['name'].lower().strip()] = w

am_reasons = {
    24611: '''Khu vực Gia Nghĩa từng bể tuyến triền miên. Phương án tách 3 BC Bắc - Đông - Nam Gia Nghĩa giúp ổn định SLA. Nhân sự tại chỗ ven (Đắc Ha) nghỉ việc 100% nếu di chuyển kho mới 25-30km. Đề xuất TẠM THỜI GIỮ NGUYÊN 100% kho trong 3-6 tháng.''',
    24615: '''Giữ nguyên phân vùng quản lý BC Nam Gia Nghĩa 2 và Bắc Gia Nghĩa trong 3-6 tháng để tránh nguy cơ gãy tuyến và biến động nhân sự.''',
    24781: '''Phường Xuân Hương sản lượng cực lớn. Địa hình bị chia cắt bởi Hồ Xuân Hương, thời tiết lạnh mưa cuối năm. Tách mới BC Xuân Hương - Đà Lạt 2 (tại Phường 10) chia tải cho BC Xuân Hương cũ (Phường 1 & 2).''',
    24784: '''Giữ nguyên 02 bưu cục (Lâm Viên 1 & 2) nằm ngay Phường 8 giúp nhân viên tập trung xử lý một phường, giảm thời gian di chuyển.''',
    24958: '''Đức Trọng có 2 BC (Đức Trọng 1 & 2) cách nhau >15km. Khó tuyển dụng mùa cà phê & mưa cuối năm, kho bãi cũ không đủ m² chứa hàng. Đề xuất giữ nguyên 2 BC.''',
    22972: '''Phan Rí Cửa & Liên Hương cách nhau 25km, 2 thị trấn sản lượng lớn nhất. Giữ nguyên 2 BC.''',
    23143: '''BTH Tân Hải cover Tân Hải và Tân Thành mới. Giữ Thuận Quý cũ cho BC Hàm Thuận Nam vì khoảng cách gần hơn và dân cư đông đúc.''',
    22528: '''Huyện Ninh Hòa cũ diện tích lớn nhất cả nước (20 xã/phường). BC Ninh Hòa 1 cover tuyến xã xa trung tâm, kéo nhân sự tại chỗ. Kho Ninh Hòa 2 diện tích quá hẹp không thể gom gộp.''',
    25000: '''BC Di Linh bán kính xa 22-45km. Tách mới BC Hàng Nhỏ / Hàng Vừa tại Đinh Trang Thượng để cover Đinh Trang Thượng, Di Linh, Phúc Thọ, Liên Đầm.''',
    24748: '''Xã Quảng Tân địa hình rộng, xa và chia cắt mạnh giữa Đắk Ngo và Quảng Tân. Kiến Đức phụ trách Đắk Ngo (109 đơn), Quảng Tín phụ trách Quảng Tân (92 đơn). Đề xuất GIỮ NGUYÊN 02 BƯU CỤC.'''
}

ward_maps = {
    24781: ('image6.png', 'Bản đồ địa bàn quy hoạch mới TP. Đà Lạt 2026'),
    24784: ('image6.png', 'Bản đồ địa bàn quy hoạch mới TP. Đà Lạt 2026'),
    25000: ('image14.png', 'Bản đồ địa bàn quy hoạch khu vực Di Linh 2026'),
    24958: ('image13.png', 'Bản đồ địa bàn quy hoạch khu vực Đức Trọng 2026'),
    24823: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    24820: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    24817: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    25084: ('image3.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Bảo Lộc & Bảo Lâm 2026'),
    22366: ('image5.png', 'Bản đồ địa bàn quy hoạch mới TP. Nha Trang 2026'),
    22402: ('image5.png', 'Bản đồ địa bàn quy hoạch mới TP. Nha Trang 2026'),
    22390: ('image5.png', 'Bản đồ địa bàn quy hoạch mới TP. Nha Trang 2026'),
    22528: ('image8.jpg', 'Bản đồ địa bàn quy hoạch mới khu vực Ninh Hòa 2026'),
    22516: ('image7.jpg', 'Bản đồ địa bàn quy hoạch khu vực Vạn Ninh - Tu Bông 2026'),
    22933: ('image10.jpg', 'Bản đồ địa bàn quy hoạch mới TP. Phan Thiết & Hàm Thắng 2026'),
    22945: ('image10.jpg', 'Bản đồ địa bàn quy hoạch mới TP. Phan Thiết 2026'),
    22960: ('image11.jpg', 'Bản đồ phân vùng quy hoạch mới Phường Bình Thuận 2026'),
    23131: ('image9.png', 'Bản đồ quy hoạch mới khu vực Phường La Gi 2026'),
    23143: ('image9.png', 'Bản đồ quy hoạch mới khu vực Xã Tân Thành 2026'),
    22972: ('image4.png', 'Bản đồ quy hoạch mới khu vực Phan Rí Cửa & Tuy Phong 2026'),
    22759: ('image12.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Phan Rang 2026'),
    22834: ('image12.png', 'Bản đồ quy hoạch mạng lưới Bưu cục khu vực Ninh Chử 2026'),
}

# Construct ward list for 36 sáp nhập wards
report_wards = []

for code, group in excel_wards:
    prov = group['Tỉnh, thành phố mới'].iloc[0]
    name = group['Tên Xã mới'].iloc[0]
    n_bcs = group['Số BC'].iloc[0]
    old_communes = group['Tên Xã cũ'].tolist()
    
    # Match with web system item
    web_item = web_ward_dict.get(name.lower().strip(), None)
    if not web_item:
        # Search by partial
        for k, v in web_ward_dict.items():
            if name.lower().strip() in k or k in name.lower().strip():
                web_item = v
                break

    if web_item:
        sys_dem = web_item.get('dem', group['TỔNG ĐƠN/NGÀY (Phường mới)'].iloc[0])
        sys_kg = web_item.get('dem_kg', group['TỔNG KG/NGÀY (Phường mới)'].iloc[0])
        sys_status = web_item.get('status', 'split')
        assigned_bc = web_item.get('assigned_bc_name', group['Đánh giá & Phương án đề xuất'].iloc[0])
        cands = web_item.get('candidates', [])
        olds_web = web_item.get('olds', [])
    else:
        sys_dem = group['TỔNG ĐƠN/NGÀY (Phường mới)'].iloc[0]
        sys_kg = group['TỔNG KG/NGÀY (Phường mới)'].iloc[0]
        sys_status = 'split' if n_bcs > 1 else 'clean'
        assigned_bc = group['Đánh giá & Phương án đề xuất'].iloc[0]
        cands = []
        olds_web = []

    bc_details = []
    if cands:
        for c in cands:
            bc_details.append({
                'name': c['bc_name'], 'id': c.get('bc', ''), 'dem': c['dem'], 'pct': c['share']
            })
    else:
        for bc_name, bc_group in group.groupby('Tên Bưu cục giao'):
            bc_total = bc_group['Sản lượng giao/ngày (đơn)'].sum() + bc_group['Sản lượng lấy/ngày (đơn)'].sum()
            pct = round((bc_total / sys_dem * 100), 1) if sys_dem > 0 else 0.0
            bc_id = bc_group['ID Bưu cục giao'].iloc[0]
            bc_details.append({
                'name': bc_name, 'id': bc_id, 'dem': bc_total, 'pct': pct
            })
    bc_details.sort(key=lambda x: x['dem'], reverse=True)

    if code == 24748:
        prop = "GIỮ NGUYÊN 02 BƯU CỤC ((DNO) Quảng Tín & (DNO) Kiến Đức)"
    else:
        prop = f"Quy hoạch Bưu cục chính: {assigned_bc}"

    if code in am_reasons:
        final_reason = am_reasons[code]
    else:
        final_reason = group['Lý do & Bố trí nhân sự'].iloc[0]

    img_info = ward_maps.get(code, (None, None))

    report_wards.append({
        'code': code, 'name': name, 'prov': prov, 'n_bcs': n_bcs,
        'old_communes': old_communes, 'olds_web': olds_web,
        'sys_dem': sys_dem, 'sys_kg': sys_kg, 'sys_status': sys_status,
        'assigned_bc': assigned_bc, 'prop': prop,
        'bc_details': bc_details, 'reason': final_reason,
        'img_file': img_info[0], 'img_caption': img_info[1]
    })

report_wards.sort(key=lambda x: x['name'])

# Build Markdown text
md = []
md.append('# BÁO CÁO TOÀN DIỆN: RÀ SOÁT VÀ QUY HOẠCH BƯU CỤC THEO ĐƠN VỊ HÀNH CHÍNH MỚI (VÙNG NTB)')
md.append('### *(Số liệu chính thức đồng bộ từ Hệ thống Web Quy hoạch Bưu cục - quyhoachbuucuc.info của Team Anh Thiên)*\n')

md.append('## I. TỔNG QUAN HIỆN TRẠNG MẠNG LƯỚI BƯU CỤC VÙNG NTB')
md.append('- **Tổng số Xã/Phường hành chính mới rà soát**: 36 Phường/Xã mới (gồm 114 xã/phường cũ sáp nhập).')
md.append(f'- **Số Phường/Xã chuẩn ranh giới 01 BC (Web status: `clean`)**: {sum(1 for w in report_wards if w["sys_status"]=="clean")} Phường/Xã.')
md.append(f'- **Số Phường/Xã bị chia cắt (Web status: `split`)**: {sum(1 for w in report_wards if w["sys_status"]=="split")} Phường/Xã.')
md.append(f'- **Tổng sản lượng giao toàn vùng NTB**: **{sum(w["sys_dem"] for w in report_wards):,.1f} đơn/ngày** ({sum(w["sys_kg"] for w in report_wards):,.1f} kg/ngày).\n')

md.append('## II. RÀ SOÁT THEO ĐƠN VỊ HÀNH CHÍNH MỚI')
md.append('### 1. Đánh giá độ phủ Bưu cục và ranh giới hành chính mới')
md.append('Sau khi sáp nhập các xã cũ thành xã/phường mới, ranh giới quản lý của các Bưu cục đang xuất hiện hiện tượng chia cắt địa bàn:')
md.append('- Các phường/xã có status `split` đang bị 2 - 3 Bưu cục cùng giao hàng, làm shipper di chuyển cắt ngang địa bàn.')
md.append('- Một số tuyến đi chéo xa ranh giới giữa các huyện/tỉnh lân cận.\n')

md.append('### 2. Danh sách 18 Tuyến giao chéo xa ranh giới cần Reassign ngay')
md.append('Bảng dưới đây liệt kê 18 tuyến xã cũ bị đi chéo ranh giới.')
md.append('> ⚠️ **ĐÁNH GIÁ VẬN HÀNH & BẤT HỢP LÝ CỦA CÁC TUYẾN CHÉO KHÁC TỈNH (ĐẶC BIỆT XÃ ĐA MI)**:')
md.append('> 1. **Trường hợp Xã Đa Mi (Bình Thuận)**: Thuật toán quét khoảng cách gợi ý chuyển về Bưu cục `(LDO) Bảo Lâm 3` (tỉnh Lâm Đồng) vì khoảng cách 15.1km. Tuy nhiên, phương án này **HOÀN TOÀN SAI VỀ MẶT HÀNH CHÍNH VÀ VẬN HÀNH**:')
md.append('>    - **Vi phạm ranh giới tỉnh**: Xã Đa Mi thuộc tỉnh Bình Thuận, việc chuyển cho BC Bảo Lâm 3 (tỉnh Lâm Đồng) phụ trách sẽ gây sai lệch luồng chia chọn kho inter-provincial, sai lệch đối soát COD và báo cáo hành chính tỉnh.')
md.append('>    - **Khác AM quản lý**: (BTH) Hàm Thuận thuộc AM Nguyễn Ngọc Khánh, trong khi (LDO) Bảo Lâm 3 thuộc AM Hồng Bích Nga.')
md.append('>    - **Kế hoạch chính thức của AM**: AM Nguyễn Ngọc Khánh đã quy hoạch **(BTH) Hàm Thuận (Mới)** giữ nguyên phụ trách Xã Đa Mi cùng cụm Đông Giang, La Dạ, Thuận Hòa, Hồng Sơn thuộc Bình Thuận.')
md.append('> 2. **Các tuyến chéo Đắk Nông & Liên tỉnh khác**: AM quản lý khu vực đề xuất **TẠM THỜI GIỮ NGUYÊN 100% PHẠM VI QUẢN LÝ THEO XÃ CŨ (trong 3-6 tháng)** để tránh rủi ro 100% nhân sự tại chỗ nghỉ việc và vỡ tuyến giao.\n')

md.append('| STT | Tên Xã/Phường cũ | Tỉnh | BC hiện tại (`from`) | Khoảng cách cũ | BC tối ưu gần nhất (`to`) | Quản lý AM tiếp nhận | Khoảng cách mới | Khoảng cách tiết kiệm | Đề xuất AM |')
md.append('|---|---|---|---|---|---|---|---|---|---|')

for r in doc2.tables[1].rows[1:]:
    cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
    xa = cells[1]
    prov = cells[2]
    bc_from = cells[3]
    bc_to = cells[5]
    
    p_from = bc_from.split(')')[0].replace('(', '').strip()[:2]
    p_to = bc_to.split(')')[0].replace('(', '').strip()[:2]
    
    if xa == "Xã Đa Mi":
        am_note = f"GIỮ NGUYÊN {bc_from} (Sai ranh giới tỉnh BTH->LDO)"
    elif p_from != p_to:
        am_note = f"GIỮ NGUYÊN {bc_from} (Tuyến khác tỉnh {p_from}->{p_to})"
    elif "Đắk Nông" in prov or "Lâm Đồng" in prov:
        am_note = f"GIỮ NGUYÊN {bc_from} (Tránh vỡ tuyến DNO/LDO)"
    else:
        am_note = "Reassign tối ưu"
        
    md.append(f'| {cells[0]} | {cells[1]} | {cells[2]} | {cells[3]} | {cells[4]} | {cells[5]} | {cells[6]} | {cells[7]} | {cells[8]} | {am_note} |')

md.append('\n## III. ĐÁNH GIÁ CHI TIẾT SẢN LƯỢNG VÀ ĐỀ XUẤT PHƯƠNG ÁN THEO SỐ LIỆU CHÍNH THỨC WEB (quyhoachbuucuc.info)\n')

for i, w in enumerate(report_wards, 1):
    md.append(f'### {i}. {w["name"]} ({w["prov"]})')
    md.append(f'- **Mã Xã mới**: `{w["code"]}`')
    md.append(f'- **SẢN LƯỢNG HỆ THỐNG WEB CHÍNH THỨC**: **{w["sys_dem"]:.1f} đơn/ngày** ({w["sys_kg"]:.1f} kg/ngày) | *Trạng thái Web: `{w["sys_status"].upper()}`*')
    
    bcs_str = ', '.join([f'"{b["name"]}"' for b in w['bc_details']])
    md.append(f'- **Các BC hiện phụ trách ({len(w["bc_details"])} BC)**: {bcs_str}')
    
    if w['olds_web']:
        old_strs = [f"{o['name']} ({o['dem']} đơn/ngày - BC: {o.get('bc_name', 'Chưa rõ')})" for o in w['olds_web']]
        md.append(f'- **Các xã cũ sáp nhập & Sản lượng từng xã (Hệ thống)**: {"; ".join(old_strs)}')
    else:
        old_str = ', '.join(w['old_communes'])
        md.append(f'- **Các xã cũ sáp nhập**: {old_str}')

    md.append('- **Tỷ lệ phân chia sản lượng thực tế (Số liệu chính thức Web)**:')
    for b in w['bc_details']:
        md.append(f'  - `{b["name"]}` (ID: {b["id"]}): **{b["dem"]:.1f} đơn/ngày** ({b["pct"]}%)')
    
    md.append(f'- **ĐỀ XUẤT PHƯƠNG ÁN HỆ THỐNG**: **{w["prop"]}**')
    md.append(f'- **LÝ DO VÀ BỐ TRÍ NHÂN SỰ CHI TIẾT (AM ghi)**:\n  {w["reason"]}')
    
    if w['img_file']:
        img_abs = os.path.join(maps_dir, w['img_file'])
        md.append(f'\n![{w["img_caption"]}]({img_abs})')
        md.append(f'*Hình {i}: {w["img_caption"]}*\n')

md.append('## IV. YÊU CẦU 3: ĐÁNH GIÁ NHU CẦU TÁCH BƯU CỤC VÀ TỐI ƯU CÔNG SUẤT KHO BÃI')
md.append('### 1. Đánh giá Bưu cục quá tải (Cần tách bớt Phường hoặc Mở rộng diện tích)')
md.append('| STT | ID Bưu cục | Tên Bưu cục | Tỉnh | Quản lý AM | Địa chỉ Bưu cục | Áp lực m² (`em2`) | Đề xuất giải pháp |')
md.append('|---|---|---|---|---|---|---|---|')

for r in doc2.tables[0].rows[1:]:
    cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
    md.append(f'| {cells[0]} | {cells[1]} | {cells[2]} | {cells[3]} | {cells[4]} | {cells[5]} | {cells[6]} | {cells[7]} |')

md.append('\n### 2. Kế hoạch Mở mới / Tách Bưu cục & Tối ưu Mạng lưới')
md.append('1. **Mở mới Bưu cục (LDO) Xuân Hương - Đà Lạt 2**: Đặt tại khu vực Phường 10 (TP. Đà Lạt), chia tải cho Bưu cục Xuân Hương cũ (phụ trách Phường 3 & Phường 10, Vol giao: 1,050 đơn/ngày, Vol lấy: 150 đơn/ngày, 8 NVPTTT + 1 NVXL).')
md.append('2. **Tách Bưu cục Hàng Nhỏ / Hàng Vừa Di Linh (Xã Đinh Trang Thượng)**: Phụ trách Đinh Trang Thượng, Di Linh, Phúc Thọ Lâm Hà, Liên Đầm để giảm bán kính di chuyển 22-45km cho BC Di Linh.')
md.append('3. **Mở mới Bưu cục Đông Hải (Tỉnh Ninh Thuận)**: Cover khu vực ven biển Đông Hải, tối ưu điểm tập kết sản lượng lấy (600 đơn giao, 250 đơn lấy, 7 NVPTTT + 1 NVXL).')
md.append('4. **Mở mới Bưu cục (LDO) B\'Lao Mới (Bảo Lộc)**: Cover Phường 1 & Phường B\'Lao (1,500 đơn giao, 1,000 đơn lấy, 15 NVPTTT + 2 NVXL), đóng cửa BC (LDO) 1 Bảo Lộc cũ.\n')

md.append('## V. TỔNG HỢP BIẾN ĐỘNG MẠNG LƯỚI BƯU CỤC NTB 2026')
md.append('- **Bưu cục Mở mới (04 BC)**: BC Xuân Hương - Đà Lạt 2, BC Di Linh Hàng Nhỏ, BC Đông Hải, BC B\'Lao Mới.')
md.append('- **Bưu cục Đóng cửa (02 BC)**: BC 1 Bảo Lộc, BC Nam Nha Trang 3.')
md.append('- **Bưu cục Gộp tuyến/Điều chỉnh phân vùng (19 BC)**: Gom các tuyến xã lẻ về BC trung tâm theo ĐVHC mới.')
md.append('- **Bưu cục Giữ nguyên vận hành (13 BC)**: Duy trì tại các khu vực đặc thù Gia Nghĩa, Đắc Nông, Đà Lạt, Đức Trọng, Phan Rí Cửa, Ninh Hòa, Xã Đa Mi (Bình Thuận)...')

full_md_text = '\n'.join(md)

# Write Markdown report to Artifact path
artifact_md_path = r'C:\Users\lap4all\.gemini\antigravity-ide\brain\d95894ad-9248-47c3-872b-72d7de66bc20\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi.md'
with open(artifact_md_path, 'w', encoding='utf-8') as f:
    f.write(full_md_text)

print(f'Saved System Thien Markdown report to: {artifact_md_path}')

# Write Word DOCX report
new_doc = docx.Document()

# Style configuration
style_normal = new_doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(10.5)

def add_styled_heading(text, level):
    h = new_doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    run = h.runs[0]
    run.font.name = 'Arial'
    if level == 0:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 51, 51)
    return h

def set_table_styling(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(r'<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        if r_idx == 0:
            trPr.append(parse_xml(r'<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="003366"/>'))
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(9)
        else:
            bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{bg_color}"/>'))
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8.5)

def insert_table_18_routes_inline(doc_table_source):
    t_doc = new_doc.add_table(rows=len(doc_table_source.rows), cols=10)
    for r_idx, r in enumerate(doc_table_source.rows):
        cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
        if r_idx == 0:
            row_data = cells + ['Đề xuất AM']
        else:
            xa = cells[1]
            prov = cells[2]
            bc_from = cells[3]
            bc_to = cells[5]
            
            p_from = bc_from.split(')')[0].replace('(', '').strip()[:2]
            p_to = bc_to.split(')')[0].replace('(', '').strip()[:2]
            
            if xa == "Xã Đa Mi":
                am_note = f"GIỮ NGUYÊN {bc_from} (Sai ranh giới tỉnh BTH->LDO)"
            elif p_from != p_to:
                am_note = f"GIỮ NGUYÊN {bc_from} (Tuyến khác tỉnh {p_from}->{p_to})"
            elif "Đắk Nông" in prov or "Lâm Đồng" in prov:
                am_note = f"GIỮ NGUYÊN {bc_from} (Tránh vỡ tuyến DNO/LDO)"
            else:
                am_note = "Reassign tối ưu"
            row_data = cells + [am_note]
        for c_idx, val in enumerate(row_data):
            t_doc.cell(r_idx, c_idx).text = val
    set_table_styling(t_doc)

def insert_table_inline(doc_table_source):
    t_doc = new_doc.add_table(rows=len(doc_table_source.rows), cols=len(doc_table_source.columns))
    for r_idx, r in enumerate(doc_table_source.rows):
        for c_idx, c in enumerate(r.cells):
            t_doc.cell(r_idx, c_idx).text = c.text.strip().replace('\n', ' ')
    set_table_styling(t_doc)

# Add Title
add_styled_heading('BÁO CÁO TOÀN DIỆN: RÀ SOÁT VÀ QUY HOẠCH BƯU CỤC THEO ĐƠN VỊ HÀNH CHÍNH MỚI (VÙNG NTB)', level=0)

i = 0
lines = md[1:]
while i < len(lines):
    line = lines[i].strip()
    if line.startswith('## '):
        add_styled_heading(line.replace('## ', ''), level=1)
    elif line.startswith('### '):
        heading_text = line.replace('### ', '')
        add_styled_heading(heading_text, level=2)
        
        if '2. Danh sách 18 Tuyến giao chéo' in heading_text:
            insert_table_18_routes_inline(doc2.tables[1])
        elif '1. Đánh giá Bưu cục quá tải' in heading_text:
            insert_table_inline(doc2.tables[0])

    elif line.startswith('#### '):
        heading_text = line.replace('#### ', '')
        add_styled_heading(heading_text, level=3)

    elif '![' in line:
        import re
        m = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if m:
            cap_text, img_p = m.group(1), m.group(2)
            if os.path.exists(img_p):
                p_img = new_doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_p, width=Inches(5.5))
                
                p_cap = new_doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_cap = p_cap.add_run(f'Hình: {cap_text}')
                r_cap.font.italic = True
                r_cap.font.size = Pt(9.5)
                r_cap.font.color.rgb = RGBColor(100, 100, 100)
    elif line.startswith('*Hình '):
        pass
    elif line.startswith('|') and '---' in line:
        pass
    elif line.startswith('|'):
        pass
    elif line.startswith('- '):
        p_b = new_doc.add_paragraph(style='List Bullet')
        txt = line.replace('- ', '')
        p_b.add_run(txt)
    elif line.startswith('  - '):
        p_b2 = new_doc.add_paragraph(style='List Bullet 2')
        txt = line.replace('  - ', '')
        p_b2.add_run(txt)
    elif line.strip():
        p_txt = new_doc.add_paragraph(line.strip())
        p_txt.paragraph_format.space_after = Pt(4)
    i += 1

docx_out_path = r'C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_Buu_Cuc_NTB_Don_Vi_Hanh_Chinh_Moi_Chuan_Web_Anh_Thien.docx'
new_doc.save(docx_out_path)
print(f'Saved OFFICIAL SYSTEM THIEN DOCX report to: {docx_out_path}')
