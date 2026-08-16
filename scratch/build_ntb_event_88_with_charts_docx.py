# -*- coding: utf-8 -*-
import sys, os, docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import pandas as pd
import numpy as np

sys.stdout.reconfigure(encoding='utf-8')

ntb_file = r'C:\Users\lap4all\Downloads\config_psbba_NTB.xlsx'
fresh_charts_dir = r'c:\Users\lap4all\Documents\Auto report\scratch\ntb_fresh_charts'

output_docx_paths = [
    r'C:\Users\lap4all\Downloads\NTB Kế hoạch Event 8.8.docx',
    r'C:\Users\lap4all\Downloads\NTB Kế hoạch Event 8.8_v2.docx',
    r'c:\Users\lap4all\Documents\Auto report\NTB Kế hoạch Event 8.8.docx',
    r'c:\Users\lap4all\Documents\Auto report\NTB Kế hoạch Event 8.8_v2.docx'
]

# 1. Load Data
df_lay = pd.read_excel(ntb_file, sheet_name='6_FC_Lay_Daily')
df_giao = pd.read_excel(ntb_file, sheet_name='7_FC_Giao_Daily')

date_cols_lay = [c for c in df_lay.columns if c not in ['Vùng', 'Tỉnh/Quận', 'ID', 'BC', 'Sàn', 'Tổng 60d']]
date_cols_giao = [c for c in df_giao.columns if c not in ['Vùng', 'Tỉnh/Quận', 'ID', 'BC', 'Sàn', 'Tổng 60d']]

days10_lay = [c for c in date_cols_lay if any(d in c for d in ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08'])][:10]
days10_giao = [c for c in date_cols_giao if any(d in c for d in ['06/08', '07/08', '08/08', '09/08', '10/08', '11/08', '12/08', '13/08', '14/08', '15/08'])][:10]

for c in days10_lay:
    df_lay[c] = pd.to_numeric(df_lay[c], errors='coerce').fillna(0)
for c in days10_giao:
    df_giao[c] = pd.to_numeric(df_giao[c], errors='coerce').fillna(0)

df_lay = df_lay.dropna(subset=['Sàn'])
df_giao = df_giao.dropna(subset=['Sàn'])

dates_header_10 = [c.split()[-1] for c in days10_lay]

# Create Document
doc = docx.Document()

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
        tblPr[0].append(borders)

def format_cell(cell, text, bold=False, italic=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=(0,0,0), bg_hex=None):
    cell.text = str(text)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        r = p.runs[0]
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(*color_rgb)
    if bg_hex:
        set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(15, 76, 129) # Ocean Cyan Blue
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0, 128, 128) # Teal Accent
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(38, 38, 38)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

def add_p(doc, text="", bold=False, italic=False, font_size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(font_size)
        r.font.bold = bold
        r.font.italic = italic
    return p

def add_centered_picture(doc, img_path, width_in_inches=6.2, caption_text=""):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_in_inches))
        if caption_text:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r_cap = p_cap.add_run(caption_text)
            r_cap.font.name = 'Arial'
            r_cap.font.size = Pt(9)
            r_cap.font.italic = True
            r_cap.font.color.rgb = RGBColor(89, 89, 89)

# TITLE
p_title = add_p(doc, "KẾ HOẠCH EVENT 8.8 - VÙNG NAM TRUNG BỘ (NTB)", bold=True, font_size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
p_title.runs[0].font.color.rgb = RGBColor(15, 76, 129)

add_p(doc, "Thời gian thực hiện: Giai đoạn Event 08/08 (Peak Forecast 06/08 - 15/08/2026)", italic=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# SECTION I
add_styled_heading(doc, "I. Mục tiêu:", level=1)

# 1. Volume Lấy
add_styled_heading(doc, "1. Volume Lấy", level=2)
add_styled_heading(doc, "Tổng quan", level=3)

# FRESH COMBO CHART FOR LAY
add_centered_picture(doc, os.path.join(fresh_charts_dir, 'fresh_combo_lay_88.png'), width_in_inches=6.4, caption_text="Biểu đồ Tổng quan FC Volume Lấy event 08.08 (Style Mới Nam Trung Bộ)")

# Table 1: Volume Lấy theo Sàn
tbl1_df = df_lay.groupby('Sàn')[days10_lay].sum().round(0).astype(int)
san_order = ['Shopee', 'Shopee-Bulky', 'Shopee-Bulky (10-15kg)', 'SME', 'SME-Bulky', 'TTS', 'TTS-Bulky']
tbl1_df = tbl1_df.reindex([s for s in san_order if s in tbl1_df.index])
tbl1_df.loc['Grand Total'] = tbl1_df.sum()

t1 = doc.add_table(rows=len(tbl1_df)+1, cols=11)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t1)

headers_t1 = ['Sàn'] + dates_header_10
for col_idx, h in enumerate(headers_t1):
    format_cell(t1.rows[0].cells[col_idx], h, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="0F4C81")

for row_idx, (san_name, row_data) in enumerate(tbl1_df.iterrows()):
    cell_row = t1.rows[row_idx+1]
    is_gt = (san_name == 'Grand Total')
    bg = "F2F2F2" if is_gt else None
    format_cell(cell_row.cells[0], san_name, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_centered_picture(doc, os.path.join(fresh_charts_dir, 'fresh_lay_san.png'), width_in_inches=6.2, caption_text="Biểu đồ 1: Chi tiết Sản lượng Lấy theo Sàn (Event 8.8 NTB)")

add_p(doc, "", space_after=6)

# Table 2: Volume Lấy theo Tỉnh
tbl2_df = df_lay.groupby('Tỉnh/Quận')[days10_lay].sum().round(0).astype(int)
tbl2_df.loc['Grand Total'] = tbl2_df.sum()

t2 = doc.add_table(rows=len(tbl2_df)+1, cols=11)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t2)

headers_t2 = ['Tỉnh/Quận'] + dates_header_10
for col_idx, h in enumerate(headers_t2):
    format_cell(t2.rows[0].cells[col_idx], h, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="0F4C81")

for row_idx, (tinh_name, row_data) in enumerate(tbl2_df.iterrows()):
    cell_row = t2.rows[row_idx+1]
    is_gt = (tinh_name == 'Grand Total')
    bg = "F2F2F2" if is_gt else None
    format_cell(cell_row.cells[0], tinh_name, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_centered_picture(doc, os.path.join(fresh_charts_dir, 'fresh_lay_tinh.png'), width_in_inches=6.2, caption_text="Biểu đồ 2: Sản lượng Lấy theo Tỉnh/Quận (Event 8.8 NTB)")

add_p(doc, "", space_after=8)

# 2. Volume Giao
add_styled_heading(doc, "2. Volume Giao", level=2)
add_styled_heading(doc, "Tổng quan & Chi tiết loại hàng", level=3)

# FRESH COMBO CHART FOR GIAO
add_centered_picture(doc, os.path.join(fresh_charts_dir, 'fresh_combo_giao_88.png'), width_in_inches=6.4, caption_text="Biểu đồ Tổng quan FC Volume Giao event 08.08 (Style Mới Nam Trung Bộ)")

# Table 3: Volume Giao theo Sàn
tbl3_df = df_giao.groupby('Sàn')[days10_giao].sum().round(0).astype(int)
tbl3_df = tbl3_df.reindex([s for s in san_order if s in tbl3_df.index])
tbl3_df.loc['Grand Total'] = tbl3_df.sum()

t3 = doc.add_table(rows=len(tbl3_df)+1, cols=11)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t3)

headers_t3 = ['Sàn'] + dates_header_10
for col_idx, h in enumerate(headers_t3):
    format_cell(t3.rows[0].cells[col_idx], h, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="0F4C81")

for row_idx, (san_name, row_data) in enumerate(tbl3_df.iterrows()):
    cell_row = t3.rows[row_idx+1]
    is_gt = (san_name == 'Grand Total')
    bg = "F2F2F2" if is_gt else None
    format_cell(cell_row.cells[0], san_name, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_centered_picture(doc, os.path.join(fresh_charts_dir, 'fresh_giao_san.png'), width_in_inches=6.2, caption_text="Biểu đồ 3: Chi tiết Sản lượng Giao theo Sàn (Event 8.8 NTB)")

add_p(doc, "", space_after=6)

# Table 4: Volume Giao theo Tỉnh
tbl4_df = df_giao.groupby('Tỉnh/Quận')[days10_giao].sum().round(0).astype(int)
tbl4_df.loc['Grand Total'] = tbl4_df.sum()

t4 = doc.add_table(rows=len(tbl4_df)+1, cols=11)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t4)

headers_t4 = ['Tỉnh/Quận'] + dates_header_10
for col_idx, h in enumerate(headers_t4):
    format_cell(t4.rows[0].cells[col_idx], h, bold=True, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=(255,255,255), bg_hex="0F4C81")

for row_idx, (tinh_name, row_data) in enumerate(tbl4_df.iterrows()):
    cell_row = t4.rows[row_idx+1]
    is_gt = (tinh_name == 'Grand Total')
    bg = "F2F2F2" if is_gt else None
    format_cell(cell_row.cells[0], tinh_name, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    for c_idx, val in enumerate(row_data):
        val_str = f"{val:,}"
        format_cell(cell_row.cells[c_idx+1], val_str, bold=is_gt, font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, bg_hex=bg)

add_centered_picture(doc, os.path.join(fresh_charts_dir, 'fresh_giao_tinh.png'), width_in_inches=6.2, caption_text="Biểu đồ 4: Sản lượng Giao theo Tỉnh/Quận (Event 8.8 NTB)")

# SECTION II
add_styled_heading(doc, "II. PHÂN TÍCH CHI TIẾT CÁC NHÓM BƯU CỤC:", level=1)

# Table 5
t5 = doc.add_table(rows=7, cols=4)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t5)

t5_data = [
    ["Đặc điểm", "Nhóm 1: Ổn định", "Nhóm 2: Cảnh báo", "Nhóm 3: Bất ổn"],
    ["Số lượng BC", "68 (82,93%)", "08 (9,76%)", "06 (7,31%)"],
    ["Mức độ rủi ro", "Thấp", "Trung bình", "Rất cao (Quá tải và cần tập trung theo dõi trong kỳ Event 8.8)"],
    ["Thực trạng Nhân sự", "- Thiếu dưới 2 NVPTTT\n- Đội ngũ ổn định", "- Thiếu hiện tại 1-4 NV\n- Đang có điều động hỗ trợ", "- Thiếu hiện tại 2-5 NV\n- Nằm trong danh sách rủi ro vận hành địa bàn đồi dốc/biển đảo"],
    ["Khả năng kiểm soát", "Tốt", "- Kiểm soát trung bình khá\n- Phụ thuộc tình hình thời tiết bão/mưa dốc", "- Khả năng kiểm soát kém nếu không can thiệp khẩn cấp"],
    ["Phương án A", "- AM theo dõi, điều hành gán, FIFO và bám sát năng suất NV.", "- Đẩy mạnh tuyển dụng bổ sung.\n- AM trực tiếp cắm BC điều hành gán & FIFO.\n- Thuê Freelance gánh ca peak 8.8.", "- AM trực tiếp tại BC điều phối 100%.\n- Thúc đẩy chính sách thưởng clear tồn.\n- Điều động nhân sự BC lân cận hỗ trợ."],
    ["Phương án B (khi PA A không thể thực hiện)", "- Tắt tuyến BC (khi vượt CAP x3)\n- Điều tiết giảm hàng KA", "- Tắt tuyến BC (khi vượt CAP x3)\n- Giữ hàng tại KTC", "- Tắt tuyến BC khẩn cấp\n- Điều tiết hàng KA giữ tại Kho TC Nha Trang / Lâm Đồng"]
]

for r_i, row in enumerate(t5_data):
    for c_i, val in enumerate(row):
        bg = "0F4C81" if r_i == 0 else ("F2F2F2" if c_i == 0 else None)
        bold = True if (r_i == 0 or c_i == 0) else False
        color = (255,255,255) if r_i == 0 else (0,0,0)
        align = WD_ALIGN_PARAGRAPH.CENTER if (r_i == 0 or c_i == 0) else WD_ALIGN_PARAGRAPH.LEFT
        format_cell(t5.rows[r_i].cells[c_i], val, bold=bold, font_size=8.5, align=align, color_rgb=color, bg_hex=bg)

add_p(doc, "", space_after=8)

# Save Main Event Document
saved = False
for p in output_docx_paths:
    try:
        doc.save(p)
        print(f"Successfully saved Main Event 8.8 Word Document: {p}")
        saved = True
    except Exception as e:
        print(f"Could not save {p}: {e}")

if saved:
    print("Main Word document updated with fresh modern charts successfully!")
