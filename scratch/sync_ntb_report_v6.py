import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

src_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_6.docx"
dst_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_7.docx"

if not os.path.exists(src_path):
    print(f"File nguồn không tồn tại: {src_path}")
    sys.exit(1)

doc = docx.Document(src_path)

print("=== BẮT ĐẦU CẬP NHẬT ĐỢT 7 ===")

# 1. Cập nhật Bảng 2 (index 1) Row 3 Col 1 (Nhân sự T7->T12)
table_2 = doc.tables[1]
table_2.rows[3].cells[1].text = "92 người (T7) → 138 người (T12) — linh hoạt theo sản lượng thực tế"
print("-> Đã sửa Nhân sự tóm tắt ở Bảng 2 thành 92 -> 138")

# 2. Cập nhật Bảng 7 (index 6) Row 8 (NV kho + QL) và Row 9 (TỔNG NHÂN SỰ)
table_7 = doc.tables[6]
table_7.rows[8].cells[0].text = "NV xử lý kho (12) + NV QL (4)"
for c_idx in range(1, 7):
    table_7.rows[8].cells[c_idx].text = "16"

# Tổng nhân sự toàn vùng (Row 9)
tot_ns = ["92", "94", "124", "126", "136", "138"]
for c_idx in range(1, 7):
    table_7.rows[9].cells[c_idx].text = tot_ns[c_idx - 1]
print("-> Đã sửa định biên kho + QL thành 16 người và tính lại Tổng nhân sự ở Bảng 7")

# 3. Cập nhật Bảng 14 (index 13) Hàng 4 (Tuyển dụng)
table_14 = doc.tables[13]
table_14.rows[4].cells[1].text = "Tuyển dụng đợt 1: 76 NV giao + 16 NV kho/QL cho T7"
print("-> Đã sửa mô tả tuyển dụng ở Bảng 14")

doc.save(dst_path)
print(f"=== ĐÃ LƯU FILE THÀNH CÔNG: {dst_path} ===")
