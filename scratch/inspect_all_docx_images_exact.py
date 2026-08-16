import docx, os, sys, hashlib

sys.stdout.reconfigure(encoding='utf-8')

src_docx = r'C:\Users\lap4all\Downloads\Quy_Hoach MANG LUOI NTB.docx'
out_dir = r'C:\Users\lap4all\Documents\Auto report\scratch\maps_exact_inspect'
os.makedirs(out_dir, exist_ok=True)

doc = docx.Document(src_docx)

extracted_maps = []
seen_hashes = set()

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    
    for run in p.runs:
        drawing_elements = run._r.xpath('.//w:drawing')
        for dr in drawing_elements:
            blip_elements = dr.xpath('.//a:blip')
            for blip in blip_elements:
                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id in doc.part.rels:
                    target_part = doc.part.rels[embed_id].target_part
                    img_bytes = target_part.blob
                    
                    h_val = hashlib.md5(img_bytes).hexdigest()
                    if h_val not in seen_hashes:
                        seen_hashes.add(h_val)
                        img_filename = f"map_{len(extracted_maps)+1}.png"
                        img_path = os.path.join(out_dir, img_filename)
                        with open(img_path, "wb") as f:
                            f.write(img_bytes)
                        
                        # Find nearby caption or text
                        nearby_text = txt
                        if not nearby_text and i > 0:
                            nearby_text = doc.paragraphs[i-1].text.strip()
                        if not nearby_text and i < len(doc.paragraphs)-1:
                            nearby_text = doc.paragraphs[i+1].text.strip()
                            
                        extracted_maps.append({
                            'index': len(extracted_maps)+1,
                            'p_index': i+1,
                            'nearby_text': nearby_text,
                            'path': img_path,
                            'hash': h_val
                        })

print(f"Extracted {len(extracted_maps)} UNIQUE map images!\n")
for m in extracted_maps:
    print(f"Image {m['index']} (at P{m['p_index']}): Nearby text = '{m['nearby_text']}'")
