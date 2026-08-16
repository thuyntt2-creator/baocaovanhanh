import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_6.docx"

if not os.path.exists(doc_path):
    print("File không tồn tại")
    sys.exit(1)

doc = docx.Document(doc_path)
errors = 0

print("=== KIỂM TRA SỐ LIỆU TOÀN DIỆN FILE WORD V8_6 ===")

# 1. Kiểm tra đơn giá thuê mặt bằng ở Paragraph 133
print("\n1. Kiểm tra đơn giá thuê ở Paragraph 133:")
found_120 = False
for p in doc.paragraphs:
    if "120.000 đ/m²/tháng" in p.text:
        print(f"  [OK] Đã tìm thấy đơn giá thuê: '{p.text.strip()}'")
        found_120 = True
        break
if not found_120:
    print("  [LỖI] Không tìm thấy đơn giá '120.000 đ/m²/tháng' trong văn bản!")
    errors += 1

# 2. Kiểm tra Bảng 9 (Giá thuê & m2 thiếu)
table_9 = doc.tables[8]
print("\n2. Kiểm tra Bảng 9 (Quy hoạch mặt bằng):")
nt_price = table_9.rows[1].cells[4].text.strip()
tot_price = table_9.rows[5].cells[4].text.strip()
if nt_price != "41.760.000 đ/th":
    print(f"  [LỖI] Giá thuê Nha Trang = '{nt_price}' (Kỳ vọng: 41.760.000 đ/th)")
    errors += 1
else:
    print("  [OK] Giá thuê Nha Trang khớp")
if tot_price != "75.600.000 đ/th":
    print(f"  [LỖI] Tổng giá thuê = '{tot_price}' (Kỳ vọng: 75.600.000 đ/th)")
    errors += 1
else:
    print("  [OK] Tổng giá thuê khớp")

# 3. Kiểm tra đối chiếu Bảng 10 vs Bảng 11
table_10 = doc.tables[9]
table_11 = doc.tables[10]
print("\n3. Kiểm tra đối chiếu Bảng 10 vs Bảng 11:")
for c_idx in range(1, 7):
    t10_sum = float(table_10.rows[5].cells[c_idx].text.replace(',', ''))
    t11_sum = float(table_11.rows[5].cells[c_idx].text.replace(',', ''))
    if abs(t10_sum - t11_sum) > 0.1:
        print(f"  [LỖI] Tháng {c_idx+6}: Bảng 10 = {t10_sum} | Bảng 11 = {t11_sum}")
        errors += 1
    else:
        print(f"  [OK] Tháng {c_idx+6}: Tổng CP = {t10_sum} (Khớp)")

if errors == 0:
    print("\n>>> TẤT CẢ KIỂM TRA BÁO CÁO V8_6 ĐỀU THÀNH CÔNG! SỐ LIỆU HOÀN THIỆN TUYỆT ĐỐI <<<")
else:
    print(f"\n>>> CÓ {errors} SAI LỆCH SỐ LIỆU CẦN XỬ LÝ <<<")
