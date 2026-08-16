# -*- coding: utf-8 -*-
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_styled_paragraph(doc, text, bold=False, italic=False, font_size=10, color_rgb=(31,31,31), space_before=2, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*color_rgb)
    return p

def format_cell(cell, text, bold=False, italic=False, font_size=9.5, color_rgb=(31,31,31), bg_hex=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    if bg_hex:
        set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    cell.text = str(text)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    if len(p.runs) > 0:
        run = p.runs[0]
        run.bold = bold
        run.italic = italic
        run.font.name = 'Calibri'
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(*color_rgb)

def generate_week_report(week_num, date_str, prev_week_num, out_filename):
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Header Title
    add_styled_paragraph(doc, f'BÁO CÁO KẾT QUẢ KINH DOANH TUẦN {week_num} – VÙNG NAM TRUNG BỘ', bold=True, font_size=16, color_rgb=(31, 56, 100), space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, f'Thời gian báo cáo: {date_str} | Kỳ so sánh: Tuần {prev_week_num}', italic=True, font_size=10.5, color_rgb=(89, 89, 89), space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    # -------------------------------------------------------------------------
    # 1. TỔNG QUAN KẾT QUẢ KINH DOANH
    # -------------------------------------------------------------------------
    add_styled_paragraph(doc, '1. TỔNG QUAN KẾT QUẢ KINH DOANH', bold=True, font_size=13, color_rgb=(31, 56, 100), space_before=6, space_after=4)
    
    # 1.1
    add_styled_paragraph(doc, '1.1. Mức độ hoàn thành kế hoạch', bold=True, font_size=11, color_rgb=(46, 117, 182), space_before=4, space_after=4)

    # Table 1
    t1 = doc.add_table(rows=5, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1_headers = ['Chỉ tiêu', '% Hoàn thành target', f'% Tăng/ giảm so với Tuần {prev_week_num}']
    
    for j, h in enumerate(t1_headers):
        format_cell(t1.cell(0, j), h, bold=True, font_size=10, color_rgb=(255, 255, 255), bg_hex='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER)

    if week_num == 28:
        t1_rows = [
            ['Tổng doanh thu LTC', '92.5%', '-0.81% (1.288 tỷ vs 1.298 tỷ VND)'],
            ['Tổng doanh thu GTTC', '94.0%', '+37.4% (1.187 tỷ vs 864.3 triệu VND)'],
            ['Doanh thu giữ cũ (KH duy trì / thăng hạng)', '93.2%', '-0.90% (1.275 tỷ vs 1.287 tỷ VND)'],
            ['Doanh thu bán mới (KH mới trong tháng)', '87.0%', '-8.93% (13.05 triệu vs 14.33 triệu VND)'],
        ]
    else: # Week 29
        t1_rows = [
            ['Tổng doanh thu LTC', '94.8%', '+2.50% (1.217 tỷ vs 1.187 tỷ VND GTTC)'],
            ['Tổng doanh thu GTTC', '96.2%', '+2.50% (1.217 tỷ vs 1.187 tỷ VND)'],
            ['Doanh thu giữ cũ (KH duy trì / thăng hạng)', '95.1%', '+2.41% (1.205 tỷ vs 1.174 tỷ VND)'],
            ['Doanh thu bán mới (KH mới trong tháng)', '84.5%', '-7.99% (12.01 triệu vs 13.05 triệu VND)'],
        ]

    for i, r_data in enumerate(t1_rows):
        bg = 'F9FBFD' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(r_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(t1.cell(i+1, j), val, font_size=9.5, bg_hex=bg, align=align)

    add_styled_paragraph(doc, '', space_after=4)

    # 1.2 Kết quả Giữ cũ
    add_styled_paragraph(doc, '1.2. Kết quả Giữ cũ', bold=True, font_size=11, color_rgb=(46, 117, 182), space_before=4, space_after=4)
    add_styled_paragraph(doc, 'Nhận định chung:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)

    if week_num == 28:
        nhan_dinh_giucu = (
            'Trong Tuần 28, doanh thu giữ cũ của Vùng Nam Trung Bộ đạt 1.275 tỷ VND (-0.9% so với Tuần 27). '
            'Cơ cấu nhóm KH duy trì tốt ở phân hạng EF (+30 KH, +9.9% doanh thu) và phân hạng G (+45 KH, +57.1% doanh thu). '
            'Tuy nhiên, tổng sản lượng nhóm KH lớn (Nhóm A & BCD) ghi nhận xu hướng sụt giảm nhẹ (-15.3% volume ở nhóm A, -4.4% volume ở nhóm BCD). '
            'Đáng chú ý, toàn Vùng ghi nhận 118 KH có dấu hiệu sụt giảm doanh thu/sản lượng >30% so với cam kết (trong đó có 2 KH nhóm A trọng điểm và 23 KH nhóm BCD).'
        )
    else: # Week 29
        nhan_dinh_giucu = (
            'Trong Tuần 29, doanh thu giữ cũ phục hồi tích cực lên 1.205 tỷ VND (ước tính GTTC +2.41% so với Tuần 28). '
            'Phân hạng BCD ghi nhận sự bứt phá ấn tượng với sản lượng đạt 10,101 đơn (+11.3% so với W28) và doanh thu nhóm duy trì đạt 41.38 triệu VND (+43.7%). '
            'Nhóm KH A giữ vững quy mô 10/10 KH có phát sinh đơn với tổng sản lượng đạt 15,836 đơn (duy trì 95.1% so với W28). '
            'Toàn Vùng ghi nhận 119 KH nằm trong danh mục cảnh báo sụt giảm >30% (tương đương W28 với 2 KH nhóm A và 23 KH nhóm BCD).'
        )
    add_styled_paragraph(doc, nhan_dinh_giucu, font_size=9.5, space_after=4)

    add_styled_paragraph(doc, 'Số liệu & Các điểm nổi bật:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)
    
    if week_num == 28:
        pts_giucu = [
            'Phân loại nhóm Khách hàng: Nhóm A đạt 10 KH (16,647 đơn, 524.4tr VND, -8.2% DT vs W27); Nhóm BCD đạt 70 KH (9,080 đơn, 260.4tr VND, -5.6% DT vs W27); Nhóm EF đạt 1,525 KH (12,199 đơn, 476.8tr VND, +9.9% DT vs W27); Nhóm G đạt 196 KH (744 đơn, 26.5tr VND, +57.1% DT vs W27).',
            'Nhóm KH nguy cơ rời bỏ / sụt giảm >30%: 118 KH (Lâm Đồng: 33 KH, Bình Thuận: 28 KH, Khánh Hòa: 27 KH, Ninh Thuận: 16 KH, Đắk Nông: 13 KH). Phân hạng: 2 KH A, 23 KH BCD, 92 KH EF, 1 KH G.',
            'Nguyên nhân chính: (1) Cạnh tranh về giá cước từ đối thủ tại khu vực Bình Thuận & Lâm Đồng; (2) Tình trạng mùa vụ kinh doanh của một số shop nông sản / hải sản khô tại Khánh Hòa & Ninh Thuận; (3) Thời gian giao hàng Ca 1 ở một số bưu cục chưa tối ưu dẫn đến shop chuyển một phần đơn qua đơn vị vận chuyển khác.',
        ]
    else: # Week 29
        pts_giucu = [
            'Phân loại nhóm Khách hàng: Nhóm A đạt 10 KH (15,836 đơn, 76.67tr VND ghi nhận log, tổng GTTC giữ ở mức cao); Nhóm BCD đạt 66 KH (10,101 đơn, +11.3% volume vs W28); Nhóm EF đạt 1,476 KH (12,225 đơn, +0.2% volume vs W28); Nhóm G đạt 178 KH (551 đơn).',
            'Nhóm KH nguy cơ rời bỏ / sụt giảm >30%: 119 KH (Lâm Đồng: 33 KH, Bình Thuận: 29 KH, Khánh Hòa: 27 KH, Ninh Thuận: 16 KH, Đắk Nông: 13 KH). Phân hạng: 2 KH A, 23 KH BCD, 92 KH EF, 2 KH G.',
            'Nguyên nhân chính: (1) Shop TIÊN HUỲNH US (Ninh Thuận) giảm sản lượng đơn hàng nhập khẩu do vướng thủ tục hải quan; (2) Shop Ny tân thành (Bình Thuận) điều chỉnh kế hoạch livestream bán hàng; (3) Tỷ lệ giao công nghiệp đúng giờ tại bưu cục Lâm Đồng đang được khắc phục giúp giữ chân nhóm EF.',
        ]

    for pt in pts_giucu:
        add_styled_paragraph(doc, f'• {pt}', font_size=9.5, space_after=3)

    # 1.3 Kết quả Bán mới
    add_styled_paragraph(doc, '1.3. Kết quả Bán mới', bold=True, font_size=11, color_rgb=(46, 117, 182), space_before=4, space_after=4)
    add_styled_paragraph(doc, 'Nhận định chung:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)

    if week_num == 28:
        nhan_dinh_banmoi = (
            'Tuần 28 ghi nhận 119 khách hàng mới phát sinh đơn LTC đầu tiên trên toàn Vùng Nam Trung Bộ, '
            'mang lại tổng doanh thu bán mới 13.05 triệu VND và tổng sản lượng 540 đơn hàng (bình quân 4.54 đơn/KH). '
            'Tỉnh Khánh Hòa tiếp tục dẫn đầu toàn Vùng về số lượng bán mới (36 KH, 3.26tr VND), tiếp theo là Ninh Thuận (20 KH, 2.25tr VND) và Lâm Đồng (20 KH, 3.08tr VND).'
        )
    else: # Week 29
        nhan_dinh_banmoi = (
            'Tuần 29 ghi nhận sự tăng trưởng về số lượng khách hàng mới với 123 KH mới (+4 KH so với Tuần 28), '
            'đạt tổng doanh thu bán mới 12.01 triệu VND và sản lượng 319 đơn hàng (bình quân 2.59 đơn/KH). '
            'Khánh Hòa vươn lên mốc 40 KH mới (2.73tr VND), Lâm Đồng bứt phá về doanh thu bán mới (28 KH, 4.16tr VND), khẳng định hiệu quả từ các chiến dịch săn shop tại địa bàn.'
        )

    add_styled_paragraph(doc, nhan_dinh_banmoi, font_size=9.5, space_after=4)
    add_styled_paragraph(doc, 'Số liệu & Các điểm nổi bật:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)

    if week_num == 28:
        pts_banmoi = [
            'Số lượng KH bán mới phát sinh trong tuần: 119 KH | Doanh thu: 13,049,122 VND | Sản lượng: 540 đơn | AOV: 4.54 đơn/KH.',
            'Khu vực / AM có kết quả bán mới tốt: AM Nguyễn Duy Long (Ninh Thuận) xuất sắc dẫn đầu Vùng với 20 KH mới (2.25tr VND, 177 đơn); AM Phan Đình Duy (Khánh Hòa) đạt 13 KH mới (1.75tr VND); AM Thái Thị Thanh Thư (Khánh Hòa) đạt 13 KH mới; AM Huỳnh Thị Kim Chi (Khánh Hòa) đạt AOV cao nhất (14.7 đơn/KH).',
            'Khu vực / AM chưa đạt tiến độ: AM Lê Văn Trường (Lâm Đồng) 1 KH mới; AM Lê Minh Đại (Đắk Nông) 1 KH mới.',
            'Các vấn đề đang ảnh hưởng: Tỷ lệ KH mới phát sinh đúng 1 đơn còn cao (~52%), cần đẩy mạnh hoạt động telesales và chăm sóc sau bán (onboarding) để biến KH mới thành KH duy trì bền vững.',
            'Dự kiến tuần tới: Đẩy mạnh tìm kiếm shop thời trang & đồ khô mùa du lịch tại Khánh Hòa & Ninh Thuận, mục tiêu phát sinh 125+ KH mới.'
        ]
    else: # Week 29
        pts_banmoi = [
            'Số lượng KH bán mới phát sinh trong tuần: 123 KH | Doanh thu: 12,006,960 VND | Sản lượng: 319 đơn | AOV: 2.59 đơn/KH.',
            'Khu vực / AM có kết quả bán mới tốt: AM Phạm Bá Thành Công (Lâm Đồng) bứt phá dẫn đầu Vùng với 17 KH mới (0.69tr VND); AM Nguyễn Duy Long (Ninh Thuận) duy trì phong độ cao với 15 KH mới (0.83tr VND); AM Phan Đình Duy (Khánh Hòa) đạt 11 KH mới; AM Hồng Bích Nga (Lâm Đồng) đạt doanh thu & sản lượng cao nhất (10 KH, 1.45tr VND, 71 đơn).',
            'Khu vực / AM chưa đạt tiến độ: AM Nguyễn Hoàng Phi (1 KH); AM Nguyễn Thanh Long (1 KH); AM Lê Thanh Nhựt (1 KH).',
            'Các vấn đề đang ảnh hưởng: Doanh thu trung bình/KH mới giảm nhẹ do nhiều shop nhỏ dùng thử; cần AM hỗ trợ tư vấn các gói cước đồng giá để kích thích shop tăng sản lượng.',
            'Dự kiến tuần tới: Triển khai gói ưu đãi cước mở rộng B2B cho 30 shop tiềm năng tại Lâm Đồng & Đắk Nông, mục tiêu đạt 130+ KH mới.'
        ]

    for pt in pts_banmoi:
        add_styled_paragraph(doc, f'• {pt}', font_size=9.5, space_after=3)

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # 2. ĐIỂM NỔI BẬT TRONG TUẦN - Khách hàng nhóm A
    # -------------------------------------------------------------------------
    add_styled_paragraph(doc, '2. ĐIỂM NỔI BẬT TRONG TUẦN - Khách hàng nhóm A', bold=True, font_size=13, color_rgb=(31, 56, 100), space_before=6, space_after=4)
    add_styled_paragraph(doc, 'Nhận định chính:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)

    if week_num == 28:
        nhan_dinh_kha = (
            'Nhóm KH A tại Vùng Nam Trung Bộ gồm 8 KH trọng điểm duy trì theo dõi trên hệ thống (tổng 10 KH A toàn Vùng). '
            '100% (10/10) KH A đều phát sinh đơn LTC trong tuần. Tuy nhiên, sản lượng ngày WTD-1 của 2 KH A sụt giảm nặng: '
            'Shop TIÊN HUỲNH US (Ninh Thuận, AM Nguyễn Duy Long) chỉ đạt 19.2% WTD-1 (146 đơn/ngày vs cam kết 3,000 đơn/tháng) '
            'và Shop Ny tân thành (Bình Thuận, AM Huỳnh Tấn Hiền) chỉ đạt 4.3% WTD-1 (93 đơn/ngày vs cam kết 2,900 đơn/tháng). '
            'Các KH A còn lại như Vận Chuyển Online (Khánh Hòa) và Công Ty Khởi Phát Thịnh (Khánh Hòa) vẫn giữ vững sản lượng lớn (>3,300 - 6,800 đơn/tuần).'
        )
    else: # Week 29
        nhan_dinh_kha = (
            'Trong Tuần 29, danh mục KH A ghi nhận sự phục hồi nhẹ ở nhóm giảm đơn: '
            'Shop Ny tân thành (Bình Thuận) tăng sản lượng lên 339 đơn/tuần (gấp 2 lần so với W28: 171 đơn); '
            'Shop TIÊN HUỲNH US (Ninh Thuận) đạt 755 đơn/tuần. '
            'Shop Vận Chuyển Online (Khánh Hòa, AM Phan Đình Duy) tiếp tục là đầu tàu tăng trưởng với sản lượng đạt 7,042 đơn/tuần (+2.86% vs W28). '
            'Vùng đang tiếp tục bám sát 55 KH BCD tiềm năng để thúc đẩy thăng hạng A trong tháng tới.'
        )

    add_styled_paragraph(doc, nhan_dinh_kha, font_size=9.5, space_after=4)

    # Table 2
    t2 = doc.add_table(rows=6, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2_headers = ['Chỉ tiêu', 'Dữ liệu', f'So sánh tăng/ giảm với Tuần {prev_week_num}', 'Ghi chú (nếu có)']
    
    for j, h in enumerate(t2_headers):
        format_cell(t2.cell(0, j), h, bold=True, font_size=10, color_rgb=(255, 255, 255), bg_hex='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER)

    if week_num == 28:
        t2_rows = [
            ['Tổng số lượng KH nhóm A đầu tháng', '10 KH', 'Bằng (0)', 'Duy trì ổn định từ đầu tháng 7'],
            ['Số KH có lên đơn đến hiện tại', '10 KH', 'Bằng (0)', '100% KH A có phát sinh đơn LTC'],
            ['Số KH A dự kiến giảm hạng', '2 KH', 'Bằng (0)', 'Shop TIÊN HUỲNH US & Ny tân thành mốc WTD-1 thấp'],
            ['Số KH A nguy cơ rời bỏ (không DT LTC)', '0 KH', 'Bằng (0)', 'Không có KH A ngừng hẳn đơn LTC'],
            ['Số KH tiềm năng lên hạng A', '55 KH', 'Bằng (0)', 'Nhóm BCD có doanh thu sát ngưỡng A'],
        ]
    else: # Week 29
        t2_rows = [
            ['Tổng số lượng KH nhóm A đầu tháng', '10 KH', 'Bằng (0)', 'Duy trì ổn định'],
            ['Số KH có lên đơn đến hiện tại', '10 KH', 'Bằng (0)', '100% KH A phát sinh đơn'],
            ['Số KH A dự kiến giảm hạng', '2 KH', 'Bằng (0)', 'Ny tân thành đang có tín hiệu phục hồi (+98% vol vs W28)'],
            ['Số KH A nguy cơ rời bỏ (không DT LTC)', '0 KH', 'Bằng (0)', 'Không có KH A dừng đơn'],
            ['Số KH tiềm năng lên hạng A', '55 KH', 'Bằng (0)', 'Theo dõi sát nhóm BCD tại Khánh Hòa & Lâm Đồng'],
        ]

    for i, r_data in enumerate(t2_rows):
        bg = 'F9FBFD' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(r_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in [0, 3] else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(t2.cell(i+1, j), val, font_size=9.5, bg_hex=bg, align=align)

    add_styled_paragraph(doc, '', space_after=4)

    # Chi tiết KH A rời bỏ / sụt giảm
    add_styled_paragraph(doc, 'CHI TIẾT KH NHÓM A NGUY CƠ SỤT GIẢM / RỜI BỎ', bold=True, font_size=11, color_rgb=(192, 0, 0), space_before=4, space_after=4)
    add_styled_paragraph(doc, 'Tổng quan nguyên nhân sụt giảm:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)
    add_styled_paragraph(doc, 'Vấn đề chính tập trung vào 2 khách hàng trọng điểm tại Ninh Thuận và Bình Thuận: Shop TIÊN HUỲNH US giảm 80.8% sản lượng do vướng đợt kiểm tra nguồn hàng nhập khẩu; Shop Ny tân thành giảm 95.7% sản lượng do tạm hoãn các phiên livestream bán hàng quy mô lớn. Cả 2 AM phụ trách đã làm việc trực tiếp và đưa ra giải pháp hỗ trợ đặc thù.', font_size=9.5, space_after=4)

    add_styled_paragraph(doc, 'Danh sách khách hàng nhóm A giảm đơn trọng điểm trong tuần:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)

    # Table 3
    t3 = doc.add_table(rows=3, cols=9)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3_headers = ['STT', 'Client_ID', 'Tên KH', 'AM phụ trách', 'Doanh thu / Cam kết', 'WTD-1 (% Cam kết)', 'Phân hạng', 'Lý do sụt giảm', 'Hành động / Hướng xử lý']
    
    for j, h in enumerate(t3_headers):
        format_cell(t3.cell(0, j), h, bold=True, font_size=9, color_rgb=(255, 255, 255), bg_hex='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER)

    if week_num == 28:
        t3_rows = [
            ['1', '5197975', 'TIÊN HUỲNH US', 'Nguyễn Duy Long', '879 đơn (CK 3,000/tháng)', '146 đơn/ngày (19.2%)', 'A', 'Vướng kiểm tra nguồn hàng nhập khẩu', 'Tư vấn quy trình kho & ưu đãi cước lưu kho'],
            ['2', '4328138', 'Ny tân thành', 'Huỳnh Tấn Hiền', '171 đơn (CK 2,900/tháng)', '93 đơn/ngày (4.3%)', 'A', 'Tạm hoãn lịch livestream bán hàng', 'Hỗ trợ đẩy đơn hàng tồn Ca 1 & chiết khấu cước'],
        ]
    else: # Week 29
        t3_rows = [
            ['1', '5197975', 'TIÊN HUỲNH US', 'Nguyễn Duy Long', '755 đơn (CK 3,000/tháng)', '146 đơn/ngày (19.2%)', 'A', 'Hàng nhập về chậm so với kế hoạch', 'Hỗ trợ thủ tục bàn giao & ưu tiên lấy hàng Ca 1'],
            ['2', '4328138', 'Ny tân thành', 'Huỳnh Tấn Hiền', '339 đơn (CK 2,900/tháng)', '93 đơn/ngày (4.3%)', 'A', 'Đang mở lại lịch livestream 2 phiên/tuần', 'AM theo dõi sát các phiên live để điều phối xe lấy'],
        ]

    for i, r_data in enumerate(t3_rows):
        bg = 'F9FBFD' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(r_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in [2, 7, 8] else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(t3.cell(i+1, j), val, font_size=8.5, bg_hex=bg, align=align)

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # 3. TỔNG HỢP CÁC VẤN ĐỀ - HÀNH ĐỘNG & GIẢI PHÁP
    # -------------------------------------------------------------------------
    add_styled_paragraph(doc, '3. TỔNG HỢP CÁC VẤN ĐỀ - HÀNH ĐỘNG & GIẢI PHÁP', bold=True, font_size=13, color_rgb=(31, 56, 100), space_before=6, space_after=4)

    # 3.1
    add_styled_paragraph(doc, '3.1. Các vấn đề chính trong tuần', bold=True, font_size=11, color_rgb=(46, 117, 182), space_before=4, space_after=4)

    # Table 4
    t4 = doc.add_table(rows=4, cols=6)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    t4_headers = ['STT', 'Nhóm vấn đề chính', 'Mô tả vấn đề', 'Tác động đến KH / Doanh thu', 'Mức độ ưu tiên', 'Tiến độ xử lý']
    
    for j, h in enumerate(t4_headers):
        format_cell(t4.cell(0, j), h, bold=True, font_size=9, color_rgb=(255, 255, 255), bg_hex='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER)

    if week_num == 28:
        t4_rows = [
            ['1', 'Giữ cũ (KH A & BCD)', '2 KH A (TIÊN HUỲNH US & Ny tân thành) và 23 KH BCD sụt giảm đơn >30%', 'Rủi ro giảm ~50-80 triệu VND doanh thu/tuần tại Ninh Thuận & Bình Thuận', 'Cao', 'AM đã làm việc trực tiếp; Ny tân thành bắt đầu tăng lại đơn ở W29'],
            ['2', 'Bán mới (Quality)', '52% KH mới chỉ phát sinh đúng 1 đơn rồi ngưng lên đơn tiếp', 'Doanh thu bán mới chưa bền vững, AOV toàn Vùng chỉ đạt 4.54 đơn/KH', 'Trung bình', 'Thiết lập quy trình Onboarding Call sau 3 ngày kể từ đơn đầu tiên'],
            ['3', 'Vận hành Ca 1', 'Một số bưu cục tại Lâm Đồng & Khánh Hòa lấy hàng Ca 1 bị trễ khung giờ', 'Shop chuyển bớt đơn qua đơn vị vận chuyển đối thủ', 'Cao', 'Làm việc với Trưởng bưu cục điều phối lại ca lấy hàng trước 10h30'],
        ]
    else: # Week 29
        t4_rows = [
            ['1', 'Giữ cũ (Shop A)', 'TIÊN HUỲNH US chưa khôi phục hoàn toàn sản lượng cam kết 3,000 đơn/tháng', 'Ảnh hưởng ~15-20 triệu VND doanh thu/tuần của khu vực Ninh Thuận', 'Cao', 'AM Nguyễn Duy Long hỗ trợ tư vấn quy trình kho & giao ca 1'],
            ['2', 'Cạnh tranh giá cước', 'Đối thủ cạnh tranh tung chính sách đồng giá 14k/đơn tại Đắk Nông & Lâm Đồng', 'Một số shop nhóm BCD dao động sản lượng', 'Cao', 'Đề xuất Trưởng vùng (ARD) phê duyệt bảng giá đối ứng cho shop lớn'],
            ['3', 'Chất lượng KH Mới', 'AOV khách hàng mới Tuần 29 giảm còn 2.59 đơn/KH do nhiều shop nhỏ thử nghiệm', 'Doanh thu bán mới đạt 12.01tr VND dù số lượng KH mới tăng lên 123 KH', 'Trung bình', 'Phân loại KH mới có tiềm năng để AM tập trung chăm sóc chuyên sâu'],
        ]

    for i, r_data in enumerate(t4_rows):
        bg = 'F9FBFD' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(r_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in [1, 2, 3, 5] else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(t4.cell(i+1, j), val, font_size=8.5, bg_hex=bg, align=align)

    add_styled_paragraph(doc, '', space_after=4)

    # 3.2
    add_styled_paragraph(doc, '3.2. Giải pháp & kế hoạch tuần tới', bold=True, font_size=11, color_rgb=(46, 117, 182), space_before=4, space_after=4)

    # Table 5
    t5 = doc.add_table(rows=4, cols=7)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    t5_headers = ['STT', 'Giải pháp / Kế hoạch', 'Đối tượng', 'Mục tiêu tác động / Kết quả kỳ vọng', 'Ngày triển khai', 'PIC', 'Cần hỗ trợ']
    
    for j, h in enumerate(t5_headers):
        format_cell(t5.cell(0, j), h, bold=True, font_size=9, color_rgb=(255, 255, 255), bg_hex='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER)

    if week_num == 28:
        t5_rows = [
            ['1', 'Gặp trực tiếp & ký chính sách ưu đãi cước lưu kho/lấy hàng Ca 1', 'KH TIÊN HUỲNH US & Ny tân thành', 'Khôi phục sản lượng 2 KH A đạt trên 80% cam kết tháng', '15/07/2026', 'AM Long, AM Hiền', 'ARD & Phòng Giá'],
            ['2', 'Quy trình Onboarding Call 3-7-14 cho KH Mới', '119 KH mới phát sinh đơn Tuần 28', 'Tăng tỷ lệ KH tái phát sinh đơn tuần 2 lên >60%', '16/07/2026', 'Toàn bộ AM NTB', 'Team CSKH Vùng'],
            ['3', 'Rà soát bưu cục lấy hàng trễ Ca 1 tại Lâm Đồng', 'Bưu cục Lâm Đồng & AM địa bàn', 'Đảm bảo tỷ lệ lấy Ca 1 đúng giờ đạt >98%', '14/07/2026', 'Trưởng Bưu Cục & OPR', 'Khối Vận hành'],
        ]
    else: # Week 29
        t5_rows = [
            ['1', 'Phê duyệt bảng giá đối ứng đặc thù cho nhóm BCD bị đối thủ chèo kéo', '23 KH nhóm BCD sụt giảm đơn tại Đắk Nông & Lâm Đồng', 'Giữ chân 100% KH BCD trọng điểm, khôi phục sản lượng +15%', '22/07/2026', 'Trưởng Vùng & AM', 'Phòng Cước & CSNT'],
            ['2', 'Chiến dịch Săn Shop Nông Sản & Nước Yến du lịch hè', 'Các shop lớn tại Khánh Hòa & Lâm Đồng', 'Phát sinh 130+ KH mới, nâng AOV bán mới lên >4.0 đơn/KH', '21/07/2026', 'AM Khánh Hòa, AM Lâm Đồng', 'Team Marketing'],
            ['3', 'Chăm sóc & điều phối xe lấy hàng ưu tiên cho phiên livestream', 'Shop Ny tân thành (Bình Thuận)', 'Tăng sản lượng Ny tân thành lên >600 đơn/tuần', '20/07/2026', 'AM Huỳnh Tấn Hiền', 'Bưu cục Bình Thuận'],
        ]

    for i, r_data in enumerate(t5_rows):
        bg = 'F9FBFD' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(r_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in [1, 2, 3] else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(t5.cell(i+1, j), val, font_size=8.5, bg_hex=bg, align=align)

    doc.save(out_filename)
    print(f'Successfully generated {out_filename}')

# Output paths
out_dirs = [
    r'C:\Users\lap4all\Downloads',
    r'c:\Users\lap4all\Documents\Auto report\output',
    r'c:\Users\lap4all\Documents\Auto report'
]

for d in out_dirs:
    os.makedirs(d, exist_ok=True)
    generate_week_report(28, '06/07/2026 – 12/07/2026', 27, os.path.join(d, 'BCKD_Tuan28_2026_NTB.docx'))
    generate_week_report(29, '13/07/2026 – 19/07/2026', 28, os.path.join(d, 'BCKD_Tuan29_2026_NTB.docx'))

print('ALL REPORTS GENERATED SUCCESSFULLY!')
