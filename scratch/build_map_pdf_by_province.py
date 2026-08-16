import docx, os, sys
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

src_docx = r'C:\Users\lap4all\Downloads\Quy_Hoach MANG LUOI NTB.docx'
out_dir = r'C:\Users\lap4all\Documents\Auto report\scratch\maps_labeled'
os.makedirs(out_dir, exist_ok=True)

doc = docx.Document(src_docx)

# Extract images in document order and match with section titles
items_with_maps = []
current_title = "VÙNG NTB"
current_province = "TỈNH LÂM ĐỒNG"

for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt: continue
    
    # Track title & province
    for n in range(1, 37):
        if txt.startswith(f'{n}. '):
            current_title = txt
            if '(Tỉnh Khánh Hòa)' in txt: current_province = "TỈNH KHÁNH HÒA"
            elif '(Tỉnh Ninh Thuận)' in txt: current_province = "TỈNH NINH THUẬN"
            elif '(Tỉnh Bình Thuận)' in txt: current_province = "TỈNH BÌNH THUẬN"
            elif '(Tỉnh Đắk Nông)' in txt or '(Tỉnh Đắc Nông)' in txt: current_province = "TỈNH ĐẮK NÔNG"
            elif '(Tỉnh Lâm Đồng)' in txt: current_province = "TỈNH LÂM ĐỒNG"
            break

    # Check images inside paragraph run elements
    for run in p.runs:
        drawing_elements = run._r.xpath('.//w:drawing')
        for dr in drawing_elements:
            blip_elements = dr.xpath('.//a:blip')
            for blip in blip_elements:
                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id in doc.part.rels:
                    target_part = doc.part.rels[embed_id].target_part
                    img_bytes = target_part.blob
                    img_filename = f"img_{len(items_with_maps)+1}.png"
                    img_path = os.path.join(out_dir, img_filename)
                    with open(img_path, "wb") as f:
                        f.write(img_bytes)
                    
                    items_with_maps.append({
                        'province': current_province,
                        'title': current_title,
                        'caption': txt if 'Hình' in txt else f"Bản đồ quy hoạch {current_title}",
                        'img_path': img_path
                    })

print(f"Matched {len(items_with_maps)} map images with exact titles and provinces!")

# Group by Province
maps_by_province = {}
for m in items_with_maps:
    prov = m['province']
    if prov not in maps_by_province:
        maps_by_province[prov] = []
    maps_by_province[prov].append(m)

# Build Word Document grouped by Province
map_doc = docx.Document()
for section in map_doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Main Title
p_t = map_doc.add_paragraph()
r_t = p_t.add_run("BỘ BẢN ĐỒ HÀNH CHÍNH QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB 2026\n(PHÂN THEO 5 TỈNH)")
r_t.font.name = "Calibri"
r_t.font.size = Pt(16)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(180, 0, 0)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(14)

for prov_name, map_list in maps_by_province.items():
    # Province Banner
    p_p = map_doc.add_paragraph()
    r_p = p_p.add_run(f"❖ QUY HOẠCH {prov_name.upper()}")
    r_p.font.name = "Calibri"
    r_p.font.size = Pt(14)
    r_p.font.bold = True
    r_p.font.color.rgb = RGBColor(0, 51, 102)
    p_p.paragraph_format.space_before = Pt(16)
    p_p.paragraph_format.space_after = Pt(8)
    
    for item in map_list:
        p_item = map_doc.add_paragraph()
        r_item = p_item.add_run(f"📍 {item['title']}")
        r_item.font.name = "Calibri"
        r_item.font.size = Pt(12)
        r_item.font.bold = True
        r_item.font.color.rgb = RGBColor(0, 102, 204)
        p_item.paragraph_format.space_before = Pt(8)
        p_item.paragraph_format.space_after = Pt(4)
        
        # Add Image
        p_img = map_doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(12)
        p_img.add_run().add_picture(item['img_path'], width=Inches(7.2))

map_docx_path = r'C:\Users\lap4all\Downloads\Ban_Do_Hanh_Chinh_Theo_Tinh_NTB_2026.docx'
map_pdf_path = r'C:\Users\lap4all\Downloads\Ban_Do_Hanh_Chinh_Theo_Tinh_NTB_2026.pdf'

map_doc.save(map_docx_path)
print(f"Saved map docx by province: {map_docx_path}")

# Convert to PDF
try:
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_win = word.Documents.Open(map_docx_path)
    doc_win.SaveAs(map_pdf_path, FileFormat=17)
    doc_win.Close()
    word.Quit()
    print(f"Successfully generated PDF: {map_pdf_path}")
except Exception as e:
    print(f"Error converting to PDF: {e}")
