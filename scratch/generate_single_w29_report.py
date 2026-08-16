# -*- coding: utf-8 -*-
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_styled_paragraph(doc, text, bold=False, italic=False, font_size=10, color_rgb=(31,31,31), space_before=2, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*color_rgb)
    return p

def format_cell(cell, text, bold=False, italic=False, font_size=9.5, color_rgb=(31,31,31), bg_hex=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    if bg_hex:
        set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    cell.text = str(text)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    if len(p.runs) > 0:
        run = p.runs[0]
        run.bold = bold
        run.italic = italic
        run.font.name = 'Calibri'
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(*color_rgb)

def generate_w29_vs_w28_report(out_filenames):
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Header Title
    add_styled_paragraph(doc, 'BÁO CÁO KẾT QUẢ KINH DOANH TUẦN 29 – VÙNG NAM TRUNG BỘ', bold=True, font_size=16, color_rgb=(31, 56, 100), space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, 'Thời gian báo cáo: Tuần 29 (13/07/2026 – 19/07/2026) | Kỳ so sánh: Tuần 28 (06/07/2026 – 12/07/2026)', italic=True, font_size=10.5, color_rgb=(89, 89, 89), space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    # -------------------------------------------------------------------------
    # 1. TỔNG QUAN KẾT QUẢ KINH DOANH
    # -------------------------------------------------------------------------
    add_styled_paragraph(doc, '1. TỔNG QUAN KẾT QUẢ KINH DOANH', bold=True, font_size=13, color_rgb=(31, 56, 100), space_before=6, space_after=4)
    
    # 1.1 Mức độ hoàn thành kế hoạch
    add_styled_paragraph(doc, '1.1. Mức độ hoàn thành kế hoạch', bold=True, font_size=11, color_rgb=(46, 117, 182), space_before=4, space_after=4)

    # Table 1
    t1 = doc.add_table(rows=5, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1_headers = ['Chỉ tiêu', '% Hoàn thành target', '% Tăng/ giảm các chỉ tiêu so với Tuần 28']
    
    for j, h in enumerate(t1_headers):
        format_cell(t1.cell(0, j), h, bold=True, font_size=10, color_rgb=(255, 255, 255), bg_hex='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER)

    t1_rows = [
        ['Tổng doanh thu LTC', '94.8%', '+2.50% (1.217 tỷ vs 1.187 tỷ VNĐ GTTC)'],
        ['Tổng doanh thu GTTC', '96.2%', '+2.50% (1.217 tỷ vs 1.187 tỷ VNĐ; 309,363 đơn)'],
        ['Doanh thu giữ cũ (KH duy trì / thăng hạng)', '95.1%', '+2.41% (1.205 tỷ vs 1.174 tỷ VNĐ)'],
        ['Doanh thu bán mới (KH mới trong tháng)', '84.5%', '-7.99% (12.01 triệu vs 13.05 triệu VNĐ)'],
    ]

    for i, r_data in enumerate(t1_rows):
        bg = 'F9FBFD' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(r_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(t1.cell(i+1, j), val, font_size=9.5, bg_hex=bg, align=align)

    add_styled_paragraph(doc, '', space_after=4)

    # 1.2 Kết quả Giữ cũ
    add_styled_paragraph(doc, '1.2. Kết quả Giữ cũ', bold=True, font_size=11, color_rgb=(46, 117, 182), space_before=4, space_after=4)
    add_styled_paragraph(doc, 'Nhận định chung:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)

    nhan_dinh_giucu = (
        'Trong Tuần 29, hoạt động giữ cũ của Vùng Nam Trung Bộ ghi nhận sự phục hồi tích cực so với Tuần 28: '
        'Tổng doanh thu GTTC giữ cũ đạt 1.205 tỷ VNĐ (tăng +2.41% so với Tuần 28: 1.174 tỷ VNĐ). '
        'Đặc biệt, phân hạng KH BCD đạt bước tiến ấn tượng với sản lượng bứt phá lên 10,101 đơn (+11.3% so với Tuần 28: 9,080 đơn) '
        'và doanh thu duy trì nhóm BCD tăng mạnh +43.7% đạt 41.38 triệu VNĐ. '
        'Nhóm KH A giữ vững quy mô 10/10 KH có phát sinh đơn LTC với sản lượng 15,836 đơn (duy trì 95.1% volume W28). '
        'Danh mục cảnh báo KH sụt giảm >30% giữ ở mốc 119 KH (so với 118 KH Tuần 28), bao gồm 2 KH nhóm A trọng điểm và 23 KH nhóm BCD.'
    )

    add_styled_paragraph(doc, nhan_dinh_giucu, font_size=9.5, space_after=4)
    add_styled_paragraph(doc, 'Số liệu & Các điểm nổi bật:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)
    
    pts_giucu = [
        'Phân loại nhóm Khách hàng (Vol & DT Tuần 29 vs Tuần 28): Nhóm A đạt 10 KH (15,836 đơn, 76.67tr VNĐ log, duy trì 95.1% vol W28); Nhóm BCD đạt 66 KH (10,101 đơn, +11.3% vol vs W28; 41.38tr VNĐ log, +43.7% DT duy trì vs W28); Nhóm EF đạt 1,476 KH (12,225 đơn, +0.2% vol vs W28; 58.25tr VNĐ log); Nhóm G đạt 178 KH (551 đơn).',
        'Nhóm KH nguy cơ rời bỏ / sụt giảm >30% (Sheet 2 Tuần 29): 119 KH (Lâm Đồng: 33 KH, Bình Thuận: 29 KH, Khánh Hòa: 27 KH, Ninh Thuận: 16 KH, Đắk Nông: 13 KH). Phân hạng: 2 KH A, 23 KH BCD, 92 KH EF, 2 KH G.',
        'Chi tiết KH Nhóm A có dấu hiệu giảm đơn: (1) Shop TIÊN HUỲNH US (Ninh Thuận, AM Nguyễn Duy Long) sản lượng W29 đạt 755 đơn (so với 879 đơn W28), WTD-1 đạt 146 đơn/ngày (19.2% cam kết 3,000 đơn/tháng) do vướng tiến độ hàng nhập khẩu; (2) Shop Ny tân thành (Bình Thuận, AM Huỳnh Tấn Hiền) có tín hiệu phục hồi rõ nét, sản lượng W29 tăng gấp 2 lần đạt 339 đơn (so với 171 đơn W28) khi mở lại 2 phiên livestream/tuần.',
        'Nguyên nhân chính: (1) Tiến độ giải phóng hàng nhập khẩu của shop TIÊN HUỲNH US bị chậm 3-5 ngày; (2) Sự cạnh tranh về cước từ đối thủ tại khu vực Đắk Nông & Lâm Đồng tác động nhẹ đến nhóm BCD; (3) Việc bám sát điều phối xe lấy hàng Ca 1 tại Bình Thuận giúp shop Ny tân thành tự tin khôi phục phiên livestream.'
    ]

    for pt in pts_giucu:
        add_styled_paragraph(doc, f'• {pt}', font_size=9.5, space_after=3)

    # 1.3 Kết quả Bán mới
    add_styled_paragraph(doc, '1.3. Kết quả Bán mới', bold=True, font_size=11, color_rgb=(46, 117, 182), space_before=4, space_after=4)
    add_styled_paragraph(doc, 'Nhận định chung:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)

    nhan_dinh_banmoi = (
        'Trong Tuần 29, công tác bán mới ghi nhận sự gia tăng về số lượng khách hàng với 123 KH mới phát sinh đơn LTC đầu tiên (+4 KH so với Tuần 28: 119 KH), '
        'mang lại tổng doanh thu bán mới 12.01 triệu VNĐ và tổng sản lượng 319 đơn hàng (bình quân 2.59 đơn/KH). '
        'Khánh Hòa tiếp tục vươn lên dẫn đầu Vùng về quy mô bán mới (40 KH mới, 2.73tr VNĐ), trong khi Lâm Đồng bứt phá về doanh thu bán mới cao nhất toàn Vùng (28 KH mới, 4.16tr VNĐ, 127 đơn).'
    )

    add_styled_paragraph(doc, nhan_dinh_banmoi, font_size=9.5, space_after=4)
    add_styled_paragraph(doc, 'Số liệu & Các điểm nổi bật:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)

    pts_banmoi = [
        'Số lượng KH bán mới phát sinh trong tuần: 123 KH (+3.36% vs W28: 119 KH) | Doanh thu: 12,006,960 VNĐ | Sản lượng: 319 đơn | AOV: 2.59 đơn/KH.',
        'Khu vực / AM có kết quả bán mới xuất sắc: AM Phạm Bá Thành Công (Lâm Đồng) dẫn đầu Vùng về số lượng KH mới (17 KH mới, 0.69tr VNĐ); AM Nguyễn Duy Long (Ninh Thuận) duy trì phong độ cao với 15 KH mới (0.83tr VNĐ); AM Phan Đình Duy (Khánh Hòa) đạt 11 KH mới (0.86tr VNĐ); AM Hồng Bích Nga (Lâm Đồng) đạt doanh thu & sản lượng bán mới cao nhất Vùng (10 KH mới, 1.45tr VNĐ, 71 đơn).',
        'Khu vực / AM chưa đạt tiến độ: AM Nguyễn Hoàng Phi (1 KH), AM Nguyễn Thanh Long (1 KH), AM Lê Thanh Nhựt (1 KH).',
        'Các vấn đề đang ảnh hưởng: AOV trung bình/KH mới Tuần 29 giảm nhẹ do nhiều shop nhỏ dùng thử; cần AM tập trung tư vấn các gói cước đồng giá để kích thích shop tăng sản lượng tái lên đơn.',
        'Dự kiến tuần tới (Tuần 30): Triển khai chương trình ưu đãi cước mở rộng B2B cho 30 shop tiềm năng tại Lâm Đồng & Đắk Nông, mục tiêu đạt 130+ KH mới.'
    ]

    for pt in pts_banmoi:
        add_styled_paragraph(doc, f'• {pt}', font_size=9.5, space_after=3)

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # 2. ĐIỂM NỔI BẬT TRONG TUẦN - Khách hàng nhóm A
    # -------------------------------------------------------------------------
    add_styled_paragraph(doc, '2. ĐIỂM NỔI BẬT TRONG TUẦN - Khách hàng nhóm A', bold=True, font_size=13, color_rgb=(31, 56, 100), space_before=6, space_after=4)
    add_styled_paragraph(doc, 'Nhận định chính:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)

    nhan_dinh_kha = (
        'Trong Tuần 29, danh mục KH A giữ vững quy mô 10/10 KH có phát sinh đơn LTC trên hệ thống. '
        'Shop Vận Chuyển Online (Khánh Hòa, AM Phan Đình Duy) tiếp tục khẳng định vị thế đầu tàu tăng trưởng với sản lượng đạt 7,042 đơn/tuần (+2.86% so với Tuần 28: 6,846 đơn). '
        'Shop Ny tân thành (Bình Thuận, AM Huỳnh Tấn Hiền) có bước chuyển mình phục hồi ấn tượng (+98.2% sản lượng so với Tuần 28). '
        'Vùng đang tiếp tục bám sát 55 KH nhóm BCD có tiềm năng thăng hạng A để thúc đẩy tăng trưởng trong nửa cuối tháng 7.'
    )

    add_styled_paragraph(doc, nhan_dinh_kha, font_size=9.5, space_after=4)

    # Table 2
    t2 = doc.add_table(rows=6, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2_headers = ['Chỉ tiêu', 'Dữ liệu Tuần 29', 'So sánh tăng/ giảm với Tuần 28', 'Ghi chú (nếu có)']
    
    for j, h in enumerate(t2_headers):
        format_cell(t2.cell(0, j), h, bold=True, font_size=10, color_rgb=(255, 255, 255), bg_hex='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER)

    t2_rows = [
        ['Tổng số lượng KH nhóm A đầu tháng', '10 KH', 'Bằng (0)', 'Duy trì ổn định từ đầu tháng 7'],
        ['Số KH có lên đơn đến hiện tại', '10 KH', 'Bằng (0)', '100% KH A có phát sinh đơn LTC'],
        ['Số KH A dự kiến giảm hạng', '2 KH', 'Bằng (0)', 'Ny tân thành đang có tín hiệu phục hồi (+98.2% vol vs W28)'],
        ['Số KH A nguy cơ rời bỏ (không DT LTC)', '0 KH', 'Bằng (0)', 'Không có KH A ngừng hẳn đơn LTC'],
        ['Số KH tiềm năng lên hạng A', '55 KH', 'Bằng (0)', 'Nhóm BCD có doanh thu sát ngưỡng A tại Khánh Hòa & Lâm Đồng'],
    ]

    for i, r_data in enumerate(t2_rows):
        bg = 'F9FBFD' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(r_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in [0, 3] else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(t2.cell(i+1, j), val, font_size=9.5, bg_hex=bg, align=align)

    add_styled_paragraph(doc, '', space_after=4)

    # Chi tiết KH A nguy cơ sụt giảm / rời bỏ
    add_styled_paragraph(doc, 'CHI TIẾT KH NHÓM A NGUY CƠ SỤT GIẢM / RỜI BỎ', bold=True, font_size=11, color_rgb=(192, 0, 0), space_before=4, space_after=4)
    add_styled_paragraph(doc, 'Tổng quan nguyên nhân sụt giảm:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)
    add_styled_paragraph(doc, 'Tình hình sụt giảm tập trung ở 2 shop nhóm A: (1) Shop TIÊN HUỲNH US (Ninh Thuận) sản lượng đạt 755 đơn ở Tuần 29 (so với 879 đơn W28) do vướng tiến độ giải phóng hàng nhập khẩu; (2) Shop Ny tân thành (Bình Thuận) sản lượng đã tăng từ 171 đơn (W28) lên 339 đơn (W29) nhờ mở lại các phiên livestream nhưng vẫn dưới mốc cam kết tháng.', font_size=9.5, space_after=4)

    add_styled_paragraph(doc, 'Danh sách khách hàng nhóm A giảm đơn trọng điểm trong tuần:', bold=True, font_size=10, color_rgb=(31, 56, 100), space_after=2)

    # Table 3
    t3 = doc.add_table(rows=3, cols=9)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3_headers = ['STT', 'Client_ID', 'Tên KH', 'AM phụ trách', 'Sản lượng Tuần 29 vs Tuần 28', 'WTD-1 (% Cam kết)', 'Phân hạng', 'Lý do sụt giảm', 'Hành động / Hướng xử lý']
    
    for j, h in enumerate(t3_headers):
        format_cell(t3.cell(0, j), h, bold=True, font_size=9, color_rgb=(255, 255, 255), bg_hex='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER)

    t3_rows = [
        ['1', '5197975', 'TIÊN HUỲNH US', 'Nguyễn Duy Long', '755 đơn (vs 879 đơn W28)', '146 đơn/ngày (19.2%)', 'A', 'Hàng nhập về chậm so với kế hoạch', 'Tư vấn quy trình kho & ưu tiên lấy hàng Ca 1'],
        ['2', '4328138', 'Ny tân thành', 'Huỳnh Tấn Hiền', '339 đơn (vs 171 đơn W28, +98.2%)', '93 đơn/ngày (4.3%)', 'A', 'Đang mở lại lịch livestream 2 phiên/tuần', 'AM theo dõi sát các phiên live để điều phối xe lấy'],
    ]

    for i, r_data in enumerate(t3_rows):
        bg = 'F9FBFD' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(r_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in [2, 7, 8] else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(t3.cell(i+1, j), val, font_size=8.5, bg_hex=bg, align=align)

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # 3. TỔNG HỢP CÁC VẤN ĐỀ - HÀNH ĐỘNG & GIẢI PHÁP
    # -------------------------------------------------------------------------
    add_styled_paragraph(doc, '3. TỔNG HỢP CÁC VẤN ĐỀ - HÀNH ĐỘNG & GIẢI PHÁP', bold=True, font_size=13, color_rgb=(31, 56, 100), space_before=6, space_after=4)

    # 3.1 Các vấn đề chính trong tuần
    add_styled_paragraph(doc, '3.1. Các vấn đề chính trong tuần', bold=True, font_size=11, color_rgb=(46, 117, 182), space_before=4, space_after=4)

    # Table 4
    t4 = doc.add_table(rows=4, cols=6)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    t4_headers = ['STT', 'Nhóm vấn đề chính', 'Mô tả vấn đề', 'Tác động đến KH / Doanh thu', 'Mức độ ưu tiên', 'Tiến độ xử lý']
    
    for j, h in enumerate(t4_headers):
        format_cell(t4.cell(0, j), h, bold=True, font_size=9, color_rgb=(255, 255, 255), bg_hex='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER)

    t4_rows = [
        ['1', 'Giữ cũ (Shop A)', 'TIÊN HUỲNH US chưa khôi phục hoàn toàn sản lượng cam kết 3,000 đơn/tháng', 'Ảnh hưởng ~15-20 triệu VNĐ doanh thu/tuần tại Ninh Thuận', 'Cao', 'AM Nguyễn Duy Long hỗ trợ tư vấn quy trình kho & giao ca 1'],
        ['2', 'Cạnh tranh giá cước', 'Đối thủ tung chính sách đồng giá cước tại Đắk Nông & Lâm Đồng', 'Một số shop nhóm BCD dao động sản lượng nhẹ', 'Cao', 'Đề xuất Trưởng vùng (ARD) phê duyệt bảng giá đối ứng cho shop lớn'],
        ['3', 'Chất lượng KH Mới', 'AOV khách hàng mới Tuần 29 đạt 2.59 đơn/KH do nhiều shop nhỏ dùng thử', 'Doanh thu bán mới đạt 12.01tr VNĐ dù số lượng KH mới tăng lên 123 KH', 'Trung bình', 'Phân loại KH mới có tiềm năng để AM tập trung chăm sóc chuyên sâu'],
    ]

    for i, r_data in enumerate(t4_rows):
        bg = 'F9FBFD' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(r_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in [1, 2, 3, 5] else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(t4.cell(i+1, j), val, font_size=8.5, bg_hex=bg, align=align)

    add_styled_paragraph(doc, '', space_after=4)

    # 3.2 Giải pháp & kế hoạch tuần tới
    add_styled_paragraph(doc, '3.2. Giải pháp & kế hoạch tuần tới', bold=True, font_size=11, color_rgb=(46, 117, 182), space_before=4, space_after=4)

    # Table 5
    t5 = doc.add_table(rows=4, cols=7)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    t5_headers = ['STT', 'Giải pháp / Kế hoạch', 'Đối tượng', 'Mục tiêu tác động / Kết quả kỳ vọng', 'Ngày triển khai', 'PIC', 'Cần hỗ trợ']
    
    for j, h in enumerate(t5_headers):
        format_cell(t5.cell(0, j), h, bold=True, font_size=9, color_rgb=(255, 255, 255), bg_hex='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER)

    t5_rows = [
        ['1', 'Phê duyệt bảng giá đối ứng đặc thù cho nhóm BCD bị đối thủ chèo kéo', '23 KH nhóm BCD sụt giảm đơn tại Đắk Nông & Lâm Đồng', 'Giữ chân 100% KH BCD trọng điểm, khôi phục sản lượng +15%', '22/07/2026', 'Trưởng Vùng & AM', 'Phòng Cước & CSNT'],
        ['2', 'Chiến dịch Săn Shop Nông Sản & Đồ Khô du lịch hè', 'Các shop lớn tại Khánh Hòa & Lâm Đồng', 'Phát sinh 130+ KH mới, nâng AOV bán mới lên >4.0 đơn/KH', '21/07/2026', 'AM Khánh Hòa, AM Lâm Đồng', 'Team Marketing'],
        ['3', 'Chăm sóc & điều phối xe lấy hàng ưu tiên cho phiên livestream', 'Shop Ny tân thành (Bình Thuận)', 'Tăng sản lượng Ny tân thành lên >600 đơn/tuần', '20/07/2026', 'AM Huỳnh Tấn Hiền', 'Bưu cục Bình Thuận'],
    ]

    for i, r_data in enumerate(t5_rows):
        bg = 'F9FBFD' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(r_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in [1, 2, 3] else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(t5.cell(i+1, j), val, font_size=8.5, bg_hex=bg, align=align)

    for out_fn in out_filenames:
        try:
            doc.save(out_fn)
            print(f'Successfully saved {out_fn}')
        except Exception as e:
            print(f'Warning could not save {out_fn}: {e}')

out_files = [
    r'C:\Users\lap4all\Downloads\BCKD_Tuan29_vs_Tuan28_NTB.docx',
    r'c:\Users\lap4all\Documents\Auto report\output\BCKD_Tuan29_vs_Tuan28_NTB.docx',
    r'c:\Users\lap4all\Documents\Auto report\BCKD_Tuan29_vs_Tuan28_NTB.docx'
]

generate_w29_vs_w28_report(out_files)
print('W29 VS W28 REPORT GENERATION COMPLETE!')
