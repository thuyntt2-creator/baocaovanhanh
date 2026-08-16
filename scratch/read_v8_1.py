import docx
import sys
import os

sys.stdout.reconfigure(encoding='utf-8')

docx_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_1.docx"

if not os.path.exists(docx_path):
    print(f"File không tồn tại: {docx_path}")
    sys.exit(1)

doc = docx.Document(docx_path)

print("=== CÁC ĐOẠN VĂN BẢN TRONG FILE ===")
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[{idx}]: {text}")

print("\n=== CÁC BẢNG TRONG FILE ===")
for idx, table in enumerate(doc.tables):
    print(f"\nBảng {idx}:")
    for r_idx, row in enumerate(table.rows):
        cells_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        print(f"  Dòng {r_idx}: {cells_text}")
