import docx, os, sys, hashlib
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

src_docx = r'C:\Users\lap4all\Downloads\Quy_Hoach MANG LUOI NTB.docx'
out_dir = r'C:\Users\lap4all\Documents\Auto report\scratch\maps_final_clean'
os.makedirs(out_dir, exist_ok=True)

doc = docx.Document(src_docx)

# Extract images and map MD5 hash to unique items
unique_maps = []
seen_hashes = set()

# Manually mapped clean titles for all unique map images in docx
map_clean_meta = [
    {
        'province': 'TỈNH LÂM ĐỒNG',
        'title': 'Bản đồ Quy hoạch Cụm Bảo Lộc & Bảo Lâm (Phường 1, Phường 2, Phường B\'Lao, Xã Bảo Lâm 2)'
    },
    {
        'province': 'TỈNH BÌNH THUẬN',
        'title': 'Bản đồ Quy hoạch Cụm TP. Phan Thiết & Hàm Thắng (Phường Bình Thuận, Phường Phan Thiết, Phường Hàm Thắng)'
    },
    {
        'province': 'TỈNH KHÁNH HÒA',
        'title': 'Bản đồ Quy hoạch TP. Nha Trang (Phường Nam Nha Trang, Phường Nha Trang, Phường Tây Nha Trang)'
    },
    {
        'province': 'TỈNH NINH THUẬN',
        'title': 'Bản đồ Quy hoạch Khu vực Phường Ninh Chử & Phường Phan Rang'
    },
    {
        'province': 'TỈNH KHÁNH HÒA',
        'title': 'Bản đồ Quy hoạch Khu vực Phường Ninh Hòa, Xã Hòa Trí & Xã Tân Định'
    },
    {
        'province': 'TỈNH LÂM ĐỒNG',
        'title': 'Bản đồ Quy hoạch TP. Đà Lạt (Phường Xuân Hương - Đà Lạt & Phường Lâm Viên - Đà Lạt)'
    },
    {
        'province': 'TỈNH LÂM ĐỒNG',
        'title': 'Bản đồ Quy hoạch Khu vực Xã Di Linh (BC CK Di Linh Mới)'
    },
    {
        'province': 'TỈNH BÌNH THUẬN',
        'title': 'Bản đồ Quy hoạch Khu vực Xã Phan Rí Cửa & Tuy Phong'
    },
    {
        'province': 'TỈNH BÌNH THUẬN',
        'title': 'Bản đồ Quy hoạch Khu vực Xã Tân Thành & Cụm Nam Thành'
    },
    {
        'province': 'TỈNH KHÁNH HÒA',
        'title': 'Bản đồ Quy hoạch Khu vực Xã Vạn Thắng & Tu Bông'
    },
    {
        'province': 'TỈNH LÂM ĐỒNG',
        'title': 'Bản đồ Quy hoạch Khu vực Xã Đức Trọng (Đức Trọng 1 & Đức Trọng 2)'
    },
    {
        'province': 'TỈNH NINH THUẬN',
        'title': 'Bản đồ Quy hoạch Khu vực Xã Phước Dinh & Bưu cục Đông Hải Mới'
    },
    {
        'province': 'TỈNH KHÁNH HÒA',
        'title': 'Bản đồ Quy hoạch Cụm TP. Cam Ranh (Bưu cục Nam Cam Ranh Mới)'
    },
    {
        'province': 'TỈNH LÂM ĐỒNG',
        'title': 'Bản đồ Quy hoạch Cụm Đơn Dương (Bưu cục Lạc Xuân Mới)'
    }
]

img_index = 0
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        drawing_elements = run._r.xpath('.//w:drawing')
        for dr in drawing_elements:
            blip_elements = dr.xpath('.//a:blip')
            for blip in blip_elements:
                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id in doc.part.rels:
                    target_part = doc.part.rels[embed_id].target_part
                    img_bytes = target_part.blob
                    
                    h_val = hashlib.md5(img_bytes).hexdigest()
                    if h_val not in seen_hashes:
                        seen_hashes.add(h_val)
                        img_filename = f"clean_map_{len(unique_maps)+1}.png"
                        img_path = os.path.join(out_dir, img_filename)
                        with open(img_path, "wb") as f:
                            f.write(img_bytes)
                        
                        meta = map_clean_meta[len(unique_maps)] if len(unique_maps) < len(map_clean_meta) else {
                            'province': 'VÙNG NTB',
                            'title': f'Bản đồ quy hoạch khu vực {len(unique_maps)+1}'
                        }
                        
                        unique_maps.append({
                            'province': meta['province'],
                            'title': meta['title'],
                            'img_path': img_path,
                            'hash': h_val
                        })

print(f"Extracted {len(unique_maps)} deduplicated unique map images!")

# Group maps by Province
maps_by_province = {}
for m in unique_maps:
    prov = m['province']
    if prov not in maps_by_province:
        maps_by_province[prov] = []
    maps_by_province[prov].append(m)

# Build Clean PDF Document
map_doc = docx.Document()
for section in map_doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Main Title (NO SUBTITLE)
p_t = map_doc.add_paragraph()
r_t = p_t.add_run("BỘ BẢN ĐỒ HÀNH CHÍNH QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB 2026")
r_t.font.name = "Calibri"
r_t.font.size = Pt(16)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(180, 0, 0)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(16)

# Province order
prov_order = ["TỈNH LÂM ĐỒNG", "TỈNH KHÁNH HÒA", "TỈNH NINH THUẬN", "TỈNH BÌNH THUẬN", "TỈNH ĐẮK NÔNG"]

for prov in prov_order:
    if prov in maps_by_province:
        # Province Header
        p_p = map_doc.add_paragraph()
        r_p = p_p.add_run(f"❖ QUY HOẠCH {prov.upper()}")
        r_p.font.name = "Calibri"
        r_p.font.size = Pt(14)
        r_p.font.bold = True
        r_p.font.color.rgb = RGBColor(0, 51, 102)
        p_p.paragraph_format.space_before = Pt(16)
        p_p.paragraph_format.space_after = Pt(8)
        
        for item in maps_by_province[prov]:
            # Clean title without STT or numbers
            p_item = map_doc.add_paragraph()
            r_item = p_item.add_run(f"📍 {item['title']}")
            r_item.font.name = "Calibri"
            r_item.font.size = Pt(11.5)
            r_item.font.bold = True
            r_item.font.color.rgb = RGBColor(0, 102, 204)
            p_item.paragraph_format.space_before = Pt(8)
            p_item.paragraph_format.space_after = Pt(4)
            
            p_img = map_doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(14)
            p_img.add_run().add_picture(item['img_path'], width=Inches(7.2))

map_docx_path = r'C:\Users\lap4all\Downloads\Ban_Do_Hanh_Chinh_Quy_Hoach_NTB_2026_Chuan.docx'
map_pdf_path = r'C:\Users\lap4all\Downloads\Ban_Do_Hanh_Chinh_Quy_Hoach_NTB_2026_Chuan.pdf'

map_doc.save(map_docx_path)
print(f"Saved clean map docx: {map_docx_path}")

# Convert to PDF
try:
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_win = word.Documents.Open(map_docx_path)
    doc_win.SaveAs(map_pdf_path, FileFormat=17)
    doc_win.Close()
    word.Quit()
    print(f"Successfully generated PDF: {map_pdf_path}")
except Exception as e:
    print(f"Error converting PDF: {e}")
