import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"C:\Users\lap4all\Downloads\AOP_Hang_Nang_NTB_T7-T12_2026_v8_4.docx"

if not os.path.exists(doc_path):
    print("File không tồn tại")
    sys.exit(1)

doc = docx.Document(doc_path)

print("=== TÌM KIẾM SỐ 50 HOẶC KHO TRONG FILE WORD V8_4 ===")

# Kiểm tra các đoạn văn
for p_idx, p in enumerate(doc.paragraphs):
    if "50" in p.text:
        print(f"[ĐOẠN VĂN] Paragraph {p_idx}: {p.text}")

# Kiểm tra các bảng
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        cells_text = [cell.text.strip() for cell in row.cells]
        for c_idx, text in enumerate(cells_text):
            if text == "50" or "50 người" in text or "kho (46)" in text:
                print(f"[BẢNG] Table {t_idx+1} ({table.rows[0].cells[0].text[:30]}), Row {r_idx}, Col {c_idx}: {cells_text}")
                break
print("=== KẾT THÚC TÌM KIẾM ===")
