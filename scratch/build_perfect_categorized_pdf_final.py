import docx, os, sys
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

img_dir = r'C:\Users\lap4all\Documents\Auto report\scratch\maps_exact_flow'

# Map index mapping to operational maps
# Index 22: Xuân Hương 2
# Index 23: Di Linh CK
# Index 24: Đông Hải
# Index 25: Nam Thành
# Index 26: Nam Cam Ranh
# Index 27: Lạc Xuân Đơn Dương
# Index 5: Nam Nha Trang / Tây Nha Trang (TP Nha Trang)
# Index 2: Bảo Lộc
# Index 4: Phan Thiết
# Index 8: Ninh Hòa
# Index 9: Phan Rang
# Index 17: Phan Rí Cửa
# Index 18: Tân Thành

sec1_openings = [
    {
        'province': 'TỈNH LÂM ĐỒNG',
        'title': 'Bản đồ Quy hoạch Bưu cục CK Di Linh (Hàng cồng kềnh - Đa năng)',
        'img_path': os.path.join(img_dir, 'flow_img_23.png')
    },
    {
        'province': 'TỈNH LÂM ĐỒNG',
        'title': 'Bản đồ Quy hoạch Bưu cục Xuân Hương - Đà Lạt 2 (Mới tại Phường 10)',
        'img_path': os.path.join(img_dir, 'flow_img_22.png')
    },
    {
        'province': 'TỈNH LÂM ĐỒNG',
        'title': 'Bản đồ Quy hoạch Bưu cục Lạc Xuân (Cụm Đơn Dương)',
        'img_path': os.path.join(img_dir, 'flow_img_27.png')
    },
    {
        'province': 'TỈNH KHÁNH HÒA',
        'title': 'Bản đồ Quy hoạch Kho mới Bưu cục Nam Nha Trang 1 Mới (TP. Nha Trang)',
        'img_path': os.path.join(img_dir, 'flow_img_5.png')
    },
    {
        'province': 'TỈNH KHÁNH HÒA',
        'title': 'Bản đồ Quy hoạch Bưu cục Nam Cam Ranh (Tách mới 6 xã Nam)',
        'img_path': os.path.join(img_dir, 'flow_img_26.png')
    },
    {
        'province': 'TỈNH NINH THUẬN',
        'title': 'Bản đồ Quy hoạch Bưu cục Đông Hải (Khu vực ven biển)',
        'img_path': os.path.join(img_dir, 'flow_img_24.png')
    },
    {
        'province': 'TỈNH BÌNH THUẬN',
        'title': 'Bản đồ Quy hoạch Bưu cục Nam Thành (Khu vực Nam Thành & Nghị Đức)',
        'img_path': os.path.join(img_dir, 'flow_img_25.png')
    }
]

sec2_relocations = [
    {
        'province': 'TỈNH KHÁNH HÒA',
        'title': 'Bản đồ Di dời mặt bằng Bưu cục Tây Nha Trang (Để nhận 2 xã & Đóng cửa BC Diên Khánh 1)',
        'img_path': os.path.join(img_dir, 'flow_img_5.png')
    },
    {
        'province': 'TỈNH KHÁNH HÒA',
        'title': 'Bản đồ Mở rộng mặt bằng Bưu cục Bắc Cam Ranh (>100m²)',
        'img_path': os.path.join(img_dir, 'flow_img_26.png')
    }
]

sec3_provinces = {
    'TỈNH LÂM ĐỒNG': [
        {'title': 'Bản đồ Quy hoạch TP. Đà Lạt (Phường Xuân Hương & Phường Lâm Viên)', 'img_path': os.path.join(img_dir, 'flow_img_12.png')},
        {'title': 'Bản đồ Quy hoạch Khu vực Xã Di Linh', 'img_path': os.path.join(img_dir, 'flow_img_14.png')},
        {'title': 'Bản đồ Quy hoạch Cụm Bảo Lộc & Bảo Lâm (Phường 1, Phường 2, Phường B\'Lao, Xã Bảo Lâm 2)', 'img_path': os.path.join(img_dir, 'flow_img_2.png')},
        {'title': 'Bản đồ Quy hoạch Cụm Đơn Dương (Bưu cục Lạc Xuân Mới)', 'img_path': os.path.join(img_dir, 'flow_img_27.png')},
        {'title': 'Bản đồ Quy hoạch Khu vực Đức Trọng', 'img_path': os.path.join(img_dir, 'flow_img_21.png')}
    ],
    'TỈNH KHÁNH HÒA': [
        {'title': 'Bản đồ Quy hoạch TP. Nha Trang (Phường Nam Nha Trang, Phường Nha Trang, Phường Tây Nha Trang)', 'img_path': os.path.join(img_dir, 'flow_img_5.png')},
        {'title': 'Bản đồ Quy hoạch Cụm TP. Cam Ranh', 'img_path': os.path.join(img_dir, 'flow_img_26.png')},
        {'title': 'Bản đồ Quy hoạch Khu vực Ninh Hòa & Xã Hòa Trí, Tân Định', 'img_path': os.path.join(img_dir, 'flow_img_8.png')},
        {'title': 'Bản đồ Quy hoạch Khu vực Vạn Ninh & Tu Bông', 'img_path': os.path.join(img_dir, 'flow_img_20.png')}
    ],
    'TỈNH NINH THUẬN': [
        {'title': 'Bản đồ Quy hoạch Khu vực Ninh Chử', 'img_path': os.path.join(img_dir, 'flow_img_7.png')},
        {'title': 'Bản đồ Quy hoạch Khu vực TP. Phan Rang', 'img_path': os.path.join(img_dir, 'flow_img_9.png')},
        {'title': 'Bản đồ Quy hoạch Khu vực Xã Phước Dinh & Bưu cục Đông Hải', 'img_path': os.path.join(img_dir, 'flow_img_24.png')}
    ],
    'TỈNH BÌNH THUẬN': [
        {'title': 'Bản đồ Quy hoạch Cụm TP. Phan Thiết & Hàm Thắng (Phường Bình Thuận, Phường Phan Thiết, Phường Hàm Thắng)', 'img_path': os.path.join(img_dir, 'flow_img_4.png')},
        {'title': 'Bản đồ Quy hoạch Khu vực Phan Rí Cửa & Tuy Phong', 'img_path': os.path.join(img_dir, 'flow_img_17.png')},
        {'title': 'Bản đồ Quy hoạch Khu vực Xã Tân Thành & Cụm Nam Thành', 'img_path': os.path.join(img_dir, 'flow_img_18.png')}
    ]
}

# Build PDF Document
map_doc = docx.Document()
for section in map_doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Main Title (NO SUBTITLE AT ALL)
p_t = map_doc.add_paragraph()
r_t = p_t.add_run("BỘ BẢN ĐỒ HÀNH CHÍNH QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB 2026")
r_t.font.name = "Calibri"
r_t.font.size = Pt(16)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(180, 0, 0)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(16)

def add_sec_title(text):
    p = map_doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(13.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)

# MỤC 1: BẢN ĐỒ CÁC BƯU CỤC MỞ MỚI / TÁCH BƯU CỤC
add_sec_title("📌 MỤC 1: BẢN ĐỒ CÁC PHƯƠNG ÁN MỞ MỚI / TÁCH BƯU CỤC (07 BC)")
for item in sec1_openings:
    p_item = map_doc.add_paragraph()
    r_item = p_item.add_run(f"📍 [{item['province']}] - {item['title']}")
    r_item.font.name = "Calibri"
    r_item.font.size = Pt(11.5)
    r_item.font.bold = True
    r_item.font.color.rgb = RGBColor(0, 102, 204)
    p_item.paragraph_format.space_before = Pt(8)
    p_item.paragraph_format.space_after = Pt(4)
    
    p_img = map_doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_after = Pt(14)
    p_img.add_run().add_picture(item['img_path'], width=Inches(7.2))

# MỤC 2: BẢN ĐỒ CÁC BƯU CỤC DI DỜI / MỞ RỘNG MB KHO
add_sec_title("🚚 MỤC 2: BẢN ĐỒ CÁC PHƯƠNG ÁN DI DỜI / MỞ RỘNG MẶT BẰNG KHO")
for item in sec2_relocations:
    p_item = map_doc.add_paragraph()
    r_item = p_item.add_run(f"📍 [{item['province']}] - {item['title']}")
    r_item.font.name = "Calibri"
    r_item.font.size = Pt(11.5)
    r_item.font.bold = True
    r_item.font.color.rgb = RGBColor(180, 100, 0)
    p_item.paragraph_format.space_before = Pt(8)
    p_item.paragraph_format.space_after = Pt(4)
    
    p_img = map_doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_after = Pt(14)
    p_img.add_run().add_picture(item['img_path'], width=Inches(7.2))

# MỤC 3: BẢN ĐỒ QUY HOẠCH ĐỊA BÀN PHÂN THEO TỈNH
add_sec_title("🗺️ MỤC 3: BẢN ĐỒ QUY HOẠCH ĐỊA BÀN SÁP NHẬP PHÂN THEO TỈNH")
for prov_name, map_list in sec3_provinces.items():
    p_p = map_doc.add_paragraph()
    r_p = p_p.add_run(f"❖ QUY HOẠCH {prov_name.upper()}")
    r_p.font.name = "Calibri"
    r_p.font.size = Pt(12.5)
    r_p.font.bold = True
    r_p.font.color.rgb = RGBColor(0, 102, 0)
    p_p.paragraph_format.space_before = Pt(12)
    p_p.paragraph_format.space_after = Pt(6)
    
    for item in map_list:
        p_item = map_doc.add_paragraph()
        r_item = p_item.add_run(f"📍 {item['title']}")
        r_item.font.name = "Calibri"
        r_item.font.size = Pt(11)
        r_item.font.bold = True
        r_item.font.color.rgb = RGBColor(51, 51, 51)
        p_item.paragraph_format.space_before = Pt(6)
        p_item.paragraph_format.space_after = Pt(4)
        
        p_img = map_doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(12)
        p_img.add_run().add_picture(item['img_path'], width=Inches(7.2))

map_docx_path = r'C:\Users\lap4all\Downloads\Ban_Do_Hanh_Chinh_Quy_Hoach_NTB_2026_Final.docx'
map_pdf_path = r'C:\Users\lap4all\Downloads\Ban_Do_Hanh_Chinh_Quy_Hoach_NTB_2026_Final.pdf'

map_doc.save(map_docx_path)
print(f"Saved final map docx: {map_docx_path}")

# Convert to PDF
try:
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_win = word.Documents.Open(map_docx_path)
    doc_win.SaveAs(map_pdf_path, FileFormat=17)
    doc_win.Close()
    word.Quit()
    print(f"Successfully generated FINAL PDF: {map_pdf_path}")
except Exception as e:
    print(f"Error converting PDF: {e}")
