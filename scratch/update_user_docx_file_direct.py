import docx, openpyxl, sys
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

# 1. Open original docx
src_path = r'C:\Users\lap4all\Downloads\Quy_Hoach MANG LUOI NTB.docx'
doc = docx.Document(src_path)

# Insert explicit Nha Trang bullets into Section IV right after item 6
sec4_p_idx = -1
sec5_p_idx = -1

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if 'IV.ĐÁNH GIÁ NHU CẦU' in txt or 'IV. ĐÁNH GIÁ NHU CẦU' in txt:
        sec4_p_idx = i
    elif 'V. TỔNG HỢP BIẾN ĐỘNG' in txt or 'V.TỔNG HỢP BIẾN ĐỘNG' in txt:
        sec5_p_idx = i

print(f"Sec IV at P{sec4_p_idx}, Sec V at P{sec5_p_idx}")

# Insert Nha Trang bullets before Section V
target_p = doc.paragraphs[sec5_p_idx]

p_nha_trang_1 = target_p.insert_paragraph_before(
    "❖ 7. Di dời mặt bằng Kho trọng điểm BC Tây Nha Trang (Tỉnh Khánh Hòa): Tìm mặt bằng mới có diện tích lớn hơn cho BC Tây Nha Trang (đây là điều kiện cần để nhận lại 2 xã Vĩnh Trung & Vĩnh Thạnh, giúp chính thức ĐÓNG CỬA BC Diên Khánh 1 và gộp xã Diên An về BC Diên Khánh 2)."
)
p_nha_trang_2 = target_p.insert_paragraph_before(
    "❖ 8. Mở mới kho diện tích lớn BC Nam Nha Trang 1 Mới (Tỉnh Khánh Hòa): Mở kho lớn cover 5 xã cũ (Vĩnh Thái, Phước Long, Phước Đồng, Phước Hải, Vĩnh Trường) và ĐÓNG CỬA BC Nam Nha Trang 2 & BC Nam Nha Trang 3."
)

# Update Section V paragraph texts
for p in doc.paragraphs[sec5_p_idx:]:
    txt = p.text.strip()
    if 'Bưu cục Di dời / Mở rộng mặt bằng kho' in txt:
        p.text = "❖ Bưu cục Di dời / Mở rộng mặt bằng kho (03 BC): BC Tây Nha Trang (Khánh Hòa - di dời MB mới), BC Bắc Cam Ranh (Khánh Hòa - mở rộng từ 100m²), BC Ninh Chử (Ninh Thuận)."
    elif 'Bưu cục Đóng cửa' in txt:
        p.text = "❖ Bưu cục Đóng cửa (03 BC): BC Diên Khánh 1 (khi BC Tây Nha Trang có MB mới), BC Nam Nha Trang 2, BC Nam Nha Trang 3 (gộp tuyến/chuyển đổi địa bàn)."

# Save new updated docx
new_doc_path = r'C:\Users\lap4all\Downloads\Quy_Hoach_MANG_LUOI_NTB_Co_Nha_Trang.docx'
doc.save(new_doc_path)
print(f"Saved new updated docx with explicit Nha Trang bullets to: {new_doc_path}")

# 2. Re-build updated Word report table docx
parsed_36 = []
current = None

for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt: continue
    
    is_heading = False
    for n in range(1, 37):
        if txt.startswith(f'{n}. '):
            is_heading = True
            break
            
    if is_heading:
        if current: parsed_36.append(current)
        raw_title = txt
        title_name = raw_title.split('(')[0].strip()
        if '. ' in title_name: title_name = title_name.split('. ', 1)[1].strip()
        prov_name = 'TỈNH LÂM ĐỒNG'
        if '(Tỉnh Khánh Hòa)' in raw_title: prov_name = 'TỈNH KHÁNH HÒA'
        elif '(Tỉnh Ninh Thuận)' in raw_title: prov_name = 'TỈNH NINH THUẬN'
        elif '(Tỉnh Bình Thuận)' in raw_title: prov_name = 'TỈNH BÌNH THUẬN'
        elif '(Tỉnh Đắc Nông)' in raw_title or '(Tỉnh Đắk Nông)' in raw_title: prov_name = 'TỈNH ĐẮK NÔNG'
        
        current = {'raw_heading': raw_title, 'title': title_name, 'province': prov_name, 'communes': [], 'proposal': '', 'reason': ''}
    elif current:
        if 'ĐỀ XUẤT PHƯƠNG ÁN' in txt:
            current['proposal'] = txt.replace('❖ ĐỀ XUẤT PHƯƠNG ÁN CỦA AM:', '').replace('❖ ĐỀ XUẤT PHƯƠNG ÁN:', '').strip()
        elif 'LÝ DO VÀ GIẢI THÍCH' in txt:
            current['reason'] = txt.replace('❖ LÝ DO VÀ GIẢI THÍCH NGUYÊN NHÂN CHI TIẾT TỪ AM:', '').replace('❖ LÝ DO VÀ GIẢI THÍCH NGUYÊN NHÂN CHI TIẾT:', '').strip()
        elif current['proposal'] and not current['reason'] and not txt.startswith('❖'):
            current['proposal'] += ' ' + txt
        elif current['reason'] and not txt.startswith('❖'):
            current['reason'] += ' ' + txt
        elif not current['proposal'] and not txt.startswith('❖'):
            if ('Giao:' in txt or 'Lấy:' in txt or 'BC:' in txt or 'Phường' in txt or 'Xã' in txt or 'Thị trấn' in txt) and not txt.startswith('TỔNG SẢN LƯỢNG') and not txt.startswith('Các BC hiện') and not txt.startswith('Mã Xã') and not txt.startswith('Tỷ lệ'):
                raw_ward = txt.split('(Giao')[0].split('(Giao :')[0].replace('•', '').replace('-', '').strip()
                if raw_ward and raw_ward not in current['communes']:
                    current['communes'].append(raw_ward)

if current: parsed_36.append(current)

provinces_dict = {}
for item in parsed_36:
    p = item['province']
    if p not in provinces_dict:
        provinces_dict[p] = []
        
    c_str = ' + '.join(item['communes']) if item['communes'] else item['title']
    prop = item['proposal'] if item['proposal'] else item['reason']
    
    buucuc_cover = "BC phụ trách theo AM"
    short_note = prop
    if 'Bưu cục' in prop or 'BC' in prop or 'GIỮ NGUYÊN' in prop or 'Gộp' in prop or 'TÁCH' in prop or 'MỞ' in prop:
        parts = prop.split('.')
        buucuc_cover = parts[0]
        if len(parts) > 1:
            short_note = '. '.join(parts[1:]).strip()
            
    provinces_dict[p].append([item['title'], c_str, buucuc_cover, short_note if short_note else buucuc_cover])

# Build Word
doc_out = docx.Document()
for section in doc_out.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_province_heading(text):
    p = doc_out.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

p_t = doc_out.add_paragraph()
r_t = p_t.add_run('TỜ TRÌNH QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB NĂM 2026\n(BẢNG TỔNG HỢP ĐÃ BỔ SUNG ĐẦY ĐỦ CẢ KẾ HOẠCH DI DỜI & MỞ KHO MỚI TẠI NHA TRANG VÀO DOC)')
r_t.font.name = 'Calibri'
r_t.font.size = Pt(15)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(180, 0, 0)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(12)

headers = ["ĐVHC Mới", "Các Xã / Phường Cũ sáp nhập", "Bưu cục Cover", "Phương án & Ghi chú"]
widths = [Inches(1.8), Inches(2.6), Inches(2.2), Inches(2.2)]

for p_name, p_rows in provinces_dict.items():
    add_province_heading(f"❖ QUY HOẠCH {p_name}")
    
    table = doc_out.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], 'F7A059')
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
            
    for r_data in p_rows:
        row_cells = table.add_row().cells
        bg_color = 'FFFFFF'
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            row_cells[c_idx].width = widths[c_idx]
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=80, right=80)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                if c_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                elif c_idx == 2:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 51, 153)

word_report_path = r'C:\Users\lap4all\Downloads\Thu_trinh_Quy_Hoach_Mang_Luoi_NTB_2026_Co_Nha_Trang.docx'
doc_out.save(word_report_path)
print(f"Saved updated Word report table to: {word_report_path}")
