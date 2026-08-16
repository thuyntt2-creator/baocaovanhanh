import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

docx_path = r"C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_ MANG LUOI NTB.docx"
doc = docx.Document(docx_path)

recording = False
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if "IV. YÊU CẦU 3" in t:
        recording = True
    elif "V. TỔNG HỢP" in t:
        recording = False
    
    if recording:
        print(f"P[{i}]: {t}")

for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i+1} ({len(table.rows)} rows, {len(table.columns)} cols) ---")
    for r in table.rows[:5]:
        print([c.text.strip().replace('\n', ' ') for c in r.cells])
