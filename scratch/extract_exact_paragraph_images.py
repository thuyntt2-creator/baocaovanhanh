import docx, os, sys

sys.stdout.reconfigure(encoding='utf-8')

src_docx = r'C:\Users\lap4all\Downloads\Quy_Hoach MANG LUOI NTB.docx'
out_dir = r'C:\Users\lap4all\Documents\Auto report\scratch\maps_exact_flow'
os.makedirs(out_dir, exist_ok=True)

doc = docx.Document(src_docx)

# Walk through paragraphs in exact document order
doc_images = []

current_heading = "VÙNG NTB"
current_province = "TỈNH LÂM ĐỒNG"

for p_idx, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    
    # Track heading & province
    for n in range(1, 37):
        if txt.startswith(f'{n}. '):
            current_heading = txt
            if '(Tỉnh Khánh Hòa)' in txt: current_province = "TỈNH KHÁNH HÒA"
            elif '(Tỉnh Ninh Thuận)' in txt: current_province = "TỈNH NINH THUẬN"
            elif '(Tỉnh Bình Thuận)' in txt: current_province = "TỈNH BÌNH THUẬN"
            elif '(Tỉnh Đắk Nông)' in txt or '(Tỉnh Đắc Nông)' in txt: current_province = "TỈNH ĐẮK NÔNG"
            elif '(Tỉnh Lâm Đồng)' in txt: current_province = "TỈNH LÂM ĐỒNG"
            break

    # Extract image directly from paragraph XML
    for run in p.runs:
        drawings = run._r.xpath('.//w:drawing')
        for dr in drawings:
            blips = dr.xpath('.//a:blip')
            for blip in blips:
                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId in doc.part.rels:
                    image_part = doc.part.rels[rId].target_part
                    img_bytes = image_part.blob
                    
                    img_name = f"flow_img_{len(doc_images)+1}.png"
                    img_path = os.path.join(out_dir, img_name)
                    with open(img_path, "wb") as f:
                        f.write(img_bytes)
                    
                    # Find caption text nearby (current p, or p before, or p after)
                    caption = txt
                    if not caption and p_idx > 0:
                        caption = doc.paragraphs[p_idx-1].text.strip()
                    if not caption and p_idx < len(doc.paragraphs)-1:
                        caption = doc.paragraphs[p_idx+1].text.strip()
                        
                    doc_images.append({
                        'index': len(doc_images)+1,
                        'p_index': p_idx+1,
                        'heading': current_heading,
                        'province': current_province,
                        'caption': caption,
                        'img_path': img_path
                    })

print(f"Extracted {len(doc_images)} images in EXACT DOCUMENT FLOW ORDER:\n")
for item in doc_images:
    print(f"Index {item['index']} (at P{item['p_index']}): [{item['province']}] {item['heading']} => Caption: '{item['caption']}'")
