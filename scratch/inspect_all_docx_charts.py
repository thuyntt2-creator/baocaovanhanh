import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

docx_path = r'C:\Users\lap4all\Downloads\NTB - Báo Cáo Tuần (HRBP - ARD) - Tuần 32 - Final.docx'
doc = docx.Document(docx_path)

out_dir = r'C:\Users\lap4all\Documents\Auto report\scratch\w32_docx_images'

image_map = {}
for rel in doc.part.rels.values():
    if 'image' in rel.target_ref:
        image_map[rel.rId] = os.path.basename(rel.target_ref)

print("=== DOCX GRAPHIC LOCATIONS ===")
for i, p in enumerate(doc.paragraphs):
    xml = p._element.xml
    if 'graphicData' in xml:
        import re
        embeds = re.findall(r'r:embed="([^"]+)"', xml)
        imgs = [image_map.get(rid, rid) for rid in embeds]
        # find nearest heading above
        heading = ""
        for j in range(i, -1, -1):
            if doc.paragraphs[j].text.strip():
                heading = doc.paragraphs[j].text.strip()
                break
        print(f"P{i:3d} | Heading: '{heading}' | Images: {imgs}")
