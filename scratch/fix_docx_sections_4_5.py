import docx
from docx.shared import Pt, RGBColor
import os

source_path = r"C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_ MANG LUOI NTB.docx"
target_path = r"C:\Users\lap4all\Downloads\Bao_Cao_Quy_Hoach_ MANG LUOI NTB_FIXED.docx"

doc = docx.Document(source_path)

# Find paragraphs where IV and V start
iv_idx = None
v_idx = None

for i, p in enumerate(doc.paragraphs):
    if "IV. YÊU CẦU 3" in p.text or "IV. YÊU CẦU" in p.text:
        iv_idx = i
    elif "V. TỔNG HỢP" in p.text or "V. TỔNG HỢP BIẾN ĐỘNG" in p.text:
        v_idx = i

print(f"Found IV at paragraph index {iv_idx}, V at {v_idx}")

# Delete paragraphs from IV onwards
if iv_idx is not None:
    for p in doc.paragraphs[iv_idx:]:
        p._element.getparent().remove(p._element)

# Add IV heading
p_iv = doc.add_paragraph()
r_iv = p_iv.add_run("IV. YÊU CẦU 3: ĐÁNH GIÁ NHU CẦU TÁCH BƯU CỤC VÀ TỐI ƯU TUYẾN CHÉO")
r_iv.bold = True
r_iv.font.size = Pt(13)
r_iv.font.color.rgb = RGBColor(0, 51, 102)

# 1. Open new / Split hubs
p = doc.add_paragraph()
r = p.add_run("1. Kế hoạch Mở mới / Tách Bưu cục & Tối ưu Mạng lưới (05 Bưu cục)")
r.bold = True
r.font.size = Pt(11)

items_iv1 = [
    ("1. Mở mới Bưu cục (LDO) Xuân Hương - Đà Lạt 2: ", "Đặt tại Phường 10 (TP. Đà Lạt), chia tải cho BC Xuân Hương cũ (phụ trách Phường 3 & Phường 10, Vol giao: 1,050 đơn/ngày, Vol lấy: 150 đơn/ngày, định biên 8 NVPTTT + 1 NVXL)."),
    ("2. Tách Bưu cục Hàng Nhỏ / Hàng Vừa Di Linh: ", "Phụ trách Xã Đinh Trang Thượng, Di Linh, Phúc Thọ Lâm Hà, Liên Đầm nhằm giảm bán kính di chuyển 22–45km cho BC Di Linh."),
    ("3. Mở mới Bưu cục Đông Hải (Tỉnh Ninh Thuận): ", "Cover khu vực ven biển Phường Đông Hải, tối ưu điểm tập kết sản lượng lấy (Vol giao: 600 đơn/ngày, Vol lấy: 250 đơn/ngày, định biên 7 NVPTTT + 1 NVXL)."),
    ("4. Mở mới Bưu cục (LDO) B'Lao Mới (Bảo Lộc): ", "Cover Phường 1 & Phường B'Lao (Vol giao: 1,500 đơn/ngày, Vol lấy: 1,000 đơn/ngày, định biên 15 NVPTTT + 2 NVXL), đóng cửa BC (LDO) 1 Bảo Lộc cũ."),
    ("5. Mở mới Bưu cục (BTH) Nam Thành (Tỉnh Bình Thuận): ", "Cover khu vực Nam Thành (250 đơn/ngày) & Nghị Đức (200 đơn/ngày), định biên 7 NVPTTT + 1 NVXL.")
]

for title, desc in items_iv1:
    p = doc.add_paragraph()
    r1 = p.add_run("❖ " + title)
    r1.bold = True
    r2 = p.add_run(desc)

# 2. Overloaded Hubs Table
p = doc.add_paragraph()
r = p.add_run("2. Danh sách 13 Bưu cục Quá tải áp lực m² kho (em2 > 20) cần Tách / Mở rộng")
r.bold = True
r.font.size = Pt(11)

table1_data = [
    ["STT", "ID BC", "Tên Bưu cục", "Tỉnh", "Quản lý AM", "Địa chỉ Bưu cục", "Áp lực em2", "Đề xuất giải pháp"],
    ["1", "1896", "(LDO) Đạ Huoai", "Lâm Đồng", "Nguyễn Lê Nguyên Vũ", "Ma Đa Guôi, Đạ Huoai, Lâm Đồng", "59.2", "Tách bớt phường / Mở rộng m²"],
    ["2", "21456000", "(BTH) Đức Linh", "Bình Thuận", "Nguyễn Ngọc Khánh", "38 DT766, xã Nam Chính, Đức Linh", "33.2", "Tách bớt phường / Mở rộng m²"],
    ["3", "21320000", "(BTH) Phú Thủy", "Bình Thuận", "Nguyễn Ngọc Khánh", "188 Trương Hán Siêu, Phan Thiết", "32.0", "Tách bớt phường / Mở rộng m²"],
    ["4", "22452000", "(BTH) Phú Quý", "Bình Thuận", "Nguyễn Ngọc Khánh", "454 Võ Văn Kiệt, Phú Quý", "28.9", "Tách bớt phường / Mở rộng m²"],
    ["5", "20543000", "(KHO) Ninh Hòa 2", "Khánh Hòa", "Phạm Bá Thành Công", "06 đường 2/4 Ninh Hiệp", "28.3", "Tách bớt phường / Mở rộng m²"],
    ["6", "2399", "(KHO) Bắc Nha Trang", "Khánh Hòa", "Phạm Bá Thành Công", "195 đường 2/4, Vĩnh Hòa, Nha Trang", "25.5", "Tách bớt phường / Mở rộng m²"],
    ["7", "20144000", "(BTH) Mũi Né", "Bình Thuận", "Nguyễn Ngọc Khánh", "16 Huỳnh Thúc Kháng, Mũi Né", "24.4", "Tách bớt phường / Mở rộng m²"],
    ["8", "21403000", "(BTH) Hàm Liêm", "Bình Thuận", "Nguyễn Ngọc Khánh", "Thôn 2, xã Hàm Liêm", "24.3", "Tách bớt phường / Mở rộng m²"],
    ["9", "2357", "(BTH) Đồng Kho", "Bình Thuận", "Nguyễn Ngọc Khánh", "158 QL55, Đức Bình, Tánh Linh", "24.2", "Tách bớt phường / Mở rộng m²"],
    ["10", "22861000", "(NTH) Phan Rang", "Ninh Thuận", "Nguyễn Duy Long", "95 Thống Nhất, Phan Rang", "23.8", "Tách bớt phường / Mở rộng m²"],
    ["11", "20648000", "(NTH) Ninh Chử", "Ninh Thuận", "Nguyễn Duy Long", "51 Phạm Ngọc Thạch, Ninh Hải", "22.1", "Tách bớt phường / Mở rộng m²"],
    ["12", "20785000", "(LDO) B'Lao", "Lâm Đồng", "Hồng Bích Nga", "95a Đội Cấn, Lộc Sơn, Bảo Lộc", "22.1", "Tách bớt phường / Mở rộng m²"],
    ["13", "20150000", "(NTH) Bảo An", "Ninh Thuận", "Nguyễn Duy Long", "254A Đường 21/8, Phan Rang", "21.7", "Tách bớt phường / Mở rộng m²"]
]

t1 = doc.add_table(rows=len(table1_data), cols=8)
for r_idx, row in enumerate(table1_data):
    for c_idx, val in enumerate(row):
        cell = t1.cell(r_idx, c_idx)
        cell.text = val
        if r_idx == 0:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

doc.add_paragraph()

# 3. Cross routes reassign table
p = doc.add_paragraph()
r = p.add_run("3. Danh sách 18 Tuyến giao chéo xa ranh giới cần Reassign ngay")
r.bold = True
r.font.size = Pt(11)

table2_data = [
    ["STT", "Tên Xã/Phường cũ", "Tỉnh", "BC hiện tại", "KC cũ", "BC tối ưu", "AM tiếp nhận", "KC mới", "Tiết kiệm"],
    ["1", "Xã Đa Mi", "Bình Thuận", "(BTH) Hàm Thuận", "40.4 km", "(LDO) Bảo Lâm 3", "Hồng Bích Nga", "15.1 km", "-25.3 km"],
    ["2", "Xã Lộc Bảo", "Lâm Đồng", "(LDO) Đạ Tẻh", "40.4 km", "(DNO) Quảng Khê", "Trần Thị Nhung", "15.6 km", "-24.8 km"],
    ["3", "Xã Đạ K' Nàng", "Lâm Đồng", "(LDO) Đam Rông 3", "31.0 km", "(LDO) Đinh Văn Lâm Hà", "Huỳnh Thị Kim Chi", "17.0 km", "-14.0 km"],
    ["4", "Xã Quảng Phú", "Đắk Nông", "(DNO) Krông Nô", "28.9 km", "(LDO) Đam Rông 3", "Huỳnh Thị Kim Chi", "15.5 km", "-13.4 km"],
    ["5", "Xã Ea Pô", "Đắk Nông", "(DNO) Cư Jút", "20.7 km", "(DLA) Ea Wer", "Ban Vận Hành", "8.8 km", "-11.9 km"],
    ["6", "Xã Thắng Hải", "Bình Thuận", "(BTH) Phước Hội", "20.5 km", "(BVT) Hòa Hiệp", "Ban Vận Hành", "7.7 km", "-12.8 km"],
    ["7", "Xã Nhân Cơ", "Đắk Nông", "(DNO) Nhân Cơ", "6.6 km", "(DNO) ĐL Nam Gia Nghĩa 2", "Trần Văn Phước", "3.2 km", "-3.4 km"],
    ["8", "Xã Phước Dinh", "Ninh Thuận", "(NTH) Phước Dinh", "15.4 km", "(NTH) Thuận Nam", "Nguyễn Duy Long", "7.7 km", "-7.7 km"],
    ["9", "Xã Đắk Gằn", "Đắk Nông", "(DNO) Đức Lập", "21.5 km", "(DNO) Krông Nô", "Trần Thị Nhung", "10.7 km", "-10.8 km"],
    ["10", "Xã Bắc Ruộng", "Bình Thuận", "(BTH) Đồng Kho", "26.4 km", "(LDO) Bảo Lâm 3", "Hồng Bích Nga", "13.1 km", "-13.3 km"],
    ["11", "Xã Công Hải", "Ninh Thuận", "(NTH) Phan Rang", "26.6 km", "(KHO) Cam Linh", "Thái Thị Thanh Thư", "15.4 km", "-11.2 km"],
    ["12", "Xã Bình Lộc", "Khánh Hòa", "(KHO) Diên Khánh 2", "8.3 km", "(KHO) Diên Thọ", "Nguyễn Hoàng Phi", "3.1 km", "-5.2 km"],
    ["13", "Phường Nghĩa Phú", "Đắk Nông", "(DNO) Bắc Gia Nghĩa", "5.2 km", "(DNO) ĐL Nam Gia Nghĩa 2", "Trần Văn Phước", "1.7 km", "-3.5 km"],
    ["14", "Xã Suối Tiên", "Khánh Hòa", "(KHO) Diên Khánh 2", "11.2 km", "(KHO) Diên Thọ", "Nguyễn Hoàng Phi", "6.5 km", "-4.7 km"],
    ["15", "Xã Đạ Tông", "Lâm Đồng", "(LDO) Đam Rông 3", "23.4 km", "(LDO) Lang Biang - Đà Lạt 1", "Lê Minh Đại", "13.5 km", "-9.9 km"],
    ["16", "Xã Phan Sơn", "Bình Thuận", "(BTH) Bắc Bình", "35.6 km", "(LDO) Ninh Gia", "Lê Văn Trường", "17.9 km", "-17.7 km"],
    ["17", "Xã Lộc Lâm", "Lâm Đồng", "(LDO) Bảo Lâm 1", "23.0 km", "(DNO) Quảng Khê", "Trần Thị Nhung", "10.8 km", "-12.2 km"],
    ["18", "Xã Mỹ Thạnh", "Bình Thuận", "(BTH) Tuyên Quang", "28.3 km", "(BTH) Đồng Kho", "Nguyễn Ngọc Khánh", "12.9 km", "-15.4 km"]
]

t2 = doc.add_table(rows=len(table2_data), cols=9)
for r_idx, row in enumerate(table2_data):
    for c_idx, val in enumerate(row):
        cell = t2.cell(r_idx, c_idx)
        cell.text = val
        if r_idx == 0:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

doc.add_paragraph()

# SECTION V
p_v = doc.add_paragraph()
r_v = p_v.add_run("V. TỔNG HỢP BIẾN ĐỘNG MẠNG LƯỚI BƯU CỤC NTB 2026")
r_v.bold = True
r_v.font.size = Pt(13)
r_v.font.color.rgb = RGBColor(0, 51, 102)

summary_items = [
    ("Bưu cục Mở mới (05 BC): ", "BC Xuân Hương - Đà Lạt 2, BC Di Linh Hàng Vừa, BC Đông Hải, BC B'Lao Mới, BC Nam Thành."),
    ("Bưu cục Đóng cửa (02 BC): ", "BC 1 Bảo Lộc, BC Nam Nha Trang 3 (gộp tuyến/chuyển đổi địa bàn)."),
    ("Bưu cục Gộp tuyến/Tối ưu phân vùng (21 BC): ", "Quy hoạch dồn sản lượng về 01 Bưu cục phụ trách chính tại các ĐVHC sáp nhập."),
    ("Bưu cục Giữ nguyên vận hành chia tuyến (13 BC): ", "Duy trì 2-3 bưu cục song song theo đề xuất AM do địa hình đặc thù, đồi dốc hoặc chia cắt địa lý (Đam Rông, Đức Trọng, Đắk Sắk, Đức An, Phan Rí Cửa, Ninh Hòa...)."),
    ("Tuyến điều chỉnh ranh giới (18 Tuyến giao chéo): ", "Reassign sang Bưu cục lân cận để tối ưu bán kính di chuyển toàn vùng NTB.")
]

for title, desc in summary_items:
    p = doc.add_paragraph()
    r1 = p.add_run("❖ " + title)
    r1.bold = True
    r2 = p.add_run(desc)

doc.save(target_path)
print(f"Successfully saved corrected DOCX to: {target_path}")

doc.save(source_path)
print(f"Successfully updated original DOCX at: {source_path}")
