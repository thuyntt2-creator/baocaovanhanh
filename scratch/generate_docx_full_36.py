import docx, sys
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

# Read original docx
doc_src = docx.Document(r'C:\Users\lap4all\Downloads\Quy_Hoach MANG LUOI NTB.docx')

parsed_36 = []
current = None

for p in doc_src.paragraphs:
    txt = p.text.strip()
    if not txt: continue
    
    is_heading = False
    for n in range(1, 37):
        if txt.startswith(f'{n}. '):
            is_heading = True
            break
            
    if is_heading:
        if current: parsed_36.append(current)
        current = {'title': txt, 'communes': '', 'proposal': '', 'reason': ''}
    elif current:
        if 'Các xã cũ sáp nhập' in txt:
            current['communes'] = txt.replace('❖ Các xã cũ sáp nhập & Sản lượng Giao/Lấy từng xã:', '').strip()
        elif 'ĐỀ XUẤT PHƯƠNG ÁN' in txt:
            current['proposal'] = txt.replace('❖ ĐỀ XUẤT PHƯƠNG ÁN CỦA AM:', '').replace('❖ ĐỀ XUẤT PHƯƠNG ÁN:', '').strip()
        elif 'LÝ DO VÀ GIẢI THÍCH' in txt:
            current['reason'] = txt.replace('❖ LÝ DO VÀ GIẢI THÍCH NGUYÊN NHÂN CHI TIẾT TỪ AM:', '').replace('❖ LÝ DO VÀ GIẢI THÍCH NGUYÊN NHÂN CHI TIẾT:', '').strip()
        elif current['communes'] and not current['proposal'] and not txt.startswith('❖'):
            current['communes'] += ' ' + txt
        elif current['proposal'] and not current['reason'] and not txt.startswith('❖'):
            current['proposal'] += ' ' + txt
        elif current['reason'] and not txt.startswith('❖'):
            current['reason'] += ' ' + txt

if current: parsed_36.append(current)

doc_out = docx.Document()

for section in doc_out.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Title
p_t = doc_out.add_paragraph()
r_t = p_t.add_run('TỜ TRÌNH QUY HOẠCH MẠNG LƯỚI BƯU CỤC VÙNG NTB NĂM 2026\n(BẢNG TỔNG HỢP FULL 36 ĐƠN VỊ HÀNH CHÍNH MỚI - CHÍNH XÁC 100% THEO FILE DOCX GỐC)')
r_t.font.name = 'Calibri'
r_t.font.size = Pt(15)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(180, 0, 0)
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t.paragraph_format.space_after = Pt(12)

table = doc_out.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

headers = ["STT & Tên ĐVHC Mới", "Các Xã / Phường Cũ sáp nhập & Sản lượng Giao/Lấy", "Bưu cục Cover (Đề xuất AM)", "Phương án Quy hoạch & Lý do chi tiết từ AM"]
widths = [Inches(1.8), Inches(2.6), Inches(1.8), Inches(2.3)]

hdr_cells = table.rows[0].cells
for i, h_text in enumerate(headers):
    hdr_cells[i].text = h_text
    hdr_cells[i].width = widths[i]
    set_cell_background(hdr_cells[i], 'F7A059') # Orange fill
    set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0, 0, 0)

for r_idx, item in enumerate(parsed_36):
    row_cells = table.add_row().cells
    bg_color = 'FFFFFF'
    
    row_cells[0].text = item['title']
    row_cells[0].width = widths[0]
    
    row_cells[1].text = item['communes']
    row_cells[1].width = widths[1]
    
    row_cells[2].text = item['proposal']
    row_cells[2].width = widths[2]
    
    row_cells[3].text = item['reason'] if item['reason'] else item['proposal']
    row_cells[3].width = widths[3]
    
    for c_idx in range(4):
        set_cell_background(row_cells[c_idx], bg_color)
        set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=80, right=80)
        p = row_cells[c_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            if c_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif c_idx == 2:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 153)

out_docx = r'C:\Users\lap4all\Downloads\Thu_trinh_Quy_Hoach_Mang_Luoi_NTB_2026_Chinh_Xac_Docx_Full.docx'
doc_out.save(out_docx)
print(f'Successfully generated docx full 36 items: {out_docx}')
